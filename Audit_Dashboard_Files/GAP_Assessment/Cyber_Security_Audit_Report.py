from flask import Blueprint, request, send_file
import os
import zipfile
import tempfile
import shutil
from datetime import datetime
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# Create blueprint for Cyber Security Audit Report
cyber_security_audit_report_bp = Blueprint('cyber_security_audit_report', __name__)

def get_ordinal_superscript_cyber(day):
    """Convert day number to superscript ordinal (1ˢᵗ, 2ⁿᵈ, 3ʳᵈ, 4ᵗʰ, etc.)"""
    superscripts = {
        'st': 'ˢᵗ',
        'nd': 'ⁿᵈ',
        'rd': 'ʳᵈ',
        'th': 'ᵗʰ'
    }
    
    if 10 <= day % 100 <= 20:
        suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    
    return f"{day:02d}{superscripts[suffix]}"

def format_date_with_superscript_cyber(date_str):
    """Format date as '21ˢᵗ October 2025'"""
    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        day = date_obj.day
        month = date_obj.strftime('%B')
        year = date_obj.year
        
        day_with_ordinal = get_ordinal_superscript_cyber(day)
        return f"{day_with_ordinal} {month} {year}"
    except:
        return date_str

def get_financial_year_cyber(date_str):
    """Get financial year for a given date (1st April to 31st March)"""
    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        
        if date_obj.month >= 4:
            fy_start_year = date_obj.year
            fy_end_year = date_obj.year + 1
        else:
            fy_start_year = date_obj.year - 1
            fy_end_year = date_obj.year
        
        start_formatted = f"{get_ordinal_superscript_cyber(1)} April {fy_start_year}"
        end_formatted = f"{get_ordinal_superscript_cyber(31)} March {fy_end_year}"
        
        return f"{start_formatted} to {end_formatted}"
    except:
        return ""

def _replace_text_in_runs_cyber(paragraph, old_text, new_text):
    """Helper function to replace text in a paragraph while preserving formatting."""
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return False

    run_formats = []
    text_positions = []
    
    start_pos = 0
    while True:
        pos = full_text.find(old_text, start_pos)
        if pos == -1:
            break
        text_positions.append((pos, pos + len(old_text)))
        start_pos = pos + 1
    
    if not text_positions:
        return False
    
    current_pos = 0
    for run in paragraph.runs:
        run_start = current_pos
        run_end = current_pos + len(run.text)
        
        for text_start, text_end in text_positions:
            if run_start < text_end and run_end > text_start:
                font_color = None
                try:
                    if run.font.color.rgb is not None:
                        font_color = run.font.color.rgb
                except:
                    font_color = None
                
                run_formats.append({
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name if run.font.name else 'Times New Roman',
                    'font_size': run.font.size,
                    'font_color': font_color
                })
                break
        
        current_pos = run_end
    
    if run_formats:
        for run in paragraph.runs:
            run.text = ""
        
        new_full_text = full_text.replace(old_text, new_text)
        
        if len(paragraph.runs) > 0:
            first_run = paragraph.runs[0]
            first_run.text = new_full_text
            
            original_format = run_formats[0]
            first_run.bold = original_format.get('bold')
            first_run.italic = original_format.get('italic')
            first_run.underline = original_format.get('underline')
            
            if original_format.get('font_name'):
                first_run.font.name = original_format.get('font_name')
            else:
                first_run.font.name = 'Times New Roman'
            
            if original_format.get('font_size'):
                first_run.font.size = original_format.get('font_size')
            else:
                first_run.font.size = Pt(12)
            
            if original_format.get('font_color') is not None:
                try:
                    first_run.font.color.rgb = original_format.get('font_color')
                except:
                    pass
            
            runs_to_remove = list(paragraph.runs[1:])
            for run in runs_to_remove:
                run._element.getparent().remove(run._element)
    else:
        for run in paragraph.runs:
            run.text = ""
        
        new_full_text = full_text.replace(old_text, new_text)
        
        if len(paragraph.runs) > 0:
            first_run = paragraph.runs[0]
            first_run.text = new_full_text
            first_run.font.name = 'Times New Roman'
            first_run.font.size = Pt(12)
    
    return True

def replace_text_in_document_cyber(doc, replacements):
    """Replace text in document while preserving formatting"""
    try:
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    _replace_text_in_runs_cyber(paragraph, old_text, new_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            if old_text in paragraph.text:
                                _replace_text_in_runs_cyber(paragraph, old_text, new_text)
        
        # Replace in headers
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        _replace_text_in_runs_cyber(paragraph, old_text, new_text)
        
        # Replace in footers
        for section in doc.sections:
            for paragraph in section.footer.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        _replace_text_in_runs_cyber(paragraph, old_text, new_text)
        
        return True
    except Exception as e:
        print(f"❌ Error replacing text: {e}")
        return False

def insert_excel_table_in_document_cyber(doc, placeholder, ws, start_col, end_col, start_row, end_row):
    """Insert Excel table into Word document with formatting"""
    try:
        print(f"\n  🔍 Searching for placeholder '{placeholder}'...")
        
        # Find placeholder paragraph
        placeholder_para = None
        
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                placeholder_para = paragraph
                print(f"    ✅ Found placeholder in paragraph!")
                break
        
        if not placeholder_para:
            print(f"  ❌ Placeholder '{placeholder}' not found in document")
            return False
        
        # Calculate dimensions
        num_rows = end_row - start_row + 1
        num_cols = end_col - start_col + 1
        
        print(f"  📏 Creating table: {num_rows} rows x {num_cols} columns")
        
        # Get parent and placeholder element
        parent = placeholder_para._element.getparent()
        placeholder_element = placeholder_para._element
        
        # Create table
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # Track merged cells to skip
        merged_cells_to_skip = set()
        
        # Handle merged cells from Excel
        print(f"  🔗 Processing merged cells...")
        merged_count = 0
        for merged_range in ws.merged_cells.ranges:
            min_row_excel = merged_range.min_row
            max_row_excel = merged_range.max_row
            min_col_excel = merged_range.min_col
            max_col_excel = merged_range.max_col
            
            # Check if merge is within our range
            if (start_row <= min_row_excel <= end_row and 
                start_col <= min_col_excel <= end_col):
                
                # Convert to table indices (0-based)
                start_row_idx = min_row_excel - start_row
                end_row_idx = max_row_excel - start_row
                start_col_idx = min_col_excel - start_col
                end_col_idx = max_col_excel - start_col
                
                # Ensure indices are within bounds
                if (0 <= start_row_idx < num_rows and 0 <= end_row_idx < num_rows and
                    0 <= start_col_idx < num_cols and 0 <= end_col_idx < num_cols):
                    
                    try:
                        start_word_cell = table.cell(start_row_idx, start_col_idx)
                        end_word_cell = table.cell(end_row_idx, end_col_idx)
                        start_word_cell.merge(end_word_cell)
                        merged_count += 1
                        
                        # Mark cells to skip
                        for r in range(start_row_idx, end_row_idx + 1):
                            for c in range(start_col_idx, end_col_idx + 1):
                                if (r, c) != (start_row_idx, start_col_idx):
                                    merged_cells_to_skip.add((r, c))
                        
                        print(f"    ✅ Merged cells: Row {start_row_idx}-{end_row_idx}, Col {start_col_idx}-{end_col_idx}")
                    except Exception as e:
                        print(f"    ⚠️ Error merging cells: {e}")
        
        print(f"  ✅ Processed {merged_count} merged cell ranges")
        
        # First pass: identify columns that should be center-aligned
        print(f"  🔍 Identifying column alignments from header row...")
        center_aligned_columns = set()
        left_aligned_columns = set()
        
        # Check first row (header row) to determine column alignment
        for col_idx in range(start_col, end_col + 1):
            table_col = col_idx - start_col
            excel_cell = ws.cell(row=start_row, column=col_idx)
            header_value = "" if excel_cell.value is None else str(excel_cell.value).strip()
            
            # Columns that should be center-aligned throughout
            if header_value in ["Sr. No.", "Complied Status (Fully Complied/Partially Complied/Not Complied/Not Applicable)", 
                               "POC Attached"]:
                center_aligned_columns.add(table_col)
                print(f"    Column {table_col} ('{header_value}'): CENTER aligned")
            
            # Columns that should be left-aligned throughout
            elif header_value in ["Requirements", "Auditor's Remark"]:
                left_aligned_columns.add(table_col)
                print(f"    Column {table_col} ('{header_value}'): LEFT aligned")
        
        print(f"  ✅ Center columns: {center_aligned_columns}")
        print(f"  ✅ Left columns: {left_aligned_columns}")
        
        print(f"  📊 Filling table with data...")
        
        # Fill table with data
        rows_filled = 0
        for row_idx in range(start_row, end_row + 1):
            for col_idx in range(start_col, end_col + 1):
                table_row = row_idx - start_row
                table_col = col_idx - start_col
                
                # Skip merged cells (except the top-left cell of merge)
                if (table_row, table_col) in merged_cells_to_skip:
                    continue
                
                try:
                    excel_cell = ws.cell(row=row_idx, column=col_idx)
                    word_cell = table.cell(table_row, table_col)
                    
                    value = "" if excel_cell.value is None else str(excel_cell.value)
                    
                    # Clear and set text
                    for p in word_cell.paragraphs:
                        p.clear()
                    word_cell.paragraphs[0].add_run(value)
                    
                    # Get the run for formatting
                    run = word_cell.paragraphs[0].runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    
                    # Check for specific content and apply special formatting
                    value_stripped = value.strip()
                    
                    # Header row with dark blue background
                    if value_stripped in ["Sr. No.", "Requirements", 
                                          "Complied Status (Fully Complied/Partially Complied/Not Complied/Not Applicable)",
                                          "Auditor's Remark", "POC Attached"]:
                        # Dark blue background, white bold text, center aligned
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        
                        # Add dark blue background
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), '00008B')  # Dark blue
                        word_cell._element.get_or_add_tcPr().append(shading_elm)
                    
                    # Apply column-based alignment (based on header)
                    elif table_col in center_aligned_columns:
                        # This column should be center-aligned
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    elif table_col in left_aligned_columns:
                        # This column should be left-aligned
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    else:
                        # Default alignment
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                except Exception as e:
                    print(f"    ⚠️ Error filling cell ({table_row}, {table_col}): {e}")
            
            rows_filled += 1
            if rows_filled % 50 == 0:
                print(f"    📝 Filled {rows_filled}/{num_rows} rows...")
        
        print(f"  ✅ Table filled with {rows_filled} rows")
        
        # Set column widths based on header content
        print(f"  📏 Setting column widths...")
        try:
            # Define width mapping based on header content
            header_width_map = {
                "Sr. No.": 20,
                "Requirements": 60,
                "Complied Status (Fully Complied/Partially Complied/Not Complied/Not Applicable)": 25,
                "Auditor's Remark": 80,
                "POC Attached": 25
            }
            
            # Get column widths from first row headers
            for col_idx in range(start_col, end_col + 1):
                table_col = col_idx - start_col
                excel_cell = ws.cell(row=start_row, column=col_idx)
                header_value = "" if excel_cell.value is None else str(excel_cell.value).strip()
                
                # Find matching width
                width = None
                for header_key, header_width in header_width_map.items():
                    if header_key in header_value:
                        width = header_width
                        break
                
                if width:
                    # Set column width
                    col = table.columns[table_col]
                    for cell in col.cells:
                        cell.width = Inches(width * 0.1)
                    print(f"    Column {table_col} ('{header_value}'): {width} units")
        
        except Exception as e:
            print(f"  ⚠️ Error setting column widths: {e}")
        
        # Insert table before placeholder
        print(f"  🔧 Inserting table into document...")
        parent.insert(parent.index(placeholder_element), table._element)
        
        # Remove placeholder
        print(f"  🗑️ Removing placeholder...")
        parent.remove(placeholder_element)
        
        print(f"  ✅ Table insertion complete!")
        return True
        
    except Exception as e:
        print(f"  ❌ Error inserting table: {e}")
        import traceback
        traceback.print_exc()
        return False

def parse_image_name_cyber(filename):
    """Parse image name to extract sorting key (first two digits only)"""
    name_without_ext = os.path.splitext(filename)[0]
    
    # Extract first number from filename
    match = re.match(r'^(\d+)', name_without_ext)
    
    if match:
        full_num_str = match.group(1)
        full_num = int(full_num_str)
        
        # Take only first two digits for sorting
        # If number has 1 digit, use as is (e.g., 1 → 1, 7 → 7)
        # If number has 2+ digits, take first two (e.g., 111 → 11, 1201 → 12, 28 → 28)
        if len(full_num_str) >= 2:
            sort_key = int(full_num_str[:2])
        else:
            sort_key = full_num
        
        # Extract remaining text after the number and underscore
        remaining = re.sub(r'^\d+_?', '', name_without_ext).strip()
        
        # Return: (sort_key for first 2 digits, full original number, remaining text, filename)
        return (sort_key, full_num, remaining, filename)
    
    return (999, 999, name_without_ext, filename)

def sort_images_cyber(image_files):
    """Sort images by first two digits only"""
    parsed_images = [parse_image_name_cyber(img) for img in image_files]
    
    # Sort by first two digits (sort_key), then by full number if needed
    sorted_images = sorted(parsed_images, key=lambda x: (x[0], x[1]))
    return [item[3] for item in sorted_images]

def insert_images_as_annexures_cyber(doc, placeholder, sorted_images, image_paths, ws):
    """Insert images as numbered annexures and update POC Attached column"""
    try:
        print(f"\n  🔍 Searching for placeholder '{placeholder}'...")
        print(f"  📋 Total images to insert: {len(sorted_images)}")
        
        # Find placeholder
        placeholder_para = None
        for i, paragraph in enumerate(doc.paragraphs):
            if placeholder in paragraph.text:
                placeholder_para = paragraph
                print(f"    ✅ Found placeholder at paragraph index {i}")
                break
        
        if not placeholder_para:
            print(f"  ❌ Placeholder '{placeholder}' not found")
            return False
        
        parent = placeholder_para._element.getparent()
        placeholder_element = placeholder_para._element
        
        # Create mapping of question numbers to annexure numbers
        question_to_annexure = {}
        annexure_num = 1
        prev_first_num = None
        
        print(f"\n  🖼️ Processing images and updating POC Attached column...")
        
        for idx, filename in enumerate(sorted_images, 1):
            print(f"\n    📷 Image {idx}/{len(sorted_images)}: {filename}")
            
            image_path = image_paths.get(filename)
            
            if not image_path:
                print(f"      ⚠️ Image path not found in dictionary")
                continue
            
            if not os.path.exists(image_path):
                print(f"      ⚠️ Image file does not exist: {image_path}")
                continue
            
            parsed = parse_image_name_cyber(filename)
            sort_key = parsed[0]  # First two digits for sorting
            full_num = parsed[1]  # Full original number
            remaining = parsed[2]  # Remaining text
            
            print(f"      📝 Sort key (first 2 digits): {sort_key}, Full number: {full_num}, Remaining: {remaining}")
            
            # Check if new annexure or continuation (based on full number)
            if prev_first_num != full_num:
                # New annexure
                print(f"      ✨ Creating new Annexure {annexure_num}")
                
                # Map this question number (full number) to annexure number
                question_to_annexure[str(full_num)] = annexure_num
                
                annexure_para = doc.add_paragraph()
                annexure_run = annexure_para.add_run(f"Annexure {annexure_num} ({remaining})")
                annexure_run.font.bold = True
                annexure_run.font.size = Pt(12)
                annexure_run.font.name = 'Times New Roman'
                parent.insert(parent.index(placeholder_element), annexure_para._element)
                
                annexure_num += 1
                prev_first_num = full_num
            else:
                # Continuation - just show remaining text
                print(f"      ➕ Continuation of same question number")
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(remaining)
                title_run.font.size = Pt(12)
                title_run.font.name = 'Times New Roman'
                title_run.font.bold = True
                parent.insert(parent.index(placeholder_element), title_para._element)
            
            # Add blank line before image
            blank_para = doc.add_paragraph()
            parent.insert(parent.index(placeholder_element), blank_para._element)
            
            # Add image with fixed height (430px) and proportional width
            try:
                print(f"      🖼️ Adding image...")
                from PIL import Image as PILImage
                
                # Get original image dimensions
                pil_img = PILImage.open(image_path)
                orig_width, orig_height = pil_img.size
                
                # Calculate proportional width for 430px height
                target_height_px = 430
                aspect_ratio = orig_width / orig_height
                target_width_px = target_height_px * aspect_ratio
                
                # Convert to inches (assuming 96 DPI)
                height_inches = target_height_px / 96
                width_inches = target_width_px / 96
                
                print(f"      📏 Original: {orig_width}x{orig_height}px, Target: {target_width_px:.0f}x{target_height_px}px")
                
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                picture = run.add_picture(image_path, height=Inches(height_inches))
                
                # Add 1pt border to image
                try:
                    inline = picture._inline
                    pic = inline.graphic.graphicData.pic
                    
                    # Get or create spPr (shape properties)
                    spPr = pic.spPr
                    if spPr is None:
                        spPr = OxmlElement('pic:spPr')
                        pic.insert(0, spPr)
                    
                    # Add line (border) element
                    ln = OxmlElement('a:ln')
                    ln.set('w', '12700')  # 1pt = 12700 EMUs
                    
                    # Add solid fill for the line
                    solidFill = OxmlElement('a:solidFill')
                    srgbClr = OxmlElement('a:srgbClr')
                    srgbClr.set('val', '000000')  # Black
                    solidFill.append(srgbClr)
                    ln.append(solidFill)
                    
                    # Add line to shape properties
                    spPr.append(ln)
                    
                    print(f"      🖼️ Added 1pt black border to image")
                except Exception as border_err:
                    print(f"      ⚠️ Could not add border: {border_err}")
                
                parent.insert(parent.index(placeholder_element), img_para._element)
                print(f"      ✅ Image added successfully")
            except Exception as e:
                print(f"      ❌ Error adding image: {e}")
                import traceback
                traceback.print_exc()
            
            # Add page break (except for last image)
            is_last_image = (idx == len(sorted_images))
            if not is_last_image:
                try:
                    from docx.enum.text import WD_BREAK
                    page_break_para = doc.add_paragraph()
                    page_break_para.add_run().add_break(WD_BREAK.PAGE)
                    parent.insert(parent.index(placeholder_element), page_break_para._element)
                    print(f"      📄 Page break added")
                except Exception as e:
                    print(f"      ⚠️ Error adding page break: {e}")
            else:
                print(f"      ℹ️ Last image - skipping page break")
        
        # Remove placeholder
        print(f"\n  🗑️ Removing placeholder...")
        parent.remove(placeholder_element)
        
        print(f"\n  ✅ Successfully inserted {len(sorted_images)} images as annexures")
        
        # Now update POC Attached column in the table
        print(f"\n📝 Updating POC Attached column in table...")
        update_poc_attached_in_table(doc, question_to_annexure)
        
        return True
        
    except Exception as e:
        print(f"  ❌ Error inserting annexures: {e}")
        import traceback
        traceback.print_exc()
        return False

def update_poc_attached_in_table(doc, question_to_annexure):
    """Update POC Attached column in the table based on question numbers"""
    try:
        # Find the table (should be the first table in document after our insertion)
        for table in doc.tables:
            # Check if this is our main table by looking for specific headers
            if table.rows and len(table.rows) > 0:
                first_row_text = []
                for cell in table.rows[0].cells:
                    first_row_text.append(cell.text)
                
                # Check if this table has our expected headers
                if "Sr. No." in first_row_text and "Requirements" in first_row_text:
                    print(f"  ✅ Found main data table")
                    
                    # Find POC Attached column index
                    poc_col_idx = None
                    for idx, text in enumerate(first_row_text):
                        if "POC Attached" in text:
                            poc_col_idx = idx
                            break
                    
                    if poc_col_idx is None:
                        print(f"  ⚠️ POC Attached column not found")
                        return False
                    
                    # Iterate through data rows and update POC Attached column
                    for row_idx, row in enumerate(table.rows[1:], 1):  # Skip header row
                        try:
                            # Get Sr. No. from first column
                            sr_no_cell = row.cells[0]
                            sr_no = sr_no_cell.text.strip()
                            
                            # Check if we have an annexure for this question number
                            if sr_no in question_to_annexure:
                                annexure_no = question_to_annexure[sr_no]
                                
                                # Update POC Attached cell
                                poc_cell = row.cells[poc_col_idx]
                                for p in poc_cell.paragraphs:
                                    p.clear()
                                poc_cell.paragraphs[0].add_run(f"Annexure {annexure_no}")
                                
                                # Format as red text
                                run = poc_cell.paragraphs[0].runs[0]
                                run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)
                                poc_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                poc_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                
                                print(f"    ✅ Row {row_idx}: Question {sr_no} → Annexure {annexure_no}")
                        
                        except Exception as e:
                            print(f"    ⚠️ Error updating POC for row {row_idx}: {e}")
                    
                    break  # Found and updated the table, exit loop
        
        return True
    
    except Exception as e:
        print(f"  ❌ Error updating POC Attached: {e}")
        import traceback
        traceback.print_exc()
        return False

@cyber_security_audit_report_bp.route('/process_cyber_security_audit_report', methods=['POST'])
def process_cyber_security_audit_report():
    """Process Cyber Security Audit Report generation"""
    try:
        print("\n" + "="*80)
        print("🚀 Processing Cyber Security Audit Report")
        print("="*80)
        
        # Get form data
        org_name = request.form.get('organizationName')
        org_name_other = request.form.get('organizationNameOther')
        city = request.form.get('city')
        city_other = request.form.get('cityOther')
        state = request.form.get('state')
        prepared_by_prefix = request.form.get('preparedByPrefix')
        prepared_by_name = request.form.get('preparedByName')
        start_date = request.form.get('startAuditDate')
        end_date = request.form.get('endAuditDate')
        submitted_to_prefix = request.form.get('submittedToPrefix')
        submitted_to_name = request.form.get('submittedToName')
        auditee_designation = request.form.get('auditeeDesignation')
        address = request.form.get('address')
        emails = request.form.getlist('email[]')
        
        final_org_name = org_name_other if org_name == 'Other' else org_name
        final_city = city_other if city == 'Other' else city
        
        excel_file = request.files.get('excelFile')
        zip_file = request.files.get('zipFile')
        
        if not excel_file or not zip_file:
            return "Missing required files", 400
        
        temp_dir = tempfile.mkdtemp()
        print(f"📁 Created temp directory: {temp_dir}")
        
        try:
            # Save files
            excel_path = os.path.join(temp_dir, 'cyber_security_audit.xlsx')
            zip_path = os.path.join(temp_dir, 'images.zip')
            
            excel_file.save(excel_path)
            zip_file.save(zip_path)
            print(f"💾 Saved uploaded files")
            
            # Load Excel
            print(f"\n📖 Loading Excel file...")
            wb = load_workbook(excel_path, data_only=True)
            ws = wb.active
            print(f"  ✅ Loaded worksheet: {ws.title}")
            
            # Find last row with data in column A
            last_row = 1
            for row_num in range(1, ws.max_row + 1):
                a_cell = ws.cell(row=row_num, column=1).value
                if a_cell is not None and str(a_cell).strip() != '':
                    last_row = row_num
            print(f"  📊 Last row with data: {last_row}")
            
            # Extract and sort images
            print(f"\n🖼️ Extracting images from ZIP...")
            images_dir = os.path.join(temp_dir, 'images')
            os.makedirs(images_dir, exist_ok=True)
            
            image_files = []
            image_paths = {}
            
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                print(f"  📦 ZIP contains {len(zip_ref.namelist())} files")
                for file_info in zip_ref.namelist():
                    if file_info.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                        zip_ref.extract(file_info, images_dir)
                        extracted_path = os.path.join(images_dir, file_info)
                        filename = os.path.basename(file_info)
                        image_files.append(filename)
                        image_paths[filename] = extracted_path
                        print(f"      ✅ Extracted image: {filename}")
            
            print(f"\n  📷 Total images extracted: {len(image_files)}")
            
            # Sort images
            print(f"\n  🔢 Sorting images...")
            sorted_images = sort_images_cyber(image_files)
            
            print(f"  📋 Sorted image order:")
            for idx, img in enumerate(sorted_images, 1):
                print(f"    {idx}. {img}")
            
            # Load Word template
            print(f"\n📄 Loading Word template...")
            template_path = os.path.join('static', 'Formats_and_Catalog', 'Cybersecurity_Audit_Report.docx')
            doc = Document(template_path)
            
            # Prepare replacements
            prepared_by = f"{prepared_by_prefix} {prepared_by_name}"
            submitted_to = f"{submitted_to_prefix} {submitted_to_name}"
            date_of_audit = f"{format_date_with_superscript_cyber(start_date)} to {format_date_with_superscript_cyber(end_date)}"
            audit_period = get_financial_year_cyber(start_date)
            full_address = f"{final_org_name}, {address}"
            email_list = ', '.join(emails)
            
            replacements = {
                'Orggganization____nameeee': final_org_name,
                'Ccccityyy': final_city,
                'Stttateeee': state,
                'Prepaaareddd_byyyyy': prepared_by,
                'Suuubmiiiiitttedddd_tooto': submitted_to,
                'Deeeesigggnation': auditee_designation,
                'eeeeeemail': email_list,
                'addddddddrrreessss': full_address,
                'Daaaaateeee_offff_auuuudit': date_of_audit,
                'audit_period_finanicialll_yeearrr': audit_period
            }
            
            # Replace text
            print(f"\n📝 Replacing placeholders...")
            for key, value in replacements.items():
                print(f"    {key} → {value[:50]}..." if len(value) > 50 else f"    {key} → {value}")
            
            replace_text_in_document_cyber(doc, replacements)
            print(f"  ✅ Replacements complete")
            
            # Add page break before "Annexures for Cyber Security Audit"
            print(f"\n📄 Adding page break before annexure section...")
            from docx.enum.text import WD_BREAK
            
            for paragraph in doc.paragraphs:
                if "Annexures for Cyber Security Audit" in paragraph.text or "Annexures For Cyber Security Audit" in paragraph.text:
                    # Insert page break before this paragraph
                    parent = paragraph._element.getparent()
                    page_break_para = doc.add_paragraph()
                    page_break_para.add_run().add_break(WD_BREAK.PAGE)
                    parent.insert(parent.index(paragraph._element), page_break_para._element)
                    print(f"  ✅ Added page break before 'Annexures for Cyber Security Audit'")
                    break
            
            # Insert table (Columns A to E)
            print(f"\n📊 Inserting table (A-E, 1-{last_row})...")
            insert_excel_table_in_document_cyber(doc, 'tttttaaablllelelel', ws, 1, 5, 1, last_row)
            
            # Insert annexures
            if sorted_images:
                print(f"\n🖼️ Inserting annexures...")
                insert_images_as_annexures_cyber(doc, 'Annnnnnnexxxurressss', sorted_images, image_paths, ws)
            
            # Save document
            output_path = os.path.join(temp_dir, 'Cyber_Security_Audit_Report.docx')
            doc.save(output_path)
            print(f"\n✅ Document saved")
            print("="*80)
            
            return send_file(
                output_path,
                as_attachment=True,
                download_name='Cyber_Security_Audit_Report.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
        finally:
            pass
    
    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return f"Error: {str(e)}", 500

