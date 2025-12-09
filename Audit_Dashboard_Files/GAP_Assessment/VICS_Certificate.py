from flask import Blueprint, request, send_file
import os
import tempfile
import shutil
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create blueprint for VICS Certificate
vics_certificate_bp = Blueprint('vics_certificate', __name__)

def _replace_text_in_runs_vics_cert(paragraph, old_text, new_text, make_bold=False):
    """
    Helper function to replace text in a paragraph, preserving paragraph formatting 
    but only applying bold to the replacement text if needed.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return False
    
    # Find the position of old_text
    start_pos = full_text.find(old_text)
    if start_pos == -1:
        return False
    
    end_pos = start_pos + len(old_text)
    
    # Split into three parts: before, replacement, after
    before_text = full_text[:start_pos]
    after_text = full_text[end_pos:]
    
    # Clear all existing runs
    for run in paragraph.runs:
        run.text = ""
    
    # Remove all runs except the first one
    while len(paragraph.runs) > 1:
        p_element = paragraph.runs[1]._element
        p_element.getparent().remove(p_element)
    
    # Get the first run to use as a template for formatting
    if len(paragraph.runs) == 0:
        paragraph.add_run()
    
    first_run = paragraph.runs[0]
    original_bold = first_run.bold
    original_font_name = first_run.font.name if first_run.font.name else 'Times New Roman'
    original_font_size = first_run.font.size if first_run.font.size else Pt(12)
    
    try:
        original_font_color = first_run.font.color.rgb
    except:
        original_font_color = None
    
    # Clear the first run
    first_run.text = ""
    
    # Add "before" text with original formatting
    if before_text:
        before_run = paragraph.add_run(before_text)
        before_run.bold = original_bold
        before_run.font.name = original_font_name
        before_run.font.size = original_font_size
        if original_font_color:
            before_run.font.color.rgb = original_font_color
    
    # Add replacement text (bold if requested, otherwise use original formatting)
    replacement_run = paragraph.add_run(new_text)
    replacement_run.bold = True if make_bold else original_bold
    replacement_run.font.name = original_font_name
    replacement_run.font.size = original_font_size
    if original_font_color:
        replacement_run.font.color.rgb = original_font_color
    
    # Add "after" text with original formatting
    if after_text:
        after_run = paragraph.add_run(after_text)
        after_run.bold = original_bold
        after_run.font.name = original_font_name
        after_run.font.size = original_font_size
        if original_font_color:
            after_run.font.color.rgb = original_font_color
    
    # Remove the empty first run
    if first_run.text == "":
        p_element = first_run._element
        p_element.getparent().remove(p_element)
    
    return True

def replace_text_in_document_vics_cert(doc, replacements, bold_replacements=None):
    """Replace text in document while preserving formatting
    
    Args:
        doc: Document object
        replacements: Dict of {old_text: new_text} for normal replacements
        bold_replacements: Dict of {old_text: new_text} for bold replacements
    """
    if bold_replacements is None:
        bold_replacements = {}
    
    try:
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            # First handle bold replacements
            for old_text, new_text in bold_replacements.items():
                if old_text in paragraph.text:
                    _replace_text_in_runs_vics_cert(paragraph, old_text, new_text, make_bold=True)
            
            # Then handle normal replacements
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    _replace_text_in_runs_vics_cert(paragraph, old_text, new_text, make_bold=False)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # First handle bold replacements
                        for old_text, new_text in bold_replacements.items():
                            if old_text in paragraph.text:
                                _replace_text_in_runs_vics_cert(paragraph, old_text, new_text, make_bold=True)
                        
                        # Then handle normal replacements
                        for old_text, new_text in replacements.items():
                            if old_text in paragraph.text:
                                _replace_text_in_runs_vics_cert(paragraph, old_text, new_text, make_bold=False)
        
        # Replace in headers
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                # First handle bold replacements
                for old_text, new_text in bold_replacements.items():
                    if old_text in paragraph.text:
                        _replace_text_in_runs_vics_cert(paragraph, old_text, new_text, make_bold=True)
                
                # Then handle normal replacements
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        _replace_text_in_runs_vics_cert(paragraph, old_text, new_text, make_bold=False)
        
        # Replace in footers
        for section in doc.sections:
            for paragraph in section.footer.paragraphs:
                # First handle bold replacements
                for old_text, new_text in bold_replacements.items():
                    if old_text in paragraph.text:
                        _replace_text_in_runs_vics_cert(paragraph, old_text, new_text, make_bold=True)
                
                # Then handle normal replacements
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        _replace_text_in_runs_vics_cert(paragraph, old_text, new_text, make_bold=False)
        
        return True
    except Exception as e:
        print(f"❌ Error replacing text: {e}")
        return False

def insert_excel_table_in_document_vics_cert(doc, placeholder, ws, start_col, end_col, start_row, end_row):
    """Insert Excel table into Word document with formatting"""
    try:
        print(f"\n  🔍 Searching for placeholder '{placeholder}'...")
        
        # Find placeholder paragraph
        placeholder_para = None
        
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                placeholder_para = paragraph
                print(f"    ✅ Found placeholder in paragraph!")
                break
        
        if not placeholder_para:
            print(f"  ❌ Placeholder '{placeholder}' not found in document")
            return False
        
        # Calculate dimensions
        num_rows = end_row - start_row + 1
        num_cols = end_col - start_col + 1
        
        print(f"  📏 Creating table: {num_rows} rows x {num_cols} columns")
        
        # Get parent and placeholder element
        parent = placeholder_para._element.getparent()
        placeholder_element = placeholder_para._element
        
        # Create table
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # Track merged cells to skip
        merged_cells_to_skip = set()
        
        # Handle merged cells from Excel
        print(f"  🔗 Processing merged cells...")
        merged_count = 0
        for merged_range in ws.merged_cells.ranges:
            min_row_excel = merged_range.min_row
            max_row_excel = merged_range.max_row
            min_col_excel = merged_range.min_col
            max_col_excel = merged_range.max_col
            
            # Check if merge is within our range
            if (start_row <= min_row_excel <= end_row and 
                start_col <= min_col_excel <= end_col):
                
                # Convert to table indices (0-based)
                start_row_idx = min_row_excel - start_row
                end_row_idx = max_row_excel - start_row
                start_col_idx = min_col_excel - start_col
                end_col_idx = max_col_excel - start_col
                
                # Ensure indices are within bounds
                if (0 <= start_row_idx < num_rows and 0 <= end_row_idx < num_rows and
                    0 <= start_col_idx < num_cols and 0 <= end_col_idx < num_cols):
                    
                    try:
                        start_word_cell = table.cell(start_row_idx, start_col_idx)
                        end_word_cell = table.cell(end_row_idx, end_col_idx)
                        start_word_cell.merge(end_word_cell)
                        merged_count += 1
                        
                        # Mark cells to skip
                        for r in range(start_row_idx, end_row_idx + 1):
                            for c in range(start_col_idx, end_col_idx + 1):
                                if (r, c) != (start_row_idx, start_col_idx):
                                    merged_cells_to_skip.add((r, c))
                        
                        print(f"    ✅ Merged cells: Row {start_row_idx}-{end_row_idx}, Col {start_col_idx}-{end_col_idx}")
                    except Exception as e:
                        print(f"    ⚠️ Error merging cells: {e}")
        
        print(f"  ✅ Processed {merged_count} merged cell ranges")
        
        # First pass: identify columns that should be center-aligned
        print(f"  🔍 Identifying column alignments from header row...")
        center_aligned_columns = set()
        left_aligned_columns = set()
        
        # Check first row (header row) to determine column alignment
        for col_idx in range(start_col, end_col + 1):
            table_col = col_idx - start_col
            excel_cell = ws.cell(row=start_row, column=col_idx)
            header_value = "" if excel_cell.value is None else str(excel_cell.value).strip()
            
            # Columns that should be center-aligned throughout
            if header_value in ["Sr. No.", "Max Marks", "Yes/No", "Marks given by the Auditor", 
                               "POC Attached", "Yes/No given by Auditor", "Marks given by Bank", 
                               "Yes/No given by the Bank", "Input"]:
                center_aligned_columns.add(table_col)
                print(f"    Column {table_col} ('{header_value}'): CENTER aligned")
            
            # Columns that should be left-aligned throughout
            elif header_value in ["Questions", "Auditor's Observation", "Auditor's Observation"]:
                left_aligned_columns.add(table_col)
                print(f"    Column {table_col} ('{header_value}'): LEFT aligned")
        
        print(f"  ✅ Center columns: {center_aligned_columns}")
        print(f"  ✅ Left columns: {left_aligned_columns}")
        
        print(f"  📊 Filling table with data...")
        
        # Fill table with data
        rows_filled = 0
        for row_idx in range(start_row, end_row + 1):
            for col_idx in range(start_col, end_col + 1):
                table_row = row_idx - start_row
                table_col = col_idx - start_col
                
                # Skip merged cells (except the top-left cell of merge)
                if (table_row, table_col) in merged_cells_to_skip:
                    continue
                
                try:
                    excel_cell = ws.cell(row=row_idx, column=col_idx)
                    word_cell = table.cell(table_row, table_col)
                    
                    value = "" if excel_cell.value is None else str(excel_cell.value)
                    
                    # Clear and set text
                    for p in word_cell.paragraphs:
                        p.clear()
                    word_cell.paragraphs[0].add_run(value)
                    
                    # Get the run for formatting
                    run = word_cell.paragraphs[0].runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    
                    # Check for specific content and apply special formatting
                    value_stripped = value.strip()
                    
                    # Header row with dark blue background
                    if value_stripped in ["Auditor's Observation", "Max Marks", "Yes/No", "Marks given by the Auditor", 
                                          "POC Attached", "Sr. No.", "Questions", "Yes/No given by Auditor",
                                          "Marks given by Bank", "Yes/No given by the Bank", "Input", 
                                          "A) Info Sec Processes & Controls", "B) Governance & Policy", 
                                          "C) Vendor Management", "D) Cyber Crisis Management"]:
                        # Dark blue background, white bold text, center aligned
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        
                        # Add dark blue background
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), '00008B')  # Dark blue
                        word_cell._element.get_or_add_tcPr().append(shading_elm)
                    
                    # Apply column-based alignment (based on header)
                    elif table_col in center_aligned_columns:
                        # This column should be center-aligned
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    elif table_col in left_aligned_columns:
                        # This column should be left-aligned
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    else:
                        # Default alignment based on content
                        if value_stripped.replace('.', '').replace('_', '').isdigit() or len(value_stripped) < 5:
                            # Short content or numbers - center
                            word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        else:
                            # Long content - left
                            word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                except Exception as e:
                    print(f"    ⚠️ Error filling cell ({table_row}, {table_col}): {e}")
            
            rows_filled += 1
            if rows_filled % 50 == 0:
                print(f"    📝 Filled {rows_filled}/{num_rows} rows...")
        
        print(f"  ✅ Table filled with {rows_filled} rows")
        
        # Set column widths based on header content
        print(f"  📏 Setting column widths...")
        try:
            # Define width mapping based on header content (in characters/units)
            header_width_map = {
                "Sr. No.": 20,
                "Max Marks": 15,
                "Yes/No": 18,
                "Marks given by the Auditor": 20,
                "Marks given by Bank": 15,
                "Yes/No given by the Bank": 20,
                "Yes/No given by Auditor": 18,
                "POC Attached": 30,
                "Questions": 50,
                "Auditor's Observation": 50,
                "Input": 20
            }
            
            # Get column widths from first row headers
            for col_idx in range(start_col, end_col + 1):
                table_col = col_idx - start_col
                excel_cell = ws.cell(row=start_row, column=col_idx)
                header_value = "" if excel_cell.value is None else str(excel_cell.value).strip()
                
                # Find matching width
                width = None
                for header_key, header_width in header_width_map.items():
                    if header_key in header_value:
                        width = header_width
                        break
                
                if width:
                    # Set column width (using twips: 1 character ≈ 120 twips)
                    col = table.columns[table_col]
                    for cell in col.cells:
                        cell.width = Inches(width * 0.1)  # Approximate conversion
                    print(f"    Column {table_col} ('{header_value}'): {width} units")
        
        except Exception as e:
            print(f"  ⚠️ Error setting column widths: {e}")
        
        # Insert table before placeholder
        print(f"  🔧 Inserting table into document...")
        parent.insert(parent.index(placeholder_element), table._element)
        
        # Remove placeholder
        print(f"  🗑️ Removing placeholder...")
        parent.remove(placeholder_element)
        
        print(f"  ✅ Table insertion complete!")
        return True
        
    except Exception as e:
        print(f"  ❌ Error inserting table: {e}")
        import traceback
        traceback.print_exc()
        return False

def insert_summary_table_vics_cert(doc, placeholder, ws):
    """Insert summary table from Excel cells B253-B258, E253-E258, F253-F258"""
    try:
        print(f"\n  🔍 Searching for placeholder '{placeholder}'...")
        
        # Find placeholder paragraph (no page break)
        placeholder_para = None
        
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                placeholder_para = paragraph
                print(f"    ✅ Found placeholder in paragraph!")
                break
        
        if not placeholder_para:
            print(f"  ❌ Placeholder '{placeholder}' not found in document")
            return False
        
        # Create 6x3 table (6 rows, 3 columns)
        print(f"  📏 Creating summary table: 6 rows x 3 columns")
        
        parent = placeholder_para._element.getparent()
        placeholder_element = placeholder_para._element
        
        table = doc.add_table(rows=6, cols=3)
        table.style = 'Table Grid'
        
        # Data ranges to copy
        rows_to_copy = [253, 254, 255, 256, 257, 258]
        cols_to_copy = [2, 5, 6]  # B=2, E=5, F=6
        
        print(f"  📊 Filling summary table with data...")
        
        for table_row_idx, excel_row in enumerate(rows_to_copy):
            for table_col_idx, excel_col in enumerate(cols_to_copy):
                try:
                    excel_cell = ws.cell(row=excel_row, column=excel_col)
                    word_cell = table.cell(table_row_idx, table_col_idx)
                    
                    value = "" if excel_cell.value is None else str(excel_cell.value)
                    
                    # Clear and set text
                    for p in word_cell.paragraphs:
                        p.clear()
                    word_cell.paragraphs[0].add_run(value)
                    
                    # Get the run for formatting
                    run = word_cell.paragraphs[0].runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    
                    value_stripped = value.strip()
                    
                    # Check if cell contains special headers that need dark blue background
                    if "Section" in value_stripped or "Total Mark given by the Bank" in value_stripped or "Total Mark given by the Auditor" in value_stripped:
                        # Header cell with dark blue background
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        
                        # Add dark blue background
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), '00008B')  # Dark blue
                        word_cell._element.get_or_add_tcPr().append(shading_elm)
                        
                        # Set column width for Section column
                        if table_col_idx == 0:
                            word_cell.width = Inches(8.0)
                    elif table_col_idx == 0:
                        # Regular cell in first column (left align)
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        word_cell.width = Inches(8.0)
                    else:
                        # Other columns (center align, middle align)
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    print(f"    ✅ Filled cell ({table_row_idx}, {table_col_idx}): {value_stripped[:30]}...")
                    
                except Exception as e:
                    print(f"    ⚠️ Error filling summary cell ({table_row_idx}, {table_col_idx}): {e}")
        
        print(f"  ✅ Summary table filled")
        
        # Insert table before placeholder
        print(f"  🔧 Inserting summary table into document...")
        parent.insert(parent.index(placeholder_element), table._element)
        
        # Remove placeholder
        print(f"  🗑️ Removing placeholder...")
        parent.remove(placeholder_element)
        
        print(f"  ✅ Summary table insertion complete!")
        return True
        
    except Exception as e:
        print(f"  ❌ Error inserting summary table: {e}")
        import traceback
        traceback.print_exc()
        return False

@vics_certificate_bp.route('/process_vics_certificate', methods=['POST'])
def process_vics_certificate():
    """Process VICS Certificate generation"""
    try:
        print("\n" + "="*80)
        print("🚀 Processing VICS Certificate")
        print("="*80)
        
        # Get form data
        org_name = request.form.get('organizationName')
        org_name_other = request.form.get('organizationNameOther')
        
        final_org_name = org_name_other if org_name == 'Other' else org_name
        
        excel_file = request.files.get('excelFile')
        
        if not excel_file:
            return "Missing Excel file", 400
        
        temp_dir = tempfile.mkdtemp()
        print(f"📁 Created temp directory: {temp_dir}")
        
        try:
            # Save Excel file
            excel_path = os.path.join(temp_dir, 'vics_cert.xlsx')
            excel_file.save(excel_path)
            print(f"💾 Saved uploaded Excel file")
            
            # Load Excel
            print(f"\n📖 Loading Excel file...")
            wb = load_workbook(excel_path, data_only=True)
            
            vics_ws = None
            
            for sheet_name in wb.sheetnames:
                if 'VICS' in sheet_name.upper():
                    vics_ws = wb[sheet_name]
                    print(f"  ✅ Found VICS worksheet: {sheet_name}")
                    break
            
            if not vics_ws:
                return "VICS worksheet not found in Excel file", 400
            
            # Get F250 value for total marks
            f250_value = vics_ws.cell(row=250, column=6).value
            total_marks = str(f250_value) if f250_value is not None else "0"
            print(f"  📊 F250 total marks: {total_marks}")
            
            # Load Word template
            print(f"\n📄 Loading Word template...")
            template_path = os.path.join('static', 'Formats_and_Catalog', 'VICS Certificate.docx')
            doc = Document(template_path)
            
            # Prepare replacements
            replacements = {}
            
            # Bold replacements (organization name and total marks)
            bold_replacements = {
                'Orrrrrrganizzzation': final_org_name,
                'ttttttttotal_marrrrks': total_marks
            }
            
            # Replace text
            print(f"\n📝 Replacing placeholders...")
            print(f"  Regular replacements: {replacements}")
            print(f"  Bold replacements: {bold_replacements}")
            
            replace_text_in_document_vics_cert(doc, replacements, bold_replacements)
            print(f"  ✅ Replacements complete")
            
            # Insert VICS table (Columns A to G, rows 1 to 250)
            if vics_ws:
                print(f"\n📊 Inserting VICS table (A-G, 1-250)...")
                insert_excel_table_in_document_vics_cert(doc, 'VICSSSSS_TABLEEEEEE', vics_ws, 1, 7, 1, 250)
            
            # Insert Summary table from B253-B258, E253-E258, F253-F258
            if vics_ws:
                print(f"\n📊 Inserting VICS Summary table...")
                insert_summary_table_vics_cert(doc, 'VICCSS_Summary', vics_ws)
            
            # Save document
            output_path = os.path.join(temp_dir, 'VICS_Certificate.docx')
            doc.save(output_path)
            print(f"\n✅ Document saved")
            print("="*80)
            
            return send_file(
                output_path,
                as_attachment=True,
                download_name='VICS_Certificate.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
        finally:
            pass
    
    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return f"Error: {str(e)}", 500

