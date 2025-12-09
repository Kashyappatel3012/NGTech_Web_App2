from flask import Blueprint, request, send_file
import os
import zipfile
import tempfile
import shutil
from datetime import datetime
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# Create blueprint for Gap Assessment Word Report with Bank Input
gap_assessment_report_bank_input_bp = Blueprint('gap_assessment_report_bank_input', __name__)

def get_ordinal_superscript_bank(day):
    """Convert day number to superscript ordinal (1ˢᵗ, 2ⁿᵈ, 3ʳᵈ, 4ᵗʰ, etc.)"""
    superscripts = {
        'st': 'ˢᵗ',
        'nd': 'ⁿᵈ',
        'rd': 'ʳᵈ',
        'th': 'ᵗʰ'
    }
    
    if 10 <= day % 100 <= 20:
        suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    
    return f"{day:02d}{superscripts[suffix]}"

def format_date_with_superscript_bank(date_str):
    """Format date as '21ˢᵗ October 2025'"""
    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        day = date_obj.day
        month = date_obj.strftime('%B')
        year = date_obj.year
        
        day_with_ordinal = get_ordinal_superscript_bank(day)
        return f"{day_with_ordinal} {month} {year}"
    except:
        return date_str

def get_financial_year_bank(date_str):
    """Get financial year for a given date (1st April to 31st March)"""
    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        
        if date_obj.month >= 4:
            fy_start_year = date_obj.year
            fy_end_year = date_obj.year + 1
        else:
            fy_start_year = date_obj.year - 1
            fy_end_year = date_obj.year
        
        start_formatted = f"{get_ordinal_superscript_bank(1)} April {fy_start_year}"
        end_formatted = f"{get_ordinal_superscript_bank(31)} March {fy_end_year}"
        
        return f"{start_formatted} to {end_formatted}"
    except:
        return ""

def check_loc_cells_bank(ws):
    """Check if specific LOC worksheet cells have content"""
    try:
        cells_to_check = {
            'b58': ws.cell(row=58, column=2).value,
            'a59': ws.cell(row=59, column=1).value,
            'b59': ws.cell(row=59, column=2).value,
            'b113': ws.cell(row=113, column=2).value,
            'a114': ws.cell(row=114, column=1).value
        }
        
        has_content = {}
        for key, value in cells_to_check.items():
            has_content[key] = value is not None and str(value).strip() != ''
        
        # Scenario 1: All 5 have content
        if all(has_content.values()):
            return "level4"
        # Scenario 2: First 3 have content, last 2 don't
        elif has_content['b58'] and has_content['a59'] and has_content['b59'] and not has_content['b113'] and not has_content['a114']:
            return "level3"
        # Scenario 3: All empty
        else:
            return "level2"
    except:
        return "level2"

def get_areas_covered_text_bank(level):
    """Get the areas covered text based on level"""
    texts = {
        "level2": """1. Network Management and Security
2. Secure Configuration
3. Application Security Life Cycle (ASLC)
4. Change Management
5. Periodic Testing
6. User Access Control/Management
7. Authentication Framework for Customers
8. Anti-Phishing
9. User/Employee/Management Awareness
10. Audit Logs
11. Incident Response and Management""",
        
        "level3": """1. Network Management and Security
2. Secure Configuration
3. Application Security Life Cycle (ASLC)
4. Change Management
5. Periodic Testing
6. User Access Control/Management
7. Authentication Framework for Customers
8. Anti-Phishing
9. User/Employee/Management Awareness
10. Audit Logs
11. Incident Response and Management
12. Network Management and Security
13. Secure Configuration
14. Application Security Life Cycle (ASLC)
15. User Access Control
16. Advance Real-time Threat Defence and Management
17. Maintenance, Monitoring and Analysis of Audit Logs
18. Incident Response and Management
19. Risk based transaction monitoring""",
        
        "level4": """1. Network Management and Security
2. Secure Configuration
3. Application Security Life Cycle (ASLC)
4. Change Management
5. Periodic Testing
6. User Access Control/Management
7. Authentication Framework for Customers
8. Anti-Phishing
9. User/Employee/Management Awareness
10. Audit Logs
11. Incident Response and Management
12. Network Management and Security
13. Secure Configuration
14. Application Security Life Cycle (ASLC)
15. User Access Control
16. Advance Real-time Threat Defence and Management
17. Maintenance, Monitoring and Analysis of Audit Logs
18. Incident Response and Management
19. Risk based transaction monitoring
20. Arrangement for continuous surveillance- Setting up of Cyber Security Operation Centre(C-SOC)
21. If Yes, CSOC in place then answer the following
22. Expectations from C-SOC
23. Steps for setting up C-SOC- Technological Aspects
24. Participation in Cyber Drills
25. Incident Response and Management
26. Forensics and Metrics
27. IT Strategy and Policy
28. IT and IS governance Framework
29. IT strategy Committee
30. IT Steering Committee
31. CISO"""
    }
    
    return texts.get(level, texts["level2"])

def _replace_text_in_runs_bank(paragraph, old_text, new_text):
    """
    Helper function to replace text in a paragraph while preserving formatting.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return False

    run_formats = []
    text_positions = []
    
    start_pos = 0
    while True:
        pos = full_text.find(old_text, start_pos)
        if pos == -1:
            break
        text_positions.append((pos, pos + len(old_text)))
        start_pos = pos + 1
    
    if not text_positions:
        return False
    
    current_pos = 0
    for run in paragraph.runs:
        run_start = current_pos
        run_end = current_pos + len(run.text)
        
        for text_start, text_end in text_positions:
            if run_start < text_end and run_end > text_start:
                font_color = None
                try:
                    if run.font.color.rgb is not None:
                        font_color = run.font.color.rgb
                except:
                    font_color = None
                
                run_formats.append({
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name if run.font.name else 'Times New Roman',
                    'font_size': run.font.size,
                    'font_color': font_color
                })
                break
        
        current_pos = run_end
    
    if run_formats:
        for run in paragraph.runs:
            run.text = ""
        
        new_full_text = full_text.replace(old_text, new_text)
        
        if len(paragraph.runs) > 0:
            first_run = paragraph.runs[0]
            first_run.text = new_full_text
            
            original_format = run_formats[0]
            first_run.bold = original_format.get('bold')
            first_run.italic = original_format.get('italic')
            first_run.underline = original_format.get('underline')
            
            # Set font name (default to Times New Roman if not specified)
            if original_format.get('font_name'):
                first_run.font.name = original_format.get('font_name')
            else:
                first_run.font.name = 'Times New Roman'
            
            # Set font size (default to 12pt if not specified)
            if original_format.get('font_size'):
                first_run.font.size = original_format.get('font_size')
            else:
                first_run.font.size = Pt(12)
            
            if original_format.get('font_color') is not None:
                try:
                    first_run.font.color.rgb = original_format.get('font_color')
                except:
                    pass
            
            runs_to_remove = list(paragraph.runs[1:])
            for run in runs_to_remove:
                run._element.getparent().remove(run._element)
    else:
        # No formatting found, create new run with default formatting
        for run in paragraph.runs:
            run.text = ""
        
        new_full_text = full_text.replace(old_text, new_text)
        
        if len(paragraph.runs) > 0:
            first_run = paragraph.runs[0]
            first_run.text = new_full_text
            first_run.font.name = 'Times New Roman'
            first_run.font.size = Pt(12)
    
    return True

def _replace_text_with_left_indent_bank(paragraph, old_text, new_text):
    """Replace text and add left indent for areas covered list"""
    if not _replace_text_in_runs_bank(paragraph, old_text, new_text):
        return False
    
    # Add left indent (0.5 inches)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    return True

def replace_text_in_document_bank(doc, replacements):
    """Replace text in document while preserving formatting"""
    try:
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    if old_text == "Aarreas_Covvered":
                        _replace_text_with_left_indent_bank(paragraph, old_text, new_text)
                    else:
                        _replace_text_in_runs_bank(paragraph, old_text, new_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            if old_text in paragraph.text:
                                if old_text == "Aarreas_Covvered":
                                    _replace_text_with_left_indent_bank(paragraph, old_text, new_text)
                                else:
                                    _replace_text_in_runs_bank(paragraph, old_text, new_text)
        
        # Replace in headers
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        _replace_text_in_runs_bank(paragraph, old_text, new_text)
        
        # Replace in footers
        for section in doc.sections:
            for paragraph in section.footer.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        _replace_text_in_runs_bank(paragraph, old_text, new_text)
        
        return True
    except Exception as e:
        print(f"❌ Error replacing text: {e}")
        return False

def get_rgb_bank(color_obj):
    """Extract RGB from Excel color object"""
    try:
        if hasattr(color_obj, 'rgb') and color_obj.rgb:
            rgb_str = color_obj.rgb
            if rgb_str.startswith('FF'):
                rgb_str = rgb_str[2:]
            return rgb_str
    except:
        pass
    return None

def insert_excel_table_in_document_bank(doc, placeholder, ws, start_col, end_col, start_row, end_row):
    """Insert Excel table into Word document with simple formatting"""
    try:
        print(f"\n  🔍 Searching for placeholder '{placeholder}'...")
        
        # Find placeholder paragraph
        placeholder_para = None
        
        for paragraph in doc.paragraphs:
            print(f"    Checking paragraph: '{paragraph.text[:50]}...'")
            if placeholder in paragraph.text:
                placeholder_para = paragraph
                print(f"    ✅ Found placeholder in paragraph!")
                break
        
        if not placeholder_para:
            print(f"  ❌ Placeholder '{placeholder}' not found in document")
            print(f"  📋 Searched {len(doc.paragraphs)} paragraphs")
            return False
        
        # Calculate dimensions
        num_rows = end_row - start_row + 1
        num_cols = end_col - start_col + 1
        
        print(f"  📏 Creating table: {num_rows} rows x {num_cols} columns")
        
        # Get parent and placeholder element
        parent = placeholder_para._element.getparent()
        placeholder_element = placeholder_para._element
        
        # Create table
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # Track merged cells to skip
        merged_cells_to_skip = set()
        
        # Handle merged cells from Excel
        print(f"  🔗 Processing merged cells...")
        merged_count = 0
        for merged_range in ws.merged_cells.ranges:
            min_row_excel = merged_range.min_row
            max_row_excel = merged_range.max_row
            min_col_excel = merged_range.min_col
            max_col_excel = merged_range.max_col
            
            # Check if merge is within our range
            if (start_row <= min_row_excel <= end_row and 
                start_col <= min_col_excel <= end_col):
                
                # Convert to table indices (0-based)
                start_row_idx = min_row_excel - start_row
                end_row_idx = max_row_excel - start_row
                start_col_idx = min_col_excel - start_col
                end_col_idx = max_col_excel - start_col
                
                # Ensure indices are within bounds
                if (0 <= start_row_idx < num_rows and 0 <= end_row_idx < num_rows and
                    0 <= start_col_idx < num_cols and 0 <= end_col_idx < num_cols):
                    
                    try:
                        start_word_cell = table.cell(start_row_idx, start_col_idx)
                        end_word_cell = table.cell(end_row_idx, end_col_idx)
                        start_word_cell.merge(end_word_cell)
                        merged_count += 1
                        
                        # Mark cells to skip
                        for r in range(start_row_idx, end_row_idx + 1):
                            for c in range(start_col_idx, end_col_idx + 1):
                                if (r, c) != (start_row_idx, start_col_idx):
                                    merged_cells_to_skip.add((r, c))
                        
                        print(f"    ✅ Merged cells: Row {start_row_idx}-{end_row_idx}, Col {start_col_idx}-{end_col_idx}")
                    except Exception as e:
                        print(f"    ⚠️ Error merging cells: {e}")
        
        print(f"  ✅ Processed {merged_count} merged cell ranges")
        
        # First pass: identify columns that should be center-aligned
        print(f"  🔍 Identifying column alignments from header row...")
        center_aligned_columns = set()
        left_aligned_columns = set()
        
        # Check first row (header row) to determine column alignment
        for col_idx in range(start_col, end_col + 1):
            table_col = col_idx - start_col
            excel_cell = ws.cell(row=start_row, column=col_idx)
            header_value = "" if excel_cell.value is None else str(excel_cell.value).strip()
            
            # Columns that should be center-aligned throughout
            if header_value in ["Sr. No.", "Max Marks", "Yes/No", "Marks given by the Auditor", 
                               "POC Attached", "Yes/No given by Auditor", "Marks given by Bank", 
                               "Yes/No given by the Bank", "Input"]:
                center_aligned_columns.add(table_col)
                print(f"    Column {table_col} ('{header_value}'): CENTER aligned")
            
            # Columns that should be left-aligned throughout
            elif header_value in ["Questions", "Auditor's Observation", "Auditor's Observation"]:
                left_aligned_columns.add(table_col)
                print(f"    Column {table_col} ('{header_value}'): LEFT aligned")
        
        print(f"  ✅ Center columns: {center_aligned_columns}")
        print(f"  ✅ Left columns: {left_aligned_columns}")
        
        print(f"  📊 Filling table with data...")
        
        # Fill table with data
        rows_filled = 0
        for row_idx in range(start_row, end_row + 1):
            for col_idx in range(start_col, end_col + 1):
                table_row = row_idx - start_row
                table_col = col_idx - start_col
                
                # Skip merged cells (except the top-left cell of merge)
                if (table_row, table_col) in merged_cells_to_skip:
                    continue
                
                try:
                    excel_cell = ws.cell(row=row_idx, column=col_idx)
                    word_cell = table.cell(table_row, table_col)
                    
                    value = "" if excel_cell.value is None else str(excel_cell.value)
                    
                    # Clear and set text
                    for p in word_cell.paragraphs:
                        p.clear()
                    word_cell.paragraphs[0].add_run(value)
                    
                    # Get the run for formatting
                    run = word_cell.paragraphs[0].runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    
                    # Check for specific content and apply special formatting
                    value_stripped = value.strip()
                    
                    # Header row with dark blue background
                    if value_stripped in ["Auditor's Observation", "Max Marks", "Yes/No", "Marks given by the Auditor", 
                                          "POC Attached", "Sr. No.", "Questions", "Yes/No given by Auditor",
                                          "Marks given by Bank", "Yes/No given by the Bank", "Input", 
                                          "A) Info Sec Processes & Controls", "B) Governance & Policy", 
                                          "C) Vendor Management", "D) Cyber Crisis Management"]:
                        # Dark blue background, white bold text, center aligned
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        
                        # Add dark blue background
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), '00008B')  # Dark blue
                        word_cell._element.get_or_add_tcPr().append(shading_elm)
                    
                    # Apply column-based alignment (based on header)
                    elif table_col in center_aligned_columns:
                        # This column should be center-aligned
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    elif table_col in left_aligned_columns:
                        # This column should be left-aligned
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    else:
                        # Default alignment based on content
                        if value_stripped.replace('.', '').replace('_', '').isdigit() or len(value_stripped) < 5:
                            # Short content or numbers - center
                            word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        else:
                            # Long content - left
                            word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                except Exception as e:
                    print(f"    ⚠️ Error filling cell ({table_row}, {table_col}): {e}")
            
            rows_filled += 1
            if rows_filled % 50 == 0:
                print(f"    📝 Filled {rows_filled}/{num_rows} rows...")
        
        print(f"  ✅ Table filled with {rows_filled} rows")
        
        # Set column widths based on header content
        print(f"  📏 Setting column widths...")
        try:
            # Define width mapping based on header content (in characters/units)
            header_width_map = {
                "Sr. No.": 20,
                "Max Marks": 15,
                "Yes/No": 18,
                "Marks given by the Auditor": 20,  # Changed to 20
                "Marks given by Bank": 15,
                "Yes/No given by the Bank": 20,    # Changed to 20
                "Yes/No given by Auditor": 18,
                "POC Attached": 30,
                "Questions": 50,
                "Auditor's Observation": 50,
                "Input": 20
            }
            
            # Get column widths from first row headers
            for col_idx in range(start_col, end_col + 1):
                table_col = col_idx - start_col
                excel_cell = ws.cell(row=start_row, column=col_idx)
                header_value = "" if excel_cell.value is None else str(excel_cell.value).strip()
                
                # Find matching width
                width = None
                for header_key, header_width in header_width_map.items():
                    if header_key in header_value:
                        width = header_width
                        break
                
                if width:
                    # Set column width (using twips: 1 character ≈ 120 twips)
                    col = table.columns[table_col]
                    for cell in col.cells:
                        cell.width = Inches(width * 0.1)  # Approximate conversion
                    print(f"    Column {table_col} ('{header_value}'): {width} units")
        
        except Exception as e:
            print(f"  ⚠️ Error setting column widths: {e}")
        
        # Insert table before placeholder
        print(f"  🔧 Inserting table into document...")
        parent.insert(parent.index(placeholder_element), table._element)
        
        # Remove placeholder
        print(f"  🗑️ Removing placeholder...")
        parent.remove(placeholder_element)
        
        # Add page break after table
        print(f"  📄 Adding page break after table...")
        try:
            from docx.enum.text import WD_BREAK
            # Find the table we just inserted and add page break after it
            for i, paragraph in enumerate(doc.paragraphs):
                # Find first paragraph after where placeholder was
                pass
            
            # Add page break by creating a new paragraph
            page_break_para = doc.add_paragraph()
            page_break_para.add_run().add_break(WD_BREAK.PAGE)
            print(f"  ✅ Page break added after table")
        except Exception as e:
            print(f"  ⚠️ Could not add page break: {e}")
        
        print(f"  ✅ Table insertion complete!")
        return True
        
    except Exception as e:
        print(f"  ❌ Error inserting table: {e}")
        import traceback
        traceback.print_exc()
        return False

def parse_image_name_bank(filename):
    """Parse image name to extract sorting key"""
    name_without_ext = os.path.splitext(filename)[0]
    
    match = re.match(r'^(\d+)_([0-9]+|[A-Z])(.*)$', name_without_ext, re.IGNORECASE)
    
    if match:
        first_num = int(match.group(1))
        second_part = match.group(2)
        remaining = match.group(3).strip(' _-')
        
        if second_part.isdigit():
            return (first_num, int(second_part), None, remaining, filename)
        else:
            return (first_num, None, second_part.upper(), remaining, filename)
    
    return (999, 999, 'ZZZ', name_without_ext, filename)

def sort_images_bank(image_files):
    """Sort images: number_number first, then number_letter"""
    parsed_images = [parse_image_name_bank(img) for img in image_files]
    
    def sort_key(item):
        first_num, second_num, letter, remaining, filename = item
        
        if second_num is not None:
            return (first_num, 0, second_num, '', remaining)
        else:
            return (first_num, 1, 0, letter if letter else '', remaining)
    
    sorted_images = sorted(parsed_images, key=sort_key)
    return [item[4] for item in sorted_images]

def insert_images_as_annexures_bank(doc, placeholder, sorted_images, image_paths):
    """Insert images as numbered annexures"""
    try:
        print(f"\n  🔍 Searching for placeholder '{placeholder}'...")
        print(f"  📋 Total images to insert: {len(sorted_images)}")
        
        # Find placeholder
        placeholder_para = None
        for i, paragraph in enumerate(doc.paragraphs):
            if placeholder in paragraph.text:
                placeholder_para = paragraph
                print(f"    ✅ Found placeholder at paragraph index {i}")
                break
        
        if not placeholder_para:
            print(f"  ❌ Placeholder '{placeholder}' not found")
            print(f"  📋 Searched {len(doc.paragraphs)} paragraphs")
            return False
        
        parent = placeholder_para._element.getparent()
        placeholder_element = placeholder_para._element
        
        annexure_num = 1
        prev_prefix = None
        images_inserted = 0
        
        print(f"\n  🖼️ Processing images...")
        
        for idx, filename in enumerate(sorted_images, 1):
            print(f"\n    📷 Image {idx}/{len(sorted_images)}: {filename}")
            
            image_path = image_paths.get(filename)
            
            if not image_path:
                print(f"      ⚠️ Image path not found in dictionary")
                continue
            
            if not os.path.exists(image_path):
                print(f"      ⚠️ Image file does not exist: {image_path}")
                continue
            
            print(f"      ✅ Image file exists")
            
            parsed = parse_image_name_bank(filename)
            first_num = parsed[0]
            second_num = parsed[1]
            letter = parsed[2]
            remaining = parsed[3]
            
            # Create prefix
            if second_num is not None:
                current_prefix = f"{first_num}_{second_num}"
                display_prefix = f"{first_num}.{second_num}"
            else:
                current_prefix = f"{first_num}_{letter}"
                display_prefix = f"{first_num}_{letter}"
            
            print(f"      📝 Prefix: {current_prefix}, Display: {display_prefix}")
            
            # Check if new annexure or continuation
            if prev_prefix != current_prefix:
                # New annexure
                print(f"      ✨ Creating new Annexure {annexure_num}")
                annexure_para = doc.add_paragraph()
                annexure_run = annexure_para.add_run(f"Annexure {annexure_num} ({display_prefix} {remaining})")
                annexure_run.font.bold = True
                annexure_run.font.size = Pt(12)
                annexure_run.font.name = 'Times New Roman'
                parent.insert(parent.index(placeholder_element), annexure_para._element)
                
                annexure_num += 1
                prev_prefix = current_prefix
            else:
                # Continuation - show only remaining text (without prefix)
                print(f"      ➕ Continuation of Annexure (same prefix)")
                title_para = doc.add_paragraph()
                # For continuation, show only the remaining text (without prefix) for both VICS and LOC
                # For VICS (number_number): show only "remaining" (e.g., "Hardware Count")
                # For LOC (number_letter): show only "remaining" (e.g., "VAPT Report")
                if second_num is not None:
                    # VICS continuation: show only remaining text (without prefix)
                    full_title = remaining if remaining else f"{display_prefix}"
                else:
                    # LOC continuation: show only remaining text (NOT the letter)
                    full_title = remaining if remaining else letter
                
                title_run = title_para.add_run(full_title)
                title_run.font.size = Pt(12)
                title_run.font.name = 'Times New Roman'
                title_run.font.bold = True
                parent.insert(parent.index(placeholder_element), title_para._element)
            
            # Add blank line before image
            blank_para = doc.add_paragraph()
            parent.insert(parent.index(placeholder_element), blank_para._element)
            
            # Add image with fixed height and proportional width
            try:
                print(f"      🖼️ Adding image...")
                from PIL import Image as PILImage
                
                # Get original image dimensions
                pil_img = PILImage.open(image_path)
                orig_width, orig_height = pil_img.size
                
                # Calculate proportional width for 400px height
                target_height_px = 440
                aspect_ratio = orig_width / orig_height
                target_width_px = target_height_px * aspect_ratio
                
                # Convert to inches (assuming 96 DPI)
                height_inches = target_height_px / 96
                width_inches = target_width_px / 96
                
                print(f"      📏 Original: {orig_width}x{orig_height}px, Target: {target_width_px:.0f}x{target_height_px}px")
                
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                picture = run.add_picture(image_path, height=Inches(height_inches))
                
                # Add 1pt border to image
                try:
                    inline = picture._inline
                    pic = inline.graphic.graphicData.pic
                    
                    # Get or create spPr (shape properties)
                    spPr = pic.spPr
                    if spPr is None:
                        spPr = OxmlElement('pic:spPr')
                        pic.insert(0, spPr)
                    
                    # Add line (border) element
                    ln = OxmlElement('a:ln')
                    ln.set('w', '12700')  # 1pt = 12700 EMUs
                    
                    # Add solid fill for the line
                    solidFill = OxmlElement('a:solidFill')
                    srgbClr = OxmlElement('a:srgbClr')
                    srgbClr.set('val', '000000')  # Black
                    solidFill.append(srgbClr)
                    ln.append(solidFill)
                    
                    # Add line to shape properties
                    spPr.append(ln)
                    
                    print(f"      🖼️ Added 1pt black border to image")
                except Exception as border_err:
                    print(f"      ⚠️ Could not add border: {border_err}")
                    import traceback
                    traceback.print_exc()
                
                parent.insert(parent.index(placeholder_element), img_para._element)
                images_inserted += 1
                print(f"      ✅ Image added successfully")
            except Exception as e:
                print(f"      ❌ Error adding image: {e}")
                import traceback
                traceback.print_exc()
            
            # Add page break (except for last image)
            is_last_image = (idx == len(sorted_images))
            if not is_last_image:
                try:
                    from docx.enum.text import WD_BREAK
                    page_break_para = doc.add_paragraph()
                    page_break_para.add_run().add_break(WD_BREAK.PAGE)
                    parent.insert(parent.index(placeholder_element), page_break_para._element)
                    print(f"      📄 Page break added")
                except Exception as e:
                    print(f"      ⚠️ Error adding page break: {e}")
            else:
                print(f"      ℹ️ Last image - skipping page break")
        
        # Remove placeholder
        print(f"\n  🗑️ Removing placeholder...")
        parent.remove(placeholder_element)
        
        print(f"\n  ✅ Successfully inserted {images_inserted}/{len(sorted_images)} images as annexures")
        return True
        
    except Exception as e:
        print(f"  ❌ Error inserting annexures: {e}")
        import traceback
        traceback.print_exc()
        return False

@gap_assessment_report_bank_input_bp.route('/process_gap_assessment_report_bank_input', methods=['POST'])
def process_gap_assessment_report_bank_input():
    """Process Gap Assessment Report generation with Bank Input"""
    try:
        print("\n" + "="*80)
        print("🚀 Processing Gap Assessment Report with Bank Input")
        print("="*80)
        
        # Get form data
        org_name = request.form.get('organizationName')
        org_name_other = request.form.get('organizationNameOther')
        city = request.form.get('city')
        city_other = request.form.get('cityOther')
        state = request.form.get('state')
        prepared_by_prefix = request.form.get('preparedByPrefix')
        prepared_by_name = request.form.get('preparedByName')
        start_date = request.form.get('startAuditDate')
        end_date = request.form.get('endAuditDate')
        auditee_prefix = request.form.get('auditeePersonPrefix')
        auditee_name = request.form.get('auditeePersonName')
        auditee_designation = request.form.get('auditeeDesignation')
        bank_address = request.form.get('bankAddress')
        emails = request.form.getlist('email[]')
        
        final_org_name = org_name_other if org_name == 'Other' else org_name
        final_city = city_other if city == 'Other' else city
        
        excel_file = request.files.get('excelFile')
        zip_file = request.files.get('zipFile')
        
        if not excel_file or not zip_file:
            return "Missing required files", 400
        
        temp_dir = tempfile.mkdtemp()
        print(f"📁 Created temp directory: {temp_dir}")
        
        try:
            # Save files
            excel_path = os.path.join(temp_dir, 'gap_assessment_bank.xlsx')
            zip_path = os.path.join(temp_dir, 'images_bank.zip')
            
            excel_file.save(excel_path)
            zip_file.save(zip_path)
            print(f"💾 Saved uploaded files")
            
            # Load Excel
            print(f"\n📖 Loading Excel file...")
            wb = load_workbook(excel_path, data_only=True)
            
            vics_ws = None
            loc_ws = None
            
            for sheet_name in wb.sheetnames:
                if 'VICS' in sheet_name.upper():
                    vics_ws = wb[sheet_name]
                    print(f"  ✅ Found VICS worksheet")
                elif 'LOC' in sheet_name.upper():
                    loc_ws = wb[sheet_name]
                    print(f"  ✅ Found LOC worksheet")
            
            # Check LOC level
            loc_level = "level2"
            if loc_ws:
                loc_level = check_loc_cells_bank(loc_ws)
                print(f"  📍 LOC Level: {loc_level}")
            
            areas_covered_text = get_areas_covered_text_bank(loc_level)
            
            # Find LOC last row
            loc_last_row = 1
            if loc_ws:
                for row_num in range(1, loc_ws.max_row + 1):
                    b_cell = loc_ws.cell(row=row_num, column=2).value
                    if b_cell is not None and str(b_cell).strip() != '':
                        loc_last_row = row_num
                print(f"  📊 LOC last row: {loc_last_row}")
            
            # Extract and sort images
            print(f"\n🖼️ Extracting images from ZIP...")
            images_dir = os.path.join(temp_dir, 'images_bank')
            os.makedirs(images_dir, exist_ok=True)
            
            image_files = []
            image_paths = {}
            
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                print(f"  📦 ZIP contains {len(zip_ref.namelist())} files")
                for file_info in zip_ref.namelist():
                    print(f"    Found in ZIP: {file_info}")
                    if file_info.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                        zip_ref.extract(file_info, images_dir)
                        extracted_path = os.path.join(images_dir, file_info)
                        filename = os.path.basename(file_info)
                        image_files.append(filename)
                        image_paths[filename] = extracted_path
                        print(f"      ✅ Extracted image: {filename} -> {extracted_path}")
            
            print(f"\n  📷 Total images extracted: {len(image_files)}")
            
            # Sort images
            print(f"\n  🔢 Sorting images...")
            sorted_images = sort_images_bank(image_files)
            
            print(f"  📋 Sorted image order:")
            for idx, img in enumerate(sorted_images, 1):
                print(f"    {idx}. {img}")
            
            # Separate VICS and LOC images
            print(f"\n  📂 Separating VICS and LOC images...")
            vics_images = []
            loc_images = []
            
            for img in sorted_images:
                parsed = parse_image_name_bank(img)
                print(f"    Analyzing: {img}")
                print(f"      Parsed: first={parsed[0]}, second_num={parsed[1]}, letter={parsed[2]}, remaining={parsed[3]}")
                
                if parsed[1] is not None:  # number_number
                    vics_images.append(img)
                    print(f"      → VICS image (number_number)")
                elif parsed[2] is not None:  # number_letter
                    loc_images.append(img)
                    print(f"      → LOC image (number_letter)")
                else:
                    print(f"      → Unrecognized pattern, skipped")
            
            print(f"\n  📊 Final categorization:")
            print(f"    VICS images ({len(vics_images)}): {vics_images}")
            print(f"    LOC images ({len(loc_images)}): {loc_images}")
            
            # Load Word template
            print(f"\n📄 Loading Word template...")
            template_path = os.path.join('static', 'Formats_and_Catalog', 'Gap_Assessment_of_Cybersecurity.docx')
            doc = Document(template_path)
            
            # Prepare replacements
            prepared_by = f"{prepared_by_prefix} {prepared_by_name}"
            date_of_audit = f"{format_date_with_superscript_bank(start_date)} to {format_date_with_superscript_bank(end_date)}"
            audit_period = get_financial_year_bank(start_date)
            submitted_to = f"{auditee_prefix} {auditee_name}"
            bank_full_address = f"{final_org_name}, {bank_address}"
            email_list = ', '.join(emails)
            
            replacements = {
                'Orrrrrganizzzzaaation_naaammee': final_org_name,
                'Ciiiiittyy': final_city,
                'Staaaate': state,
                'Preeeepaared_byy': prepared_by,
                'Daaaateee_of_audddit': date_of_audit,
                'auuuuudit_period': audit_period,
                'Subbbmitteeed_tooo': submitted_to,
                'Desiggggnation': auditee_designation,
                'Baaank_addrreessss': bank_full_address,
                'emaaaill': email_list,
                'Aarreas_Covvered': areas_covered_text
            }
            
            # Replace text
            print(f"\n📝 Replacing placeholders...")
            for key, value in replacements.items():
                print(f"    {key} → {value[:50]}..." if len(value) > 50 else f"    {key} → {value}")
            
            replace_text_in_document_bank(doc, replacements)
            print(f"  ✅ Replacements complete")
            
            # Add page break before "Annexures For VICS" if it exists
            print(f"\n📄 Adding page breaks before annexure sections...")
            from docx.enum.text import WD_BREAK
            
            for paragraph in doc.paragraphs:
                if "Annexures For VICS" in paragraph.text or "Annexures for VICS" in paragraph.text:
                    # Insert page break before this paragraph
                    parent = paragraph._element.getparent()
                    page_break_para = doc.add_paragraph()
                    page_break_para.add_run().add_break(WD_BREAK.PAGE)
                    parent.insert(parent.index(paragraph._element), page_break_para._element)
                    print(f"  ✅ Added page break before 'Annexures For VICS'")
                    break
            
            for paragraph in doc.paragraphs:
                if "Annexures For LOC" in paragraph.text or "Annexures for LOC" in paragraph.text:
                    # Insert page break before this paragraph
                    parent = paragraph._element.getparent()
                    page_break_para = doc.add_paragraph()
                    page_break_para.add_run().add_break(WD_BREAK.PAGE)
                    parent.insert(parent.index(paragraph._element), page_break_para._element)
                    print(f"  ✅ Added page break before 'Annexures For LOC'")
                    break
            
            # Insert VICS table (Columns A to H)
            if vics_ws:
                print(f"\n📊 Inserting VICS table (A-H)...")
                insert_excel_table_in_document_bank(doc, 'VIIICSSSS_Table', vics_ws, 1, 8, 1, 250)
            
            # Insert LOC table (Columns A to F)
            if loc_ws:
                print(f"\n📊 Inserting LOC table (A-F)...")
                insert_excel_table_in_document_bank(doc, 'LOOCCC_table', loc_ws, 1, 6, 1, loc_last_row)
            
            # Insert VICS annexures
            if vics_images:
                print(f"\n🖼️ Inserting VICS annexures...")
                insert_images_as_annexures_bank(doc, 'Annnnnnnnnnnnnnneessurer', vics_images, image_paths)
            
            # Insert LOC annexures
            if loc_images:
                print(f"\n🖼️ Inserting LOC annexures...")
                insert_images_as_annexures_bank(doc, 'LLLLOCCCC_Annexuuuuree', loc_images, image_paths)
            
            # Remove everything after the marker word
            print(f"\n🗑️ Removing content after marker...")
            marker_word = "alasskdkfndnxkcmskalevndkslalmsnckdkslalakscndkslalskdncksdj"
            
            paragraphs_to_remove = []
            found_marker = False
            
            for i, paragraph in enumerate(doc.paragraphs):
                if found_marker:
                    paragraphs_to_remove.append(paragraph)
                elif marker_word in paragraph.text:
                    # Found the marker, remove this and all following paragraphs
                    found_marker = True
                    paragraphs_to_remove.append(paragraph)
                    print(f"  ✅ Found marker at paragraph {i}")
            
            # Remove paragraphs in reverse order to avoid index issues
            for paragraph in reversed(paragraphs_to_remove):
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
            
            if found_marker:
                print(f"  ✅ Removed {len(paragraphs_to_remove)} paragraphs after marker")
            else:
                print(f"  ℹ️ Marker word not found in document")
            
            # Save document
            output_path = os.path.join(temp_dir, 'Gap_Assessment_Report_with_Bank_Input.docx')
            doc.save(output_path)
            print(f"\n✅ Document saved")
            print("="*80)
            
            return send_file(
                output_path,
                as_attachment=True,
                download_name='Gap_Assessment_Report_with_Bank_Input.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
        finally:
            pass
    
    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return f"Error: {str(e)}", 500

