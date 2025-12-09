import copy
import io
import os
import re
import tempfile
import traceback
import zipfile
import shutil
from datetime import datetime, timedelta
from lxml import etree
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor
import pandas as pd
from openpyxl import load_workbook
import openpyxl
from PIL import Image as PILImage
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import qn
from docx.oxml.ns import nsdecls
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from flask import Blueprint, request, send_file, jsonify


word_report_bp = Blueprint('word_report', __name__)

# Helper function to convert NaN values to "NA"
def clean_value(value):
    """Convert NaN, None, empty strings to 'NA'"""
    if pd.isna(value) or value is None or str(value).lower() in ['nan', 'none', '']:
        return "NA"
    return str(value)

# Helper to convert openpyxl color to RGB
def get_rgb(color):
    if color is None or color.type != "rgb":
        return None
    if len(color.rgb) > 6:
        return color.rgb[2:]
    return color.rgb

def apply_1pt_border_to_picture(picture):
    """
    Apply a 1pt solid black border to a picture object.
    Works for all images added via add_picture().
    
    Args:
        picture: The picture object returned by run.add_picture()
    """
    try:
        # Access the picture element
        pic = picture._inline.graphic.graphicData.pic
        spPr = pic.spPr
        
        # Remove any existing line borders
        for ln in spPr.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}ln'):
            spPr.remove(ln)
        
        # Create 1pt border with the width attribute directly in the tag
        from docx.oxml import parse_xml
        border = parse_xml(
            '<a:ln xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" w="12700">'
            '<a:solidFill><a:srgbClr val="000000"/></a:solidFill>'
            '<a:prstDash val="solid"/>'
            '</a:ln>'
        )
        spPr.append(border)
        
        print(f"✅ Applied 1pt border (12700 EMUs)")
        
    except Exception as e:
        print(f"❌ Error applying border: {e}")
        import traceback
        traceback.print_exc()

def add_borders_to_cell(cell):
    """Add borders to a table cell"""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Create border elements
        borders = ['top', 'left', 'bottom', 'right']
        for border_name in borders:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcPr.append(border)
    except Exception as e:
        print(f"Warning: Could not add borders to cell: {e}")

def add_borders_to_row(row):
    """Add borders to all cells in a table row"""
    for cell in row.cells:
        add_borders_to_cell(cell)

from datetime import datetime, timedelta

def convert_to_dd_mm_yyyy(date_str):
    """Convert various date formats to DD.MM.YYYY format"""
    date_formats = ["%d.%m.%Y", "%Y-%m-%d", "%m/%d/%Y", "%d-%m-%Y", "%B %Y", "%b %Y"]
    for fmt in date_formats:
        try:
            date_obj = datetime.strptime(date_str, fmt)
            return date_obj.strftime("%d.%m.%Y")
        except ValueError:
            continue
    try:
        year = int(date_str)
        return f"01.01.{year}"
    except ValueError:
        raise ValueError(f"Could not parse date: {date_str}")

def generate_vapt_timeline(start_date_str, end_date_str):
    """Generate VAPT timeline between start_date and end_date, skipping weekends and bank holidays"""
    
    BANK_HOLIDAYS = [
        "01.01.2025",  # New Year's Day
        "15.08.2025",  # Independence Day
        "02.10.2025",  # Gandhi Jayanti
    ]
    
    def is_working_day(date):
        date_str = date.strftime("%d.%m.%Y")
        return date.weekday() < 5 and date_str not in BANK_HOLIDAYS

    # Parse dates
    try:
        start_date = datetime.strptime(convert_to_dd_mm_yyyy(start_date_str), "%d.%m.%Y")
        end_date = datetime.strptime(convert_to_dd_mm_yyyy(end_date_str), "%d.%m.%Y")
    except Exception as e:
        print(f"Error parsing dates: {e}")
        today = datetime.now()
        start_date = today.replace(day=1)
        end_date = today

    # Collect working days
    all_dates = [d for d in (start_date + timedelta(days=i) for i in range((end_date - start_date).days + 1)) if is_working_day(d)]
    total_working_days = len(all_dates)
    if total_working_days == 0:
        return []

    phase_order = [
        "Planning",
        "Environment Setup",
        "Pre-Assessment Preparation",
        "VAPT Execution",
        "Reporting"
    ]

    timeline = []
    date_index = 0

    if total_working_days < len(phase_order):
        # Less days than phases: assign phases sequentially, sharing days at the end
        for i, phase in enumerate(phase_order):
            # Assign the next available day, reuse last day if we run out
            assigned_day = all_dates[min(date_index, total_working_days - 1)]
            start = assigned_day.strftime("%d.%m.%Y")
            end = start
            timeline.append({"Phase": phase, "Dates": f"{start} - {end}", "Start": start, "End": end})
            date_index += 1
    else:
        # Enough days: allocate normally
        remaining_days = total_working_days
        phase_allocation = {
            "Planning": 1,
            "Environment Setup": 1,
            "Pre-Assessment Preparation": 2,
            "Reporting": 1
        }
        allocated = sum(phase_allocation.values())
        phase_allocation["VAPT Execution"] = remaining_days - allocated

        for phase in phase_order:
            days_needed = phase_allocation.get(phase, 0)
            if days_needed <= 0 or date_index >= len(all_dates):
                continue
            end_index = min(date_index + days_needed, len(all_dates))
            phase_dates = all_dates[date_index:end_index]
            start = phase_dates[0].strftime("%d.%m.%Y")
            end = phase_dates[-1].strftime("%d.%m.%Y")
            timeline.append({"Phase": phase, "Dates": f"{start} - {end}", "Start": start, "End": end})
            date_index = end_index

    return timeline


def find_and_print_metadata(file_path):
    """
    Extract metadata from Excel file's Meta_Data worksheet
    """
    try:
        # Load the workbook
        workbook = openpyxl.load_workbook(file_path)
        
        # Check if 'Meta_Data' worksheet exists
        if 'Meta_Data' not in workbook.sheetnames:
            print("Error: 'Meta_Data' worksheet not found!")
            return {}
        
        # Get the Meta_Data worksheet
        sheet = workbook['Meta_Data']
        
        # Dictionary to store the values we're looking for
        target_values = {
            "Organization Name": None,
            "City": None,
            "State": None,
            "Start Date": None,
            "End Date": None
        }
        
        # Variables to store the special values
        report_prepared_by_value = None
        auditee_details_value1 = None  # First value (diagonal cell)
        auditee_details_value2 = None  # Second value (two steps down and right)
        bank_email_addresses = []      # List to store all email addresses
        auditing_team_members = []     # List to store all auditing team member data
        
        # Iterate through all cells in the worksheet
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value is not None:
                    # Check if cell value matches any of our target keys
                    for key in target_values.keys():
                        if str(cell.value).strip() == key:
                            # Get the adjacent cell (next column, same row)
                            adjacent_cell = sheet.cell(row=cell.row, column=cell.column + 1)
                            if adjacent_cell.value is not None:
                                target_values[key] = adjacent_cell.value
                            break
                    
                    # Special handling for "REPORT PREPARED BY" (separate from other keys)
                    if str(cell.value).strip() == "REPORT PREPARED BY":
                        # Get the cell below and to the right (next column, next row)
                        diagonal_cell = sheet.cell(row=cell.row + 1, column=cell.column + 1)
                        if diagonal_cell.value is not None:
                            report_prepared_by_value = diagonal_cell.value
                    
                    # Special handling for "AUDITEE DETAILS"
                    if str(cell.value).strip() == "AUDITEE DETAILS":
                        # Get the first value: cell below and to the right (next column, next row)
                        diagonal_cell1 = sheet.cell(row=cell.row + 1, column=cell.column + 1)
                        if diagonal_cell1.value is not None:
                            auditee_details_value1 = diagonal_cell1.value
                        
                        # Get the second value: two steps down and to the right (next column, row+2)
                        diagonal_cell2 = sheet.cell(row=cell.row + 2, column=cell.column + 1)
                        if diagonal_cell2.value is not None:
                            auditee_details_value2 = diagonal_cell2.value
                    
                    # Special handling for "BANK EMAIL ADDRESSES"
                    if str(cell.value).strip() == "BANK EMAIL ADDRESSES":
                        # Start from the cell below and to the right (next column, next row)
                        current_row = cell.row + 1
                        current_col = cell.column + 1
                        
                        # Keep reading cells vertically until an empty cell is found
                        while True:
                            email_cell = sheet.cell(row=current_row, column=current_col)
                            if email_cell.value is not None and str(email_cell.value).strip() != "":
                                bank_email_addresses.append(str(email_cell.value).strip())
                                current_row += 1  # Move to next row
                            else:
                                break  # Stop when empty cell is found
                    
                    # Special handling for "AUDITING TEAM MEMBER" entries
                    if "AUDITING TEAM MEMBER" in str(cell.value).strip():
                        member_data = {}
                        member_number = str(cell.value).strip().split()[-1]  # Get the number (1, 2, 3, etc.)
                        
                        # Define the labels for each row
                        labels = [
                            f"Team Member {member_number} - Name",
                            f"Team Member {member_number} - Designation", 
                            f"Team Member {member_number} - Email",
                            f"Team Member {member_number} - Qualification",
                            f"Team Member {member_number} - Certified"
                        ]
                        
                        # Start from the cell below and to the right (next column, next row)
                        current_row = cell.row + 1
                        current_col = cell.column + 1
                        label_index = 0
                        
                        # Read up to 5 cells or until empty cell is found
                        for i in range(5):
                            member_cell = sheet.cell(row=current_row + i, column=current_col)
                            if member_cell.value is not None and str(member_cell.value).strip() != "":
                                if label_index < len(labels):
                                    member_data[labels[label_index]] = str(member_cell.value).strip()
                                    label_index += 1
                            else:
                                break  # Stop when empty cell is found
                        
                        # Add member data to the list if we found any values
                        if member_data:
                            auditing_team_members.append(member_data)
        
        # Create result dictionary
        result = {
            "organization_name": target_values.get("Organization Name"),
            "city": target_values.get("City"),
            "state": target_values.get("State"),
            "start_date": target_values.get("Start Date"),
            "end_date": target_values.get("End Date"),
            "report_prepared_by": report_prepared_by_value,
            "auditee_details_1": auditee_details_value1,
            "auditee_details_2": auditee_details_value2,
            "bank_email_addresses": bank_email_addresses,
            "auditing_team_members": auditing_team_members
        }
        
        # Close the workbook
        workbook.close()
        
        return result
        
    except FileNotFoundError:
        print(f"Error: File '{file_path}' not found!")
        return {}
    except Exception as e:
        print(f"An error occurred: {str(e)}")
        return {}

# NEW FUNCTION: Highlights cells with specific keywords
def highlight_keywords(word_cell, value):
    """
    Checks if a cell's text contains specific keywords and changes its background color to blue.
    """
    if value.strip().upper() in ["HOST", "PORT", "SERVICE"]:
        blue_rgb = "#1376d1"  # Hex code for blue
        tc = word_cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), blue_rgb)
        tcPr.append(shd)

def create_nmap_table_from_excel(doc, excel_file):
    """
    Create a formatted table from the 'Nmap Files' worksheet in Excel
    Returns the table object without adding it to document
    """
    try:
        # Load Excel workbook and get the 'Nmap Files' worksheet
        wb = load_workbook(excel_file)
        
        # Try to get the 'Nmap Files' worksheet, fallback to active sheet if not found
        try:
            ws = wb['Nmap Files']
        except KeyError:
            print("Worksheet 'Nmap Files' not found, using active sheet")
            ws = wb.active
        
        # Create a table with Excel size (but don't add it to doc yet)
        table = doc.add_table(rows=ws.max_row, cols=ws.max_column)
        table.style = 'Table Grid'

        # A set to keep track of merged cells to avoid overwriting them
        merged_cells_to_skip = set()

        # First, handle merged cells
        for merged_range in ws.merged_cells:
            merged_range_coords = list(merged_range.cells)
            
            top_left_cell_ref = merged_range_coords[0]
            start_row = top_left_cell_ref[0] - 1
            start_col = top_left_cell_ref[1] - 1
            
            bottom_right_cell_ref = merged_range_coords[-1]
            end_row = bottom_right_cell_ref[0] - 1
            end_col = bottom_right_cell_ref[1] - 1
            
            start_word_cell = table.cell(start_row, start_col)
            end_word_cell = table.cell(end_row, end_col)
            
            start_word_cell.merge(end_word_cell)
            
            for r_idx in range(start_row, end_row + 1):
                for c_idx in range(start_col, end_col + 1):
                    if (r_idx, c_idx) != (start_row, start_col):
                        merged_cells_to_skip.add((r_idx, c_idx))

        # Fill Word table with Excel data and formatting
        for i, row in enumerate(ws.iter_rows(values_only=False)):
            for j, cell in enumerate(row):
                if (i, j) in merged_cells_to_skip:
                    continue
                    
                word_cell = table.cell(i, j)
                value = "" if cell.value is None else str(cell.value)
                word_cell.text = value

                # Format font
                if cell.font:
                    # Ensure we have at least one run in the paragraph
                    if not word_cell.paragraphs[0].runs:
                        word_cell.paragraphs[0].add_run(value)
                    
                    run = word_cell.paragraphs[0].runs[0]
                    run.font.name = cell.font.name or "Calibri"
                    run.font.size = Pt(12)
                    run.font.bold = cell.font.b
                    run.font.italic = cell.font.i
                    if cell.font.color:
                        rgb = get_rgb(cell.font.color)
                        if rgb:
                            run.font.color.rgb = RGBColor.from_string(rgb)

                # Set horizontal and vertical alignment
                word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Call the function to highlight keywords
                highlight_keywords(word_cell, value)

        print("✅ Nmap table created successfully from Excel")
        return table
        
    except Exception as e:
        print(f"Error creating Nmap table from Excel: {e}")
        traceback.print_exc()
        return None
def normalize_vulnerability_name_for_filename(vuln_name):
    """
    Normalize vulnerability name to be compatible with file names by replacing
    invalid characters with dashes.
    
    Args:
        vuln_name (str): Original vulnerability name
        
    Returns:
        str: Normalized vulnerability name safe for file names
    """
    if not vuln_name:
        return vuln_name
    
    # Characters that are not allowed in file names
    invalid_chars = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
    
    normalized_name = str(vuln_name)
    
    # Replace invalid characters with dashes
    for char in invalid_chars:
        normalized_name = normalized_name.replace(char, '-')
    
    # Remove multiple consecutive dashes and trim
    normalized_name = re.sub(r'-+', '-', normalized_name)
    normalized_name = normalized_name.strip('-')
    
    return normalized_name

def extract_poc_images_from_excel(excel_file):
    """
    Extract images from Excel POC column and associate them with vulnerabilities.

    Args:
        excel_file (str): Path to Excel file

    Returns:
        tuple: A tuple containing a dictionary mapping vulnerability names to a list of image data,
               and the pandas DataFrame of the Excel data.
    """
    try:
        wb = load_workbook(excel_file)
        ws = wb['Infra_VAPT']
    except KeyError:
        print("Worksheet 'Infra_VAPT' not found in Excel file. Returning empty data.")
        return {}, pd.DataFrame()
    except Exception as e:
        print(f"Error loading Excel file: {e}")
        return {}, pd.DataFrame()

    df = pd.read_excel(excel_file, sheet_name='Infra_VAPT')
    vulnerability_images = extract_poc_images_from_worksheet(ws, df)

    return vulnerability_images, df

def extract_poc_images_from_worksheet(worksheet, df):
    """
    Extract images from the POC columns (L, M, N, O, P, Q, R) and associate them with vulnerabilities.
    """
    vulnerability_images = {}
    
    # Find POC column range by looking for merged cells in row 1
    poc_col_start = None
    poc_col_end = None
    
    for merged_range in worksheet.merged_cells.ranges:
        if merged_range.min_row == 1 and merged_range.max_row == 1:
            # Check if this merged cell contains "POC"
            first_cell = worksheet.cell(row=1, column=merged_range.min_col)
            if first_cell.value and "POC" in str(first_cell.value).upper():
                poc_col_start = merged_range.min_col
                poc_col_end = merged_range.max_col
                break
    
    # Fallback: if no merged cell found, look for single column with "POC"
    if poc_col_start is None or poc_col_end is None:
        for col_idx, cell in enumerate(worksheet[1], 1):
            if cell.value and "POC" in str(cell.value).upper():
                poc_col_start = col_idx
                poc_col_end = col_idx
                break
    
    if poc_col_start is None or poc_col_end is None:
        print("POC column(s) not found in worksheet")
        return vulnerability_images
    
    print(f"Found POC columns from {poc_col_start} to {poc_col_end}")
    images_by_cell = extract_images_by_cell_coordinates(worksheet, poc_col_start, poc_col_end)
    
    for row_idx, image_data_list in images_by_cell.items():
        if row_idx <= 1:
            continue
            
        df_row_index = row_idx - 2
        
        if 0 <= df_row_index < len(df):
            vuln_name_cell = df.iloc[df_row_index]['Name of Vulnerability']
            if pd.notna(vuln_name_cell):
                vuln_name = str(vuln_name_cell)
                # Normalize the vulnerability name for filename compatibility
                normalized_vuln_name = normalize_vulnerability_name_for_filename(vuln_name)
                
                # Store both original and normalized names for flexible matching
                if vuln_name not in vulnerability_images:
                    vulnerability_images[vuln_name] = []
                if normalized_vuln_name not in vulnerability_images:
                    vulnerability_images[normalized_vuln_name] = []
                
                # Associates images with both original and normalized vulnerability names
                vulnerability_images[vuln_name].extend(image_data_list)
                vulnerability_images[normalized_vuln_name].extend(image_data_list)
        
    return vulnerability_images

def extract_images_by_cell_coordinates(worksheet, poc_col_start, poc_col_end):
    """
    Extract all images from worksheet POC columns and group them by cell coordinates.
    Now handles multiple POC columns (L, M, N, O, P, Q, R).
    """
    images_by_cell = {}
    
    if hasattr(worksheet, '_drawing') and worksheet._drawing:
        try:
            drawing = worksheet._drawing
            for shape in drawing._shapes:
                if hasattr(shape, 'anchor') and hasattr(shape.anchor, '_from'):
                    cell_ref = shape.anchor._from
                    row_idx = cell_ref.row + 1
                    col_idx = cell_ref.col + 1
                    
                    # Check if column is within POC column range
                    if poc_col_start <= col_idx <= poc_col_end:
                        if hasattr(shape, '_blip') and shape._blip:
                            image_data = shape._blip._blob
                            if image_data:
                                if row_idx not in images_by_cell:
                                    images_by_cell[row_idx] = []
                                images_by_cell[row_idx].append(image_data)
        except Exception as e:
            pass
    
    if hasattr(worksheet, '_images'):
        for img_obj in worksheet._images:
            try:
                row_idx = img_obj.anchor._from.row + 1
                col_idx = img_obj.anchor._from.col + 1
                
                # Check if column is within POC column range
                if poc_col_start <= col_idx <= poc_col_end:
                    image_data = None
                    if hasattr(img_obj, '_data'):
                        try:
                            image_data = img_obj._data() if callable(img_obj._data) else img_obj._data
                        except:
                            pass
                    
                    if image_data:
                        if row_idx not in images_by_cell:
                            images_by_cell[row_idx] = []
                        images_by_cell[row_idx].append(image_data)
                        
            except Exception as e:
                pass
    
    return images_by_cell

def copy_table_with_formatting(source_table):
    """
    Create a deep copy of a table while ensuring all formatting is preserved.
    """
    copied_table = copy.deepcopy(source_table)
    return copied_table

def _replace_text_in_runs_with_bold(paragraph, old_text, new_text, make_bold=False):
    """
    Helper function to replace text in a paragraph while preserving formatting.
    This handles cases where a single placeholder is split across multiple runs.
    If make_bold is True, the new text will be bold.
    """
    # Combine all run texts to a single string for easy replacement
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return False

    # Get the formatting of the runs before we delete them
    run_formats = []
    for run in paragraph.runs:
        run_formats.append({
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': 'calibri',
            'font_size': run.font.size,
            'font_color': run.font.color.rgb
        })

    # Clear the paragraph
    for run in paragraph.runs:
        run.clear()
    
    # Replace the text
    new_full_text = full_text.replace(old_text, new_text)

    # Rebuild the runs with the new text and old formatting
    new_runs_texts = re.split(r'(' + re.escape(new_text) + r')', new_full_text)
    
    # This is a simplified approach, a more complex one would be needed
    # for perfect preservation of multiple run formatting within a replaced block.
    # For now, we apply the formatting of the first run.
    first_run_format = run_formats[0] if run_formats else {}
    
    for segment in new_runs_texts:
        if segment:
            new_run = paragraph.add_run(segment)
            if segment == new_text and make_bold:
                # Make the replaced text bold
                new_run.bold = True
                new_run.italic = first_run_format.get('italic')
                new_run.underline = first_run_format.get('underline')
                new_run.font.name = first_run_format.get('font_name')
                new_run.font.size = first_run_format.get('font_size')
                new_run.font.color.rgb = first_run_format.get('font_color')
            else:
                # Apply original formatting
                new_run.bold = first_run_format.get('bold')
                new_run.italic = first_run_format.get('italic')
                new_run.underline = first_run_format.get('underline')
                new_run.font.name = first_run_format.get('font_name')
                new_run.font.size = first_run_format.get('font_size')
                new_run.font.color.rgb = first_run_format.get('font_color')

    return True

def _replace_text_in_runs(paragraph, old_text, new_text):
    """
    Helper function to replace text in a paragraph while preserving formatting.
    This handles cases where a single placeholder is split across multiple runs.
    """
    # Combine all run texts to a single string for easy replacement
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return False

    # Get the formatting of the runs before we delete them
    run_formats = []
    for run in paragraph.runs:
        # Safely get font color to avoid None reference errors
        font_color = None
        try:
            if run.font.color.rgb is not None:
                font_color = run.font.color.rgb
        except:
            font_color = None
            
        run_formats.append({
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': 'calibri',
            'font_size': run.font.size,
            'font_color': font_color
        })

    # Clear the paragraph safely
    try:
        # Clear runs one by one to avoid corruption
        runs_to_clear = list(paragraph.runs)  # Create a copy to avoid modification during iteration
        for run in runs_to_clear:
            run.clear()
    except Exception as e:
        print(f"Warning: Error clearing paragraph runs: {e}")
        # Fallback: clear all runs at once
        for run in paragraph.runs:
            run.clear()
    
    # Replace the text
    new_full_text = full_text.replace(old_text, new_text)

    # Rebuild the runs with the new text and old formatting
    new_runs_texts = re.split(r'(' + re.escape(new_text) + r')', new_full_text)
    
    # This is a simplified approach, a more complex one would be needed
    # for perfect preservation of multiple run formatting within a replaced block.
    # For now, we apply the formatting of the first run.
    first_run_format = run_formats[0] if run_formats else {}
    
    for segment in new_runs_texts:
        if segment:
            new_run = paragraph.add_run(segment)
            new_run.bold = first_run_format.get('bold')
            new_run.italic = first_run_format.get('italic')
            new_run.underline = first_run_format.get('underline')
            new_run.font.name = first_run_format.get('font_name')
            new_run.font.size = first_run_format.get('font_size')
            # Safely set font color
            if first_run_format.get('font_color') is not None:
                try:
                    new_run.font.color.rgb = first_run_format.get('font_color')
                except:
                    pass  # Skip color setting if it fails

    return True

def replace_text_in_table(table, old_text, new_text):
    """
    Replace text in table cells, handling placeholders that might be split across runs.
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_text_in_runs(paragraph, old_text, new_text)

def format_date_for_month_year(date_str):
    """
    Format date string to 'Month Year' format (e.g., 'September 2025', 'March 2024')
    """
    if not date_str:
        return ""
    
    try:
        from datetime import datetime
        # Try different date formats
        date_formats = [
            '%Y-%m-%d', '%d-%m-%Y', '%m/%d/%Y', '%d/%m/%Y',
            '%Y/%m/%d', '%d-%m-%y', '%m-%d-%Y', '%d.%m.%Y',
            '%Y.%m.%d', '%d %m %Y', '%Y %m %d'
        ]
        
        date_obj = None
        for fmt in date_formats:
            try:
                date_obj = datetime.strptime(str(date_str).strip(), fmt)
                break
            except ValueError:
                continue
        
        if date_obj:
            return date_obj.strftime('%B %Y')
        else:
            return str(date_str)
    except Exception as e:
        print(f"Error formatting date '{date_str}': {e}")
        return str(date_str)

def format_audit_date_period(start_date_str, end_date_str):
    """
    Format audit date period to 'DD.MM.YYYY - DD.MM.YYYY' format
    """
    if not start_date_str or not end_date_str:
        return ""
    
    try:
        formatted_start = convert_to_dd_mm_yyyy(str(start_date_str).strip())
        formatted_end = convert_to_dd_mm_yyyy(str(end_date_str).strip())
        return f"{formatted_start} - {formatted_end}"
    except Exception as e:
        print(f"Error formatting audit date period '{start_date_str}' to '{end_date_str}': {e}")
        return f"{start_date_str} - {end_date_str}"

def format_date_for_range(start_date_str, end_date_str):
    """
    Format date range to 'DDth MMMM YYYY to DD MMMM YYYY' format (e.g., '12th September 2025 to 30 September 2025')
    """
    if not start_date_str or not end_date_str:
        return ""
    
    try:
        from datetime import datetime
        
        # Try different date formats
        date_formats = [
            '%Y-%m-%d', '%d-%m-%Y', '%m/%d/%Y', '%d/%m/%Y',
            '%Y/%m/%d', '%d-%m-%y', '%m-%d-%Y', '%d.%m.%Y',
            '%Y.%m.%d', '%d %m %Y', '%Y %m %d'
        ]
        
        start_date_obj = None
        end_date_obj = None
        
        # Parse start date
        for fmt in date_formats:
            try:
                start_date_obj = datetime.strptime(str(start_date_str).strip(), fmt)
                break
            except ValueError:
                continue
        
        # Parse end date
        for fmt in date_formats:
            try:
                end_date_obj = datetime.strptime(str(end_date_str).strip(), fmt)
                break
            except ValueError:
                continue
        
        if start_date_obj and end_date_obj:
            # Format day with leading zero and superscript ordinal suffix
            def get_ordinal_suffix(day):
                if 10 <= day % 100 <= 20:
                    suffix = 'th'
                else:
                    suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
                
                # Convert suffix to superscript using Unicode characters
                superscript_map = {
                    's': 'ˢ', 't': 'ᵗ', 'n': 'ⁿ', 'd': 'ᵈ', 'r': 'ʳ', 'h': 'ʰ'
                }
                superscript_suffix = ''.join(superscript_map.get(c, c) for c in suffix)
                
                # Add leading zero for single digit days
                day_str = f"{day:02d}"
                return f"{day_str}{superscript_suffix}"
            
            start_day = get_ordinal_suffix(start_date_obj.day)
            end_day = get_ordinal_suffix(end_date_obj.day)
            
            start_formatted = f"{start_day} {start_date_obj.strftime('%B %Y')}"
            end_formatted = f"{end_day} {end_date_obj.strftime('%B %Y')}"
            
            return f"{start_formatted} to {end_formatted}"
        else:
            return f"{start_date_str} to {end_date_str}"
    except Exception as e:
        print(f"Error formatting date range '{start_date_str}' to '{end_date_str}': {e}")
        return f"{start_date_str} to {end_date_str}"

def replace_auditor_placeholders_and_add_rows(doc, metadata):
    """
    Replace auditor placeholders and add additional rows for multiple team members
    """
    try:
        auditing_team_members = metadata.get("auditing_team_members", [])
        
        if not auditing_team_members:
            print("No auditing team members found")
            return
        
        # Find the table containing auditor placeholders
        target_table = None
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Auditorrrrr" in cell.text:
                        target_table = table
                        break
                if target_table:
                    break
            if target_table:
                break
        
        if not target_table:
            print("Table with auditor placeholders not found")
            return
        
        # Find the row with auditor placeholders
        target_row = None
        column_mapping = {}  # Store which column contains which placeholder
        
        for row in target_table.rows:
            # Check if this row contains auditor placeholders
            has_auditor_placeholders = False
            for cell in row.cells:
                if "Auditorrrrr" in cell.text:
                    has_auditor_placeholders = True
                    target_row = row
                    break
            
            if has_auditor_placeholders:
                # Map all columns in this row
                for j, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        if "Auditorrrrr_Desi" in paragraph.text:
                            column_mapping[j] = "designation"
                        elif "Auditorrrrr_email" in paragraph.text:
                            column_mapping[j] = "email"
                        elif "Auditorrrrr_Qua" in paragraph.text:
                            column_mapping[j] = "qualification"
                        elif "Auditorrrrr_Cert" in paragraph.text:
                            column_mapping[j] = "certified"
                        elif "Auditorrrrr" in paragraph.text and "Desi" not in paragraph.text and "email" not in paragraph.text and "Qua" not in paragraph.text and "Cert" not in paragraph.text:
                            column_mapping[j] = "name"
                break
        
        if not target_row:
            print("Row with auditor placeholders not found")
            return
        
        print(f"📋 Column mapping detected: {column_mapping}")
        
        # Replace placeholders in the first row with fixed values
        fixed_auditor_data = {
            "name": "Niraj Goyal",
            "designation": "CEO / Director", 
            "email": "admin@ngtech.co.in",
            "qualification": "CA, CEH, DISA",
            "certified": "Yes"
        }
        
        for j, cell in enumerate(target_row.cells):
            if j in column_mapping:
                column_type = column_mapping[j]
                for paragraph in cell.paragraphs:
                    # Replace Team Member 1 placeholders with fixed values
                    if column_type == "name":
                        _replace_text_in_runs(paragraph, "Auditorrrrr", fixed_auditor_data["name"])
                    elif column_type == "designation":
                        _replace_text_in_runs(paragraph, "Auditorrrrr_Desi", fixed_auditor_data["designation"])
                    elif column_type == "email":
                        _replace_text_in_runs(paragraph, "Auditorrrrr_email", fixed_auditor_data["email"])
                    elif column_type == "qualification":
                        _replace_text_in_runs(paragraph, "Auditorrrrr_Qua", fixed_auditor_data["qualification"])
                    elif column_type == "certified":
                        _replace_text_in_runs(paragraph, "Auditorrrrr_Cert", fixed_auditor_data["certified"])
        
        # Add additional rows for Team Member 2, 3, etc. (limit to 8 additional rows)
        # Since first row is fixed, we start from index 0 of auditing_team_members for Team Member 2
        max_additional_rows = min(8, len(auditing_team_members))
        
        for i in range(max_additional_rows):
            member_data = auditing_team_members[i]
            member_number = i + 2  # Start from Team Member 2
            
            print(f"🔍 Processing Team Member {member_number} (Excel index {i})")
            print(f"📊 Available keys in member_data: {list(member_data.keys())}")
            
            # Create a new row by copying the target row
            new_row = target_table.add_row()
            
            # Copy the structure from the target row
            for j, cell in enumerate(new_row.cells):
                if j < len(target_row.cells):
                    # Safely clear the new cell content
                    try:
                        # Remove all existing paragraphs to avoid extra spacing
                        while len(cell.paragraphs) > 0:
                            p = cell.paragraphs[0]
                            p._element.getparent().remove(p._element)
                    except Exception as e:
                        print(f"Warning: Could not clear cell content: {e}")
                        # Fallback: clear the entire cell element
                        cell._element.clear()
                    
                    # Copy formatting and add content
                    original_cell = target_row.cells[j]
                    
                    # Determine content based on column mapping
                    new_content = ""
                    if j in column_mapping:
                        column_type = column_mapping[j]
                        # The Excel data is stored with the actual member number from Excel (1, 2, 3, etc.)
                        # We need to find the correct key in the member_data
                        excel_member_number = None
                        for key in member_data.keys():
                            if key.startswith("Team Member ") and key.endswith(" - Name"):
                                # Extract the number from the key
                                try:
                                    excel_member_number = key.split("Team Member ")[1].split(" - Name")[0]
                                    break
                                except:
                                    continue
                        
                        if excel_member_number:
                            if column_type == "name":
                                new_content = member_data.get(f"Team Member {excel_member_number} - Name", "")
                            elif column_type == "designation":
                                new_content = member_data.get(f"Team Member {excel_member_number} - Designation", "")
                            elif column_type == "email":
                                new_content = member_data.get(f"Team Member {excel_member_number} - Email", "")
                            elif column_type == "qualification":
                                new_content = member_data.get(f"Team Member {excel_member_number} - Qualification", "")
                            elif column_type == "certified":
                                certified_value = member_data.get(f"Team Member {excel_member_number} - Certified", "")
                                # Convert to proper case (Yes/No)
                                if certified_value.lower() == "yes":
                                    new_content = "Yes"
                                elif certified_value.lower() == "no":
                                    new_content = "No"
                                else:
                                    new_content = certified_value  # Keep original if not yes/no
                            
                            print(f"📝 {column_type}: '{new_content}' (from key: Team Member {excel_member_number} - {column_type.title()})")
                        else:
                            print(f"⚠️  No Excel member number found for Team Member {member_number}")
                    else:
                        # Check if this is the first column (Sr. No. column)
                        if j == 0 and original_cell.text.strip().isdigit():
                            # This is the Sr. No. column, update the number
                            new_content = str(member_number)
                        else:
                            # Keep original text for other non-auditor columns
                            new_content = original_cell.text
                    
                    # Add content with clean formatting
                    if new_content:
                        new_paragraph = cell.add_paragraph()
                        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Remove any extra spacing - set to exactly 0
                        new_paragraph.paragraph_format.space_before = Pt(0)
                        new_paragraph.paragraph_format.space_after = Pt(0)
                        new_paragraph.paragraph_format.line_spacing = 1.0
                        new_paragraph.paragraph_format.first_line_indent = Pt(0)
                        new_paragraph.paragraph_format.left_indent = Pt(0)
                        new_paragraph.paragraph_format.right_indent = Pt(0)
                        
                        # Ensure no extra spacing is applied
                        new_paragraph.paragraph_format.space_before_auto = False
                        new_paragraph.paragraph_format.space_after_auto = False
                        
                        # Set consistent formatting for all new rows
                        new_run = new_paragraph.add_run(new_content)
                        new_run.font.name = 'Calibri (Body)'
                        new_run.font.size = Pt(12)
                        new_run.font.bold = False
                        new_run.font.italic = False
                        new_run.font.underline = False
                        
                        # Set cell vertical alignment
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        
                        # Set cell padding (top and bottom 1px equivalent)
                        try:
                            from docx.oxml.shared import qn
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            
                            # Set cell margins for padding - use minimal margins to avoid extra spacing
                            from docx.oxml import parse_xml
                            cell_margin_xml = f'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:w="0" w:type="dxa"/><w:bottom w:w="0" w:type="dxa"/></w:tcMar>'
                            cell_margin = parse_xml(cell_margin_xml)
                            tcPr.append(cell_margin)
                        except Exception as e:
                            print(f"Warning: Could not set cell margins: {e}")
                            # Continue without setting margins to avoid corruption
        
        print(f"✅ Added {max_additional_rows + 1} auditing team member(s) to table (max 8 additional rows)")
        
    except Exception as e:
        print(f"❌ Error replacing auditor placeholders: {str(e)}")
        traceback.print_exc()

def replace_metadata_placeholders(doc, metadata):
    """
    Replace metadata placeholders in the Word document with extracted values
    """
    try:
        # Prepare bank email addresses as a single string
        bank_emails = ""
        if metadata.get("bank_email_addresses"):
            bank_emails = "\n".join(metadata.get("bank_email_addresses", []))
        
        # Format dates
        month_year = format_date_for_month_year(metadata.get("end_date", ""))
        date_range = format_date_for_range(metadata.get("start_date", ""), metadata.get("end_date", ""))
        
        # Define placeholder mappings
        replacements = {
            "ORGANIZATION_NAMEE": metadata.get("organization_name", ""),
            "Organization_Namee": metadata.get("organization_name", ""),
            "CITYY": metadata.get("city", ""),
            "Stateee": metadata.get("state", ""),
            "Audit_Date_Period": format_audit_date_period(metadata.get('start_date', ''), metadata.get('end_date', '')),
            "Maker_Name_R": metadata.get("report_prepared_by", ""),
            "Organization_Personn": metadata.get("auditee_details_1", ""),
            "Designationn": metadata.get("auditee_details_2", ""),
            "Auditee_Email_Adresss": bank_emails,
            "Monthh Yearr": month_year,
            "00rd Month Year to 00th Month Year": date_range
        }
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    # Special handling for date_range to make it bold
                    if old_text == "00rd Month Year to 00th Month Year":
                        _replace_text_in_runs_with_bold(paragraph, old_text, new_text, make_bold=True)
                    else:
                        _replace_text_in_runs(paragraph, old_text, new_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            if old_text in paragraph.text:
                                # Special handling for date_range to make it bold
                                if old_text == "00rd Month Year to 00th Month Year":
                                    _replace_text_in_runs_with_bold(paragraph, old_text, new_text, make_bold=True)
                                else:
                                    _replace_text_in_runs(paragraph, old_text, new_text)
        
        # Generate and replace timeline placeholders
        print("🔍 Generating VAPT timeline...")
        start_date = metadata.get("start_date", "")
        end_date = metadata.get("end_date", "")
        
        if start_date and end_date:
            timeline = generate_vapt_timeline(start_date, end_date)
            if timeline:
                print("📅 Generated timeline:")
                for phase in timeline:
                    print(f"  {phase['Phase']}: {phase['Start']} - {phase['End']}")
                
                # Replace timeline placeholders
                replace_timeline_placeholders(doc, timeline)
            else:
                print("⚠️ No timeline generated - using default placeholders")
        else:
            print("⚠️ Missing start_date or end_date - skipping timeline generation")
        
        print("✅ Metadata placeholders replaced successfully")
        
    except Exception as e:
        print(f"❌ Error replacing metadata placeholders: {str(e)}")
        traceback.print_exc()

def replace_timeline_placeholders(doc, timeline):
    """
    Replace timeline placeholders in the Word document with generated timeline data
    """
    try:
        # Create mapping of placeholders to phase names
        phase_placeholders = {
            "Planning": {
                "start": "11.11.1111",
                "end": "22.22.2222"
            },
            "Environment Setup": {
                "start": "33.33.3333", 
                "end": "44.44.4444"
            },
            "Pre-Assessment Preparation": {
                "start": "55.55.5555",
                "end": "66.66.6666"
            },
            "VAPT Execution": {
                "start": "77.77.7777",
                "end": "88.88.8888"
            },
            "Reporting": {
                "start": "99.99.9999",
                "end": "00.00.0000"
            }
        }

        # Extract dates from timeline
        phase_dates = {}
        for phase in timeline:
            phase_name = phase['Phase']
            phase_dates[phase_name] = {
                'start': phase['Start'],
                'end': phase['End']
            }

        # Create replacements dictionary
        timeline_replacements = {}
        for phase_name, placeholders in phase_placeholders.items():
            if phase_name in phase_dates:
                timeline_replacements[placeholders['start']] = phase_dates[phase_name]['start']
                timeline_replacements[placeholders['end']] = phase_dates[phase_name]['end']
            else:
                # Use default values if phase not found in timeline
                timeline_replacements[placeholders['start']] = placeholders['start']
                timeline_replacements[placeholders['end']] = placeholders['end']

        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for old_text, new_text in timeline_replacements.items():
                if old_text in paragraph.text:
                    _replace_text_in_runs(paragraph, old_text, new_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in timeline_replacements.items():
                            if old_text in paragraph.text:
                                _replace_text_in_runs(paragraph, old_text, new_text)
        
        print("✅ Timeline placeholders replaced successfully")
        
    except Exception as e:
        print(f"❌ Error replacing timeline placeholders: {str(e)}")
        traceback.print_exc()

def replace_text_in_table_with_risk_colors(table, old_text, new_text):
    """
    Replace text in a table while preserving formatting and applying risk-based background colors
    
    Args:
        table: The table object from python-docx
        old_text: Text to be replaced
        new_text: Text to replace with (should be the risk factor)
    """
    
    # Define risk level colors (using hex values for better compatibility)
    risk_colors = {
        'CRITICAL': '8B0000',      # Dark red
        'HIGH': 'FF0000',          # Red  
        'MEDIUM': 'FFA500',        # Orange
        'LOW': '008000'            # Green
    }
    
    # Get the appropriate background color
    risk_level = str(new_text).strip().upper()
    bg_color_hex = risk_colors.get(risk_level, None)
    
    for row in table.rows:
        for cell in row.cells:
            # Replace in cell paragraphs with bold formatting
            if _replace_text_in_runs_with_bold(cell.paragraphs[0], old_text, new_text, make_bold=True):
                # Set font color to white for better contrast with background
                if bg_color_hex:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)  # White color
                
                # Apply background color to the cell
                if bg_color_hex:
                    try:
                        # Method 1: Direct XML manipulation
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        
                        # Remove existing shading if any
                        existing_shd = tcPr.find(qn('w:shd'))
                        if existing_shd is not None:
                            tcPr.remove(existing_shd)
                        
                        # Create shading element
                        shd_xml = f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{bg_color_hex}"/>'
                        shd = parse_xml(shd_xml)
                        tcPr.append(shd)
                        
                    except Exception as e:
                        # Fallback for older docx versions or complex formatting
                        try:
                            tc = cell._tc
                            if tc.tcPr is None:
                                tcPr = parse_xml('<w:tcPr/>')
                                tc.append(tcPr)
                            else:
                                tcPr = tc.tcPr
                            
                            shd = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:color="auto" w:fill="{bg_color_hex}"/>')
                            tcPr.append(shd)
                        except Exception as e2:
                            pass
            
            # Recursively check for nested tables
            for nested_table in cell.tables:
                replace_text_in_table_with_risk_colors(nested_table, old_text, new_text)

def is_ip_address(text):
    """
    Check if a text string is a valid IP address
    """
    # Simple IP address pattern
    ip_pattern = r'^(\d{1,3}\.){3}\d{1,3}$'
    if re.match(ip_pattern, text):
        # Additional validation to check if each octet is 0-255
        parts = text.split('.')
        try:
            return all(0 <= int(part) <= 255 for part in parts)
        except ValueError:
            return False
    return False

def get_vulnerability_counts(excel_file):
    """
    Extract vulnerability counts by risk level from Excel file
    Handles case-insensitive matching for risk levels (Critical, HIGH, medium, Low, etc.)
    
    Args:
        excel_file (str): Path to Excel file
        
    Returns:
        tuple: (categories, values) for chart data
    """
    try:
        df = pd.read_excel(excel_file, sheet_name='Infra_VAPT')
        
        if df.empty:
            print("No data found in 'Infra_VAPT' worksheet")
            return [], []
        
        # Count vulnerabilities by risk level (case-insensitive)
        # Show original values first
        original_risk_counts = df['Risk Factor'].value_counts()
        print(f"📊 Original risk counts from Excel: {dict(original_risk_counts)}")
        
        # Convert to uppercase for consistent comparison, handle NaN values
        df['Risk_Factor_Upper'] = df['Risk Factor'].fillna('UNKNOWN').str.upper().str.strip()
        risk_counts = df['Risk_Factor_Upper'].value_counts()
        print(f"📊 Normalized risk counts (uppercase): {dict(risk_counts)}")
        
        # Define the order we want for categories - ALWAYS include all categories
        category_order = ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW']
        
        categories = []
        values = []
        
        # Always include all predefined categories, even if count is 0
        for category in category_order:
            categories.append(category)
            if category in risk_counts:
                values.append(int(risk_counts[category]))
                print(f"📊 {category}: {int(risk_counts[category])}")
            else:
                values.append(0)  # Set to 0 if no vulnerabilities found for this risk level
                print(f"📊 {category}: 0 (not found in data)")
        
        # Add any other risk levels not in our predefined order
        for category, count in risk_counts.items():
            if category not in category_order:
                categories.append(category)
                values.append(int(count))
        
        print(f"📊 Vulnerability counts: {dict(zip(categories, values))}")
        return categories, values
        
    except Exception as e:
        print(f"❌ Error extracting vulnerability counts: {e}")
        return [], []

def update_chart_in_docx(docx_path, excel_data, chart_file=None):
    """Update the chart in the DOCX file directly - Only Y-axis values and scaling"""
    try:
        # Step 1: Get vulnerability data from Excel
        categories, values = get_vulnerability_counts(excel_data)
        if not categories or not values:
            print("❌ Failed to extract vulnerability data from Excel")
            return False

        print(f"📊 Extracted vulnerability data: {dict(zip(categories, values))}")

        # Create a temporary copy to work with
        temp_dir = tempfile.mkdtemp()
        temp_docx = os.path.join(temp_dir, "temp.docx")
        shutil.copy2(docx_path, temp_docx)

        # Step 2: Read the DOCX as a ZIP file and find chart files
        with zipfile.ZipFile(temp_docx, 'r') as zin:
            file_list = zin.namelist()
            
            # Find chart files if not specified
            if chart_file is None:
                chart_files = [f for f in file_list if f.startswith('word/charts/') and f.endswith('.xml')]
                if not chart_files:
                    print("❌ No chart files found in Word document")
                    return False
                chart_file = chart_files[0]  # Use the first chart found
                print(f"📊 Using chart file: {chart_file}")
            
            try:
                with zin.open(chart_file) as chart_file_obj:
                    chart_xml = chart_file_obj.read()
            except KeyError:
                print(f"❌ Chart file not found: {chart_file}")
                return False

        # Step 3: Parse and update the chart XML
        tree = etree.fromstring(chart_xml)

        ns = {
            'c': 'http://schemas.openxmlformats.org/drawingml/2006/chart',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
        }

        # Update Y-axis values only
        val_pts = tree.findall('.//c:val//c:numCache//c:pt', namespaces=ns)
        
        print(f"📊 Found {len(val_pts)} data points in chart, updating with {len(values)} values")
        
        # Only update Y-axis values, leave X-axis categories unchanged
        for i, val_pt in enumerate(val_pts):
            if i < len(values):
                val_elem = val_pt.find('c:v', namespaces=ns)
                if val_elem is not None:
                    val_elem.text = str(values[i])
                    print(f"📊 Updated data point {i}: {values[i]}")
                else:
                    print(f"⚠️ No value element found for data point {i}")
            else:
                print(f"⚠️ More chart data points ({len(val_pts)}) than values ({len(values)})")
                break

        # Y-axis scaling logic (1,2 or 2,4 or 4,8 pattern)
        max_value = max(values) if values else 1
        
        # Handle case where all values are 0
        if max_value == 0:
            max_value = 1
            print("📊 All values are 0, setting max to 1 for chart display")
        
        # Determine Y-axis max and interval using power-of-2 logic
        if max_value <= 10:
            max_axis = 10
            interval = 1
        elif max_value <= 20:
            max_axis = 20
            interval = 2
        elif max_value <= 40:
            max_axis = 40
            interval = 4
        elif max_value <= 80:
            max_axis = 80
            interval = 8
        elif max_value <= 120:
            max_axis = 120
            interval = 12
        else:
            # For larger values, find next power of 2
            max_axis = 1
            while max_axis < max_value:
                max_axis *= 2
            interval = max_axis // 4  # Quarter intervals

        # Update Y-axis scaling
        val_ax = tree.find('.//c:valAx', namespaces=ns)
        scaling_elem = val_ax.find('.//c:scaling', namespaces=ns)
        if scaling_elem is not None:
            max_elem = scaling_elem.find('.//c:max', namespaces=ns)
            if max_elem is None:
                max_elem = etree.SubElement(scaling_elem, '{http://schemas.openxmlformats.org/drawingml/2006/chart}max')
            max_elem.set('val', str(max_axis))
        
        # Update major unit (interval)
        major_unit_elem = tree.find('.//c:majorUnit', namespaces=ns)
        if major_unit_elem is None:
            axis_elem = tree.find('.//c:valAx', namespaces=ns)
            if axis_elem is not None:
                major_unit_elem = etree.SubElement(axis_elem, '{http://schemas.openxmlformats.org/drawingml/2006/chart}majorUnit')
        if major_unit_elem is not None:
            major_unit_elem.set('val', str(interval))

        # Serialize and update the DOCX
        updated_chart_xml = etree.tostring(tree, pretty_print=True, xml_declaration=True, encoding='UTF-8')

        with zipfile.ZipFile(temp_docx, 'w') as zout:
            for file in file_list:
                if file == chart_file:
                    zout.writestr(file, updated_chart_xml)
                else:
                    with zipfile.ZipFile(docx_path, 'r') as zin:
                        with zin.open(file) as src_file:
                            zout.writestr(file, src_file.read())

        shutil.move(temp_docx, docx_path)
        shutil.rmtree(temp_dir)

        print(f"✅ Y-axis updated successfully in: {docx_path}")
        print(f"📈 Y-axis scaled to max: {max_axis}, interval: {interval}")
        return True

    except Exception as e:
        print(f"❌ Error updating Y-axis data: {str(e)}")
        import traceback
        traceback.print_exc()
        if 'temp_dir' in locals():
            shutil.rmtree(temp_dir, ignore_errors=True)
        return False

def replace_affected_systems_with_table(table, placeholder, affected_systems_text):
    """
    Replace the placeholder with an actual Word table containing affected systems data
    """
    target_cell = None
    for row in table.rows:
        for cell in row.cells:
            if placeholder in cell.text:
                target_cell = cell
                break
        if target_cell:
            break
    
    if not target_cell:
        return
    
    if not affected_systems_text or affected_systems_text.strip() == '':
        target_cell.text = "No affected systems"
        return
    
    lines = affected_systems_text.strip().split('\n')
    table_data = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        if 'Branch' in line or not is_ip_address(line):
            branch_name = line
            ip_addresses = []
            
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if not next_line:
                    j += 1
                    continue
                elif is_ip_address(next_line):
                    ip_addresses.append(next_line)
                    j += 1
                else:
                    break
            
            if ip_addresses:
                if len(ip_addresses) == 1:
                    table_data.append([branch_name, ip_addresses[0]])
                else:
                    table_data.append([branch_name, ip_addresses[0]])
                    for ip in ip_addresses[1:]:
                        table_data.append(["", ip])
            
            i = j
        else:
            table_data.append(["", line])
            i += 1
    
    if not table_data:
        target_cell.text = "No affected systems"
        return
    
    target_cell.text = ""
    
    for paragraph in target_cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    
    new_table = target_cell.add_table(rows=len(table_data), cols=2)
    
    new_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    for row in new_table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(2.0)
    
    for row in new_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            for border_name in ['top', 'left', 'bottom', 'right']:
                existing_border = tcPr.find(qn(f'w:{border_name}'))
                if existing_border is not None:
                    tcPr.remove(existing_border)
            
            for border_name in ['top', 'left', 'bottom', 'right']:
                border_xml = f'<w:{border_name} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                border = parse_xml(border_xml)
                tcPr.append(border)
    
    for i, (branch_name, ip_address) in enumerate(table_data):
        row = new_table.rows[i]
        row.cells[0].text = branch_name
        row.cells[1].text = ip_address
        
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 0
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Inches(0.17)

def add_images_at_end_of_document(doc, vulnerability_images, df):
    """
    Add images at the end of the document with titles from Name of Vulnerability column.
    """
    try:
        if not any(vulnerability_images.values()):
            return

        doc.add_page_break()
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run("Proof of Concept (POC) Images")
        title_run.font.size = Pt(16)
        title_run.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        image_count = 0
        for idx, row in enumerate(df.iterrows(), start=1):
            index, row_data = row
            vuln_name = str(row_data['Name of Vulnerability'])
            
            # Check both original and normalized vulnerability names for image matching
            normalized_vuln_name = normalize_vulnerability_name_for_filename(vuln_name)
            has_images = False
            
            if vuln_name in vulnerability_images and vulnerability_images[vuln_name]:
                has_images = True
                image_data_list = vulnerability_images[vuln_name]
            elif normalized_vuln_name in vulnerability_images and vulnerability_images[normalized_vuln_name]:
                has_images = True
                image_data_list = vulnerability_images[normalized_vuln_name]
            
            if has_images:
                image_count += 1
                vuln_title = doc.add_paragraph()
                vuln_run = vuln_title.add_run(f"{image_count}. {vuln_name}")
                vuln_run.font.size = Pt(14)
                vuln_run.bold = True
                
                for i, image_data in enumerate(image_data_list, start=1):
                    try:
                        pil_image = PILImage.open(io.BytesIO(image_data))
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                            pil_image.save(tmp, format='PNG')
                            temp_image_path = tmp.name
                        caption_para = doc.add_paragraph()
                        caption_run = caption_para.add_run(f"Image {i}")
                        caption_run.italic = True
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        image_para = doc.add_paragraph()
                        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = image_para.add_run()
                        picture = run.add_picture(temp_image_path, width=Inches(6.0))
                        apply_1pt_border_to_picture(picture)  # Apply 1pt border
                        os.unlink(temp_image_path)
                        doc.add_paragraph()
                    except Exception as e:
                        error_para = doc.add_paragraph()
                        error_para.add_run(f"Error loading image {i}: {str(e)}")
                
                if image_count < len([v for v in vulnerability_images.values() if v]):
                    doc.add_page_break()
    except Exception as e:
        traceback.print_exc()

def replace_vulnerability_details_with_images(doc, excel_file):
    """
    Replaces vulnerability details and creates new tables with POC placeholders.
    """
    try:
        # Check if A2 cell is empty in Infra_VAPT worksheet
        try:
            wb = load_workbook(excel_file, data_only=True)
            if 'Infra_VAPT' in wb.sheetnames:
                ws = wb['Infra_VAPT']
                a2_value = ws.cell(row=2, column=1).value  # A2 is row 2, column 1 (1-indexed)
                # Check if A2 is empty or None
                if a2_value is None or str(a2_value).strip() == '':
                    print("A2 cell is empty - removing table with a.1.1 and Vulnerability___Name, replacing Tables_Heree with message")
                    
                    # Find the target table (the one containing "a.1.1")
                    target_table = None
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if "a.1.1" in cell.text.lower():
                                    target_table = table
                                    break
                            if target_table:
                                break
                        if target_table:
                            break
                    
                    if target_table:
                        # Remove the entire table
                        target_table._element.getparent().remove(target_table._element)
                        print("✅ Removed a.1.1 table")
                    else:
                        print("⚠️ Could not find table containing 'a.1.1'")
                    
                    # Remove all tables containing "Recommendation___" (in case they're on different pages)
                    tables_to_remove = []
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if "Recommendation___" in cell.text:
                                    tables_to_remove.append(table)
                                    break
                            if table in tables_to_remove:
                                break
                    
                    for table in tables_to_remove:
                        try:
                            table._element.getparent().remove(table._element)
                            print(f"✅ Removed table containing 'Recommendation___'")
                        except Exception as e:
                            print(f"⚠️ Could not remove table: {e}")
                    
                    # Replace "Tables_Heree" with green bold message
                    for paragraph in doc.paragraphs:
                        if "Tables_Heree" in paragraph.text:
                            # Clear existing runs
                            paragraph.clear()
                            # Add line break before the message
                            break_run = paragraph.add_run()
                            break_run.add_break(WD_BREAK.LINE)
                            # Add new run with green bold text
                            run = paragraph.add_run("Note: No vulnerabilities were identified during the Infrastructure VAPT Audit.")
                            run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
                            run.font.bold = True
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            print("✅ Replaced 'Tables_Heree' with green bold message")
                            break
                    
                    # Also check in runs
                    for paragraph in doc.paragraphs:
                        for idx, run in enumerate(paragraph.runs):
                            if "Tables_Heree" in run.text:
                                # Create a new run with line break and insert it before the current run
                                break_run = paragraph.add_run()
                                break_run.add_break(WD_BREAK.LINE)
                                # Get the XML element for the break run and move it before the current run
                                run_element = run._element
                                break_element = break_run._element
                                run_element.getparent().remove(break_element)
                                run_element.getparent().insert(run_element.getparent().index(run_element), break_element)
                                
                                run.text = "Note: No vulnerabilities were identified during the Infrastructure VAPT Audit."
                                run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
                                run.font.bold = True
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                print("✅ Replaced 'Tables_Heree' in run with green bold message")
                                break
                    
                    wb.close()
                    return
            wb.close()
        except Exception as e:
            print(f"⚠️ Error checking A2 cell: {e}")
            # Continue with normal processing if check fails
        
        vulnerability_images, df = extract_poc_images_from_excel(excel_file)
        
        if df.empty:
            return

        template_table = None
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Vulnerability___Name" in cell.text:
                        template_table = table
                        break
                if template_table:
                    break
            if template_table:
                break
        if not template_table:
            return

        marker_paragraph = None
        for paragraph in doc.paragraphs:
            if "Tables_Heree" in paragraph.text:
                marker_paragraph = paragraph
                break
        if not marker_paragraph:
            return

        parent = marker_paragraph._element.getparent()
        marker_index = parent.index(marker_paragraph._element)
        parent.remove(template_table._element)
        for run in marker_paragraph.runs:
            if "Tables_Heree" in run.text:
                run.text = run.text.replace("Tables_Heree", "")

        for idx, row in enumerate(df.iterrows(), start=1):
            index, row_data = row
            new_table = copy_table_with_formatting(template_table)
            
            vuln_name = str(row_data['Name of Vulnerability'])
            replace_text_in_table(new_table, "Vulnerability___Name", vuln_name)
            
            risk_factor = str(row_data['Risk Factor'])
            replace_text_in_table_with_risk_colors(new_table, "Riskkkk", risk_factor)
            
            replace_text_in_table(new_table, "CWE_ID__", str(row_data['CVE ID']) if pd.notna(row_data['CVE ID']) else "NA")
            replace_text_in_table(new_table, "CVSS__", str(row_data['CVSS']))
            
            # Use the new function to create a table for affected systems
            replace_affected_systems_with_table(new_table, "Affected_URL___", str(row_data['Affected Systems']))
            
            # Simple text replacements for the remaining fields
            replace_text_in_table(new_table, "Observation___", str(row_data['Audit Observation']))
            replace_text_in_table(new_table, "Impact___", str(row_data['Impact']))
            replace_text_in_table(new_table, "Recommendation___", str(row_data['Recommendation/Countermeasure']))
            replace_text_in_table(new_table, "Reference_Link___", str(row_data['Reference Link']))
            replace_text_in_table(new_table, "New Observationnnn", "New Observation")
            replace_text_in_table(new_table, "Not Applicableeee", "Not Applicable")
            
            replace_text_in_table(new_table, "POCss", f"POCss{idx}11")
            
            for table_row in new_table.rows:
                for cell in table_row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        pattern = r'\b1\.([1-9])\b'
                        
                        def replace_numbering(match):
                            return f"{idx}.{match.group(1)}"
                        
                        new_text = re.sub(pattern, replace_numbering, original_text)
                        
                        if new_text != original_text:
                            font_props = {}
                            if paragraph.runs:
                                run = paragraph.runs[0]
                                font_props['name'] = 'calibri'
                                font_props['size'] = run.font.size
                                font_props['color_rgb'] = run.font.color.rgb
                                font_props['bold'] = run.bold
                                font_props['italic'] = run.italic
                                font_props['underline'] = run.underline

                            paragraph.text = ""
                            new_run = paragraph.add_run(new_text)
                            if font_props.get('name') is not None:
                                new_run.font.name = font_props['name']
                            if font_props.get('size') is not None:
                                new_run.font.size = font_props['size']
                            if font_props.get('color_rgb') is not None:
                                new_run.font.color.rgb = font_props['color_rgb']
                            if font_props.get('bold') is not None:
                                new_run.bold = font_props['bold']
                            if font_props.get('italic') is not None:
                                new_run.italic = font_props['italic']
                            if font_props.get('underline') is not None:
                                new_run.underline = font_props['underline']

            parent.insert(marker_index + 1, new_table._element)
            marker_index += 1
            
            if idx < len(df):
                p = doc.add_paragraph()
                p.add_run().add_break(WD_BREAK.PAGE)
                parent.insert(marker_index + 1, p._element)
                marker_index += 1
        
        add_images_at_end_of_document(doc, vulnerability_images, df)

        print("✅ Vulnerability details with POC links updated successfully")

    except Exception as e:
        traceback.print_exc()

def replace_poc_placeholders_with_images(doc, excel_file):
    """
    Finds POC<table number> placeholders and replaces them with images from Excel.
    If no images are found, replaces the placeholder with "NIL".
    Images will be centered, have width 6.5 inches, and a 1px solid black border.
    """
    try:
        wb = load_workbook(excel_file)
        ws = wb['Infra_VAPT']
        
        images_by_excel_row = {}
        if hasattr(ws, "_images"):
            for img in ws._images:
                try:
                    row_idx = img.anchor._from.row + 1
                    img_data = img._data() if callable(img._data) else img._data
                    if row_idx not in images_by_excel_row:
                        images_by_excel_row[row_idx] = []
                    images_by_excel_row[row_idx].append(img_data)
                except Exception as e:
                    print(f"Error processing image in Excel: {e}")
                    continue

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    
                    found_placeholder = False
                    for paragraph in cell.paragraphs:
                        text_content = paragraph.text
                        match = re.search(r'(POCss(\d+))', text_content)

                        if match:
                            full_placeholder = match.group(1)
                            table_number_str = match.group(2)
                            
                            try:
                                vuln_index = int(table_number_str[0])
                                excel_row = vuln_index + 1
                                
                                if excel_row in images_by_excel_row and images_by_excel_row[excel_row]:
                                    image_data_list = images_by_excel_row[excel_row]
                                    
                                    # Clear existing content in the cell first
                                    for p in list(cell.paragraphs):
                                        if "POCss" in p.text:
                                            p._element.getparent().remove(p._element)
                                    
                                    # Add images with single line break after each image
                                    for i, image_data in enumerate(image_data_list):
                                        try:
                                            # Create a new paragraph for the image
                                            image_paragraph = cell.add_paragraph()
                                            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            
                                            image_stream = io.BytesIO(image_data)
                                            run = image_paragraph.add_run()
                                            
                                            # Insert image with fixed width
                                            picture = run.add_picture(image_stream, width=Inches(6.5))
                                            
                                            # Apply 1pt border to the image
                                            apply_1pt_border_to_picture(picture)
                                            
                                            # Add ONE line break after each image
                                            line_break_para = cell.add_paragraph()
                                            line_break_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            # line_break_para.add_run().add_break()
                                            
                                        except Exception as e:
                                            print(f"Error adding image from Excel row {excel_row}: {e}")
                                            error_para = cell.add_paragraph()
                                            error_para.add_run(f" [Error loading image: {str(e)}] ")

                                    print(f"✅ Replaced '{full_placeholder}' with {len(image_data_list)} image(s) from Excel row {excel_row}.")
                                    found_placeholder = True
                                else:
                                    # No images found - replace with "NIL"
                                    for p in list(cell.paragraphs):
                                        if "POCss" in p.text:
                                            if full_placeholder in p.text:
                                                # Store formatting from first run (optional)
                                                original_runs = p.runs
                                                font_props = {}
                                                if original_runs:
                                                    first_run = original_runs[0]
                                                    font_props = {
                                                        # 'bold': first_run.bold,
                                                        'italic': first_run.italic,
                                                        'underline': first_run.underline,
                                                        'font_name': 'calibri',
                                                        'font_size': first_run.font.size,
                                                        'font_color': first_run.font.color.rgb
                                                    }

                                                # Clear entire cell
                                                for paragraph in cell.paragraphs:
                                                    p_el = paragraph._element
                                                    p_el.getparent().remove(p_el)

                                                for tbl in cell.tables:
                                                    tbl_el = tbl._element
                                                    tbl_el.getparent().remove(tbl_el)

                                                # Add one line break before NIL
                                                br_before = cell.add_paragraph()
                                                # br_before.add_run().add_break()

                                                # Add NIL centered
                                                nil_para = cell.add_paragraph()
                                                nil_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                nil_run = nil_para.add_run("NIL")
                                                if font_props.get('bold') is not None:
                                                    nil_run.bold = font_props['bold']
                                                if font_props.get('italic') is not None:
                                                    nil_run.italic = font_props['italic']
                                                if font_props.get('underline') is not None:
                                                    nil_run.underline = font_props['underline']
                                                if font_props.get('font_name') is not None:
                                                    nil_run.font.name = font_props['font_name']
                                                if font_props.get('font_size') is not None:
                                                    nil_run.font.size = font_props['font_size']
                                                if font_props.get('font_color') is not None:
                                                    nil_run.font.color.rgb = font_props['font_color']

                                                # Add one line break after NIL
                                                br_after = cell.add_paragraph()
                                                # br_after.add_run().add_break()

                                                print(f"✅ Replaced '{full_placeholder}' with 'NIL' (no images found for Excel row {excel_row})")
                                                found_placeholder = True
                                                break

                            except (ValueError, IndexError):
                                continue

    except Exception as e:
        print(f"Error processing image in Excel: {e}")
        traceback.print_exc()
        
def remove_content_after_marker(doc, marker_text):
    """
    Safely removes all content (paragraphs, tables, etc.) after a specific marker text in a Word document,
    including the marker text itself.

    Args:
        doc: The python-docx Document object.
        marker_text: The string to search for as the end marker (will be removed).
    """
    marker_found = False

    # Iterate over paragraphs and tables in body in document order
    body_elements = list(doc.element.body)

    for element in body_elements:
        if element.tag.endswith('p'):  # Paragraph
            para = next((p for p in doc.paragraphs if p._element == element), None)
            if para and marker_text in para.text:
                marker_found = True
                # Find text before the marker and keep only that
                para_text = para.text
                marker_index = para_text.find(marker_text)
                
                if marker_index > 0:
                    # Keep text before the marker
                    text_to_keep = para_text[:marker_index].strip()
                    para.clear()
                    if text_to_keep:
                        para.add_run(text_to_keep)
                else:
                    # Marker is at beginning or no text before it, remove the entire paragraph
                    element.getparent().remove(element)
                continue

            if marker_found:
                element.getparent().remove(element)

        elif element.tag.endswith('tbl'):  # Table
            tbl = next((t for t in doc.tables if t._element == element), None)
            if tbl:
                marker_in_table = False
                for row in tbl.rows:
                    for cell in row.cells:
                        if marker_text in cell.text:
                            marker_in_table = True
                            # Remove marker text from cell
                            cell_text = cell.text
                            marker_index = cell_text.find(marker_text)
                            if marker_index > 0:
                                cell.text = cell_text[:marker_index].strip()
                            else:
                                cell.text = ""
                            break
                    if marker_in_table:
                        break
                
                if marker_in_table:
                    marker_found = True
                    continue

            if marker_found:
                element.getparent().remove(element)

    return doc

def replace_open_ports_with_nmap_table(doc, excel_file, placeholder="OPEN_PORTSSSS"):
    """
    Replace OPEN_PORTSSSS placeholder with Nmap table from Excel
    """
    try:
        # Create the Nmap table
        nmap_table = create_nmap_table_from_excel(doc, excel_file)
        if not nmap_table:
            print("❌ Failed to create Nmap table")
            return False
        
        # Search for the placeholder in paragraphs
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                # Store the parent and index for insertion
                parent = paragraph._element.getparent()
                index = parent.index(paragraph._element)
                
                # Remove the placeholder paragraph
                parent.remove(paragraph._element)
                
                # Insert the table at the same position
                parent.insert(index, nmap_table._element)
                print(f"✅ Replaced '{placeholder}' with Nmap table")
                return True
        
        # If not found in paragraphs, check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            # Clear the cell and add the table
                            cell._element.clear()
                            cell._element.append(nmap_table._element)
                            print(f"✅ Replaced '{placeholder}' with Nmap table in table cell")
                            return True
        
        print(f"❌ Placeholder '{placeholder}' not found in document")
        return False
        
    except Exception as e:
        print(f"Error replacing OPEN_PORTSSSS with Nmap table: {e}")
        traceback.print_exc()
        return False

def read_scope_worksheet_data(excel_file):
    """
    Read data from the Scope worksheet in Excel file
    Handles cases where branches may have multiple IP addresses
    
    Args:
        excel_file (str): Path to Excel file
        
    Returns:
        list: List of dictionaries containing branch names and IP addresses
    """
    try:
        # Load Excel workbook
        wb = load_workbook(excel_file)
        
        # Check if 'Scope' worksheet exists
        if 'Scope' not in wb.sheetnames:
            print("Warning: 'Scope' worksheet not found in Excel file")
            return []
        
        # Get the Scope worksheet
        ws = wb['Scope']
        
        scope_data = []
        current_branch = None
        
        # Read data from the worksheet (assuming data starts from row 2, with headers in row 1)
        for row in ws.iter_rows(min_row=2, values_only=True):
            sr_no = row[0] if row[0] is not None else ""
            branch_name = str(row[1]).strip() if row[1] else ""
            host_ip = str(row[2]).strip() if row[2] else ""
            
            # If we have a branch name (not empty), it's a new branch
            if branch_name:
                current_branch = branch_name
                if host_ip:  # If there's also an IP in the same row
                    scope_data.append({
                        'branch_name': branch_name,
                        'host_ip': host_ip
                    })
            # If branch name is empty but we have an IP and a current branch, it's an additional IP
            elif not branch_name and host_ip and current_branch:
                scope_data.append({
                    'branch_name': current_branch,
                    'host_ip': host_ip
                })
        
        wb.close()
        print(f"✅ Read {len(scope_data)} entries from Scope worksheet")
        return scope_data
        
    except Exception as e:
        print(f"❌ Error reading Scope worksheet: {e}")
        return []

def replace_scope_placeholders_with_data(doc, excel_file):
    """
    Find and replace 'Branch Namee or Server Namee' and 'IP Addressesss' placeholders
    with data from the Scope worksheet, creating new rows for each entry
    """
    try:
        # Read scope data from Excel
        scope_data = read_scope_worksheet_data(excel_file)
        
        if not scope_data:
            print("No scope data found, skipping placeholder replacement")
            return
        
        # Find tables containing the placeholders
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        if "Branch Namee or Server Namee" in paragraph.text or "IP Addressesss" in paragraph.text:
                            # Found the target row, now replace and add new rows
                            replace_scope_placeholders_in_table(table, row_idx, cell_idx, scope_data)
                            return
        
        print("⚠️ Scope placeholders not found in any table")
        
    except Exception as e:
        print(f"❌ Error replacing scope placeholders: {e}")
        traceback.print_exc()

def replace_scope_placeholders_in_table(table, target_row_idx, target_cell_idx, scope_data):
    """
    Replace scope placeholders in a specific table row and add new rows for additional data
    Groups IP addresses by branch name for better organization
    """
    try:
        target_row = table.rows[target_row_idx]
        
        # Group scope data by branch name
        grouped_data = {}
        for entry in scope_data:
            branch_name = entry['branch_name']
            if branch_name not in grouped_data:
                grouped_data[branch_name] = []
            grouped_data[branch_name].append(entry['host_ip'])
        
        # Convert to list of entries with combined IP addresses
        processed_data = []
        for branch_name, ip_list in grouped_data.items():
            # Combine multiple IPs with line breaks
            combined_ips = '\n'.join(ip_list)
            processed_data.append({
                'branch_name': branch_name,
                'host_ip': combined_ips
            })
        
        # Replace placeholders in the first row with first processed entry
        if processed_data:
            first_entry = processed_data[0]
            
            # Replace Branch Namee or Server Namee
            for paragraph in target_row.cells[target_cell_idx].paragraphs:
                if "Branch Namee or Server Namee" in paragraph.text:
                    _replace_text_in_runs(paragraph, "Branch Namee or Server Namee", first_entry['branch_name'])
            
            # Replace IP Addressesss (two columns after Branch Name column - A to C)
            ip_cell_idx = target_cell_idx + 2
            if ip_cell_idx < len(target_row.cells):
                for paragraph in target_row.cells[ip_cell_idx].paragraphs:
                    if "IP Addressesss" in paragraph.text:
                        _replace_text_in_runs(paragraph, "IP Addressesss", first_entry['host_ip'])
            
            # Add borders to the target row
            add_borders_to_row(target_row)
        
        # Add new rows for remaining processed entries
        for i in range(1, len(processed_data)):
            entry = processed_data[i]
            
            # Create a new row by copying the target row structure
            new_row = table.add_row()
            
            # Copy the structure and content from the target row
            for j, cell in enumerate(new_row.cells):
                if j < len(target_row.cells):
                    # Clear the new cell
                    cell._element.clear()
                    
                    # Copy formatting from target cell
                    target_cell = target_row.cells[j]
                    
                    # Determine content based on column position
                    if j == target_cell_idx:  # Branch Name column
                        new_content = entry['branch_name']
                    elif j == target_cell_idx + 2:  # IP Address column (two columns after Branch Name - A to C)
                        new_content = entry['host_ip']
                    elif j == 0:  # Sr.No column
                        new_content = str(i + 1)  # Continue numbering
                    else:
                        # Keep original content for other columns
                        new_content = target_cell.text
                    
                    # Add content with proper formatting
                    if new_content:
                        new_paragraph = cell.add_paragraph()
                        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Set consistent formatting
                        new_run = new_paragraph.add_run(new_content)
                        new_run.font.name = 'Calibri (Body)'
                        new_run.font.size = Pt(12)
                        new_run.font.bold = False
                        new_run.font.italic = False
                        new_run.font.underline = False
                        
                        # Set cell vertical alignment
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Add borders to the new row
            add_borders_to_row(new_row)
        
        print(f"✅ Replaced scope placeholders and added {len(processed_data)} rows of data")
        
    except Exception as e:
        print(f"❌ Error replacing scope placeholders in table: {e}")
        traceback.print_exc()

def replace_a11_row_with_vulnerabilities(doc, excel_file):
    """Find the table containing a.1.1 and add Excel data with proper columns"""
    try:
        # Read Excel data
        df = pd.read_excel(excel_file, sheet_name='Infra_VAPT')
        
        # Check if A2 cell is empty (row 2, column 1, index 0 for row, index 0 for column)
        # A2 in Excel corresponds to the first data row (index 0) in pandas dataframe
        a2_empty = False
        if not df.empty and len(df) >= 1:
            # Check if the first data row (index 0) has empty value in first column (Sr.No)
            first_cell_value = df.iloc[0, 0] if len(df.columns) > 0 else None
            # Also check if it contains the note message
            note_message = "Note: No vulnerabilities were identified during the Infrastructure VAPT Audit."
            if pd.isna(first_cell_value) or str(first_cell_value).strip() == '' or note_message in str(first_cell_value):
                a2_empty = True
        elif df.empty:
            a2_empty = True
        
        # Find the target table (the one containing "a.1.1")
        target_table = None
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "a.1.1" in cell.text.lower():
                        target_table = table
                        break
                if target_table:
                    break
            if target_table:
                break

        if not target_table:
            print("Could not find table containing 'a.1.1'")
            return

        # If A2 is empty, delete the a.1.1 table
        if a2_empty:
            print("A2 is empty - removing a.1.1 table...")
            # Remove the entire table
            target_table._element.getparent().remove(target_table._element)
            print("✅ Removed a.1.1 table")
            
            # Remove all tables containing "Recommendation___" (in case they're on different pages)
            tables_to_remove = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "Recommendation___" in cell.text:
                            tables_to_remove.append(table)
                            break
                    if table in tables_to_remove:
                        break
            
            for table in tables_to_remove:
                try:
                    table._element.getparent().remove(table._element)
                    print(f"✅ Removed table containing 'Recommendation___'")
                except Exception as e:
                    print(f"⚠️ Could not remove table: {e}")
            
            # Replace "Tables_Heree" with note message
            for paragraph in doc.paragraphs:
                if "Tables_Heree" in paragraph.text:
                    # Clear existing runs
                    paragraph.clear()
                    # Add line break before the message
                    break_run = paragraph.add_run()
                    break_run.add_break(WD_BREAK.LINE)
                    # Add note message
                    run = paragraph.add_run("Note: No vulnerabilities were identified during the Infrastructure VAPT Audit.")
                    run.font.name = 'Calibri'
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
                    run.font.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print("✅ Replaced 'Tables_Heree' with note message")
                    break
            
            # Also check in runs
            for paragraph in doc.paragraphs:
                for idx, run in enumerate(paragraph.runs):
                    if "Tables_Heree" in run.text:
                        # Create a new run with line break and insert it before the current run
                        break_run = paragraph.add_run()
                        break_run.add_break(WD_BREAK.LINE)
                        # Get the XML element for the break run and move it before the current run
                        run_element = run._element
                        break_element = break_run._element
                        run_element.getparent().remove(break_element)
                        run_element.getparent().insert(run_element.getparent().index(run_element), break_element)
                        
                        run.text = "Note: No vulnerabilities were identified during the Infrastructure VAPT Audit."
                        run.font.name = 'Calibri'
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
                        run.font.bold = True
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print("✅ Replaced 'Tables_Heree' in run with note message")
                        break
            
            return
        
        if df.empty:
            print("No data found in 'Infra_VAPT' worksheet")
            return

        print("Found target table - updating content...")

        # Clear all existing rows except header (assuming first row is header)
        while len(target_table.rows) > 1:
            target_table._tbl.remove(target_table.rows[1]._tr)

        # Add Excel data with proper columns
        for idx in range(len(df)):
            new_row = target_table.add_row()
            
            # Ensure we have at least 11 columns (adjust based on your Word table)
            if len(new_row.cells) < 11:
                continue

            for cell in new_row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER    
            
            # Column 1: Sr.No (index + 1)
            first_cell = new_row.cells[0]
            for paragraph in first_cell.paragraphs:
                paragraph.clear()
            p = first_cell.paragraphs[0]
            p.text = str(idx + 1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 2: Affected System (Excel column I, index 8)
            second_cell = new_row.cells[1]
            for paragraph in second_cell.paragraphs:
                paragraph.clear()
            affected_system = str(df.iloc[idx, 9]) if pd.notna(df.iloc[idx, 1]) else "NA"
            p = second_cell.paragraphs[0]
            p.text = affected_system
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 3: Vulnerability Description (Excel column B, index 1)
            third_cell = new_row.cells[2]
            for paragraph in third_cell.paragraphs:
                paragraph.clear()
            vulnerability = str(df.iloc[idx, 2]) if pd.notna(df.iloc[idx, 2]) else "NA"
            p = third_cell.paragraphs[0]
            p.text = vulnerability
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 4: CVE ID (Excel column D, index 3)
            fourth_cell = new_row.cells[3]
            for paragraph in fourth_cell.paragraphs:
                paragraph.clear()
            # Get the value and handle both NaN and string "nan"
            raw_value = df.iloc[idx, 4]
            if pd.isna(raw_value) or str(raw_value).lower() == 'nan':
                cve_id = "NA"
            else:
                cve_id = str(raw_value)
            p = fourth_cell.paragraphs[0]
            p.text = cve_id
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)
                
            # Column 5: Risk Factor (Excel column C, index 2)
            fifth_cell = new_row.cells[7]
            for paragraph in fifth_cell.paragraphs:
                paragraph.clear()
            risk_factor = str(df.iloc[idx, 3]) if pd.notna(df.iloc[idx, 7]) else "NA"
            p = fifth_cell.paragraphs[0]
            p.text = risk_factor
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(255, 255, 255) 
                run.bold = True

            # Apply risk-based background colors
            risk_factor_lower = risk_factor.lower()
            if 'critical' in risk_factor_lower:
                fifth_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="8B0000"/>'.format(nsdecls('w'))))  # Dark Red
            elif 'high' in risk_factor_lower:
                fifth_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w'))))  # Red
            elif 'medium' in risk_factor_lower:
                fifth_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FFA500"/>'.format(nsdecls('w'))))  # Orange
            elif 'low' in risk_factor_lower:
                fifth_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="008000"/>'.format(nsdecls('w'))))  # Green
            
            # Column 6: CVSS (Excel column E, index 4)
            sixth_cell = new_row.cells[4]
            for paragraph in sixth_cell.paragraphs:
                paragraph.clear()
            p = sixth_cell.paragraphs[0]
            p.text = "NA"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 7: Audit Observation (Excel column F, index 5)
            seventh_cell = new_row.cells[5]
            for paragraph in seventh_cell.paragraphs:
                paragraph.clear()
            p = seventh_cell.paragraphs[0]
            p.text = "NA"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 8: Impact (Excel column G, index 6)
            eighth_cell = new_row.cells[6]
            for paragraph in eighth_cell.paragraphs:
                paragraph.clear()
            p = eighth_cell.paragraphs[0]
            p.text = "NA"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 9: Recommendation (Excel column H, index 7)
            ninth_cell = new_row.cells[8]
            for paragraph in ninth_cell.paragraphs:
                paragraph.clear()
            recommendation_w = clean_value(df.iloc[idx, 8])
            p = ninth_cell.paragraphs[0]
            p.text = recommendation_w
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 10: Reference Link (Excel column J, index 9)
            tenth_cell = new_row.cells[9]
            for paragraph in tenth_cell.paragraphs:
                paragraph.clear()
            reference = clean_value(df.iloc[idx, 10])
            p = tenth_cell.paragraphs[0]
            p.text = reference
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 11: Status (New/Old)
            eleven_cell = new_row.cells[10]
            for paragraph in eleven_cell.paragraphs:
                paragraph.clear()
            p = eleven_cell.paragraphs[0]
            p.text = "New"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

        print(f"✅ Successfully updated a.1.1 table with {len(df)} rows of data")

    except Exception as e:
        print(f"Error updating a.1.1 table: {str(e)}")
        traceback.print_exc()

def generate_word_report_from_excel(excel_file_path):
    """
    Generate a Word report from an Excel file containing audit data using the comprehensive logic
    """
    try:
        # Get the template path from static folder
        template_path = os.path.join('static', 'Formats_and_Catalog', 'Infrastructure_VAPT_First_Audit_Report.docx')
        
        if not os.path.exists(template_path):
            raise Exception(f"Template file not found at: {template_path}")
        
        # Load the Word template
        doc = Document(template_path)
        
        # Extract metadata from Excel file
        print("🔍 Extracting metadata from Excel file...")
        metadata = find_and_print_metadata(excel_file_path)
        
        # Print extracted metadata for debugging
        if metadata:
            print("📋 Extracted metadata:")
            for key, value in metadata.items():
                if value:
                    print(f"  {key}: {value}")
        
        # Replace metadata placeholders in the document
        if metadata:
            print("🔍 Replacing metadata placeholders...")
            replace_metadata_placeholders(doc, metadata)
            
            # Replace auditor placeholders and add additional rows for multiple team members
            print("🔍 Processing auditing team members...")
            replace_auditor_placeholders_and_add_rows(doc, metadata)
        else:
            print("⚠️ No metadata found in Excel file")
        
        # Replace vulnerability details with images
        replace_vulnerability_details_with_images(doc, excel_file_path)
        replace_poc_placeholders_with_images(doc, excel_file_path)
        
        # Replace OPEN_PORTSSSS placeholder with Nmap table
        replace_open_ports_with_nmap_table(doc, excel_file_path, "OPEN_PORTSSSS")
        
        # Replace scope placeholders with data from Scope worksheet
        print("🔍 Replacing scope placeholders with data from Scope worksheet...")
        replace_scope_placeholders_with_data(doc, excel_file_path)
        
        # Replace a.1.1 table with vulnerabilities data
        replace_a11_row_with_vulnerabilities(doc, excel_file_path)
        
        # Remove content after marker
        marker_to_remove_after = "CEH Cert. No: 12345678901234567890"
        remove_content_after_marker(doc, marker_to_remove_after)
        
        # Save the document temporarily to update chart
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_docx:
            doc.save(temp_docx.name)
            temp_docx_path = temp_docx.name
        
        # Update chart in Word document with vulnerability data from Excel
        try:
            print("Updating chart in Word document with vulnerability data...")
            chart_update_success = update_chart_in_docx(temp_docx_path, excel_file_path)
            if chart_update_success:
                print("✅ Chart updated successfully in Word document")
            else:
                print("⚠️ Chart update failed, but continuing with document generation")
        except Exception as e:
            print(f"❌ Error updating chart: {e}")
            print("⚠️ Continuing with document generation without chart update")
        
        # Reload the document with updated chart
        doc = Document(temp_docx_path)
        
        # Clean up temporary file
        os.unlink(temp_docx_path)
        
        print(f"✅ Word report generated successfully from Excel file")
        return doc
        
    except Exception as e:
        raise Exception(f"Error generating Word report: {str(e)}")

@word_report_bp.route('/process_first_audit_word_report', methods=['POST'])
def process_first_audit_word_report():
    """
    Process the uploaded Excel file and generate a Word report
    """
    try:
        # Check if file was uploaded
        if 'excelFile' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['excelFile']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Check file extension
        if not file.filename.lower().endswith('.xlsx'):
            return jsonify({'error': 'Please upload a valid Excel file (.xlsx)'}), 400
        
        # Note: Using static filename for Word report instead of Excel filename
        
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as temp_file:
            file.save(temp_file.name)
            temp_file_path = temp_file.name
        
        try:
            print("🔍 Starting Word report generation...")
            # Generate Word report using the comprehensive logic
            doc = generate_word_report_from_excel(temp_file_path)
            print("🔍 Word report generated, saving to temporary file...")
            
            # Save Word document to temporary file and send it
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as word_temp_file:
                doc.save(word_temp_file.name)
                word_temp_file_path = word_temp_file.name
                print(f"🔍 Word file saved to: {word_temp_file_path}")
            
            print("🔍 Attempting to send file to client...")
            # Return the Word file with automatic cleanup
            try:
                response = send_file(
                word_temp_file_path,
                as_attachment=True,
                download_name='Infrastructure_VAPT_First_Audit_Report.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
                print("🔍 File sent successfully!")
                
                # Clean up Excel file immediately
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                
                # Schedule Word file cleanup after response is sent
                import atexit
                def cleanup_word_file():
                    try:
                        if os.path.exists(word_temp_file_path):
                            os.unlink(word_temp_file_path)
                            print(f"🧹 Cleaned up Word file: {word_temp_file_path}")
                    except:
                        pass
                
                # Register cleanup function
                atexit.register(cleanup_word_file)
                
                return response
            except Exception as send_error:
                print(f"❌ Error sending file: {str(send_error)}")
                # Clean up Word file on error
                if os.path.exists(word_temp_file_path):
                    os.unlink(word_temp_file_path)
                raise send_error
            
        except Exception as e:
            print(f"❌ Error in processing: {str(e)}")
            # Clean up temporary files on error
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
            raise e
                
    except Exception as e:
        print(f"❌ Error in Flask route: {str(e)}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500
