from flask import Blueprint, request, jsonify
import os
import re
from docx import Document
from docx.shared import Pt
from datetime import datetime

# Create blueprint for Infrastructure VAPT Compliance Certificate
infrastructure_vapt_compliance_certificate_bp = Blueprint('infrastructure_vapt_compliance_certificate', __name__)

def format_date_with_suffix(date_obj):
    """Format date like '2nd September 2025' with superscript suffix"""
    day = date_obj.day
    if 4 <= day <= 20 or 24 <= day <= 30:
        suffix = "th"
    else:
        suffix = ["st", "nd", "rd"][day % 10 - 1]
    
    month = date_obj.strftime('%B')
    year = date_obj.year
    
    return f"{day}{suffix} {month} {year}", suffix

def create_infrastructure_vapt_compliance_certificate(form_data):
    """Create Infrastructure VAPT Compliance Certificate Word document"""
    try:
        print(f"🔍 Starting Infrastructure VAPT Compliance Certificate creation...")
        
        # Get form data
        org_name = form_data.get('organizationName', '')
        if org_name == 'Other':
            org_name = form_data.get('organizationNameOther', '')
        
        address = form_data.get('address', '')
        designation = form_data.get('designation', '')
        
        financial_year = form_data.get('financialYear', '')
        if financial_year == 'Other':
            financial_year = form_data.get('financialYearOther', '')
        
        branch_number = form_data.get('branchNumber', '')
        report_id = form_data.get('reportId', '')
        report_date = form_data.get('reportDate', '')
        
        # Format dates
        try:
            report_date_obj = datetime.strptime(report_date, '%Y-%m-%d')
            report_date_formatted, suffix = format_date_with_suffix(report_date_obj)
            report_date_dd_mm_yyyy = report_date_obj.strftime('%d.%m.%Y')
        except:
            report_date_formatted = report_date
            report_date_dd_mm_yyyy = report_date
        
        # Load template
        template_path = os.path.join('static', 'Formats_and_Catalog', 'Infrastructure VAPT Compliance Certificate.docx')
        if not os.path.exists(template_path):
            raise Exception(f"Template not found: {template_path}")
        
        print(f"   📄 Loading template: {template_path}")
        doc = Document(template_path)
        
        # Replace placeholders
        replacements = {
            'dddddessignation': designation,
            'ooooorgannization': org_name,
            'aaddrrrreessss': address,
            'fiiiiinanical_year': financial_year,
            'brrrranch_nuuumber': branch_number,
            'RRRRRRRRid': report_id,
            'RRRRRRDATE': report_date_formatted,
            'RRRDateformat': report_date_dd_mm_yyyy
        }
        
        print(f"   🔄 Replacing placeholders in paragraphs...")
        for key, value in replacements.items():
            for para in doc.paragraphs:
                if key in para.text:
                    # Special handling for RRRRRRDATE with superscript
                    if key == 'RRRRRRDATE':
                        for run in para.runs:
                            if key in run.text:
                                # Split date into day+number and suffix
                                match = re.match(r'(\d+)(st|nd|rd|th)(\s+\w+\s+\d+)', str(value))
                                if match:
                                    day_num = match.group(1)
                                    suffix = match.group(2)
                                    rest = match.group(3)
                                    
                                    # Clear the run
                                    run.text = run.text.replace(key, '')
                                    
                                    # Add day number
                                    run.text = day_num
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(12)
                                    run.font.superscript = False
                                    
                                    # Add suffix as superscript
                                    suffix_run = para.add_run(suffix)
                                    suffix_run.font.name = 'Times New Roman'
                                    suffix_run.font.size = Pt(12)
                                    suffix_run.font.superscript = True
                                    
                                    # Add rest of date
                                    rest_run = para.add_run(rest)
                                    rest_run.font.name = 'Times New Roman'
                                    rest_run.font.size = Pt(12)
                                    rest_run.font.superscript = False
                                    
                                    print(f"      ✅ Replaced '{key}' with '{value}' (with superscript)")
                                else:
                                    run.text = run.text.replace(key, str(value))
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(12)
                    else:
                        for run in para.runs:
                            if key in run.text:
                                run.text = run.text.replace(key, str(value))
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)
                                print(f"      ✅ Replaced '{key}' with '{value}'")
        
        # Replace placeholders in tables
        print(f"   🔄 Replacing placeholders in tables...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in para.text:
                                # Special handling for RRRRRRDATE with superscript
                                if key == 'RRRRRRDATE':
                                    for run in para.runs:
                                        if key in run.text:
                                            match = re.match(r'(\d+)(st|nd|rd|th)(\s+\w+\s+\d+)', str(value))
                                            if match:
                                                day_num = match.group(1)
                                                suffix = match.group(2)
                                                rest = match.group(3)
                                                
                                                run.text = run.text.replace(key, '')
                                                run.text = day_num
                                                run.font.name = 'Times New Roman'
                                                run.font.size = Pt(12)
                                                run.font.superscript = False
                                                
                                                suffix_run = para.add_run(suffix)
                                                suffix_run.font.name = 'Times New Roman'
                                                suffix_run.font.size = Pt(12)
                                                suffix_run.font.superscript = True
                                                
                                                rest_run = para.add_run(rest)
                                                rest_run.font.name = 'Times New Roman'
                                                rest_run.font.size = Pt(12)
                                                rest_run.font.superscript = False
                                                
                                                print(f"      ✅ Replaced '{key}' with '{value}' (with superscript) in table")
                                            else:
                                                run.text = run.text.replace(key, str(value))
                                                run.font.name = 'Times New Roman'
                                                run.font.size = Pt(12)
                                else:
                                    for run in para.runs:
                                        if key in run.text:
                                            run.text = run.text.replace(key, str(value))
                                            run.font.name = 'Times New Roman'
                                            run.font.size = Pt(12)
                                            print(f"      ✅ Replaced '{key}' with '{value}' in table")
        
        print(f"   ✅ Infrastructure VAPT Compliance Certificate created successfully")
        return doc
        
    except Exception as e:
        print(f"   ❌ Error creating certificate: {e}")
        import traceback
        traceback.print_exc()
        raise

@infrastructure_vapt_compliance_certificate_bp.route('/process_infrastructure_vapt_compliance_certificate', methods=['POST'])
def process_infrastructure_vapt_compliance_certificate():
    """Process Infrastructure VAPT Compliance Certificate form submission"""
    try:
        print("\n" + "="*80)
        print("🚀 Processing Infrastructure VAPT Compliance Certificate")
        print("="*80)
        
        # Get form data
        form_data = {
            'organizationName': request.form.get('organizationName'),
            'organizationNameOther': request.form.get('organizationNameOther'),
            'address': request.form.get('address'),
            'designation': request.form.get('designation'),
            'financialYear': request.form.get('financialYear'),
            'financialYearOther': request.form.get('financialYearOther'),
            'branchNumber': request.form.get('branchNumber'),
            'reportId': request.form.get('reportId'),
            'reportDate': request.form.get('reportDate')
        }
        
        # Create the certificate
        doc = create_infrastructure_vapt_compliance_certificate(form_data)
        
        # Save the Word document with same name as template
        filename = 'Infrastructure_VAPT_Compliance_Certificate.docx'
        filepath = os.path.join('static', 'uploads', filename)
        
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        doc.save(filepath)
        
        print(f"\n✅ Infrastructure VAPT Compliance Certificate created: {filename}")
        print("="*80)
        
        return jsonify({
            'success': True,
            'filename': filename,
            'download_url': f'/static/uploads/{filename}'
        })
    
    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@infrastructure_vapt_compliance_certificate_bp.route('/cleanup_infrastructure_vapt_compliance_certificate', methods=['POST'])
def cleanup_infrastructure_vapt_compliance_certificate():
    """Clean up Infrastructure VAPT Compliance Certificate file after download"""
    try:
        data = request.get_json()
        filename = data.get('filename')
        
        files_deleted = []
        
        if filename:
            file_path = os.path.join('static', 'uploads', filename)
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                    files_deleted.append(file_path)
                    print(f"✅ Deleted: {file_path}")
                except Exception as e:
                    print(f"❌ Error deleting {file_path}: {e}")
        
        print(f"📊 Cleanup summary: {len(files_deleted)} files deleted")
        
        return jsonify({
            'success': True,
            'files_deleted': len(files_deleted)
        })
    except Exception as e:
        print(f"❌ Error cleaning up: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

