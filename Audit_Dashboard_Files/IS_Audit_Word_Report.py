import os
import shutil
import tempfile
import re
import io
from datetime import datetime
from flask import Blueprint, request, send_file, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
import threading
import pandas as pd
from openpyxl import load_workbook
from PIL import Image as PILImage

is_audit_word_report_bp = Blueprint('is_audit_word_report', __name__)

def extract_images_from_excel_cells(worksheet, row_idx, column_letters):
    """
    Extract images from specific cells in an Excel row.
    
    Args:
        worksheet: openpyxl worksheet object
        row_idx: Row number (1-based)
        column_letters: List of column letters (e.g., ['J', 'K', 'I'])
    
    Returns:
        list: List of image data (binary)
    """
    images = []
    
    # Convert column letters to indices
    column_indices = []
    for col_letter in column_letters:
        col_idx = 0
        for char in col_letter.upper():
            col_idx = col_idx * 26 + (ord(char) - ord('A') + 1)
        column_indices.append(col_idx)
    
    try:
        # Extract images from worksheet
        if hasattr(worksheet, '_images'):
            for img in worksheet._images:
                try:
                    img_row = img.anchor._from.row + 1
                    img_col = img.anchor._from.col + 1
                    
                    # Check if image is in target row and columns
                    if img_row == row_idx and img_col in column_indices:
                        img_data = img._data() if callable(img._data) else img._data
                        if img_data:
                            images.append(img_data)
                except Exception as e:
                    print(f"Error extracting image: {e}")
                    continue
    except Exception as e:
        print(f"Error accessing worksheet images: {e}")
    
    return images

def extract_worksheet_names_from_excel(excel_file_path):
    """
    Extract all worksheet names from the Excel file.
    
    Returns:
        list: List of worksheet names
    """
    try:
        wb = load_workbook(excel_file_path, read_only=True)
        sheet_names = wb.sheetnames
        wb.close()
        return sheet_names
    except Exception as e:
        print(f"❌ Error extracting worksheet names: {e}")
        return []


def extract_audit_data_from_excel(excel_file_path):
    """
    Extract audit data from all worksheets in the Excel file.
    
    Returns:
        list: List of dictionaries containing worksheet data
    """
    try:
        # Load workbook
        wb = load_workbook(excel_file_path)
        
        all_worksheets_data = []
        
        for sheet_name in wb.sheetnames:
            print(f"\n📋 Processing worksheet: {sheet_name}")
            ws = wb[sheet_name]
            
            # Read data using pandas for easier manipulation
            df = pd.read_excel(excel_file_path, sheet_name=sheet_name)
            
            # Extract observations from each row
            observations = []
            has_any_data = False  # Track if there's any data beyond headers
            
            # If worksheet is completely empty or has no data, mark as compliant
            if df.empty:
                print(f"  ℹ️ Worksheet '{sheet_name}' is completely empty - marking as compliant")
                worksheet_data = {
                    'sheet_name': sheet_name,
                    'observations': [],
                    'is_compliant': True
                }
                all_worksheets_data.append(worksheet_data)
                continue
            
            # Start from row 2 (index 1 in pandas, row 2 in Excel)
            for idx, row in df.iterrows():
                # Check if row is completely empty (blank row)
                if row.isna().all():
                    print(f"  ℹ️ Found blank row at index {idx}, stopping for this worksheet")
                    break
                
                # Mark that we have at least some data in rows
                has_any_data = True
                
                # Extract data from columns
                observation_short = row.get('Observation (Short/Brief)', '')
                risk_factor = row.get('Risk Factor', '')
                observation_full = row.get('Observation', '')
                impact = row.get('Impact', '')
                recommendation = row.get('Recommendation', '')
                
                # Skip if observation short is empty/NaN (but still mark that data exists)
                if pd.isna(observation_short) or str(observation_short).strip() == '':
                    print(f"  ℹ️ Row {idx} has empty observation, skipping this row")
                    continue
                
                # Extract images from columns J, K, I (POC columns)
                # Row index in Excel = pandas index + 2 (header row + 0-based to 1-based)
                excel_row_idx = idx + 2
                images = extract_images_from_excel_cells(ws, excel_row_idx, ['J', 'K', 'I'])
                
                observation_data = {
                    'observation_short': str(observation_short).strip() if pd.notna(observation_short) else '',
                    'risk_factor': str(risk_factor).strip() if pd.notna(risk_factor) else '',
                    'observation_full': str(observation_full).strip() if pd.notna(observation_full) else '',
                    'impact': str(impact).strip() if pd.notna(impact) else '',
                    'recommendation': str(recommendation).strip() if pd.notna(recommendation) else '',
                    'images': images
                }
                
                observations.append(observation_data)
                print(f"  ✅ Extracted observation {len(observations)}: {observation_short[:50]}...")
            
            # Determine if worksheet is compliant (only headers, no data rows)
            # If has_any_data is False, it means only header row exists (no data rows below header)
            is_compliant = not has_any_data
            
            # Add worksheet data even if no observations (only header row)
            worksheet_data = {
                'sheet_name': sheet_name,
                'observations': observations,
                'is_compliant': is_compliant  # True if only header row, no data rows at all
            }
            all_worksheets_data.append(worksheet_data)
            
            if observations:
                print(f"  ✅ Total observations extracted from '{sheet_name}': {len(observations)}")
            elif is_compliant:
                print(f"  ✅ Worksheet '{sheet_name}' has only headers (no data rows) - marking as compliant")
            else:
                print(f"  ⚠️ Worksheet '{sheet_name}' has data rows but no valid observations")
        
        print(f"\n✅ Total worksheets processed: {len(all_worksheets_data)}")
        return all_worksheets_data
        
    except Exception as e:
        print(f"❌ Error extracting data from Excel: {str(e)}")
        import traceback
        traceback.print_exc()
        return []

def replace_assets_review_with_excel_data(doc, excel_file_path, placeholder="Asssssssssettttsss_Reevieeww"):
    """
    Replace the placeholder with formatted audit observations from all worksheets.
    
    Args:
        doc: Document object
        excel_file_path: Path to the Excel file
        placeholder: Text to replace
    """
    try:
        print(f"\n🔄 Replacing '{placeholder}' with Excel data...")
        
        # Extract all worksheet data
        worksheets_data = extract_audit_data_from_excel(excel_file_path)
        
        if not worksheets_data:
            print("⚠️ No data extracted from Excel, skipping replacement")
            return
        
        # Find the paragraph containing the placeholder
        placeholder_paragraph = None
        placeholder_index = None
        
        for idx, paragraph in enumerate(doc.paragraphs):
            if placeholder in paragraph.text:
                placeholder_paragraph = paragraph
                placeholder_index = idx
                print(f"✅ Found placeholder at paragraph index {idx}")
                break
        
        if not placeholder_paragraph:
            print(f"⚠️ Placeholder '{placeholder}' not found in document")
            return
        
        # Get the parent element to insert content
        parent = placeholder_paragraph._element.getparent()
        placeholder_element = placeholder_paragraph._element
        
        # Process each worksheet
        for sheet_idx, worksheet_data in enumerate(worksheets_data, 1):
            sheet_name = worksheet_data['sheet_name']
            observations = worksheet_data['observations']
            is_compliant = worksheet_data.get('is_compliant', False)
            
            print(f"\n📝 Adding worksheet {sheet_idx}: {sheet_name}")
            
            # Add worksheet title (e.g., "1. Network Review")
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{sheet_idx}. {sheet_name}")
            title_run.font.size = Pt(16)
            title_run.font.name = 'Times New Roman'
            title_run.bold = True
            
            # Move the title paragraph to correct position
            parent.insert(parent.index(placeholder_element), title_para._element)
            
            # Add blank line after title
            blank_para = doc.add_paragraph()
            parent.insert(parent.index(placeholder_element), blank_para._element)
            
            # If worksheet is compliant (only headers, no data)
            if is_compliant:
                print(f"  ✅ Worksheet '{sheet_name}' is compliant - adding compliant message")
                
                # Add compliant message in green and bold
                compliant_para = doc.add_paragraph()
                compliant_run = compliant_para.add_run("All points are compliant with the auditor's observations.")
                compliant_run.font.size = Pt(12)
                compliant_run.font.name = 'Times New Roman'
                compliant_run.bold = True
                compliant_run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
                
                parent.insert(parent.index(placeholder_element), compliant_para._element)
                
                # Add blank line after compliant message
                blank_after = doc.add_paragraph()
                parent.insert(parent.index(placeholder_element), blank_after._element)
                
                # Add page break if not the last worksheet
                if sheet_idx < len(worksheets_data):
                    page_break_para = doc.add_paragraph()
                    page_break_run = page_break_para.add_run()
                    page_break_run.add_break(WD_BREAK.PAGE)
                    parent.insert(parent.index(placeholder_element), page_break_para._element)
                
                continue  # Skip to next worksheet
            
            # Process each observation in this worksheet
            for obs_idx, obs in enumerate(observations, 1):
                print(f"  Adding observation {sheet_idx}.{obs_idx}: {obs['observation_short'][:50]}...")
                
                # 1. Observation Short (e.g., "1.1 Inbound TCP Keep-Alives enabled.")
                obs_short_para = doc.add_paragraph()
                obs_short_run = obs_short_para.add_run(f"{sheet_idx}.{obs_idx} {obs['observation_short']}")
                obs_short_run.font.size = Pt(12)
                obs_short_run.font.name = 'Times New Roman'
                obs_short_run.bold = True
                parent.insert(parent.index(placeholder_element), obs_short_para._element)
                
                # Add blank line
                blank1 = doc.add_paragraph()
                parent.insert(parent.index(placeholder_element), blank1._element)
                
                # 2. Risk Factor
                if obs['risk_factor']:
                    risk_para = doc.add_paragraph()
                    
                    # Add "Risk Factor:" label in bold
                    label_run = risk_para.add_run("Risk Factor: ")
                    label_run.font.size = Pt(12)
                    label_run.font.name = 'Times New Roman'
                    label_run.bold = True
                    
                    # Add risk factor value with color coding
                    value_run = risk_para.add_run(obs['risk_factor'])
                    value_run.font.size = Pt(12)
                    value_run.font.name = 'Times New Roman'
                    value_run.bold = True
                    
                    # Apply color based on risk level
                    risk_value = obs['risk_factor'].strip().lower()
                    if 'critical' in risk_value:
                        value_run.font.color.rgb = RGBColor(139, 0, 0)  # Dark Red
                    elif 'high' in risk_value:
                        value_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                    elif 'medium' in risk_value:
                        value_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
                    elif 'low' in risk_value:
                        value_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                    
                    parent.insert(parent.index(placeholder_element), risk_para._element)
                    
                    # Add blank line
                    blank2 = doc.add_paragraph()
                    parent.insert(parent.index(placeholder_element), blank2._element)
                
                # 3. Observation Full
                if obs['observation_full']:
                    obs_full_para = doc.add_paragraph()
                    
                    # Add "Observation:" label in bold
                    obs_label_run = obs_full_para.add_run("Observation: ")
                    obs_label_run.font.size = Pt(12)
                    obs_label_run.font.name = 'Times New Roman'
                    obs_label_run.bold = True
                    
                    # Add observation text (not bold)
                    obs_value_run = obs_full_para.add_run(obs['observation_full'])
                    obs_value_run.font.size = Pt(12)
                    obs_value_run.font.name = 'Times New Roman'
                    
                    parent.insert(parent.index(placeholder_element), obs_full_para._element)
                    
                    # Add blank line
                    blank3 = doc.add_paragraph()
                    parent.insert(parent.index(placeholder_element), blank3._element)
                
                # 4. Impact
                if obs['impact']:
                    impact_para = doc.add_paragraph()
                    
                    # Add "Impact:" label in bold
                    impact_label_run = impact_para.add_run("Impact: ")
                    impact_label_run.font.size = Pt(12)
                    impact_label_run.font.name = 'Times New Roman'
                    impact_label_run.bold = True
                    
                    # Add impact text (not bold)
                    impact_value_run = impact_para.add_run(obs['impact'])
                    impact_value_run.font.size = Pt(12)
                    impact_value_run.font.name = 'Times New Roman'
                    
                    parent.insert(parent.index(placeholder_element), impact_para._element)
                    
                    # Add blank line
                    blank4 = doc.add_paragraph()
                    parent.insert(parent.index(placeholder_element), blank4._element)
                
                # 5. POC Images
                if obs['images']:
                    # If images exist, add "POC:" on separate line
                    poc_para = doc.add_paragraph()
                    poc_run = poc_para.add_run("POC:")
                    poc_run.font.size = Pt(12)
                    poc_run.font.name = 'Times New Roman'
                    poc_run.bold = True
                    parent.insert(parent.index(placeholder_element), poc_para._element)
                    
                    # Add images
                    for img_idx, img_data in enumerate(obs['images'], 1):
                        try:
                            # Get image dimensions to calculate width based on aspect ratio
                            pil_img = PILImage.open(io.BytesIO(img_data))
                            img_width, img_height = pil_img.size
                            
                            # Target height is 300px, calculate width to maintain aspect ratio
                            target_height_px = 350
                            aspect_ratio = img_width / img_height
                            target_width_px = target_height_px * aspect_ratio
                            
                            # Convert pixels to inches (assuming 96 DPI)
                            target_height_inches = target_height_px / 96
                            
                            # Create centered paragraph for image
                            img_para = doc.add_paragraph()
                            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Add the image
                            img_stream = io.BytesIO(img_data)
                            img_run = img_para.add_run()
                            picture = img_run.add_picture(img_stream, height=Inches(target_height_inches))
                            
                            print(f"    ✅ Added image {img_idx}")
                            
                            # Insert the image paragraph
                            parent.insert(parent.index(placeholder_element), img_para._element)
                            
                        except Exception as e:
                            print(f"    ⚠️ Error adding image {img_idx}: {e}")
                            import traceback
                            traceback.print_exc()
                    
                    # Add blank line after all images
                    blank5 = doc.add_paragraph()
                    parent.insert(parent.index(placeholder_element), blank5._element)
                else:
                    # No images - add "POC: NIL" on same line
                    poc_nil_para = doc.add_paragraph()
                    poc_run = poc_nil_para.add_run("POC: ")
                    poc_run.font.size = Pt(12)
                    poc_run.font.name = 'Times New Roman'
                    poc_run.bold = True
                    nil_run = poc_nil_para.add_run("NIL")
                    nil_run.font.size = Pt(12)
                    nil_run.font.name = 'Times New Roman'
                    parent.insert(parent.index(placeholder_element), poc_nil_para._element)
                    
                    # Add blank line after POC: NIL
                    blank5 = doc.add_paragraph()
                    parent.insert(parent.index(placeholder_element), blank5._element)
                
                # 6. Recommendation
                if obs['recommendation']:
                    rec_para = doc.add_paragraph()
                    
                    # Add "Recommendation:" label in bold
                    rec_label_run = rec_para.add_run("Recommendation: ")
                    rec_label_run.font.size = Pt(12)
                    rec_label_run.font.name = 'Times New Roman'
                    rec_label_run.bold = True
                    
                    # Add recommendation text (not bold)
                    rec_value_run = rec_para.add_run(obs['recommendation'])
                    rec_value_run.font.size = Pt(12)
                    rec_value_run.font.name = 'Times New Roman'
                    
                    parent.insert(parent.index(placeholder_element), rec_para._element)
                    
                    # Add blank line
                    blank6 = doc.add_paragraph()
                    parent.insert(parent.index(placeholder_element), blank6._element)
                
                # Add page break after each observation (including the last one)
                page_break_para = doc.add_paragraph()
                page_break_run = page_break_para.add_run()
                page_break_run.add_break(WD_BREAK.PAGE)
                parent.insert(parent.index(placeholder_element), page_break_para._element)
        
        # Remove the placeholder paragraph
        parent.remove(placeholder_element)
        print(f"\n✅ Successfully replaced '{placeholder}' with data from {len(worksheets_data)} worksheet(s)")
        
    except Exception as e:
        print(f"❌ Error replacing assets review: {str(e)}")
        import traceback
        traceback.print_exc()

def set_table_font_size(table, font_size):
    """
    Set font size for all text in a table
    """
    try:
        print(f"  🔤 Setting font size to {font_size}pt for all table text...")
        
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)
                        run.font.name = 'Times New Roman'
        
        print(f"  ✅ Set font size to {font_size}pt for all table text")
        
    except Exception as e:
        print(f"  ❌ Error setting table font size: {e}")

def ensure_table_alignment(table):
    """
    Ensure all cells in the table have proper alignment
    """
    try:
        print(f"  📐 Ensuring proper alignment for all table cells...")
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # Ensure cell has at least one paragraph
                if not cell.paragraphs:
                    cell.add_paragraph()
                
                # Get the first paragraph
                first_para = cell.paragraphs[0]
                
                # Set default alignment if not already set
                if first_para.alignment is None:
                    first_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Set default vertical alignment if not already set
                if cell.vertical_alignment is None:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                
                # Ensure all paragraphs in the cell have the same alignment
                for para in cell.paragraphs:
                    if para.alignment is None:
                        para.alignment = first_para.alignment
        
        print(f"  ✅ Ensured proper alignment for all table cells")
        
    except Exception as e:
        print(f"  ❌ Error ensuring table alignment: {e}")

def update_table_column_widths(table):
    """
    Update column widths for a table after it's been inserted into the document
    """
    try:
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn
        
        actual_col_count = len(table.columns)
        print(f"  🔧 Updating column widths for {actual_col_count} columns...")
        
        # Define width percentages based on actual column count
        if actual_col_count == 5:
            # For 5 columns: SR.NO, BRANCH CHECK POINTS, Risk Factor, Compliance, REMARKS
            width_percentages = [5, 20, 10, 10, 55]  # Total = 100% (F column = 55%)
            print(f"  📏 Using 5-column layout (SR.NO=5%, CHECK POINTS=20%, RISK=10%, COMPLIANCE=10%, REMARKS=55%)")
        elif actual_col_count == 6:
            # For 6 columns: A=5%, B=10%, C=10%, D=10%, E=5%, F=60%
            width_percentages = [5, 10, 10, 10, 5, 60]  # Total = 100% (F column = 60%)
            print(f"  📏 Using 6-column layout (A=5%, B=10%, C=10%, D=10%, E=5%, F=60%)")
        else:
            # For other column counts, distribute equally
            equal_percentage = 100 / actual_col_count
            width_percentages = [equal_percentage] * actual_col_count
            print(f"  📏 Using equal widths ({equal_percentage:.1f}% each for {actual_col_count} columns)")
        
        # Get the table element and its grid
        tbl = table._tbl
        tblGrid = tbl.find(qn('w:tblGrid'))
        
        if tblGrid is not None:
            # Get existing grid columns
            existing_grid_cols = tblGrid.findall(qn('w:gridCol'))
            print(f"  🔍 Found {len(existing_grid_cols)} existing grid columns")
            
            # Calculate widths
            total_width_twips = 9360  # 6.5 inches in twips (1440 twips per inch)
            
            # Update existing grid columns or create new ones
            for col_idx in range(actual_col_count):
                if col_idx < len(width_percentages):
                    percentage = width_percentages[col_idx]
                    width_twips = int(total_width_twips * percentage / 100)
                else:
                    # For extra columns, use 5% each
                    width_twips = int(total_width_twips * 5 / 100)
                
                if col_idx < len(existing_grid_cols):
                    # Update existing grid column
                    existing_grid_cols[col_idx].set(qn('w:w'), str(width_twips))
                    print(f"    Column {col_idx + 1}: {width_percentages[col_idx] if col_idx < len(width_percentages) else 5}% = {width_twips} twips (updated existing)")
                else:
                    # Create new grid column
                    gridCol = OxmlElement('w:gridCol')
                    gridCol.set(qn('w:w'), str(width_twips))
                    tblGrid.append(gridCol)
                    print(f"    Column {col_idx + 1}: {width_percentages[col_idx] if col_idx < len(width_percentages) else 5}% = {width_twips} twips (created new)")
            
            # Remove any extra grid columns if we have fewer columns now
            while len(tblGrid.findall(qn('w:gridCol'))) > actual_col_count:
                last_grid_col = tblGrid.findall(qn('w:gridCol'))[-1]
                tblGrid.remove(last_grid_col)
                print(f"    Removed extra grid column")
            
            # Set table properties to use fixed column widths
            tblPr = tbl.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # Add table width property
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblW.set(qn('w:w'), str(total_width_twips))
                tblW.set(qn('w:type'), 'dxa')  # dxa = twips
                tblPr.append(tblW)
            else:
                tblW.set(qn('w:w'), str(total_width_twips))
                tblW.set(qn('w:type'), 'dxa')
            
            # Apply AutoFit to Content (like right-click → AutoFit → AutoFit to Content)
            # This is equivalent to the manual "AutoFit to Content" option
            
            print(f"  🔧 Applying AutoFit to Content...")
            
            # Remove any existing table width constraints
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is not None:
                tblPr.remove(tblW)
                print(f"    ✅ Removed fixed table width constraint")
            
            # Set table layout to autofit to content (not autofit to window)
            tblLayout = tblPr.find(qn('w:tblLayout'))
            if tblLayout is None:
                tblLayout = OxmlElement('w:tblLayout')
                # Remove the type attribute to enable AutoFit to Content
                tblPr.append(tblLayout)
                print(f"    ✅ Set table layout to AutoFit to Content")
            else:
                # Remove type attribute to enable AutoFit to Content
                if tblLayout.get(qn('w:type')):
                    tblLayout.attrib.pop(qn('w:type'))
                print(f"    ✅ Updated table layout to AutoFit to Content")
            
            # Remove fixed column widths from grid columns to allow autofit
            grid_cols = tblGrid.findall(qn('w:gridCol'))
            for i, gridCol in enumerate(grid_cols):
                if gridCol.get(qn('w:w')):
                    gridCol.attrib.pop(qn('w:w'), None)
                    print(f"    ✅ Removed fixed width from column {i+1}")
            
            # Set table style to allow autofit
            tbl.set(qn('w:tblStyle'), 'TableGrid')
            
            # Add specific properties for AutoFit to Content
            # This ensures it's AutoFit to Content, not AutoFit to Window
            tblLook = tblPr.find(qn('w:tblLook'))
            if tblLook is None:
                tblLook = OxmlElement('w:tblLook')
                tblLook.set(qn('w:val'), '04A0')
                tblLook.set(qn('w:firstRow'), '1')
                tblLook.set(qn('w:lastRow'), '0')
                tblLook.set(qn('w:firstColumn'), '1')
                tblLook.set(qn('w:lastColumn'), '0')
                tblLook.set(qn('w:noHBand'), '0')
                tblLook.set(qn('w:noVBand'), '1')
                tblPr.append(tblLook)
            
            # Enable autofit on the table object
            table.allow_autofit = True
            
            # Also remove any fixed widths from python-docx column objects
            for col in table.columns:
                col.width = None
            
            # Force AutoFit to Content by setting table to use content-based sizing
            # This is the key difference between AutoFit to Window vs AutoFit to Content
            tbl.set(qn('w:tblStyle'), 'TableGrid')
            
            # Add table width as "auto" to enable content-based sizing
            tblW_auto = OxmlElement('w:tblW')
            tblW_auto.set(qn('w:w'), '0')
            tblW_auto.set(qn('w:type'), 'auto')
            tblPr.append(tblW_auto)
            
            # Set font size to 12pt for all text in the table
            set_table_font_size(table, 12)
            
            # Ensure all cells have proper alignment
            ensure_table_alignment(table)
            
            print(f"  ✅ Successfully applied AutoFit to Content (equivalent to right-click → AutoFit → AutoFit to Content)")
            
        else:
            print(f"  ⚠️ Could not find table grid, applying AutoFit to Content using fallback method...")
            # Fallback: Apply AutoFit to Content using python-docx methods
            
            # Remove any fixed widths from columns to allow autofit
            for col in table.columns:
                # Reset column width to None to allow autofit
                col.width = None
                print(f"    Reset column width to allow autofit")
            
            # Enable autofit on the table object
            table.allow_autofit = True
            
            # Set font size to 12pt for all text in the table
            set_table_font_size(table, 12)
            
            # Ensure all cells have proper alignment
            ensure_table_alignment(table)
            
            print(f"  ✅ Applied AutoFit to Content using fallback method (equivalent to right-click → AutoFit → AutoFit to Content)")
            
    except Exception as e:
        print(f"  ❌ Error updating column widths: {e}")
        import traceback
        traceback.print_exc()

def get_excel_alignment(cell):
    """
    Get alignment properties from Excel cell
    Returns horizontal and vertical alignment
    """
    horizontal_align = WD_ALIGN_PARAGRAPH.LEFT  # default
    vertical_align = WD_ALIGN_VERTICAL.TOP  # default
    
    if cell.alignment:
        # Horizontal alignment
        if cell.alignment.horizontal == 'center':
            horizontal_align = WD_ALIGN_PARAGRAPH.CENTER
        elif cell.alignment.horizontal == 'right':
            horizontal_align = WD_ALIGN_PARAGRAPH.RIGHT
        elif cell.alignment.horizontal == 'justify':
            horizontal_align = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:  # left or general
            horizontal_align = WD_ALIGN_PARAGRAPH.LEFT
        
        # Vertical alignment
        if cell.alignment.vertical == 'center':
            vertical_align = WD_ALIGN_VERTICAL.CENTER
        elif cell.alignment.vertical == 'bottom':
            vertical_align = WD_ALIGN_VERTICAL.BOTTOM
        else:  # top or general
            vertical_align = WD_ALIGN_VERTICAL.TOP
    
    return horizontal_align, vertical_align


def clean_cell_simple(cell, horizontal_align=WD_ALIGN_PARAGRAPH.LEFT, vertical_align=WD_ALIGN_VERTICAL.TOP):
    """
    Clean cell content with comprehensive alignment preservation
    """
    try:
        # Ensure cell has at least one paragraph
        if not cell.paragraphs:
            cell.add_paragraph()
        
        # Keep only the first paragraph
        while len(cell.paragraphs) > 1:
            last_para = cell.paragraphs[-1]
            last_para._p.getparent().remove(last_para._p)
        
        # Clear and reset the first paragraph with alignment
        first_para = cell.paragraphs[0]
        cell_text = cell.text.strip() if cell.text else ""
        
        # Remove all runs
        for run in first_para.runs[:]:
            first_para._p.remove(run._r)
        
        # Add a clean run with the text
        if cell_text:
            first_para.add_run(cell_text)
        
        # Apply horizontal alignment to the paragraph
        first_para.alignment = horizontal_align
        
        # Apply vertical alignment to the cell
        cell.vertical_alignment = vertical_align
        
        # Debug: Print alignment info
        print(f"    📐 Applied alignment: H={horizontal_align}, V={vertical_align}")
        
    except Exception as e:
        print(f"    ⚠️ Alignment preservation failed: {e}")
        # Fallback: just set basic alignment
        if cell.paragraphs:
            cell.paragraphs[0].alignment = horizontal_align
        cell.vertical_alignment = vertical_align


def make_specific_words_bold_in_table(table):
    """
    Make specific words bold in the table with case sensitivity
    """
    # List of words to make bold (case-sensitive)
    words_to_bold = [
        "SR. NO.",
        "SR. NO",
        "BRANCH CODE",
        "BRANCH NAME",
        "DATE OF VISITS",
        "ASSISTED BY",
        "AUDITED BY",
        "AUDITOR NAME",
        "BRANCH CHECK POINTS",
        "RISK FACTOR",
        "COMPLIANCE/NON-COMPLIANCE",
        "REMARKS",
        "SYSTEM IP'S",
        "CBS ACCESS CONTROL",
        "PHYSICAL AND ENVIRONMENTAL SECURITY",
        "POWER BACK UP",
        "USER AWARENESS",
        "MAINTENANCE AND BUSINESS CONTINUITY CONTROLS",
        "PATCH MANAGEMENT",
        "NETWORK SECURITY",
        "ENDPOINTS VULNERABILITY",
        "ATM MACHINE ROOM",
        "EMAIL-SECURITY",
        "REMOTE ACCESS",
        "UNAUTHORIZED APPLICATIONS / PERSONAL DATA",
        "Important Note"
    ]
    
    print(f"  🔤 Making specific words bold (case-sensitive)...")
    bold_count = 0
    
    try:
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    cell_text = paragraph.text
                    
                    # Check if any of the words to bold are in this cell
                    for word in words_to_bold:
                        if word in cell_text:
                            # Found a match - need to make it bold
                            print(f"    Found '{word}' in Row {row_idx+1}, Col {col_idx+1}")
                            
                            # Split the text around the word
                            parts = cell_text.split(word)
                            
                            # Clear the paragraph runs
                            for run in paragraph.runs[:]:
                                paragraph._p.remove(run._r)
                            
                            # Reconstruct with bold word
                            for i, part in enumerate(parts):
                                if part:
                                    # Add regular text
                                    run = paragraph.add_run(part)
                                    run.font.size = Pt(12)
                                    run.font.name = 'Times New Roman'
                                
                                # Add bold word (except after last part)
                                if i < len(parts) - 1:
                                    run = paragraph.add_run(word)
                                    run.font.size = Pt(12)
                                    run.font.name = 'Times New Roman'
                                    run.font.bold = True
                                    bold_count += 1
                            
                            # Update cell_text for next iteration
                            cell_text = paragraph.text
        
        print(f"  ✅ Made {bold_count} word(s) bold in the table")
        
    except Exception as e:
        print(f"  ❌ Error making words bold: {e}")
        import traceback
        traceback.print_exc()

def replace_branch_review_with_excel_tables(doc, excel_file_path, placeholder="Braaanch_Revviewww"):
    """
    Replace the placeholder with tables from all worksheets in the Excel file.
    Uses optimized approach with alignment preservation.
    
    Args:
        doc: Document object
        excel_file_path: Path to the Excel file
        placeholder: Text to replace
    """
    try:
        print(f"\n🔄 Replacing '{placeholder}' with Excel tables (with alignment)...")
        
        # Load the Excel workbook
        workbook = load_workbook(excel_file_path)
        
        if not workbook.sheetnames:
            print("⚠️ No worksheets found in Excel file")
            return
        
        print(f"✅ Found {len(workbook.sheetnames)} worksheet(s) in Excel file")
        
        # Find the paragraph containing the placeholder
        placeholder_paragraph = None
        
        for idx, paragraph in enumerate(doc.paragraphs):
            if placeholder in paragraph.text:
                placeholder_paragraph = paragraph
                print(f"✅ Found placeholder at paragraph index {idx}")
                break
        
        if not placeholder_paragraph:
            print(f"⚠️ Placeholder '{placeholder}' not found in document")
            return
        
        # Get the parent element to insert content
        parent = placeholder_paragraph._element.getparent()
        placeholder_element = placeholder_paragraph._element
        
        # Process each worksheet
        for sheet_idx, sheet_name in enumerate(workbook.sheetnames, 1):
            print(f"\n📋 Processing worksheet {sheet_idx}: {sheet_name}")
            worksheet = workbook[sheet_name]
            
            max_row = worksheet.max_row
            max_col = worksheet.max_column
            
            if max_row == 0 or max_col == 0:
                print(f"  ⚠️ Worksheet '{sheet_name}' is empty, skipping...")
                continue
            
            print(f"  Creating table with {max_row} rows and {max_col} columns")
            
            merged_ranges = list(worksheet.merged_cells.ranges)
            
            # Create table
            table = doc.add_table(rows=max_row, cols=max_col)
            table.style = 'Table Grid'
            
            # Store alignment information for each cell
            alignment_data = {}
            
            # First pass: collect alignment data from Excel
            print(f"  📊 Collecting alignment data for {max_row} rows and {max_col} columns...")
            for row in range(1, max_row + 1):
                for col in range(1, max_col + 1):
                    excel_cell = worksheet.cell(row=row, column=col)
                    horizontal_align, vertical_align = get_excel_alignment(excel_cell)
                    alignment_data[(row, col)] = (horizontal_align, vertical_align)
                    print(f"    Row {row}, Col {col}: H={horizontal_align}, V={vertical_align}")
            
            # Second pass: copy data and apply alignment
            for row in range(1, max_row + 1):
                for col in range(1, max_col + 1):
                    word_cell = table.rows[row-1].cells[col-1]
                    excel_cell = worksheet.cell(row=row, column=col)
                    
                    # Get alignment for this cell
                    horizontal_align, vertical_align = alignment_data[(row, col)]
                    
                    # Check if this cell is part of a merged range (but not the main cell)
                    is_secondary_merged = False
                    main_cell_alignment = None
                    
                    for merged_range in merged_ranges:
                        if (merged_range.min_row <= row <= merged_range.max_row and
                            merged_range.min_col <= col <= merged_range.max_col):
                            if not (row == merged_range.min_row and col == merged_range.min_col):
                                is_secondary_merged = True
                            else:
                                # This is the main cell of merged range
                                main_cell_alignment = alignment_data[(merged_range.min_row, merged_range.min_col)]
                            break
                    
                    if is_secondary_merged and main_cell_alignment:
                        # For secondary merged cells, use the main cell's alignment
                        horizontal_align, vertical_align = main_cell_alignment
                        word_cell.text = ""  # Secondary merged cells should be empty
                        print(f"    Row {row}, Col {col}: Secondary merged - using main cell alignment")
                    elif not is_secondary_merged:
                        # This is either a normal cell or main cell of merged range
                        word_cell.text = str(excel_cell.value) if excel_cell.value else ""
                        print(f"    Row {row}, Col {col}: Normal cell - applying alignment H={horizontal_align}, V={vertical_align}")
                    
                    # Apply cleaning with alignment for each cell
                    clean_cell_simple(word_cell, horizontal_align, vertical_align)
            
            # Third pass: Apply merged cells
            for merged_range in merged_ranges:
                min_row = merged_range.min_row
                min_col = merged_range.min_col
                max_row_range = merged_range.max_row
                max_col_range = merged_range.max_col
                
                try:
                    if (min_row <= len(table.rows) and min_col <= len(table.rows[0].cells)):
                        main_cell = table.cell(min_row-1, min_col-1)
                        
                        # Get the main cell's alignment
                        main_horizontal_align, main_vertical_align = alignment_data[(min_row, min_col)]
                        
                        # Clean before merging with main cell alignment
                        print(f"    Merged cell main: Row {min_row}, Col {min_col} - applying alignment H={main_horizontal_align}, V={main_vertical_align}")
                        clean_cell_simple(main_cell, main_horizontal_align, main_vertical_align)
                        
                        # Merge cells
                        for row in range(min_row, max_row_range + 1):
                            for col in range(min_col, max_col_range + 1):
                                if not (row == min_row and col == min_col):
                                    cell_to_merge = table.cell(row-1, col-1)
                                    # Clean with main cell alignment before merging
                                    print(f"    Merged cell secondary: Row {row}, Col {col} - applying alignment H={main_horizontal_align}, V={main_vertical_align}")
                                    clean_cell_simple(cell_to_merge, main_horizontal_align, main_vertical_align)
                                    main_cell.merge(cell_to_merge)
                                    
                except Exception as e:
                    print(f"    ⚠️ Merge error: {e}")
            
            # Insert the table before the placeholder first
            parent.insert(parent.index(placeholder_element), table._element)
            
            print(f"  ✅ Added table with {max_row} rows and {max_col} columns")
            
            # Now update column widths after table is inserted
            update_table_column_widths(table)
            
            # Make specific words bold in the table (case-sensitive)
            make_specific_words_bold_in_table(table)
            
            # Add page break after each table (including the last one)
            page_break_para = doc.add_paragraph()
            page_break_run = page_break_para.add_run()
            page_break_run.add_break(WD_BREAK.PAGE)
            parent.insert(parent.index(placeholder_element), page_break_para._element)
            
            print(f"  ✅ Added page break after table {sheet_idx}")
        
        # Remove the placeholder paragraph
        parent.remove(placeholder_element)
        print(f"\n✅ Successfully replaced '{placeholder}' with {len(workbook.sheetnames)} table(s)")
        
    except Exception as e:
        print(f"❌ Error replacing branch review: {str(e)}")
        import traceback
        traceback.print_exc()


def cleanup_temp_file(file_path):
    """Clean up temporary file after a delay"""
    try:
        import time
        time.sleep(5)  # Wait for file to be downloaded
        if os.path.exists(file_path):
            os.remove(file_path)
            print(f"Cleaned up temporary file: {file_path}")
    except Exception as e:
        print(f"Error cleaning up temporary file: {e}")

def reduce_specific_column_widths(doc):
    """
    Find tables with specific words and reduce their column widths with case sensitivity
    """
    from docx.shared import Pt, Inches
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    
    # Define words and their target widths (in pixels)
    column_widths = {
        "SR. NO.": 45,  # 40px
        "System IPs": 140,  # 100px
        "MAC Address": 160,  # 100px
        "Employee Name": 130,  # 100px
        "COMPLIANCE/NON-COMPLIANCE": 100,  # 100px
        "REMARKS": 285  # 200px
    }
    
    print(f"\n🔍 Searching for specific words to reduce column widths...")
    tables_modified = 0
    
    try:
        for table_idx, table in enumerate(doc.tables, 1):
            print(f"  📋 Processing Table {table_idx} with {len(table.columns)} columns")
            
            # Check each column for specific words
            for col_idx in range(len(table.columns)):
                # Check all cells in this column for specific words
                found_word = None
                target_width = None
                
                for row in table.rows:
                    try:
                        cell = row.cells[col_idx]
                        cell_text = cell.text.strip()
                        
                        # Check for exact case-sensitive matches
                        for word, width in column_widths.items():
                            if word in cell_text:
                                found_word = word
                                target_width = width
                                print(f"    ✅ Found '{word}' in Table {table_idx}, Column {col_idx + 1}")
                                break
                        
                        if found_word:
                            break
                    except Exception as e:
                        continue
                
                if found_word and target_width:
                    # Reduce this column's width using multiple methods
                    try:
                        print(f"    🔧 Reducing Column {col_idx + 1} width to {target_width}px...")
                        
                        # Convert pixels to inches and twips
                        # 1 inch = 96 pixels, so target_width pixels = target_width/96 inches
                        width_inches = target_width / 96.0
                        width_twips = int(target_width * 15)  # 1 pixel = 15 twips
                        
                        # Method 1: Set python-docx column width
                        table.columns[col_idx].width = Inches(width_inches)
                        
                        # Method 2: Direct XML manipulation for more control
                        tbl = table._tbl
                        tblPr = tbl.find(qn('w:tblPr'))
                        
                        if tblPr is not None:
                            # Set table layout to fixed
                            tblLayout = tblPr.find(qn('w:tblLayout'))
                            if tblLayout is None:
                                tblLayout = OxmlElement('w:tblLayout')
                                tblPr.append(tblLayout)
                            tblLayout.set(qn('w:type'), 'fixed')
                            
                            # Set table width to auto to allow column control
                            tblW = tblPr.find(qn('w:tblW'))
                            if tblW is not None:
                                tblPr.remove(tblW)
                            
                            tblW = OxmlElement('w:tblW')
                            tblW.set(qn('w:w'), '0')
                            tblW.set(qn('w:type'), 'auto')
                            tblPr.append(tblW)
                        
                        # Method 3: Update grid column width in XML
                        tblGrid = tbl.find(qn('w:tblGrid'))
                        if tblGrid is not None:
                            grid_cols = tblGrid.findall(qn('w:gridCol'))
                            if col_idx < len(grid_cols):
                                grid_cols[col_idx].set(qn('w:w'), str(width_twips))
                                print(f"      📏 Set XML grid column width to {width_twips} twips ({target_width}px)")
                        
                        # Method 4: Set cell widths directly for this column
                        for row in table.rows:
                            try:
                                cell = row.cells[col_idx]
                                # Set cell width using XML
                                tc = cell._tc
                                tcPr = tc.find(qn('w:tcPr'))
                                if tcPr is None:
                                    tcPr = OxmlElement('w:tcPr')
                                    tc.insert(0, tcPr)
                                
                                tcW = tcPr.find(qn('w:tcW'))
                                if tcW is not None:
                                    tcPr.remove(tcW)
                                
                                tcW = OxmlElement('w:tcW')
                                tcW.set(qn('w:w'), str(width_twips))
                                tcW.set(qn('w:type'), 'dxa')
                                tcPr.append(tcW)
                                
                            except Exception as e:
                                print(f"      ⚠️ Error setting cell width: {e}")
                        
                        print(f"    ✅ Successfully reduced Column {col_idx + 1} to {target_width}px")
                        tables_modified += 1
                        
                    except Exception as e:
                        print(f"    ❌ Error setting column width: {e}")
                        import traceback
                        traceback.print_exc()
        
        if tables_modified > 0:
            print(f"  ✅ Modified {tables_modified} column(s) with specific words")
        else:
            print(f"  ℹ️ No specific word columns found")
        
    except Exception as e:
        print(f"  ❌ Error reducing specific column widths: {e}")
        import traceback
        traceback.print_exc()

def reduce_sr_no_column_width(doc):
    """
    Find tables with "SR. NO." and reduce that column's width to 15px
    """
    from docx.shared import Pt, Inches
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    
    print(f"\n🔍 Searching for 'SR. NO.' columns to reduce width...")
    tables_modified = 0
    
    try:
        for table_idx, table in enumerate(doc.tables, 1):
            print(f"  📋 Processing Table {table_idx} with {len(table.columns)} columns")
            
            # Check each column for "SR. NO."
            for col_idx in range(len(table.columns)):
                # Check all cells in this column for "SR. NO."
                column_has_sr_no = False
                
                for row in table.rows:
                    try:
                        cell = row.cells[col_idx]
                        cell_text = cell.text.strip()
                        
                        if "SR. NO." in cell_text:
                            column_has_sr_no = True
                            print(f"    ✅ Found 'SR. NO.' in Table {table_idx}, Column {col_idx + 1}")
                            break
                    except Exception as e:
                        continue
                
                if column_has_sr_no:
                    # Reduce this column's width to 15px using multiple methods
                    try:
                        print(f"    🔧 Reducing Column {col_idx + 1} width to 15px...")
                        
                        # Method 1: Set python-docx column width
                        table.columns[col_idx].width = Inches(0.708)  # 15px ≈ 0.208 inches
                        
                        # Method 2: Direct XML manipulation for more control
                        tbl = table._tbl
                        tblPr = tbl.find(qn('w:tblPr'))
                        
                        if tblPr is not None:
                            # Set table layout to fixed
                            tblLayout = tblPr.find(qn('w:tblLayout'))
                            if tblLayout is None:
                                tblLayout = OxmlElement('w:tblLayout')
                                tblPr.append(tblLayout)
                            tblLayout.set(qn('w:type'), 'fixed')
                            
                            # Set table width to auto to allow column control
                            tblW = tblPr.find(qn('w:tblW'))
                            if tblW is not None:
                                tblPr.remove(tblW)
                            
                            tblW = OxmlElement('w:tblW')
                            tblW.set(qn('w:w'), '0')
                            tblW.set(qn('w:type'), 'auto')
                            tblPr.append(tblW)
                        
                        # Method 3: Update grid column width in XML
                        tblGrid = tbl.find(qn('w:tblGrid'))
                        if tblGrid is not None:
                            grid_cols = tblGrid.findall(qn('w:gridCol'))
                            if col_idx < len(grid_cols):
                                # 15px = 300 twips (1 inch = 1440 twips)
                                grid_cols[col_idx].set(qn('w:w'), '300')
                                print(f"      📏 Set XML grid column width to 300 twips (15px)")
                        
                        # Method 4: Set cell widths directly for this column
                        for row in table.rows:
                            try:
                                cell = row.cells[col_idx]
                                # Set cell width using XML
                                tc = cell._tc
                                tcPr = tc.find(qn('w:tcPr'))
                                if tcPr is None:
                                    tcPr = OxmlElement('w:tcPr')
                                    tc.insert(0, tcPr)
                                
                                tcW = tcPr.find(qn('w:tcW'))
                                if tcW is not None:
                                    tcPr.remove(tcW)
                                
                                tcW = OxmlElement('w:tcW')
                                tcW.set(qn('w:w'), '300')  # 15px = 300 twips
                                tcW.set(qn('w:type'), 'dxa')
                                tcPr.append(tcW)
                                
                            except Exception as e:
                                print(f"      ⚠️ Error setting cell width: {e}")
                        
                        print(f"    ✅ Successfully reduced Column {col_idx + 1} to 15px")
                        tables_modified += 1
                        
                    except Exception as e:
                        print(f"    ❌ Error setting column width: {e}")
                        import traceback
                        traceback.print_exc()
        
        if tables_modified > 0:
            print(f"  ✅ Modified {tables_modified} column(s) with 'SR. NO.'")
        else:
            print(f"  ℹ️ No 'SR. NO.' columns found")
        
    except Exception as e:
        print(f"  ❌ Error reducing SR. NO. column width: {e}")
        import traceback
        traceback.print_exc()

def add_page_borders_to_document(doc):
    """
    Add page borders to all pages in the document.
    Top: 1pt black
    Right: 1pt black  
    Left: 1pt black
    Bottom: 1/16pt white (very thin white line)
    """
    try:
        print("\n🖼️ Adding page borders to document...")
        
        for section_idx, section in enumerate(doc.sections):
            section_element = section._sectPr
            
            # Remove any existing page borders first
            existing_borders = section_element.find(qn('w:pgBorders'))
            if existing_borders is not None:
                section_element.remove(existing_borders)
            
            # Create the page borders element
            # 1/16pt = 0.5 half-points (w:sz="1")
            pg_borders_xml = '''
                <w:pgBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                             w:offsetFrom="page">
                    <w:top w:val="single" w:sz="8" w:space="24" w:color="000000"/>
                    <w:left w:val="single" w:sz="8" w:space="24" w:color="000000"/>
                    <w:bottom w:val="single" w:sz="1" w:space="24" w:color="FFFFFF"/>
                    <w:right w:val="single" w:sz="8" w:space="24" w:color="000000"/>
                </w:pgBorders>
            '''
            
            # Parse the XML and add it to the section
            pg_borders = parse_xml(pg_borders_xml)
            
            # Insert pgBorders at the correct position in sectPr
            inserted = False
            for i, child in enumerate(section_element):
                tag = child.tag
                if tag in [qn('w:pgSz'), qn('w:pgMar'), qn('w:cols'), qn('w:docGrid')]:
                    section_element.insert(i, pg_borders)
                    inserted = True
                    break
            
            if not inserted:
                section_element.append(pg_borders)
            
            print(f"  ✅ Added borders to section {section_idx + 1}")
        
        print(f"✅ Successfully added page borders to all {len(doc.sections)} section(s)")
        print("   Border specifications:")
        print("   - Top: 1pt black")
        print("   - Left: 1pt black")
        print("   - Right: 1pt black")
        print("   - Bottom: 1/16pt white (very thin)")
        print("   - Offset: 24pt from page edge")
        
        return True
        
    except Exception as e:
        print(f"❌ Error adding page borders: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def format_low_risk_text(doc):
    """
    Find all instances of "Low risk:" in the document and make them bold and underlined.
    Only formats the "Low risk:" text, not the entire paragraph/run.
    """
    print("\n📝 Formatting 'Low risk:' text in document...")
    formatted_count = 0
    
    # Helper function to format "Low risk:" in a paragraph
    def format_in_paragraph(para):
        nonlocal formatted_count
        full_text = para.text
        
        if 'Low risk:' in full_text or 'low risk:' in full_text:
            # Find the position of "Low risk:" or "low risk:"
            search_terms = ['Low risk:', 'low risk:']
            
            for search_term in search_terms:
                if search_term in full_text:
                    # Clear existing runs
                    for run in para.runs:
                        run.text = ""
                    
                    # Split text at "Low risk:" or "low risk:"
                    parts = full_text.split(search_term, 1)
                    
                    if len(parts) == 2:
                        # Add text before "Low risk:"
                        if parts[0]:
                            before_run = para.add_run(parts[0])
                            before_run.font.size = Pt(12)
                            before_run.font.name = 'Times New Roman'
                        
                        # Add "Low risk:" with bold and underline
                        low_risk_run = para.add_run('Low risk:')
                        low_risk_run.font.size = Pt(12)
                        low_risk_run.font.name = 'Times New Roman'
                        low_risk_run.bold = True
                        low_risk_run.underline = True
                        
                        # Add text after "Low risk:"
                        if parts[1]:
                            after_run = para.add_run(parts[1])
                            after_run.font.size = Pt(12)
                            after_run.font.name = 'Times New Roman'
                        
                        formatted_count += 1
                        print(f"  ✅ Formatted 'Low risk:' (only the text, not the whole line)")
                    
                    # Only process first occurrence
                    break
    
    # Search in all paragraphs
    for paragraph in doc.paragraphs:
        format_in_paragraph(paragraph)
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    format_in_paragraph(paragraph)
    
    print(f"✅ Successfully formatted {formatted_count} instance(s) of 'Low risk:'")
    print("   Formatting: Bold + Underline (only 'Low risk:' text)")
    return formatted_count


def add_borders_to_all_images(doc):
    """
    Add 1pt black border to all images in the document.
    """
    print("\n🖼️ Adding borders to all images in document...")
    image_count = 0
    
    # Iterate through all paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Check if run contains an inline shape (image)
            if run._element.xml.find('pic:pic') != -1:
                try:
                    # Get the inline shape
                    inline = run._element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
                    
                    if inline is not None:
                        # Get the picture element
                        graphic = inline.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}graphic')
                        
                        if graphic is not None:
                            graphicData = graphic.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}graphicData')
                            
                            if graphicData is not None:
                                pic = graphicData.find('.//{http://schemas.openxmlformats.org/drawingml/2006/picture}pic')
                                
                                if pic is not None:
                                    spPr = pic.find('.//{http://schemas.openxmlformats.org/drawingml/2006/picture}spPr')
                                    
                                    if spPr is not None:
                                        # Remove existing border if any
                                        existing_ln = spPr.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}ln')
                                        if existing_ln is not None:
                                            spPr.remove(existing_ln)
                                        
                                        # Create 1pt black border (12700 EMUs = 1pt)
                                        border_xml = r"""
                                            <a:ln xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" w="16867">
                                                <a:solidFill>
                                                    <a:srgbClr val="000000"/>
                                                </a:solidFill>
                                                <a:prstDash val="solid"/>
                                            </a:ln>
                                        """
                                        
                                        # Parse and add the border
                                        ln_element = parse_xml(border_xml)
                                        spPr.append(ln_element)
                                        
                                        image_count += 1
                                        print(f"  ✅ Added 1pt black border to image {image_count}")
                                        
                except Exception as e:
                    print(f"  ⚠️ Error adding border to image: {e}")
                    continue
    
    # Also check images in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run._element.xml.find('pic:pic') != -1:
                            try:
                                inline = run._element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
                                
                                if inline is not None:
                                    graphic = inline.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}graphic')
                                    
                                    if graphic is not None:
                                        graphicData = graphic.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}graphicData')
                                        
                                        if graphicData is not None:
                                            pic = graphicData.find('.//{http://schemas.openxmlformats.org/drawingml/2006/picture}pic')
                                            
                                            if pic is not None:
                                                spPr = pic.find('.//{http://schemas.openxmlformats.org/drawingml/2006/picture}spPr')
                                                
                                                if spPr is not None:
                                                    # Remove existing border if any
                                                    existing_ln = spPr.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}ln')
                                                    if existing_ln is not None:
                                                        spPr.remove(existing_ln)
                                                    
                                                    # Create 1pt black border
                                                    border_xml = r"""
                                                        <a:ln xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" w="12700">
                                                            <a:solidFill>
                                                                <a:srgbClr val="000000"/>
                                                            </a:solidFill>
                                                            <a:prstDash val="solid"/>
                                                        </a:ln>
                                                    """
                                                    
                                                    ln_element = parse_xml(border_xml)
                                                    spPr.append(ln_element)
                                                    
                                                    image_count += 1
                                                    print(f"  ✅ Added 1pt black border to image {image_count} (in table)")
                                                    
                            except Exception as e:
                                print(f"  ⚠️ Error adding border to image in table: {e}")
                                continue
    
    print(f"✅ Successfully added borders to {image_count} image(s)")
    print("   Border specification: 1pt solid black")
    return image_count


def _replace_text_in_runs(paragraph, old_text, new_text):
    """
    Helper function to replace text in a paragraph while preserving formatting.
    This handles cases where a single placeholder is split across multiple runs.
    Preserves the exact formatting of the original placeholder text.
    """
    # Combine all run texts to a single string for easy replacement
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return False

    # Find which run(s) contain the old text and preserve their formatting
    run_formats = []
    text_positions = []
    
    # Track where the old text appears in the combined text
    start_pos = 0
    while True:
        pos = full_text.find(old_text, start_pos)
        if pos == -1:
            break
        text_positions.append((pos, pos + len(old_text)))
        start_pos = pos + 1
    
    if not text_positions:
        return False
    
    # Get detailed formatting information for each run
    current_pos = 0
    for run in paragraph.runs:
        run_start = current_pos
        run_end = current_pos + len(run.text)
        
        # Check if this run contains part of the old text
        for text_start, text_end in text_positions:
            if run_start < text_end and run_end > text_start:
                # This run contains part of the old text
                # Safely get font color to avoid None reference errors
                font_color = None
                try:
                    if run.font.color.rgb is not None:
                        font_color = run.font.color.rgb
                except:
                    font_color = None
                
                run_formats.append({
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name if run.font.name else 'Times New Roman',
                    'font_size': run.font.size,
                    'font_color': font_color,
                    'start_pos': run_start,
                    'end_pos': run_end
                })
                break
        
        current_pos = run_end
    
    # If we found formatting info, use the first one (most common case)
    if run_formats:
        # Clear existing runs by setting text to empty (safer than run.clear())
        for run in paragraph.runs:
            run.text = ""
        
        # Replace the text
        new_full_text = full_text.replace(old_text, new_text)
        
        # Use the first run to add all the new text with preserved formatting
        if len(paragraph.runs) > 0:
            first_run = paragraph.runs[0]
            first_run.text = new_full_text
            
            # Apply the formatting from the original run that contained the placeholder
            original_format = run_formats[0]
            first_run.bold = original_format.get('bold')
            first_run.italic = original_format.get('italic')
            first_run.underline = original_format.get('underline')
            
            # Set font name and size
            if original_format.get('font_name'):
                first_run.font.name = original_format.get('font_name')
            if original_format.get('font_size'):
                first_run.font.size = original_format.get('font_size')
            
            # Safely set font color
            if original_format.get('font_color') is not None:
                try:
                    first_run.font.color.rgb = original_format.get('font_color')
                except:
                    pass  # Skip color setting if it fails
            
            # Remove extra empty runs (but keep at least one)
            runs_to_remove = list(paragraph.runs[1:])
            for run in runs_to_remove:
                run._element.getparent().remove(run._element)
        else:
            # If no runs exist, create a new one with preserved formatting
            new_run = paragraph.add_run(new_full_text)
            original_format = run_formats[0]
            new_run.bold = original_format.get('bold')
            new_run.italic = original_format.get('italic')
            new_run.underline = original_format.get('underline')
            new_run.font.name = original_format.get('font_name', 'Times New Roman')
            if original_format.get('font_size'):
                new_run.font.size = original_format.get('font_size')
    else:
        # Fallback: simple replacement if we can't find formatting info
        if len(paragraph.runs) == 1 and old_text in paragraph.runs[0].text:
            paragraph.runs[0].text = paragraph.runs[0].text.replace(old_text, new_text)
        else:
            # Multiple runs case - use first run's formatting
            if paragraph.runs:
                first_run = paragraph.runs[0]
                original_text = first_run.text
                first_run.text = original_text.replace(old_text, new_text)

    return True

def convert_to_dd_mm_yyyy(date_str):
    """Convert various date formats to DD.MM.YYYY format"""
    date_formats = ["%d.%m.%Y", "%Y-%m-%d", "%m/%d/%Y", "%d-%m-%Y", "%d/%m/%Y", "%Y/%m/%d"]
    for fmt in date_formats:
        try:
            date_obj = datetime.strptime(date_str, fmt)
            return date_obj.strftime("%d.%m.%Y")
        except ValueError:
            continue
    raise ValueError(f"Could not parse date: {date_str}")

def get_ordinal_suffix(day):
    """Get the ordinal suffix for a day (st, nd, rd, th)"""
    if 10 <= day % 100 <= 20:
        suffix = 'ᵗʰ'
    else:
        suffix = {1: 'ˢᵗ', 2: 'ⁿᵈ', 3: 'ʳᵈ'}.get(day % 10, 'ᵗʰ')
    return f"{day}{suffix}"

def calculate_financial_year(date_str):
    """
    Calculate financial year based on the date.
    Financial year runs from 1st April to 31st March.
    If date is between 1 April to 31 March, return "1ˢᵗ April YYYY to 31ˢᵗ March YYYY+1"
    """
    try:
        # Parse the date
        date_formats = ["%Y-%m-%d", "%d-%m-%Y", "%m/%d/%Y", "%d/%m/%Y", "%Y/%m/%d", "%d.%m.%Y"]
        date_obj = None
        for fmt in date_formats:
            try:
                date_obj = datetime.strptime(date_str, fmt)
                break
            except ValueError:
                continue
        
        if not date_obj:
            return ""
        
        # Determine financial year
        # If month is Jan-Mar (1-3), financial year started previous year
        # If month is Apr-Dec (4-12), financial year started current year
        if date_obj.month >= 4:  # April to December
            fy_start_year = date_obj.year
            fy_end_year = date_obj.year + 1
        else:  # January to March
            fy_start_year = date_obj.year - 1
            fy_end_year = date_obj.year
        
        return f"1ˢᵗ April {fy_start_year} to 31ˢᵗ March {fy_end_year}"
    
    except Exception as e:
        print(f"Error calculating financial year for date '{date_str}': {e}")
        return ""

def format_date_range(start_date_str, end_date_str):
    """
    Format date range to 'DD.MM.YYYY To DD.MM.YYYY' format
    """
    try:
        start_formatted = convert_to_dd_mm_yyyy(start_date_str)
        end_formatted = convert_to_dd_mm_yyyy(end_date_str)
        return f"{start_formatted} To {end_formatted}"
    except Exception as e:
        print(f"Error formatting date range: {e}")
        return f"{start_date_str} To {end_date_str}"

def replace_text_in_document(doc, replacements):
    """
    Replace text in the entire document including paragraphs and tables.
    Handles placeholders that might be split across multiple runs.
    Special handling for Application_Nameeeeeesssssssss to ensure left alignment.
    """
    try:
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    if old_text == "Application_Nameeeeeesssssssss":
                        # Special handling for applications list - ensure left alignment
                        _replace_text_in_runs_with_alignment(paragraph, old_text, new_text, "left")
                    else:
                        _replace_text_in_runs(paragraph, old_text, new_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            if old_text in paragraph.text:
                                if old_text == "Application_Nameeeeeesssssssss":
                                    # Special handling for applications list - ensure left alignment
                                    _replace_text_in_runs_with_alignment(paragraph, old_text, new_text, "left")
                                else:
                                    _replace_text_in_runs(paragraph, old_text, new_text)
        
        # Replace in headers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        if old_text == "Application_Nameeeeeesssssssss":
                            _replace_text_in_runs_with_alignment(paragraph, old_text, new_text, "left")
                        else:
                            _replace_text_in_runs(paragraph, old_text, new_text)
            
            # Replace in header tables
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for old_text, new_text in replacements.items():
                                if old_text in paragraph.text:
                                    if old_text == "Application_Nameeeeeesssssssss":
                                        _replace_text_in_runs_with_alignment(paragraph, old_text, new_text, "left")
                                    else:
                                        _replace_text_in_runs(paragraph, old_text, new_text)
        
        # Replace in footers
        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        if old_text == "Application_Nameeeeeesssssssss":
                            _replace_text_in_runs_with_alignment(paragraph, old_text, new_text, "left")
                        else:
                            _replace_text_in_runs(paragraph, old_text, new_text)
            
            # Replace in footer tables
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for old_text, new_text in replacements.items():
                                if old_text in paragraph.text:
                                    if old_text == "Application_Nameeeeeesssssssss":
                                        _replace_text_in_runs_with_alignment(paragraph, old_text, new_text, "left")
                                    else:
                                        _replace_text_in_runs(paragraph, old_text, new_text)
        
        print("✅ All replacements completed successfully")
        return True
    
    except Exception as e:
        print(f"❌ Error replacing text in document: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def _replace_text_in_runs_with_alignment(paragraph, old_text, new_text, alignment="left"):
    """
    Helper function to replace text in a paragraph while preserving formatting and setting alignment.
    Special function for Application_Nameeeeeesssssssss to ensure left alignment.
    """
    # Combine all run texts to a single string for easy replacement
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return False

    # Find which run(s) contain the old text and preserve their formatting
    run_formats = []
    text_positions = []
    
    # Track where the old text appears in the combined text
    start_pos = 0
    while True:
        pos = full_text.find(old_text, start_pos)
        if pos == -1:
            break
        text_positions.append((pos, pos + len(old_text)))
        start_pos = pos + 1
    
    if not text_positions:
        return False
    
    # Get detailed formatting information for each run
    current_pos = 0
    for run in paragraph.runs:
        run_start = current_pos
        run_end = current_pos + len(run.text)
        
        # Check if this run contains part of the old text
        for text_start, text_end in text_positions:
            if run_start < text_end and run_end > text_start:
                # This run contains part of the old text
                # Safely get font color to avoid None reference errors
                font_color = None
                try:
                    if run.font.color.rgb is not None:
                        font_color = run.font.color.rgb
                except:
                    font_color = None
                
                run_formats.append({
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name if run.font.name else 'Times New Roman',
                    'font_size': run.font.size,
                    'font_color': font_color,
                    'start_pos': run_start,
                    'end_pos': run_end
                })
                break
        
        current_pos = run_end
    
    # If we found formatting info, use the first one (most common case)
    if run_formats:
        # Clear existing runs by setting text to empty (safer than run.clear())
        for run in paragraph.runs:
            run.text = ""
        
        # Replace the text
        new_full_text = full_text.replace(old_text, new_text)
        
        # Use the first run to add all the new text with preserved formatting
        if len(paragraph.runs) > 0:
            first_run = paragraph.runs[0]
            first_run.text = new_full_text
            
            # Apply the formatting from the original run that contained the placeholder
            original_format = run_formats[0]
            first_run.bold = original_format.get('bold')
            first_run.italic = original_format.get('italic')
            first_run.underline = original_format.get('underline')
            
            # Set font name and size
            if original_format.get('font_name'):
                first_run.font.name = original_format.get('font_name')
            if original_format.get('font_size'):
                first_run.font.size = original_format.get('font_size')
            
            # Safely set font color
            if original_format.get('font_color') is not None:
                try:
                    first_run.font.color.rgb = original_format.get('font_color')
                except:
                    pass  # Skip color setting if it fails
            
            # Remove extra empty runs (but keep at least one)
            runs_to_remove = list(paragraph.runs[1:])
            for run in runs_to_remove:
                run._element.getparent().remove(run._element)
        else:
            # If no runs exist, create a new one with preserved formatting
            new_run = paragraph.add_run(new_full_text)
            original_format = run_formats[0]
            new_run.bold = original_format.get('bold')
            new_run.italic = original_format.get('italic')
            new_run.underline = original_format.get('underline')
            new_run.font.name = original_format.get('font_name', 'Times New Roman')
            if original_format.get('font_size'):
                new_run.font.size = original_format.get('font_size')
        
        # Set paragraph alignment to left for applications list
        if alignment == "left":
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            print(f"✅ Set left alignment for Application_Nameeeeeesssssssss replacement")
    else:
        # Fallback: simple replacement if we can't find formatting info
        if len(paragraph.runs) == 1 and old_text in paragraph.runs[0].text:
            paragraph.runs[0].text = paragraph.runs[0].text.replace(old_text, new_text)
            if alignment == "left":
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Multiple runs case - use first run's formatting
            if paragraph.runs:
                first_run = paragraph.runs[0]
                original_text = first_run.text
                first_run.text = original_text.replace(old_text, new_text)
                if alignment == "left":
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    return True

@is_audit_word_report_bp.route('/process_is_audit_word_report', methods=['POST'])
def process_is_audit_word_report():
    """
    Process IS Audit Word Report form submission.
    Replaces placeholders in the Word template with user-provided data.
    """
    try:
        print("\n" + "="*80)
        print("IS AUDIT WORD REPORT - FORM DATA RECEIVED")
        print("="*80 + "\n")
        
        # Organization Details
        print("--- ORGANIZATION DETAILS ---")
        organization_name = request.form.get('organizationName')
        organization_name_other = request.form.get('organizationNameOther')
        final_organization = organization_name_other if organization_name == 'Other' else organization_name
        print(f"Organization Name: {final_organization}")
        
        city = request.form.get('city')
        city_other = request.form.get('cityOther')
        final_city = city_other if city == 'Other' else city
        print(f"City: {final_city}")
        
        state = request.form.get('state')
        print(f"State: {state}")
        
        # Auditor Details
        print("\n--- AUDITOR DETAILS ---")
        prepared_by_prefix = request.form.get('preparedByPrefix')
        prepared_by_name = request.form.get('preparedByName')
        prepared_by = f"{prepared_by_prefix} {prepared_by_name}"
        print(f"Prepared By: {prepared_by}")
        
        start_audit_date = request.form.get('startAuditDate')
        print(f"Start Audit Date: {start_audit_date}")
        
        end_audit_date = request.form.get('endAuditDate')
        print(f"End Audit Date: {end_audit_date}")
        
        # Calculate financial year based on start date
        financial_year = calculate_financial_year(start_audit_date)
        print(f"Financial Year: {financial_year}")
        
        # Format date range
        date_range = format_date_range(start_audit_date, end_audit_date)
        print(f"Date Range: {date_range}")
        
        # Auditee Details
        print("\n--- AUDITEE DETAILS ---")
        auditee_person_prefix = request.form.get('auditeePersonPrefix')
        auditee_person_name = request.form.get('auditeePersonName')
        auditee_person = f"{auditee_person_prefix} {auditee_person_name}"
        print(f"Auditee Organization Person Name: {auditee_person}")
        
        auditee_designation = request.form.get('auditeeDesignation')
        print(f"Auditee Designation: {auditee_designation}")
        
        address = request.form.get('address')
        print(f"Address: {address}")
        
        emails = request.form.getlist('email[]')
        # Join all emails with comma and space
        all_emails = ", ".join(emails)
        print(f"Email Addresses ({len(emails)} total): {all_emails}")
        
        # Banking Details
        print("\n--- BANKING DETAILS ---")
        bank_under = request.form.get('bankUnder')
        print(f"Bank Under: {bank_under}")
        
        number_of_branches = request.form.get('numberOfBranches')
        print(f"Number of Branches: {number_of_branches}")
        
        cbs_provider_name = request.form.get('cbsProviderName')
        print(f"CBS Provider Name: {cbs_provider_name}")
        
        business_activity_model = request.form.get('businessActivityModel')
        business_activity_model_other = request.form.get('businessActivityModelOther')
        final_business_model = business_activity_model_other if business_activity_model == 'Other' else business_activity_model
        print(f"Business Activity Model: {final_business_model}")
        
        # Infrastructure Details
        print("\n--- INFRASTRUCTURE DETAILS ---")
        cbs_datacenter_managed = request.form.get('cbsDataCenterManaged')
        print(f"CBS Data Center Managed by: {cbs_datacenter_managed}")
        
        dc_location = request.form.get('dcLocation')
        dc_location_other = request.form.get('dcLocationOther')
        final_dc_location = dc_location_other if dc_location == 'Other' else dc_location
        print(f"DC Location: {final_dc_location}")
        
        dr_location = request.form.get('drLocation')
        dr_location_other = request.form.get('drLocationOther')
        final_dr_location = dr_location_other if dr_location == 'Other' else dr_location
        print(f"DR Location: {final_dr_location}")
        
        primary_connectivity = request.form.get('primaryConnectivity')
        print(f"Primary Connectivity: {primary_connectivity}")
        
        secondary_connectivity = request.form.get('secondaryConnectivity')
        print(f"Secondary Connectivity: {secondary_connectivity}")
        
        # New conditional fields
        has_other_services = request.form.get('hasOtherServices')
        print(f"Has Other Services: {has_other_services}")
        
        other_services = request.form.get('otherServices', '')
        print(f"Other Services: {other_services if other_services else 'N/A'}")
        
        other_services_dc = request.form.get('otherServicesDC', '')
        other_services_dc_other = request.form.get('otherServicesDCOther', '')
        final_other_services_dc = other_services_dc_other if other_services_dc == 'Other' else other_services_dc
        print(f"Other Services Host DC Location: {final_other_services_dc if final_other_services_dc else 'N/A'}")
        
        other_services_dr = request.form.get('otherServicesDR', '')
        other_services_dr_other = request.form.get('otherServicesDROther', '')
        final_other_services_dr = other_services_dr_other if other_services_dr == 'Other' else other_services_dr
        print(f"Other Services Host DR Location: {final_other_services_dr if final_other_services_dr else 'N/A'}")
        
        is_direct_member = request.form.get('isDirectMember')
        print(f"Is Direct Member of RBI and NPCI: {is_direct_member}")
        
        number_of_submember_banks = request.form.get('numberOfSubmemberBanks', '')
        print(f"Number of Sub-member Banks: {number_of_submember_banks if number_of_submember_banks else 'N/A'}")
        
        sponsor_bank_name = request.form.get('sponsorBankName', '')
        sponsor_bank_name_other = request.form.get('sponsorBankNameOther', '')
        final_sponsor_bank = sponsor_bank_name_other if sponsor_bank_name == 'Other' else sponsor_bank_name
        print(f"Sponsor Bank Name: {final_sponsor_bank if final_sponsor_bank else 'N/A'}")
        
        # Software & Security Details
        print("\n--- SOFTWARE & SECURITY DETAILS ---")
        hardware_vendor_name = request.form.get('hardwareVendorName')
        print(f"Hardware Vendor Name: {hardware_vendor_name}")
        
        antivirus_name = request.form.get('antivirusName')
        print(f"Antivirus Name: {antivirus_name}")
        
        has_antimalware = request.form.get('hasAntimalware')
        antimalware_name = request.form.get('antimalwareName', '')
        print(f"Antimalware Solution: {has_antimalware}")
        if has_antimalware == 'Yes':
            print(f"Antimalware Name: {antimalware_name}")
        
        mail_messaging_service = request.form.get('mailMessagingService')
        print(f"Mail and Messaging Service: {mail_messaging_service}")
        
        has_ad = request.form.get('hasAD')
        print(f"Bank has AD: {has_ad}")
        
        has_pam = request.form.get('hasPAM')
        pam_name = request.form.get('pamName', '')
        print(f"Bank has PAM: {has_pam}")
        if has_pam == 'Yes':
            print(f"PAM Name: {pam_name}")
        
        applications = request.form.getlist('applications[]')
        print(f"CBS and Non-CBS Applications ({len(applications)} total):")
        for idx, app in enumerate(applications, 1):
            print(f"  {idx}. {app}")
        
        # Upload Files
        print("\n--- UPLOAD FILES ---")
        excel_file_1 = request.files.get('excelFile1')
        excel_file_2 = request.files.get('excelFile2')
        
        excel_file_1_path = None
        excel_file_2_path = None
        
        # Save Excel File 1 to temporary location
        if excel_file_1:
            print(f"Excel File 1: {excel_file_1.filename}")
            excel_file_1_filename = secure_filename(excel_file_1.filename)
            temp_dir = tempfile.gettempdir()
            excel_file_1_path = os.path.join(temp_dir, f"temp_excel1_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{excel_file_1_filename}")
            excel_file_1.save(excel_file_1_path)
            print(f"  ✅ Excel File 1 saved to: {excel_file_1_path}")
        
        # Save Excel File 2 to temporary location (if needed in future)
        if excel_file_2:
            print(f"Excel File 2: {excel_file_2.filename}")
            excel_file_2_filename = secure_filename(excel_file_2.filename)
            temp_dir = tempfile.gettempdir()
            excel_file_2_path = os.path.join(temp_dir, f"temp_excel2_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{excel_file_2_filename}")
            excel_file_2.save(excel_file_2_path)
            print(f"  ✅ Excel File 2 saved to: {excel_file_2_path}")
        
        print("\n" + "="*80)
        print("PROCESSING WORD DOCUMENT")
        print("="*80 + "\n")
        
        # Load the Word template file
        template_path = os.path.join('static', 'Formats_and_Catalog', 'IS_Audit_Report.docx')
        
        if not os.path.exists(template_path):
            return jsonify({'error': 'Template file not found'}), 404
        
        # Load the document
        doc = Document(template_path)
        print(f"✅ Template loaded: {template_path}")
        
        # Format applications list with numbers (left-aligned)
        formatted_applications = ""
        if applications:
            for idx, app in enumerate(applications, 1):
                formatted_applications += f"{idx}. {app}\n"
            formatted_applications = formatted_applications.rstrip('\n')  # Remove last newline
        else:
            formatted_applications = "N/A"
        
        # Create replacement dictionary
        replacements = {
            "Organizationnnnnn_Nameeeee": final_organization,
            "Cityyyyyy": final_city,
            "Stateeeeeee": state,
            "Audit_Period_Financial_Yearrrrr": financial_year,
            "Preparedddd_Byyyy": prepared_by,
            "Date_of_Audittttt": date_range,
            "Submitteddd_to": auditee_person,
            "Submitted_to_Designationnn": auditee_designation,
            "Addressssssss": address,
            "Bank_Email_IDDDDD": all_emails,
            "Nummmber_of_Branchessss": number_of_branches,
            "CBSS_Vendor_Nameeeee": cbs_provider_name,
            "Business_model_nameeeee": final_business_model,
            "CBS_DC_managed_byyy": cbs_datacenter_managed,
            "DC_Locationnnnn": final_dc_location,
            "DR_Locationnnnn": final_dr_location,
            "Primary_Connectivityyyy": primary_connectivity,
            "Secondry_Connectivityyyyyy": secondary_connectivity,
            "Hardware_Vendor_Nameeeee": hardware_vendor_name,
            "Antivirusssss": antivirus_name,
            "Mail_Service_Providerrrrrr": mail_messaging_service,
            "Application_Nameeeeeesssssssss": formatted_applications
        }
        
        # Add conditional replacements based on bank under selection
        if bank_under == "NABARD":
            print("\n--- NABARD SPECIFIC REPLACEMENTS ---")
            replacements["Bankkkkk_memberrr_offffff"] = f"{final_organization} was established in accordance with the provisions of the NABARD."
            replacements["Regulateryyyyyyyy_Guidelineeeeeee"] = "Conduct a System audit based on NABARD Information Security Audit guidelines."
            replacements["Ciircularr_Nameeeee"] = "NABARD circular no NB.DoS.HO.POL/3634/J-1/2014-15 dated 25/02/2015."
            replacements["Regulateeeeeeryyyyy"] = "NABARD"
            print("✅ Added NABARD-specific replacements")
            
        elif bank_under == "RBI":
            print("\n--- RBI SPECIFIC REPLACEMENTS ---")
            replacements["Bankkkkk_memberrr_offffff"] = f"{final_organization} was established in accordance with the provisions of the Reserve Bank of India (RBI)."
            replacements["Regulateryyyyyyyy_Guidelineeeeeee"] = "Conduct a System Audit based on RBI Information Security Audit guidelines."
            replacements["Ciircularr_Nameeeee"] = "RBI's circular no RBI/2010-11/494 DBS.CO.ITC.BC. No.6 / 31.02.008 / 2010-11."
            replacements["Regulateeeeeeryyyyy"] = "RBI"
            print("✅ Added RBI-specific replacements")
        
        # Add conditional replacements based on yes/no selections
        print("\n--- CONDITIONAL REPLACEMENTS ---")
        
        # Handle has_other_services
        if has_other_services == "No":
            replacements["otther_Servicesss"] = ""  # Remove the word and line breaks
            print("✅ Removed other services section (No selected)")
        elif has_other_services == "Yes":
            other_services_text = f"For the services other than CBS, Bank is operating the services ({other_services}) hosted in {final_city} and having its near DR at {final_other_services_dc} and Far DR at {final_other_services_dr}."
            replacements["otther_Servicesss"] = other_services_text
            print("✅ Added other services section (Yes selected)")
        
        # Handle is_direct_member
        if is_direct_member == "No":
            replacements["Iss_Diirect_memmmber"] = ""  # Remove the word and line breaks
            print("✅ Removed direct member section (No selected)")
        elif is_direct_member == "Yes":
            direct_member_text = f"Bank is direct member of RBI and NPCI for some services. For the same bank act as sponsor bank for {number_of_submember_banks} sub-member banks for processing transactions on the behalf of them."
            replacements["Iss_Diirect_memmmber"] = direct_member_text
            print("✅ Added direct member section (Yes selected)")
        
        # Handle hasAntimalware
        if has_antimalware == "No":
            replacements["hassss_antimalware"] = ""  # Remove the word and line breaks
            print("✅ Removed antimalware section (No selected)")
        elif has_antimalware == "Yes":
            antimalware_text = f"{antimalware_name} Anti-malware solution is installed in Head office and branch systems to prevent bank from malware attacks."
            replacements["hassss_antimalware"] = antimalware_text
            print("✅ Added antimalware section (Yes selected)")
        
        # Handle hasAD
        if has_ad == "No":
            replacements["hasssss_addddddd"] = ""  # Remove the word and line breaks
            print("✅ Removed AD section (No selected)")
        elif has_ad == "Yes":
            ad_text = "All Bank systems are managed through a centralized active directory."
            replacements["hasssss_addddddd"] = ad_text
            print("✅ Added AD section (Yes selected)")
        
        # Handle hasPAM
        if has_pam == "No":
            replacements["hasPAMMMMM"] = ""  # Remove the word and line breaks
            print("✅ Removed PAM section (No selected)")
        elif has_pam == "Yes":
            pam_text = f"Access to all servers is managed through PAM (Privilege Access Management) solution by {pam_name}."
            replacements["hasPAMMMMM"] = pam_text
            print("✅ Added PAM section (Yes selected)")
        
        print("\n📝 Replacements to be made:")
        for old, new in replacements.items():
            print(f"  '{old}' → '{new}'")
        
        # Replace text in the document
        replace_text_in_document(doc, replacements)
        
        # Replace "Allllll_Asssetsssssssss" with numbered list of worksheet names if Excel file is provided
        if excel_file_1_path and os.path.exists(excel_file_1_path):
            print("\n" + "="*80)
            print("EXTRACTING WORKSHEET NAMES FOR ASSET LIST")
            print("="*80)
            
            # Extract worksheet names
            worksheet_names = extract_worksheet_names_from_excel(excel_file_1_path)
            
            if worksheet_names:
                print(f"✅ Found {len(worksheet_names)} worksheet(s):")
                for idx, name in enumerate(worksheet_names, 1):
                    print(f"  {idx}. {name}")
                
                # Find and replace the placeholder with formatted list
                placeholder_found = False
                for paragraph in doc.paragraphs:
                    if "Allllll_Asssetsssssssss" in paragraph.text:
                        # Get the parent element
                        parent = paragraph._element.getparent()
                        placeholder_element = paragraph._element
                        
                        # Add each worksheet name as a numbered list with left margin
                        for idx, sheet_name in enumerate(worksheet_names, 1):
                            asset_para = doc.add_paragraph()
                            asset_run = asset_para.add_run(f"{idx}. {sheet_name}")
                            asset_run.font.size = Pt(12)
                            asset_run.font.name = 'Times New Roman'
                            
                            # Add left margin (indent)
                            asset_para.paragraph_format.left_indent = Inches(0.5)
                            
                            # Insert before the placeholder
                            parent.insert(parent.index(placeholder_element), asset_para._element)
                        
                        # Remove the placeholder paragraph
                        parent.remove(placeholder_element)
                        placeholder_found = True
                        print("✅ Successfully replaced 'Allllll_Asssetsssssssss' with worksheet list")
                        break
                
                if not placeholder_found:
                    print("⚠️ Placeholder 'Allllll_Asssetsssssssss' not found in document")
            else:
                print("⚠️ No worksheets found in Excel file")
        
        # Replace Assets Review placeholder with Excel data if Excel file is provided
        if excel_file_1_path and os.path.exists(excel_file_1_path):
            print("\n" + "="*80)
            print("PROCESSING EXCEL DATA FOR ASSETS REVIEW")
            print("="*80)
            replace_assets_review_with_excel_data(doc, excel_file_1_path, "Asssssssssettttsss_Reevieeww")
        
        # Replace Branch Review placeholder with Excel tables if Excel file 2 is provided
        if excel_file_2_path and os.path.exists(excel_file_2_path):
            print("\n" + "="*80)
            print("PROCESSING EXCEL TABLES FOR BRANCH REVIEW")
            print("="*80)
            replace_branch_review_with_excel_tables(doc, excel_file_2_path, "Braaanch_Revviewww")
        
        # Add page borders just before saving
        add_page_borders_to_document(doc)
        
        # Add borders to all images just before saving
        add_borders_to_all_images(doc)
        
        # Format "Low risk:" text (bold + underline)
        format_low_risk_text(doc)
        
        # Reduce specific column widths based on content
        reduce_specific_column_widths(doc)
        
        # Create a temporary file with fixed name
        output_filename = "Is Audit Report.docx"
        
        # Create temp directory if it doesn't exist
        temp_dir = tempfile.gettempdir()
        temp_file_path = os.path.join(temp_dir, output_filename)
        
        # Save the modified document
        doc.save(temp_file_path)
        print(f"✅ Word file saved successfully: {temp_file_path}")
        
        print("\n" + "="*80)
        print("END OF PROCESSING")
        print("="*80 + "\n")
        
        # Cleanup temporary Excel files
        if excel_file_1_path and os.path.exists(excel_file_1_path):
            try:
                os.remove(excel_file_1_path)
                print(f"🗑️ Cleaned up temporary Excel File 1: {excel_file_1_path}")
            except Exception as e:
                print(f"⚠️ Error cleaning up Excel File 1: {e}")
        
        if excel_file_2_path and os.path.exists(excel_file_2_path):
            try:
                os.remove(excel_file_2_path)
                print(f"🗑️ Cleaned up temporary Excel File 2: {excel_file_2_path}")
            except Exception as e:
                print(f"⚠️ Error cleaning up Excel File 2: {e}")
        
        # Schedule cleanup for Word file
        cleanup_thread = threading.Thread(target=cleanup_temp_file, args=(temp_file_path,))
        cleanup_thread.daemon = True
        cleanup_thread.start()
        
        # Send the file for download
        return send_file(
            temp_file_path,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"❌ Error processing IS Audit Word Report: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500
