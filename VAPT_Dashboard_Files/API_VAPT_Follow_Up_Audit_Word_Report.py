import copy
import io
import os
import re
import tempfile
import traceback
import zipfile
import shutil
from datetime import datetime, timedelta
from lxml import etree
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor
import pandas as pd
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
import openpyxl
from PIL import Image as PILImage
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import qn
from docx.oxml.ns import nsdecls
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from flask import Blueprint, request, send_file, jsonify

from VAPT_Dashboard_Files.Web_Application_First_Audit_Word_Report import (
    replace_scope_placeholders_with_data as web_app_first_replace_scope_placeholders_with_data,
    copy_table_with_formatting,
    replace_text_in_table,
    replace_text_in_table_with_risk_colors,
    replace_affected_systems_with_table
)

api_follow_up_word_report_bp = Blueprint('api_follow_up_word_report', __name__)


def apply_1pt_border_to_picture(picture):
    """
    Apply a 1pt solid black border to a picture object.
    Works for all images added via add_picture().

    Args:
        picture: The picture object returned by run.add_picture()
    """
    try:
        pic = picture._inline.graphic.graphicData.pic
        spPr = pic.spPr

        for ln in spPr.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}ln'):
            spPr.remove(ln)

        border = parse_xml(
            '<a:ln xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" w="12700">'
            '<a:solidFill><a:srgbClr val="000000"/></a:solidFill>'
            '<a:prstDash val="solid"/>'
            '</a:ln>'
        )
        spPr.append(border)

        print(f"✅ Applied 1pt border (12700 EMUs)")

    except Exception as e:
        print(f"❌ Error applying border: {e}")
        traceback.print_exc()


def clean_value(value):
    """Convert NaN, None, empty strings to 'NA'"""
    if pd.isna(value) or value is None or str(value).lower() in ['nan', 'none', '']:
        return "NA"
    return str(value)


def get_rgb(color):
    if color is None or color.type != "rgb":
        return None
    if len(color.rgb) > 6:
        return color.rgb[2:]
    return color.rgb


def add_borders_to_cell(cell):
    """Add borders to a table cell"""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        borders = ['top', 'left', 'bottom', 'right']
        for border_name in borders:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcPr.append(border)
    except Exception as e:
        print(f"Warning: Could not add borders to cell: {e}")


def add_borders_to_row(row):
    """Add borders to all cells in a table row"""
    for cell in row.cells:
        add_borders_to_cell(cell)


def _clear_cell_contents(cell):
    """Remove all paragraphs and nested tables from a cell."""
    for paragraph in list(cell.paragraphs):
        p_element = paragraph._element
        p_parent = p_element.getparent()
        if p_parent is not None:
            p_parent.remove(p_element)
    for nested_table in list(cell.tables):
        tbl_element = nested_table._element
        tbl_parent = tbl_element.getparent()
        if tbl_parent is not None:
            tbl_parent.remove(tbl_element)


def _set_cell_to_nil(cell):
    """Clear cell content and insert centered 'NIL' text."""
    _clear_cell_contents(cell)
    nil_paragraph = cell.add_paragraph()
    nil_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nil_paragraph.paragraph_format.space_before = Pt(0)
    nil_paragraph.paragraph_format.space_after = Pt(0)
    nil_run = nil_paragraph.add_run("NIL")
    nil_run.font.name = 'Calibri'
    nil_run.font.size = Pt(12)

def replace_text_preserving_format(paragraph, old_text, new_text):
    """
    Replace text within existing runs so original formatting (font name, size, color, etc.)
    is preserved as much as possible. Returns True if a replacement was made.
    """
    replaced = False
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            replaced = True
    return replaced


def add_ordinal_suffix(day):
    """Add ordinal suffix to day with leading zero and superscript (01ˢᵗ, 02ⁿᵈ, 03ʳᵈ, 04ᵗʰ, etc.)"""
    if 10 <= day % 100 <= 20:
        suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')

    superscript_map = {
        's': 'ˢ', 't': 'ᵗ', 'n': 'ⁿ', 'd': 'ᵈ', 'r': 'ʳ', 'h': 'ʰ'
    }
    superscript_suffix = ''.join(superscript_map.get(c, c) for c in suffix)

    day_str = f"{day:02d}"
    return f"{day_str}{superscript_suffix}"


def convert_to_dd_mm_yyyy(date_str):
    """Convert various date formats to DD.MM.YYYY format"""
    date_formats = ["%d.%m.%Y", "%Y-%m-%d", "%m/%d/%Y", "%d-%m-%Y", "%B %Y", "%b %Y"]
    for fmt in date_formats:
        try:
            date_obj = datetime.strptime(date_str, fmt)
            return date_obj.strftime("%d.%m.%Y")
        except ValueError:
            continue
    try:
        year = int(date_str)
        return f"01.01.{year}"
    except ValueError:
        raise ValueError(f"Could not parse date: {date_str}")


def format_date_for_month_year(date_str):
    """Format date string to 'Month YYYY' format"""
    try:
        if not date_str:
            return ""

        date_formats = ["%d.%m.%Y", "%Y-%m-%d", "%m/%d/%Y", "%d-%m-%Y"]
        for fmt in date_formats:
            try:
                date_obj = datetime.strptime(str(date_str), fmt)
                return date_obj.strftime("%B %Y")
            except ValueError:
                continue

        year_match = re.search(r'\b(20\d{2})\b', str(date_str))
        if year_match:
            year = year_match.group(1)
            return f"January {year}"

        return str(date_str)
    except Exception as e:
        print(f"Error formatting date: {e}")
        return str(date_str)


def format_date_for_range(start_date_str, end_date_str):
    """Format date range to 'DD Month YYYY to DD Month YYYY' format"""
    try:
        if not start_date_str or not end_date_str:
            return ""

        start_formatted = format_date_for_dd_month_yyyy(start_date_str)
        end_formatted = format_date_for_dd_month_yyyy(end_date_str)

        return f"{start_formatted} to {end_formatted}"
    except Exception as e:
        print(f"Error formatting date range: {e}")
        return ""


def format_date_for_dd_month_yyyy(date_str):
    """Format date string to 'DD Month YYYY' format with ordinal suffix"""
    try:
        if not date_str:
            return ""

        date_formats = ["%d.%m.%Y", "%Y-%m-%d", "%m/%d/%Y", "%d-%m-%Y"]
        for fmt in date_formats:
            try:
                date_obj = datetime.strptime(str(date_str), fmt)
                day = date_obj.day
                month = date_obj.strftime("%B")
                year = date_obj.year
                ordinal_day = add_ordinal_suffix(day)
                return f"{ordinal_day} {month} {year}"
            except ValueError:
                continue

        return str(date_str)
    except Exception as e:
        print(f"Error formatting date: {e}")
        return str(date_str)


def format_audit_date_period(start_date_str, end_date_str):
    """Format audit date period"""
    try:
        if not start_date_str or not end_date_str:
            return ""

        start_formatted = convert_to_dd_mm_yyyy(start_date_str)
        end_formatted = convert_to_dd_mm_yyyy(end_date_str)

        return f"{start_formatted} - {end_formatted}"
    except Exception as e:
        print(f"Error formatting audit date period: {e}")
        return ""


def generate_followup_vapt_timeline(start_date_str, end_date_str):
    """
    Generate Follow-Up Audit VAPT timeline with specific phase allocations
    - First Validation Cycle: 25% of dates
    - Gap: 40% of dates (not included in timeline)
    - Second Validation Cycle: 20%
    - Reporting: 15%
    """

    BANK_HOLIDAYS = [
        "01.01.2025",
        "15.08.2025",
        "02.10.2025",
    ]

    def is_working_day(date):
        date_str = date.strftime("%d.%m.%Y")
        return date.weekday() < 5 and date_str not in BANK_HOLIDAYS

    try:
        start_date = datetime.strptime(convert_to_dd_mm_yyyy(start_date_str), "%d.%m.%Y")
        end_date = datetime.strptime(convert_to_dd_mm_yyyy(end_date_str), "%d.%m.%Y")
    except Exception as e:
        print(f"Error parsing dates: {e}")
        return []

    all_dates = [d for d in (start_date + timedelta(days=i) for i in range((end_date - start_date).days + 1)) if is_working_day(d)]
    total_working_days = len(all_dates)

    if total_working_days == 0:
        return []

    first_validation_days = max(1, int(total_working_days * 0.25))
    gap_days = int(total_working_days * 0.40)
    second_validation_days = max(1, int(total_working_days * 0.20))
    reporting_days = total_working_days - first_validation_days - gap_days - second_validation_days

    if reporting_days < 1:
        reporting_days = 1
        if second_validation_days > 1:
            second_validation_days -= 1

    timeline = []
    date_index = 0

    if date_index < len(all_dates):
        end_index = min(date_index + first_validation_days, len(all_dates))
        phase_dates = all_dates[date_index:end_index]
        if phase_dates:
            start = phase_dates[0].strftime("%d.%m.%Y")
            end = phase_dates[-1].strftime("%d.%m.%Y")
            timeline.append({
                "Phase": "Follow-up Audit First Validation Cycle",
                "Dates": f"{start} - {end}",
                "Start": start,
                "End": end
            })
            date_index = end_index

    date_index += gap_days

    if date_index < len(all_dates):
        end_index = min(date_index + second_validation_days, len(all_dates))
        phase_dates = all_dates[date_index:end_index]
        if phase_dates:
            start = phase_dates[0].strftime("%d.%m.%Y")
            end = phase_dates[-1].strftime("%d.%m.%Y")
            timeline.append({
                "Phase": "Follow-up Audit Second Validation Cycle",
                "Dates": f"{start} - {end}",
                "Start": start,
                "End": end
            })
            date_index = end_index

    if date_index < len(all_dates):
        end_index = min(date_index + reporting_days, len(all_dates))
        phase_dates = all_dates[date_index:end_index]
        if phase_dates:
            start = phase_dates[0].strftime("%d.%m.%Y")
            end = phase_dates[-1].strftime("%d.%m.%Y")
            timeline.append({
                "Phase": "Reporting",
                "Dates": f"{start} - {end}",
                "Start": start,
                "End": end
            })

    return timeline


def find_and_print_metadata(file_path):
    """
    Extract metadata from Excel file's Meta_Data worksheet
    """
    try:
        workbook = openpyxl.load_workbook(file_path)

        if 'Meta_Data' not in workbook.sheetnames:
            print("Error: 'Meta_Data' worksheet not found!")
            return {}

        sheet = workbook['Meta_Data']

        target_values = {
            "Organization Name": None,
            "City": None,
            "State": None,
            "Start Date": None,
            "End Date": None
        }

        report_prepared_by_value = None
        auditee_details_value1 = None
        auditee_details_value2 = None
        bank_email_addresses = []
        auditing_team_members = []

        for row in sheet.iter_rows():
            for cell in row:
                if cell.value is not None:
                    for key in target_values.keys():
                        if str(cell.value).strip() == key:
                            adjacent_cell = sheet.cell(row=cell.row, column=cell.column + 1)
                            if adjacent_cell.value is not None:
                                target_values[key] = adjacent_cell.value
                            break

                    cell_text = str(cell.value).strip()
                    normalized_text = cell_text.upper()

                    if cell_text == "REPORT PREPARED BY":
                        diagonal_cell = sheet.cell(row=cell.row + 1, column=cell.column + 1)
                        if diagonal_cell.value is not None:
                            report_prepared_by_value = diagonal_cell.value

                    if cell_text == "AUDITEE DETAILS":
                        diagonal_cell1 = sheet.cell(row=cell.row + 1, column=cell.column + 1)
                        if diagonal_cell1.value is not None:
                            auditee_details_value1 = diagonal_cell1.value

                        diagonal_cell2 = sheet.cell(row=cell.row + 2, column=cell.column + 1)
                        if diagonal_cell2.value is not None:
                            auditee_details_value2 = diagonal_cell2.value

                    if "BANK EMAIL ADDRESS" in normalized_text or "ORGANIZATION EMAIL ADDRESS" in normalized_text:
                        adjacent_cell = sheet.cell(row=cell.row, column=cell.column + 1)
                        if adjacent_cell.value:
                            adjacent_value = str(adjacent_cell.value)
                            for part in re.split(r'[,;\n]+', adjacent_value):
                                cleaned = part.strip()
                                if cleaned:
                                    bank_email_addresses.append(cleaned)

                        current_row = cell.row + 1
                        current_col = cell.column + 1

                        while True:
                            email_cell = sheet.cell(row=current_row, column=current_col)
                            if email_cell.value is not None and str(email_cell.value).strip() != "":
                                bank_email_addresses.append(str(email_cell.value).strip())
                                current_row += 1
                            else:
                                break

                    if "AUDITING TEAM MEMBER" in cell_text:
                        member_data = {}
                        member_number = cell_text.strip().split()[-1]

                        labels = [
                            f"Team Member {member_number} - Name",
                            f"Team Member {member_number} - Designation",
                            f"Team Member {member_number} - Email",
                            f"Team Member {member_number} - Qualification",
                            f"Team Member {member_number} - Certified"
                        ]

                        current_row = cell.row + 1
                        current_col = cell.column + 1
                        label_index = 0

                        for i in range(5):
                            member_cell = sheet.cell(row=current_row + i, column=current_col)
                            if member_cell.value is not None and str(member_cell.value).strip() != "":
                                if label_index < len(labels):
                                    member_data[labels[label_index]] = str(member_cell.value).strip()
                                    label_index += 1
                            else:
                                break

                        if member_data:
                            auditing_team_members.append(member_data)

        internal_external_value = "NA"
        try:
            cell_b11 = sheet['B11']
            if cell_b11 and cell_b11.value is not None:
                internal_external_value = clean_value(cell_b11.value)
        except KeyError:
            pass

        try:
            first_audit_report_id = clean_value(sheet['B15'].value) if sheet['B15'].value is not None else "NA"
        except KeyError:
            first_audit_report_id = "NA"
        try:
            first_audit_report_date = clean_value(sheet['B16'].value) if sheet['B16'].value is not None else "NA"
        except KeyError:
            first_audit_report_date = "NA"

        result = {
            "organization_name": target_values.get("Organization Name"),
            "city": target_values.get("City"),
            "state": target_values.get("State"),
            "start_date": target_values.get("Start Date"),
            "end_date": target_values.get("End Date"),
            "report_prepared_by": report_prepared_by_value,
            "auditee_details_1": auditee_details_value1,
            "auditee_details_2": auditee_details_value2,
            "bank_email_addresses": bank_email_addresses,
            "auditing_team_members": auditing_team_members,
            "first_audit_report_id": first_audit_report_id,
            "first_audit_report_date": first_audit_report_date,
            "internal_external": internal_external_value
        }

        workbook.close()

        return result

    except FileNotFoundError:
        print(f"Error: File '{file_path}' not found!")
        return {}
    except Exception as e:
        print(f"An error occurred: {str(e)}")
        return {}


def replace_scope_placeholders_with_data(doc, excel_file, form_metadata=None):
    """
    Replace scope placeholders using form metadata assets or Excel file.
    If form_metadata contains asset data, use that; otherwise delegate to Excel logic.
    """
    try:
        if form_metadata and 'asset_data' in form_metadata:
            # Use form asset data
            asset_data = form_metadata['asset_data']
            if asset_data:
                print(f"🔍 Using {len(asset_data)} assets from form data")
                replace_scope_placeholders_with_form_assets(doc, asset_data)
            else:
                print("⚠️ No asset data in form metadata, using Excel data")
                web_app_first_replace_scope_placeholders_with_data(doc, excel_file)
        else:
            # Fallback to Excel logic
            web_app_first_replace_scope_placeholders_with_data(doc, excel_file)
    except Exception as e:
        print(f"❌ Error replacing scope placeholders: {e}")
        traceback.print_exc()


def replace_scope_placeholders_with_form_assets(doc, asset_data):
    """
    Replace asset placeholders using form asset data (similar to auditing team logic).
    """
    try:
        if not asset_data:
            print("No asset data provided")
            return

        # Find tables containing the placeholders
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        if "Branch Namee or Server Namee" in paragraph.text or "HAAAASH" in paragraph.text:
                            # Found the target row, now replace and add new rows
                            replace_scope_placeholders_in_table_with_assets(table, row_idx, cell_idx, asset_data)
                            replace_asset_placeholders_globally_with_assets(doc, asset_data)
                            return

        print("⚠️ Asset placeholders not found in any table")

    except Exception as e:
        print(f"❌ Error replacing asset placeholders with form data: {e}")
        traceback.print_exc()


def _asset_placeholder_mapping_from_form(asset, asset_index):
    """
    Create placeholder mapping for an asset from form data.
    """
    hash_value = asset.get('hash_value', 'NA') or 'NA'
    
    return {
        "Branch Namee or Server Namee": asset.get('description', 'NA'),
        "HAAAASH": hash_value,
        "Criticallllll": asset.get('criticality', 'NA'),
        "URRRRLLLLLLLLL____": asset.get('url', 'NA'),
        "VEEERSSSS_": asset.get('version', 'NA'),
        f"Asset {asset_index + 1} - Description": asset.get('description', 'NA'),
        f"Asset {asset_index + 1} - Criticality": asset.get('criticality', 'NA'),
        f"Asset {asset_index + 1} - Hash Value": hash_value,
        f"Asset {asset_index + 1} - URL": asset.get('url', 'NA'),
        f"Asset {asset_index + 1} - Version": asset.get('version', 'NA')
    }


def _fill_asset_placeholders_in_row_from_form(row, asset, asset_index):
    """
    Replace placeholders in a given table row with asset data from form.
    """
    replacements = _asset_placeholder_mapping_from_form(asset, asset_index)
    
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    _replace_text_in_runs(paragraph, placeholder, value)
        # Fallback replacement if any placeholder remains
        for placeholder, value in replacements.items():
            if placeholder in cell.text:
                cell.text = cell.text.replace(placeholder, value)
    
    # Ensure Sr. No column (first cell) has sequential numbering
    if row.cells:
        row.cells[0].text = str(asset_index + 1)


def replace_scope_placeholders_in_table_with_assets(table, target_row_idx, target_cell_idx, asset_data):
    """
    Replace asset placeholders in a specific table row and add new rows for additional assets.
    Similar to replace_scope_placeholders_in_table but using form asset data.
    """
    try:
        from VAPT_Dashboard_Files.Web_Application_First_Audit_Word_Report import replace_scope_placeholders_in_table, _fill_asset_placeholders_in_row
        from VAPT_Dashboard_Files.Web_Application_First_Audit_Word_Report import _asset_placeholder_mapping
        
        target_row = table.rows[target_row_idx]
        template_texts = [cell.text for cell in target_row.cells]
        template_alignments = [[p.alignment for p in cell.paragraphs] for cell in target_row.cells]
        template_vertical_alignment = [cell.vertical_alignment for cell in target_row.cells]

        def restore_row_template(row):
            for j, cell in enumerate(row.cells):
                cell.text = template_texts[j]
                cell.vertical_alignment = template_vertical_alignment[j]

                # Ensure at least one paragraph exists
                if not cell.paragraphs:
                    cell.add_paragraph()

                # Apply stored alignments where possible
                for idx, paragraph in enumerate(cell.paragraphs):
                    if j < len(template_alignments) and idx < len(template_alignments[j]):
                        paragraph.alignment = template_alignments[j][idx]
                    else:
                        paragraph.alignment = template_alignments[j][0] if template_alignments[j] else paragraph.alignment
        
        def clone_row_structure():
            new_row = table.add_row()
            restore_row_template(new_row)
            return new_row
        
        for idx, asset in enumerate(asset_data):
            if idx == 0:
                row = target_row
                restore_row_template(row)
            else:
                row = clone_row_structure()
            _fill_asset_placeholders_in_row_from_form(row, asset, idx)
            # Set font size to 12 for all text in the row
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
            # Add borders to all cells in the row (important for newly added rows)
            add_borders_to_row(row)
        
        print(f"✅ Replaced asset placeholders with {len(asset_data)} assets")
        
    except Exception as e:
        print(f"❌ Error replacing asset placeholders in table: {e}")
        traceback.print_exc()


def replace_asset_placeholders_globally_with_assets(doc, asset_data):
    """
    Replace remaining asset placeholders (outside the main table) using the first asset entry.
    """
    if not asset_data:
        return
    
    # Use first asset for global replacements
    first_asset = asset_data[0] if asset_data else None
    if not first_asset:
        return
    
    replacements = _asset_placeholder_mapping_from_form(first_asset, 0)
    
    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                _replace_text_in_runs(paragraph, placeholder, value)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in replacements.items():
                        if placeholder in paragraph.text:
                            _replace_text_in_runs(paragraph, placeholder, value)


def _replace_text_in_runs(paragraph, old_text, new_text):
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return False

    for run in list(paragraph.runs):
        run.clear()

    paragraph.add_run(full_text.replace(old_text, new_text))
    return True


def _replace_text_in_runs_with_bold(paragraph, old_text, new_text, make_bold=False):
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return False

    run_formats = []
    for run in paragraph.runs:
        run_formats.append({
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'font_color': run.font.color.rgb if run.font.color else None
        })

    for run in list(paragraph.runs):
        run.clear()

    new_full_text = full_text.replace(old_text, new_text)
    segments = new_full_text.split(new_text)

    for i, segment in enumerate(segments):
        if segment:
            seg_run = paragraph.add_run(segment)
            seg_run.bold = run_formats[0]['bold'] if run_formats else False
            seg_run.italic = run_formats[0]['italic'] if run_formats else False
            seg_run.underline = run_formats[0]['underline'] if run_formats else False
            if run_formats and run_formats[0]['font_name']:
                seg_run.font.name = run_formats[0]['font_name']
            if run_formats and run_formats[0]['font_size']:
                seg_run.font.size = run_formats[0]['font_size']
            if run_formats and run_formats[0]['font_color']:
                try:
                    seg_run.font.color.rgb = run_formats[0]['font_color']
                except Exception:
                    pass

        if i != len(segments) - 1:
            highlighted_run = paragraph.add_run(new_text)
            if make_bold:
                highlighted_run.bold = True
            else:
                highlighted_run.bold = run_formats[0]['bold'] if run_formats else False
            highlighted_run.italic = run_formats[0]['italic'] if run_formats else False
            highlighted_run.underline = run_formats[0]['underline'] if run_formats else False
            # Set font to Calibri for replaced text
            highlighted_run.font.name = 'Calibri'
            if run_formats and run_formats[0]['font_size']:
                highlighted_run.font.size = run_formats[0]['font_size']
            if run_formats and run_formats[0]['font_color']:
                try:
                    highlighted_run.font.color.rgb = run_formats[0]['font_color']
                except Exception:
                    pass

    return True


def replace_metadata_placeholders(doc, metadata):
    try:
        bank_emails = ""
        if metadata.get("bank_email_addresses"):
            bank_emails = "\n".join(metadata.get("bank_email_addresses", []))

        month_year = format_date_for_month_year(metadata.get("end_date", ""))
        date_range = format_date_for_range(metadata.get("start_date", ""), metadata.get("end_date", ""))

        # Format first_audit_report_date to DD.MM.YYYY
        first_audit_report_date = metadata.get("first_audit_report_date", "")
        if first_audit_report_date:
            try:
                # Try to parse and format the date
                date_obj = datetime.strptime(first_audit_report_date, "%Y-%m-%d")
                first_audit_report_date = date_obj.strftime("%d.%m.%Y")
            except (ValueError, TypeError):
                # If parsing fails, try to convert using convert_to_dd_mm_yyyy
                try:
                    first_audit_report_date = convert_to_dd_mm_yyyy(first_audit_report_date)
                except:
                    pass  # Keep original if all conversions fail

        replacements = {
            "ORGANIZATION_NAMEE": metadata.get("organization_name", ""),
            "Organization_Namee": metadata.get("organization_name", ""),
            "CITYY": metadata.get("city", ""),
            "Stateee": metadata.get("state", ""),
            "First_Audit_Report_Idd": metadata.get("first_audit_report_id", ""),
            "First_Audit_Report_Datee": first_audit_report_date,
            "INTEERNALLLOREXTERRNAL": metadata.get("internal_external", ""),
            "Audit_Date_Period": format_audit_date_period(metadata.get('start_date', ''), metadata.get('end_date', '')),
            "Maker_Name_R": metadata.get("report_prepared_by", ""),
            "Organization_Personn": metadata.get("auditee_details_1", ""),
            "Designationn": metadata.get("auditee_details_2", ""),
            "Auditee_Email_Adresss": bank_emails,
            "Monthh Yearr": month_year,
            "00rd Month Year to 00th Month Year": date_range
        }

        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    if old_text == "00rd Month Year to 00th Month Year":
                        _replace_text_in_runs_with_bold(paragraph, old_text, new_text, make_bold=True)
                    elif not replace_text_preserving_format(paragraph, old_text, new_text):
                        _replace_text_in_runs(paragraph, old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            if old_text in paragraph.text:
                                if old_text == "00rd Month Year to 00th Month Year":
                                    _replace_text_in_runs_with_bold(paragraph, old_text, new_text, make_bold=True)
                                elif not replace_text_preserving_format(paragraph, old_text, new_text):
                                    _replace_text_in_runs(paragraph, old_text, new_text)

        print("✅ Metadata placeholders replaced successfully")

    except Exception as e:
        print(f"❌ Error replacing metadata placeholders: {str(e)}")
        traceback.print_exc()


def replace_auditor_placeholders_and_add_rows(doc, metadata):
    try:
        auditing_team_members = metadata.get("auditing_team_members", [])

        if not auditing_team_members:
            print("No auditing team members found")
            return

        target_table = None
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Auditorrrrr" in cell.text:
                        target_table = table
                        break
                if target_table:
                    break
            if target_table:
                break

        if not target_table:
            print("Table with auditor placeholders not found")
            return

        target_row = None
        column_mapping = {}

        for row in target_table.rows:
            has_auditor_placeholders = False
            for cell in row.cells:
                if "Auditorrrrr" in cell.text:
                    has_auditor_placeholders = True
                    target_row = row
                    break

            if has_auditor_placeholders:
                for j, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        if "Auditorrrrr_Desi" in paragraph.text:
                            column_mapping[j] = "designation"
                        elif "Auditorrrrr_email" in paragraph.text:
                            column_mapping[j] = "email"
                        elif "Auditorrrrr_Qua" in paragraph.text:
                            column_mapping[j] = "qualification"
                        elif "Auditorrrrr_Cert" in paragraph.text:
                            column_mapping[j] = "certified"
                        elif "Auditorrrrr" in paragraph.text and all(keyword not in paragraph.text for keyword in ["Desi", "email", "Qua", "Cert"]):
                            column_mapping[j] = "name"
                break

        if not target_row:
            print("Row with auditor placeholders not found")
            return

        print(f"📋 Column mapping detected: {column_mapping}")

        fixed_auditor_data = {
            "name": "Niraj Goyal",
            "designation": "CEO / Director",
            "email": "admin@ngtech.co.in",
            "qualification": "CA, CEH, DISA",
            "certified": "Yes"
        }

        for j, cell in enumerate(target_row.cells):
            if j in column_mapping:
                column_type = column_mapping[j]
                for paragraph in cell.paragraphs:
                    if column_type == "name":
                        _replace_text_in_runs(paragraph, "Auditorrrrr", fixed_auditor_data["name"])
                    elif column_type == "designation":
                        _replace_text_in_runs(paragraph, "Auditorrrrr_Desi", fixed_auditor_data["designation"])
                    elif column_type == "email":
                        _replace_text_in_runs(paragraph, "Auditorrrrr_email", fixed_auditor_data["email"])
                    elif column_type == "qualification":
                        _replace_text_in_runs(paragraph, "Auditorrrrr_Qua", fixed_auditor_data["qualification"])
                    elif column_type == "certified":
                        _replace_text_in_runs(paragraph, "Auditorrrrr_Cert", fixed_auditor_data["certified"])

        max_additional_rows = min(8, len(auditing_team_members))

        for i in range(max_additional_rows):
            member_data = auditing_team_members[i]
            member_number = i + 2

            new_row = target_table.add_row()

            for j, cell in enumerate(new_row.cells):
                if j < len(target_row.cells):
                    while len(cell.paragraphs) > 0:
                        p = cell.paragraphs[0]
                        p._element.getparent().remove(p._element)

                    original_cell = target_row.cells[j]

                    new_content = ""
                    if j in column_mapping:
                        column_type = column_mapping[j]
                        excel_member_number = None
                        for key in member_data.keys():
                            if key.startswith("Team Member ") and key.endswith(" - Name"):
                                try:
                                    excel_member_number = key.split("Team Member ")[1].split(" - Name")[0]
                                    break
                                except Exception:
                                    continue

                        if excel_member_number:
                            if column_type == "name":
                                new_content = member_data.get(f"Team Member {excel_member_number} - Name", "")
                            elif column_type == "designation":
                                new_content = member_data.get(f"Team Member {excel_member_number} - Designation", "")
                            elif column_type == "email":
                                new_content = member_data.get(f"Team Member {excel_member_number} - Email", "")
                            elif column_type == "qualification":
                                new_content = member_data.get(f"Team Member {excel_member_number} - Qualification", "")
                            elif column_type == "certified":
                                certified_value = member_data.get(f"Team Member {excel_member_number} - Certified", "")
                                if isinstance(certified_value, str):
                                    value_lower = certified_value.strip().lower()
                                    if value_lower == "yes":
                                        new_content = "Yes"
                                    elif value_lower == "no":
                                        new_content = "No"
                                    else:
                                        new_content = certified_value
                                else:
                                    new_content = certified_value
                        else:
                            print(f"⚠️  No Excel member number found for Team Member {member_number}")
                    else:
                        if j == 0 and original_cell.text.strip().isdigit():
                            new_content = str(member_number)
                        else:
                            new_content = original_cell.text

                    if new_content:
                        new_paragraph = cell.add_paragraph()
                        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        new_paragraph.paragraph_format.space_before = Pt(0)
                        new_paragraph.paragraph_format.space_after = Pt(0)
                        new_paragraph.paragraph_format.line_spacing = 1.0
                        new_paragraph.paragraph_format.first_line_indent = Pt(0)
                        new_paragraph.paragraph_format.left_indent = Pt(0)
                        new_paragraph.paragraph_format.right_indent = Pt(0)
                        new_paragraph.paragraph_format.space_before_auto = False
                        new_paragraph.paragraph_format.space_after_auto = False

                        new_run = new_paragraph.add_run(new_content)
                        new_run.font.name = 'Calibri (Body)'
                        new_run.font.size = Pt(12)
                        new_run.font.bold = False
                        new_run.font.italic = False
                        new_run.font.underline = False

                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                        try:
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            cell_margin_xml = f'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:w="0" w:type="dxa"/><w:bottom w:w="0" w:type="dxa"/></w:tcMar>'
                            cell_margin = parse_xml(cell_margin_xml)
                            tcPr.append(cell_margin)
                        except Exception as e:
                            print(f"Warning: Could not set cell margins: {e}")

        print(f"✅ Added {max_additional_rows + 1} auditing team member(s) to table (max 8 additional rows)")

    except Exception as e:
        print(f"❌ Error replacing auditor placeholders: {str(e)}")
        traceback.print_exc()


def highlight_keywords(word_cell, value):
    if value.strip().upper() in ["HOST", "PORT", "SERVICE"]:
        blue_rgb = "#1376d1"
        tc = word_cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), blue_rgb)
        tcPr.append(shd)


def create_nmap_table_from_excel(doc, excel_file):
    try:
        wb = load_workbook(excel_file)

        try:
            ws = wb['Nmap Files']
        except KeyError:
            print("Worksheet 'Nmap Files' not found, using active sheet")
            ws = wb.active

        table = doc.add_table(rows=ws.max_row, cols=ws.max_column)
        table.style = 'Table Grid'

        merged_cells_to_skip = set()

        for merged_range in ws.merged_cells:
            merged_range_coords = list(merged_range.cells)

            top_left_cell_ref = merged_range_coords[0]
            start_row = top_left_cell_ref[0] - 1
            start_col = top_left_cell_ref[1] - 1

            bottom_right_cell_ref = merged_range_coords[-1]
            end_row = bottom_right_cell_ref[0] - 1
            end_col = bottom_right_cell_ref[1] - 1

            start_word_cell = table.cell(start_row, start_col)
            end_word_cell = table.cell(end_row, end_col)

            start_word_cell.merge(end_word_cell)

            for r_idx in range(start_row, end_row + 1):
                for c_idx in range(start_col, end_col + 1):
                    if (r_idx, c_idx) != (start_row, start_col):
                        merged_cells_to_skip.add((r_idx, c_idx))

        for i, row in enumerate(ws.iter_rows(values_only=False)):
            for j, cell in enumerate(row):
                if (i, j) in merged_cells_to_skip:
                    continue

                word_cell = table.cell(i, j)
                value = "" if cell.value is None else str(cell.value)
                word_cell.text = value

                if cell.font:
                    if not word_cell.paragraphs[0].runs:
                        word_cell.paragraphs[0].add_run(value)

                    run = word_cell.paragraphs[0].runs[0]
                    run.font.name = cell.font.name or "Calibri"
                    run.font.size = Pt(12)
                    run.font.bold = cell.font.b
                    run.font.italic = cell.font.i
                    if cell.font.color:
                        rgb = get_rgb(cell.font.color)
                        if rgb:
                            run.font.color.rgb = RGBColor.from_string(rgb)

                word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                highlight_keywords(word_cell, value)

        print("✅ Nmap table created successfully from Excel")
        return table

    except Exception as e:
        print(f"Error creating Nmap table from Excel: {e}")
        traceback.print_exc()
        return None


def read_scope_worksheet_data(excel_file):
    """
    Reuse scope worksheet reader from first audit logic to ensure consistency.
    """
    try:
        from VAPT_Dashboard_Files.Web_Application_First_Audit_Word_Report import read_scope_worksheet_data as read_scope_data_web
        return read_scope_data_web(excel_file)
    except ImportError:
        print("⚠️ Could not import API scope reader, falling back to empty list.")
        return []


def replace_followup_vulnerability_details_with_images(doc, excel_file):
    """
    Build the vulnerability section for the API follow-up report.
    This mirrors the follow-up logic but reads from the 'API VAPT' sheet.
    """
    try:
        poc_images, old_poc_images, df, remarks_dict = extract_followup_poc_images_from_excel(excel_file)

        if df.empty:
            print("⚠️ No vulnerability data found")
            return
        
        # Load workbook to get status from Status columns with dates
        from openpyxl import load_workbook
        from openpyxl.utils import get_column_letter
        wb = load_workbook(excel_file, data_only=True)
        ws = None
        for sheet_name in wb.sheetnames:
            if "api vapt" in sheet_name.lower() or "api" in sheet_name.lower():
                ws = wb[sheet_name]
                break
        
        if not ws:
            print("⚠️ 'API VAPT' worksheet not found")
            wb.close()
            return
        
        # Find all Status columns with dates
        status_columns = find_status_columns_with_dates(ws, header_row=1)
        latest_status_col = status_columns[-1] if status_columns else None

        template_table = None
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Vulnerability___Name" in cell.text:
                        template_table = table
                        break
                if template_table:
                    break
            if template_table:
                break

        if not template_table:
            print("⚠️ Template table with 'Vulnerability___Name' not found")
            return

        marker_paragraph = None
        for paragraph in doc.paragraphs:
            if "Tables_Heree" in paragraph.text:
                marker_paragraph = paragraph
                break

        if not marker_paragraph:
            print("⚠️ Marker 'Tables_Heree' not found")
            return

        parent = marker_paragraph._element.getparent()
        marker_index = parent.index(marker_paragraph._element)
        parent.remove(template_table._element)

        for run in marker_paragraph.runs:
            if "Tables_Heree" in run.text:
                run.text = run.text.replace("Tables_Heree", "")

        for idx, row in enumerate(df.iterrows(), start=1):
            _, row_data = row
            new_table = copy_table_with_formatting(template_table)

            vuln_name = str(row_data.get('Name of Vulnerability', f'Vulnerability {idx}'))
            risk_factor = str(row_data.get('Risk Factor', 'NA'))
            # Handle both old and new column names for backward compatibility
            cve_id = str(row_data.get('CVE/CWE ID', row_data.get('CVE ID', 'NA'))) if pd.notna(row_data.get('CVE/CWE ID', row_data.get('CVE ID'))) else "NA"
            cvss_value = str(row_data.get('CVSS', 'NA')) if pd.notna(row_data.get('CVSS')) else "NA"
            affected_systems = str(row_data.get('Affected Systems', ''))
            observation = str(row_data.get('Audit Observation', 'NA')) if pd.notna(row_data.get('Audit Observation')) else "NA"
            impact = str(row_data.get('Impact', 'NA')) if pd.notna(row_data.get('Impact')) else "NA"
            recommendation = str(row_data.get('Recommendation / Countermeasure', 'NA')) if pd.notna(row_data.get('Recommendation / Countermeasure')) else "NA"
            reference_link = str(row_data.get('Reference Link', 'NA')) if pd.notna(row_data.get('Reference Link')) else "NA"

            # Get status from the latest Status column (or closed Status column if exists)
            excel_row_idx = idx + 1  # Excel row (row 1 is header, so row 2 = index 0)
            status_value = 'New'  # Default
            
            if status_columns:
                # First check for closed status column
                closed_status_col = find_closed_status_column_for_vulnerability(ws, status_columns, excel_row_idx, idx - 1)
                target_status_col = closed_status_col if closed_status_col else latest_status_col
                
                if target_status_col:
                    status_cell = ws.cell(row=excel_row_idx, column=target_status_col['col'])
                    if status_cell.value:
                        status_value = str(status_cell.value).strip()
            
            status_lower = str(status_value).lower().strip()
            print(f"DEBUG: status_value='{status_value}', status_lower='{status_lower}'")

            replace_text_in_table(new_table, "Vulnerability___Name", vuln_name)
            replace_text_in_table_with_risk_colors(new_table, "Riskkkk", risk_factor)
            replace_text_in_table(new_table, "CWE_ID__", cve_id)
            replace_text_in_table(new_table, "CVSS__", cvss_value)
            replace_affected_systems_with_table(new_table, "Affected_URL___", affected_systems)
            replace_text_in_table(new_table, "Observation___", observation)
            replace_text_in_table(new_table, "Impact___", impact)
            replace_text_in_table(new_table, "Recommendation___", recommendation)
            replace_text_in_table(new_table, "Reference_Link___", reference_link)

            # Determine Statussssss value based on status (similar to Newwww_or_Repeatttt but with "Open" instead of "Repeat Observation")
            # Check for "closed with exception" first (substring match to handle variations)
            if "closed with exception" in status_lower:
                status_display = "Closed With Exception"
            elif status_lower == "open":
                status_display = "Open"
            elif status_lower in ["closed", "close", "na"]:
                status_display = "Closed"
            elif status_lower == "new":
                # "New" status typically means it's a new vulnerability, treat as "Open"
                status_display = "Open"
            else:
                # Default to Closed if unknown status
                print(f"WARNING: Unknown status '{status_lower}', defaulting to 'Closed'")
                status_display = "Closed"
            replace_text_in_table(new_table, "Statussssss", status_display)

            if "closed" in status_lower or "close" in status_lower or "closed with exception" in status_lower:
                new_or_repeat = "-"
            elif "open" in status_lower:
                new_or_repeat = "Repeat Observation"
            else:
                new_or_repeat = "-"
            replace_text_in_table(new_table, "Newwww_or_Repeatttt", new_or_repeat)

            # Get follow-up remark from remarks_dict (extracted from Excel based on Status logic)
            vuln_name = str(row_data.get('Name of Vulnerability', f'Vulnerability {idx}'))
            follow_up_remark = remarks_dict.get(vuln_name, '')
            
            if not follow_up_remark or not str(follow_up_remark).strip():
                # Fallback to default if not found in remarks_dict
                if "closed" in status_lower or "close" in status_lower or "closed with exception" in status_lower:
                    follow_up_remark = "This vulnerability has been mitigated by the Bank."
                else:
                    follow_up_remark = "This vulnerability has not been mitigated by the Bank."
            else:
                follow_up_remark = str(follow_up_remark).strip()
            
            replace_text_in_table(new_table, "Follow_up_Remarksssss", follow_up_remark)

            replace_text_in_table(new_table, "Not Applicableeee", "Not Applicable")
            replace_text_in_table(new_table, "New Observationnnn", "New Observation")

            replace_text_in_table(new_table, "POCsss", f"POCsss_{idx}_11")
            replace_text_in_table(new_table, "New_POCCsssss", f"New_POCCsssss_{idx}_11")

            for table_row in new_table.rows:
                for cell in table_row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text

                        def renumber(match):
                            return f"{idx}.{match.group(1)}"

                        updated_text = re.sub(r'\b1\.(\d+)\b', renumber, original_text)
                        if updated_text != original_text:
                            style = {}
                            if paragraph.runs:
                                ref_run = paragraph.runs[0]
                                style = {
                                    "name": ref_run.font.name,
                                    "size": ref_run.font.size,
                                    "color": ref_run.font.color.rgb if ref_run.font.color else None,
                                    "bold": ref_run.bold,
                                    "italic": ref_run.italic,
                                    "underline": ref_run.underline,
                                }

                            paragraph.text = ""
                            new_run = paragraph.add_run(updated_text)
                            if style.get("name"):
                                new_run.font.name = style["name"]
                            if style.get("size"):
                                new_run.font.size = style["size"]
                            if style.get("color"):
                                try:
                                    new_run.font.color.rgb = style["color"]
                                except Exception:
                                    pass
                            new_run.bold = style.get("bold")
                            new_run.italic = style.get("italic")
                            new_run.underline = style.get("underline")

            parent.insert(marker_index + 1, new_table._element)
            marker_index += 1

            page_break_para = doc.add_paragraph()
            page_break_para.add_run().add_break(WD_BREAK.PAGE)
            parent.insert(marker_index + 1, page_break_para._element)
            marker_index += 1
        
        # Close workbook after processing all rows
        wb.close()

        add_followup_images_at_end_of_document(doc, poc_images, old_poc_images, df)

    except Exception as e:
        print(f"❌ Error replacing follow-up vulnerability details: {e}")
        traceback.print_exc()


def find_status_columns_with_dates(ws, header_row=1):
    """
    Find all Status columns with dates in format "Status - DD.MM.YYYY"
    Returns list of dicts: [{'col': col_num, 'header': header_text, 'date': date_obj, 'col_index': col_idx}]
    """
    status_columns = []
    
    for col in range(1, ws.max_column + 1):
        cell_value = ws.cell(row=header_row, column=col).value
        if cell_value:
            cell_str = str(cell_value).strip()
            if cell_str.lower().startswith('status'):
                # Try to extract date from "Status - DD.MM.YYYY" format
                date_match = re.search(r'(\d{1,2})\.(\d{1,2})\.(\d{4})', cell_str)
                if date_match:
                    day = int(date_match.group(1))
                    month = int(date_match.group(2))
                    year = int(date_match.group(3))
                    date_obj = datetime(year, month, day)
                    status_columns.append({
                        'col': col,
                        'header': cell_str,
                        'date': date_obj,
                        'col_index': col - 1  # 0-based for pandas
                    })
    
    # Sort by date ascending (oldest first)
    status_columns.sort(key=lambda x: x['date'])
    
    return status_columns


def find_closed_status_column_for_vulnerability(ws, status_columns, row_idx, df_row_idx):
    """
    Find which Status column has "Closed" for a specific vulnerability row
    Returns: status_column dict or None
    """
    for status_col_info in status_columns:
        status_col = status_col_info['col']
        cell_value = ws.cell(row=row_idx, column=status_col).value
        if cell_value:
            status_value = str(cell_value).strip().upper()
            if status_value == 'CLOSED' or status_value == 'CLOSED WITH EXCEPTION':
                return status_col_info
    return None


def get_poc_image_columns_from_status_column(status_col):
    """
    Given a Status column number, return the 7 image columns in sequence: V, W, X, Y, Z, AA, U
    Pattern: Status+3, Status+4, Status+5, Status+6, Status+7, Status+8, Status+2
    Example: If Status is at S (19), return [22, 23, 24, 25, 26, 27, 21] = [V, W, X, Y, Z, AA, U]
    """
    # Calculate the 7 columns: +3, +4, +5, +6, +7, +8, +2 (in that order)
    poc_cols = [
        status_col + 3,  # V (first)
        status_col + 4,  # W (second)
        status_col + 5,  # X (third)
        status_col + 6,  # Y (fourth)
        status_col + 7,  # Z (fifth)
        status_col + 8,  # AA (sixth)
        status_col + 2,  # U (seventh)
    ]
    return poc_cols


def get_remark_column_from_status_column(status_col):
    """
    Given a Status column number, return the Remarks column (next column after Status)
    """
    return status_col + 1


def extract_followup_poc_images_from_excel(excel_file):
    """
    Extract POC images from the API follow-up Excel file.
    For each vulnerability:
    - Checks all Status columns for "Closed" status
    - If "Closed" found, extracts images from the 7 columns after that Status column
    - If no "Closed", uses latest Status column and extracts from the 7 columns after it
    Returns: (poc_images, old_poc_images, df, remarks_dict)
    - poc_images: Images for New_POCCsssss placeholder (from Status columns with logic above)
    - old_poc_images: Old POC images from columns L-R (12-18)
    - remarks_dict: Follow-up remarks mapped by vulnerability name
    """
    try:
        wb = load_workbook(excel_file)
        
        # Find "API VAPT" worksheet
        ws = None
        for sheet_name in wb.sheetnames:
            if "api vapt" in sheet_name.lower() or "api" in sheet_name.lower():
                ws = wb[sheet_name]
                break
        
        if not ws:
            print("❌ 'API VAPT' worksheet not found")
            wb.close()
            return {}, {}, pd.DataFrame()

        df = pd.read_excel(excel_file, sheet_name=ws.title)

        poc_images = {}
        old_poc_images = {}
        remarks_dict = {}

        # Find all Status columns with dates
        status_columns = find_status_columns_with_dates(ws, header_row=1)
        
        if not status_columns:
            print("⚠️ No Status columns with dates found")
            wb.close()
            return {}, {}, pd.DataFrame(), {}
        
        print(f"📋 Found {len(status_columns)} Status columns with dates:")
        for status_col_info in status_columns:
            print(f"  - {status_col_info['header']} at column {get_column_letter(status_col_info['col'])} ({status_col_info['col']})")
        
        # Get latest Status column (by date)
        latest_status_col = status_columns[-1] if status_columns else None
        if latest_status_col:
            print(f"📅 Latest Status column: {latest_status_col['header']} at column {get_column_letter(latest_status_col['col'])}")

        # All columns L-R (12-18) are Old POC images
        old_poc_col_start = 12  # L
        old_poc_col_end = 18    # R
        
        # Column order for insertion: M, N, O, P, Q, R, L (13, 14, 15, 16, 17, 18, 12)
        old_poc_column_order = [13, 14, 15, 16, 17, 18, 12]  # M, N, O, P, Q, R, L
        
        # Process each vulnerability row
        for df_row_idx in range(len(df)):
            vuln_name = str(df.iloc[df_row_idx]['Name of Vulnerability'])
            excel_row_idx = df_row_idx + 2  # Excel row (row 1 is header, so row 2 = index 0)
            
            # Find which Status column has "Closed" for this vulnerability
            closed_status_col = find_closed_status_column_for_vulnerability(ws, status_columns, excel_row_idx, df_row_idx)
            
            # Determine which Status column to use for POC images and remarks
            target_status_col = closed_status_col if closed_status_col else latest_status_col
            
            if target_status_col:
                status_col_num = target_status_col['col']
                status_type = "Closed" if closed_status_col else "Latest"
                print(f"📊 Vulnerability '{vuln_name}' (row {excel_row_idx}): Using {status_type} Status column '{target_status_col['header']}' at column {get_column_letter(status_col_num)}")
                
                # Get POC image columns from this Status column
                poc_image_cols = get_poc_image_columns_from_status_column(status_col_num)
                print(f"  📸 POC image columns: {[get_column_letter(col) for col in poc_image_cols]}")
                
                # Get remark column
                remark_col = get_remark_column_from_status_column(status_col_num)
                remark_col_letter = get_column_letter(remark_col)
                
                # Extract remark
                remark_cell = ws.cell(row=excel_row_idx, column=remark_col)
                remark_value = None
                if remark_cell.value:
                    remark_value = str(remark_cell.value).strip()
                    remarks_dict[vuln_name] = remark_value
                    remark_preview = remark_value[:50] + "..." if len(remark_value) > 50 else remark_value
                    print(f"  📝 Remark from column {remark_col_letter}: '{remark_preview}'")
                
                # Store POC image column info for this vulnerability (will extract images later)
                if vuln_name not in poc_images:
                    poc_images[vuln_name] = {
                        'image_columns': poc_image_cols,
                        'status_col': status_col_num,
                        'status_type': status_type,
                        'images': []  # Will store actual image data
                    }

        if hasattr(ws, "_images"):
            print(f"📸 Found {len(ws._images)} images in API sheet")
            
            for img_idx, img in enumerate(ws._images):
                try:
                    # Get image position (openpyxl uses 0-based row/col in anchor)
                    row_idx_0based = img.anchor._from.row  # 0-based
                    col_idx_0based = img.anchor._from.col  # 0-based
                    
                    # Convert to 1-based for comparison
                    row_idx = row_idx_0based + 1  # Excel row (1-based)
                    col_idx = col_idx_0based + 1  # Excel column (1-based)
                    
                    col_letter = get_column_letter(col_idx)
                    
                    # Get image data
                    img_data = None
                    if hasattr(img, '_data'):
                        if callable(img._data):
                            try:
                                img_data = img._data()
                            except Exception as e:
                                print(f"  ⚠️ Error calling img._data() for image {img_idx} at row {row_idx}, col {col_letter}: {e}")
                                continue
                        else:
                            img_data = img._data
                    
                    if img_data is None:
                        print(f"  ⚠️ No image data found for image {img_idx} at row {row_idx}, col {col_letter}")
                        continue
                    
                    # Check if image is valid (has some data)
                    if len(img_data) < 100:  # Very small images might be invalid
                        print(f"  ⚠️ Image {img_idx} at row {row_idx}, col {col_letter} appears too small ({len(img_data)} bytes)")
                        continue

                    # Convert Excel row to dataframe index (assuming row 1 is header)
                    excel_row_idx = row_idx - 2  # Row 2 in Excel = index 0 in dataframe
                    
                    if not (0 <= excel_row_idx < len(df)):
                        print(f"  ⚠️ Row {row_idx} (df index {excel_row_idx}) out of range for dataframe (length: {len(df)})")
                        continue

                    # Get vulnerability name
                    if 'Name of Vulnerability' not in df.columns:
                        print("  ⚠️ 'Name of Vulnerability' column not found in dataframe")
                        print(f"  Available columns: {list(df.columns)}")
                        continue
                    
                    vuln_name = str(df.iloc[excel_row_idx]['Name of Vulnerability'])
                    
                    print(f"  🔍 Processing image {img_idx}: row {row_idx}, col {col_letter} ({col_idx}), vulnerability: '{vuln_name}'")
                    
                    # Check if image is in POC image columns for this vulnerability
                    if vuln_name in poc_images:
                        poc_info = poc_images[vuln_name]
                        if col_idx in poc_info['image_columns']:
                            # Store image with column index for sorting
                            poc_info['images'].append((col_idx, img_data))
                            print(f"  ✅ Added POC (New) image for '{vuln_name}' at row {row_idx}, column {col_letter} ({col_idx})")
                    
                    # All images in columns L-R (12-18) are Old POC images
                    if old_poc_col_start <= col_idx <= old_poc_col_end:
                        # Store image with column index for sorting later
                        if vuln_name not in old_poc_images:
                            old_poc_images[vuln_name] = []  # Will store tuples: (col_idx, img_data)
                        old_poc_images[vuln_name].append((col_idx, img_data))
                        print(f"  ✅ Added Old POC image for '{vuln_name}' at row {row_idx}, column {col_letter} ({col_idx})")
                        
                except Exception as e:
                    print(f"  ⚠️ Error processing image {img_idx}: {e}")
                    traceback.print_exc()
        else:
            print("⚠️ No images attribute found in worksheet")

        wb.close()
        
        # Process POC images: sort by column order within each vulnerability
        poc_images_final = {}
        for vuln_name, poc_info in poc_images.items():
            if 'images' in poc_info and poc_info['images']:
                # Sort images by column order (same as image_columns order)
                image_cols = poc_info['image_columns']
                images_with_cols = poc_info['images']
                sorted_images = sorted(images_with_cols, key=lambda x: (
                    image_cols.index(x[0]) if x[0] in image_cols else 999
                ))
                # Extract just the image data in the sorted order
                poc_images_final[vuln_name] = [img_data for _, img_data in sorted_images]
                print(f"  - POC (New): '{vuln_name}': {len(poc_images_final[vuln_name])} images from {poc_info['status_type']} Status column")
        
        # Sort Old POC images by column order: M, N, O, P, Q, R, L
        for vuln_name in old_poc_images:
            # Sort images by column index according to the desired order
            images_with_cols = old_poc_images[vuln_name]
            sorted_images = sorted(images_with_cols, key=lambda x: (
                old_poc_column_order.index(x[0]) if x[0] in old_poc_column_order else 999
            ))
            # Extract just the image data in the sorted order
            old_poc_images[vuln_name] = [img_data for _, img_data in sorted_images]
            print(f"  - Old POC: '{vuln_name}': {len(old_poc_images[vuln_name])} images (sorted: M, N, O, P, Q, R, L)")
        
        # Print summary
        print(f"📊 API extracted {len(poc_images_final)} vulnerabilities with POC (New) images")
        print(f"📊 API extracted {len(old_poc_images)} vulnerabilities with Old POC images")
        print(f"📊 API extracted {len(remarks_dict)} vulnerabilities with remarks")

        return poc_images_final, old_poc_images, df, remarks_dict

    except Exception as e:
        print(f"❌ Error extracting API follow-up images: {e}")
        traceback.print_exc()
        return {}, {}, pd.DataFrame(), {}


def add_followup_images_at_end_of_document(doc, poc_images, old_poc_images, df):
    """
    Append old and new POC images for API follow-up audit at the end of the document.
    """
    try:
        has_poc = any(poc_images.values())
        has_old_poc = any(old_poc_images.values())

        if not has_poc and not has_old_poc:
            print("⚠️ No API POC images to append")
            return

        doc.add_page_break()
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("Proof of Concept (POC) Images")
        title_run.font.size = Pt(16)
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        image_count = 0
        for index, row in enumerate(df.iterrows(), start=1):
            _, row_data = row
            vuln_name = str(row_data.get('Name of Vulnerability', f'Vulnerability {index}'))

            has_images = (vuln_name in poc_images and poc_images[vuln_name]) or (
                vuln_name in old_poc_images and old_poc_images[vuln_name])
            if not has_images:
                continue

            image_count += 1
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(f"{image_count}. {vuln_name}")
            heading_run.font.size = Pt(14)
            heading_run.bold = True

            if vuln_name in old_poc_images and old_poc_images[vuln_name]:
                for i, image_data in enumerate(old_poc_images[vuln_name], start=1):
                    try:
                        pil_image = PILImage.open(io.BytesIO(image_data))
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                            pil_image.save(tmp, format='PNG')
                            temp_path = tmp.name

                        image_para = doc.add_paragraph()
                        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        image_para.paragraph_format.space_before = Pt(0)
                        image_para.paragraph_format.space_after = Pt(4 if i < len(old_poc_images[vuln_name]) else 2)
                        run = image_para.add_run()
                        picture = run.add_picture(temp_path, width=Inches(6.0))
                        apply_1pt_border_to_picture(picture)
                        os.unlink(temp_path)
                    except Exception as e:
                        err_para = doc.add_paragraph()
                        err_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        err_para.paragraph_format.space_before = Pt(0)
                        err_para.paragraph_format.space_after = Pt(2)
                        err_para.add_run(f"Error loading Old POC image {i}: {str(e)}")

            if vuln_name in poc_images and poc_images[vuln_name]:
                for i, image_data in enumerate(poc_images[vuln_name], start=1):
                    try:
                        pil_image = PILImage.open(io.BytesIO(image_data))
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                            pil_image.save(tmp, format='PNG')
                            temp_path = tmp.name

                        image_para = doc.add_paragraph()
                        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        image_para.paragraph_format.space_before = Pt(0)
                        image_para.paragraph_format.space_after = Pt(4 if i < len(poc_images[vuln_name]) else 2)
                        run = image_para.add_run()
                        picture = run.add_picture(temp_path, width=Inches(6.0))
                        apply_1pt_border_to_picture(picture)
                        os.unlink(temp_path)
                    except Exception as e:
                        err_para = doc.add_paragraph()
                        err_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        err_para.paragraph_format.space_before = Pt(0)
                        err_para.paragraph_format.space_after = Pt(2)
                        err_para.add_run(f"Error loading POC image {i}: {str(e)}")

            sections_with_images = sum(1 for name in df['Name of Vulnerability']
                                       if (name in poc_images and poc_images[name]) or
                                       (name in old_poc_images and old_poc_images[name]))
            if image_count < sections_with_images:
                doc.add_page_break()

        print(f"✅ Added {image_count} vulnerability sections with API POC images")

    except Exception as e:
        print(f"❌ Error adding API POC images: {e}")
        traceback.print_exc()


def insert_images_into_cell(cell, images):
    """
    Helper to insert multiple images into a Word cell while minimizing extra whitespace.
    """
    total = len(images)
    for i, image_data in enumerate(images, start=1):
        try:
            pil_image = PILImage.open(io.BytesIO(image_data))
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                pil_image.save(tmp, format='PNG')
                temp_path = tmp.name
            
            img_para = cell.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.paragraph_format.space_before = Pt(0)
            img_para.paragraph_format.space_after = Pt(4 if i < total else 0)
            run = img_para.add_run()
            picture = run.add_picture(temp_path, width=Inches(6.0))
            apply_1pt_border_to_picture(picture)
            os.unlink(temp_path)

        except Exception as e:
            error_para = cell.add_paragraph()
            error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            error_para.paragraph_format.space_before = Pt(0)
            error_para.paragraph_format.space_after = Pt(0)
            error_para.add_run(f"Error loading image {i}: {str(e)}")


def replace_followup_poc_placeholders_with_images(doc, excel_file):
    """
    Replace POC placeholders (POCsss / New_POCCsssss) in the API follow-up report with images.
    """
    try:
        poc_images, old_poc_images, df, remarks_dict = extract_followup_poc_images_from_excel(excel_file)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in list(cell.paragraphs):
                        text = paragraph.text
                        if not text:
                            continue

                        new_match = re.search(r'(New_POCCsssss_(\d+)_11)', text)
                        old_match = re.search(r'(POCsss_(\d+)_11)', text)

                        if new_match or old_match:
                            placeholder_element = paragraph._element
                            parent_element = placeholder_element.getparent()
                            if parent_element is not None:
                                parent_element.remove(placeholder_element)

                            handled = False

                            if old_match:
                                _, index_str = old_match.groups()
                                idx = int(index_str)
                                vuln_name = str(df.iloc[idx - 1]['Name of Vulnerability']) if idx - 1 < len(df) else ""
                                images = old_poc_images.get(vuln_name, [])
                                if images:
                                    insert_images_into_cell(cell, images)
                                    handled = True

                            if new_match:
                                _, index_str = new_match.groups()
                                idx = int(index_str)
                                vuln_name = str(df.iloc[idx - 1]['Name of Vulnerability']) if idx - 1 < len(df) else ""
                                images = poc_images.get(vuln_name, [])
                                if images:
                                    insert_images_into_cell(cell, images)
                                    handled = True

                            if not handled:
                                _set_cell_to_nil(cell)

                            break

    except Exception as e:
        print(f"❌ Error replacing Web follow-up POC placeholders: {e}")
        traceback.print_exc()


def replace_open_ports_with_nmap_table(doc, excel_file, placeholder="OPEN_PORTSSSS"):
    nmap_table = create_nmap_table_from_excel(doc, excel_file)
    if not nmap_table:
        print("❌ Failed to create Web Nmap table")
        return False

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            parent = paragraph._element.getparent()
            idx = parent.index(paragraph._element)
            parent.remove(paragraph._element)
            parent.insert(idx, nmap_table._element)
            print(f"✅ Replaced '{placeholder}' with Web Nmap table")
            return True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        cell._element.clear()
                        cell._element.append(nmap_table._element)
                        print(f"✅ Replaced '{placeholder}' with Web Nmap table inside table cell")
                        return True

    print(f"⚠️ Placeholder '{placeholder}' not found in Web document")
    return False


def replace_a11_row_with_vulnerabilities(doc, excel_file):
    """Find the table containing a.1.1 and add Excel data with proper columns including Status column"""
    try:
        # Load workbook to find latest Status column
        wb = load_workbook(excel_file)
        
        # Find "API VAPT" worksheet
        ws = None
        for sheet_name in wb.sheetnames:
            if "api vapt" in sheet_name.lower() or "api" in sheet_name.lower():
                ws = wb[sheet_name]
                break
        
        if not ws:
            print("❌ 'API VAPT' worksheet not found")
            wb.close()
            return
        
        # Find all Status columns with dates
        status_columns = []
        header_row = 1
        
        for col in range(1, ws.max_column + 1):
            cell_value = ws.cell(row=header_row, column=col).value
            if cell_value:
                cell_str = str(cell_value).strip()
                if cell_str.lower().startswith('status'):
                    # Try to extract date from "Status - DD.MM.YYYY" format
                    date_match = re.search(r'(\d{1,2})\.(\d{1,2})\.(\d{4})', cell_str)
                    if date_match:
                        day = int(date_match.group(1))
                        month = int(date_match.group(2))
                        year = int(date_match.group(3))
                        date_obj = datetime(year, month, day)
                        status_columns.append({
                            'col': col,
                            'header': cell_str,
                            'date': date_obj,
                            'col_index': col - 1  # 0-based index for pandas
                        })
        
        # Find latest Status column
        latest_status_col_name = None
        latest_status_col_index = None
        if status_columns:
            status_columns.sort(key=lambda x: x['date'], reverse=True)
            latest_status = status_columns[0]
            latest_status_col_name = latest_status['header']
            latest_status_col_index = latest_status['col_index']
            print(f"📊 Found latest Status column: {latest_status_col_name} at column {latest_status['col']} (index {latest_status_col_index})")
        else:
            print("⚠️ No Status columns with dates found, will try to find Status column by name")
        
        wb.close()
        
        # Read Excel with pandas
        df = pd.read_excel(excel_file, sheet_name='API VAPT')

        if df.empty:
            print("No data found in 'API VAPT' worksheet")
            return
        
        # Find the Status column in pandas dataframe by column name
        status_col_pandas_index = None
        if latest_status_col_name:
            # Try to find column by exact name or partial match
            for col_idx, col_name in enumerate(df.columns):
                if latest_status_col_name in str(col_name) or str(col_name) in latest_status_col_name:
                    status_col_pandas_index = col_idx
                    print(f"📊 Found Status column in pandas dataframe: '{col_name}' at index {col_idx}")
                    break
        
        # Fallback: try to find any Status column if not found by name
        if status_col_pandas_index is None:
            for col_idx, col_name in enumerate(df.columns):
                col_str = str(col_name).strip().lower()
                if col_str.startswith('status'):
                    # Check if it has a date
                    if re.search(r'\d{1,2}\.\d{1,2}\.\d{4}', str(col_name)):
                        status_col_pandas_index = col_idx
                        print(f"📊 Found Status column by fallback: '{col_name}' at index {col_idx}")
                        break
        
        # Final fallback: use the openpyxl column index if available
        if status_col_pandas_index is None and latest_status_col_index is not None:
            if latest_status_col_index < len(df.columns):
                status_col_pandas_index = latest_status_col_index
                print(f"📊 Using openpyxl column index {latest_status_col_index} as pandas index")
            else:
                print("⚠️ Could not find Status column, using default column 11")
                status_col_pandas_index = 11 if df.shape[1] > 11 else None

        target_table = None
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "a.1.1" in cell.text.lower():
                        target_table = table
                        break
                if target_table:
                    break
            if target_table:
                break

        if not target_table:
            print("Could not find table containing 'a.1.1'")
            return

        print("Found target table - updating content...")

        while len(target_table.rows) > 1:
            target_table._tbl.remove(target_table.rows[1]._tr)

        for idx in range(len(df)):
            new_row = target_table.add_row()

            if len(new_row.cells) < 12:
                continue

            for cell in new_row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            first_cell = new_row.cells[0]
            for paragraph in first_cell.paragraphs:
                paragraph.clear()
            p = first_cell.paragraphs[0]
            p.text = str(idx + 1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            second_cell = new_row.cells[1]
            for paragraph in second_cell.paragraphs:
                paragraph.clear()
            affected_system = str(df.iloc[idx, 10]) if df.shape[1] > 10 and pd.notna(df.iloc[idx, 10]) else "NA"
            p = second_cell.paragraphs[0]
            p.text = affected_system
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            third_cell = new_row.cells[2]
            for paragraph in third_cell.paragraphs:
                paragraph.clear()
            vulnerability = str(df.iloc[idx, 2]) if pd.notna(df.iloc[idx, 2]) else "NA"
            p = third_cell.paragraphs[0]
            p.text = vulnerability
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            fourth_cell = new_row.cells[3]
            for paragraph in fourth_cell.paragraphs:
                paragraph.clear()
            raw_value = df.iloc[idx, 4]
            cve_id = "NA" if pd.isna(raw_value) or str(raw_value).lower() == 'nan' else str(raw_value)
            p = fourth_cell.paragraphs[0]
            p.text = cve_id
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            fifth_cell = new_row.cells[7]
            for paragraph in fifth_cell.paragraphs:
                paragraph.clear()
            risk_factor = str(df.iloc[idx, 3]) if pd.notna(df.iloc[idx, 3]) else "NA"
            p = fifth_cell.paragraphs[0]
            p.text = risk_factor
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True

            risk_factor_lower = risk_factor.lower()
            if 'critical' in risk_factor_lower:
                fifth_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="8B0000"/>'.format(nsdecls('w'))))
            elif 'high' in risk_factor_lower:
                fifth_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w'))))
            elif 'medium' in risk_factor_lower:
                fifth_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FFA500"/>'.format(nsdecls('w'))))
            elif 'low' in risk_factor_lower:
                fifth_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="008000"/>'.format(nsdecls('w'))))

            sixth_cell = new_row.cells[4]
            for paragraph in sixth_cell.paragraphs:
                paragraph.clear()
            p = sixth_cell.paragraphs[0]
            p.text = "NA"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            seventh_cell = new_row.cells[5]
            for paragraph in seventh_cell.paragraphs:
                paragraph.clear()
            p = seventh_cell.paragraphs[0]
            p.text = "NA"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            eighth_cell = new_row.cells[6]
            for paragraph in eighth_cell.paragraphs:
                paragraph.clear()
            p = eighth_cell.paragraphs[0]
            p.text = "NA"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            ninth_cell = new_row.cells[8]
            for paragraph in ninth_cell.paragraphs:
                paragraph.clear()
            recommendation_w = clean_value(df.iloc[idx, 8]) if df.shape[1] > 8 else "NA"
            p = ninth_cell.paragraphs[0]
            p.text = recommendation_w
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            tenth_cell = new_row.cells[9]
            for paragraph in tenth_cell.paragraphs:
                paragraph.clear()
            reference = clean_value(df.iloc[idx, 9]) if df.shape[1] > 9 else "NA"
            p = tenth_cell.paragraphs[0]
            p.text = reference
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Get status from latest Status column
            status_value = "NA"
            if status_col_pandas_index is not None and df.shape[1] > status_col_pandas_index:
                status_value = clean_value(df.iloc[idx, status_col_pandas_index])
            
            status_lower = str(status_value).lower().strip()
            
            # Determine values for twelve_cell and eleven_cell based on latest Status column
            if status_lower == "open":
                twelve_cell_value = "Open"
                eleven_cell_value = "Repeat"
            elif status_lower in ["closed with exception"]:
                twelve_cell_value = "Closed With Exception"
                eleven_cell_value = "-"
            elif status_lower in ["closed", "close", "na"]:
                twelve_cell_value = "Closed"
                eleven_cell_value = "-"
            else:
                # Default to Closed if unknown status
                twelve_cell_value = "Closed"
                eleven_cell_value = "-"
            
            # Update twelve_cell (column 11, index 11)
            twelve_cell = new_row.cells[11]
            for paragraph in twelve_cell.paragraphs:
                paragraph.clear()
            p = twelve_cell.paragraphs[0]
            p.text = twelve_cell_value
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)
            
            # Update eleven_cell (column 10, index 10)
            eleven_cell = new_row.cells[10]
            for paragraph in eleven_cell.paragraphs:
                paragraph.clear()
            p = eleven_cell.paragraphs[0]
            p.text = eleven_cell_value
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

        print(f"✅ Successfully updated a.1.1 table with {len(df)} rows of data")

    except ValueError as e:
        print(f"Error updating a.1.1 table: {e}")
    except Exception as e:
        print(f"Error updating a.1.1 table: {str(e)}")
        traceback.print_exc()


def replace_followup_timeline_placeholders(doc, timeline):
    phase_placeholders = {
        "Follow-up Audit First Validation Cycle": {
            "start": "11.11.1111",
            "end": "22.22.2222"
        },
        "Follow-up Audit Second Validation Cycle": {
            "start": "33.33.3333",
            "end": "44.44.4444"
        },
        "Reporting": {
            "start": "55.55.5555",
            "end": "66.66.6666"
        }
    }

    replacements = {}
    for phase in timeline:
        name = phase.get("Phase")
        if name in phase_placeholders:
            placeholders = phase_placeholders[name]
            replacements[placeholders['start']] = phase.get('Start', '')
            replacements[placeholders['end']] = phase.get('End', '')

    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                _replace_text_in_runs(paragraph, placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in replacements.items():
                        if placeholder in paragraph.text:
                            _replace_text_in_runs(paragraph, placeholder, value)


def read_revalidation_cycle_dates_from_excel(excel_file_path):
    """
    Read REVALIDATION CYCLE dates from Excel Meta_Data worksheet.
    Returns a dictionary mapping cycle numbers to start and end dates.
    """
    try:
        workbook = openpyxl.load_workbook(excel_file_path)
        
        if 'Meta_Data' not in workbook.sheetnames:
            print("⚠️ Meta_Data worksheet not found")
            return {}
        
        sheet = workbook['Meta_Data']
        cycles = {}
        
        # Define mappings: cycle number -> (start_placeholder, end_placeholder)
        cycle_mappings = {
            1: ("11.11.1111", "22.22.2222"),
            2: ("33.33.3333", "44.44.4444"),
            3: ("55.55.5555", "66.66.6666"),
            4: ("77.77.7777", "88.88.8888"),
            5: ("99.99.9999", "aa.aa.aaaa"),
            6: ("bb.bb.bbbb", "cc.cc.cccc")
        }
        
        # Find all REVALIDATION CYCLE entries
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    cell_text = str(cell.value).strip()
                    
                    # Check for REVALIDATION CYCLE - N pattern
                    match = re.match(r'REVALIDATION CYCLE\s*-\s*(\d+)', cell_text, re.IGNORECASE)
                    if match:
                        cycle_num = int(match.group(1))
                        
                        if cycle_num in cycle_mappings:
                            start_placeholder, end_placeholder = cycle_mappings[cycle_num]
                            
                            # Search for Start Date and End Date in column A after this cell
                            cycle_row = cell.row
                            start_date = None
                            end_date = None
                            
                            # Look for Start Date and End Date in the next 10 rows
                            for search_row in range(cycle_row + 1, min(cycle_row + 11, sheet.max_row + 1)):
                                a_cell = sheet.cell(row=search_row, column=1)  # Column A
                                if a_cell.value:
                                    a_value = str(a_cell.value).strip()
                                    
                                    if 'start date' in a_value.lower() and not start_date:
                                        b_cell = sheet.cell(row=search_row, column=2)  # Column B
                                        if b_cell.value:
                                            start_date = str(b_cell.value).strip()
                                            start_date = convert_to_dd_mm_yyyy(start_date) if start_date else None
                                    
                                    elif 'end date' in a_value.lower() and not end_date:
                                        b_cell = sheet.cell(row=search_row, column=2)  # Column B
                                        if b_cell.value:
                                            end_date = str(b_cell.value).strip()
                                            end_date = convert_to_dd_mm_yyyy(end_date) if end_date else None
                            
                            if start_date and end_date:
                                cycles[cycle_num] = {
                                    'start': start_date,
                                    'end': end_date,
                                    'start_placeholder': start_placeholder,
                                    'end_placeholder': end_placeholder
                                }
                                print(f"✅ Found REVALIDATION CYCLE - {cycle_num}: {start_date} to {end_date}")
        
        workbook.close()
        return cycles
        
    except Exception as e:
        print(f"❌ Error reading revalidation cycle dates from Excel: {e}")
        traceback.print_exc()
        return {}


def replace_timeline_from_excel_metadata(doc, excel_file_path, metadata=None):
    """
    Replace timeline placeholders from Excel Meta_Data worksheet.
    Also handles row deletion if cycle is not found.
    """
    try:
        cycles = read_revalidation_cycle_dates_from_excel(excel_file_path)
        
        if not cycles:
            print("⚠️ No revalidation cycles found in Excel")
            return
        
        # Create replacements dictionary
        replacements = {}
        cycles_to_remove = []
        
        for cycle_num in range(1, 7):  # Cycles 1-6
            if cycle_num in cycles:
                cycle_data = cycles[cycle_num]
                replacements[cycle_data['start_placeholder']] = cycle_data['start']
                replacements[cycle_data['end_placeholder']] = cycle_data['end']
            else:
                # Mark cycles not found for row removal
                cycle_mappings = {
                    1: ("11.11.1111", "22.22.2222"),
                    2: ("33.33.3333", "44.44.4444"),
                    3: ("55.55.5555", "66.66.6666"),
                    4: ("77.77.7777", "88.88.8888"),
                    5: ("99.99.9999", "aa.aa.aaaa"),
                    6: ("bb.bb.bbbb", "cc.cc.cccc")
                }
                if cycle_num in cycle_mappings:
                    start_ph, end_ph = cycle_mappings[cycle_num]
                    cycles_to_remove.append(start_ph)
                    cycles_to_remove.append(end_ph)
                    print(f"⚠️ REVALIDATION CYCLE - {cycle_num} not found, will remove rows with {start_ph} or {end_ph}")
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    _replace_text_in_runs(paragraph, placeholder, value)
        
        # Replace placeholders in tables and remove rows if needed
        for table in doc.tables:
            rows_to_remove = []
            for row_idx, row in enumerate(table.rows):
                row_should_remove = False
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # Replace placeholders
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                _replace_text_in_runs(paragraph, placeholder, value)
                        
                        # Check if row should be removed
                        for placeholder in cycles_to_remove:
                            if placeholder in paragraph.text or placeholder in cell.text:
                                row_should_remove = True
                                break
                    
                    if row_should_remove:
                        break
                
                if row_should_remove:
                    rows_to_remove.append(row_idx)
            
            # Remove rows in reverse order to maintain indices
            for row_idx in reversed(rows_to_remove):
                if row_idx < len(table.rows):
                    tr = table.rows[row_idx]._tr
                    tr.getparent().remove(tr)
                    print(f"✅ Removed row {row_idx + 1} (contained placeholder from missing cycle)")
        
        print(f"✅ Replaced timeline placeholders from Excel Meta_Data ({len(cycles)} cycles found)")
        
    except Exception as e:
        print(f"❌ Error replacing timeline from Excel metadata: {e}")
        traceback.print_exc()


def calculate_working_days_backward(end_date_str, days=2):
    """
    Calculate a date that is N working days (excluding Saturday and Sunday) before the end date.
    Returns date in DD.MM.YYYY format.
    """
    try:
        # Parse end date
        if isinstance(end_date_str, str):
            # Try different date formats
            date_formats = ["%Y-%m-%d", "%d.%m.%Y", "%d/%m/%Y"]
            end_date = None
            for fmt in date_formats:
                try:
                    end_date = datetime.strptime(end_date_str, fmt)
                    break
                except ValueError:
                    continue
            
            if not end_date:
                print(f"⚠️ Could not parse date: {end_date_str}")
                return end_date_str
        else:
            end_date = end_date_str
        
        current_date = end_date
        working_days_counted = 0
        
        # Go backward until we've counted N working days
        while working_days_counted < days:
            current_date = current_date - timedelta(days=1)
            # Check if it's not Saturday (5) or Sunday (6)
            if current_date.weekday() < 5:  # Monday = 0, Friday = 4
                working_days_counted += 1
        
        return current_date.strftime("%d.%m.%Y")
        
    except Exception as e:
        print(f"❌ Error calculating working days backward: {e}")
        traceback.print_exc()
        # Return original date if error occurs
        if isinstance(end_date_str, str):
            try:
                # Try to format it
                date_formats = ["%Y-%m-%d", "%d.%m.%Y", "%d/%m/%Y"]
                for fmt in date_formats:
                    try:
                        end_date_obj = datetime.strptime(end_date_str, fmt)
                        return end_date_obj.strftime("%d.%m.%Y")
                    except ValueError:
                        continue
            except:
                pass
        return str(end_date_str)


def replace_finish_date_placeholders(doc, finish_date):
    """
    Replace dd.dd.dddd (Finish Date - 2 working days) and ee.ee.eeee (Finish Date).
    """
    try:
        # Format finish date to DD.MM.YYYY
        if isinstance(finish_date, str):
            date_formats = ["%Y-%m-%d", "%d.%m.%Y", "%d/%m/%Y"]
            finish_date_obj = None
            for fmt in date_formats:
                try:
                    finish_date_obj = datetime.strptime(finish_date, fmt)
                    break
                except ValueError:
                    continue
            
            if finish_date_obj:
                ee_ee_eeee = finish_date_obj.strftime("%d.%m.%Y")
                dd_dd_dddd = calculate_working_days_backward(finish_date_obj, days=2)
            else:
                # Try to convert using existing function
                try:
                    ee_ee_eeee = convert_to_dd_mm_yyyy(finish_date)
                    finish_date_obj = datetime.strptime(finish_date, "%Y-%m-%d") if finish_date else None
                    dd_dd_dddd = calculate_working_days_backward(finish_date_obj, days=2) if finish_date_obj else finish_date
                except:
                    ee_ee_eeee = finish_date
                    dd_dd_dddd = finish_date
        else:
            ee_ee_eeee = finish_date.strftime("%d.%m.%Y") if hasattr(finish_date, 'strftime') else str(finish_date)
            dd_dd_dddd = calculate_working_days_backward(finish_date, days=2)
        
        replacements = {
            "dd.dd.dddd": dd_dd_dddd,
            "ee.ee.eeee": ee_ee_eeee
        }
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    _replace_text_in_runs(paragraph, placeholder, value)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                _replace_text_in_runs(paragraph, placeholder, value)
        
        print(f"✅ Replaced finish date placeholders: dd.dd.dddd = {dd_dd_dddd}, ee.ee.eeee = {ee_ee_eeee}")
        
    except Exception as e:
        print(f"❌ Error replacing finish date placeholders: {e}")
        traceback.print_exc()


def remove_content_after_marker(doc, marker_text):
    from VAPT_Dashboard_Files.Infra_VAPT_Follow_Up_Audit_Word_Report import remove_content_after_marker as infra_remove_marker
    return infra_remove_marker(doc, marker_text)


def get_followup_vulnerability_counts(excel_file):
    """
    Extract vulnerability counts from "API VAPT" worksheet:
    - First Audit: Count all vulnerabilities by Risk Factor (column D)
    - Follow-up Audit: Count vulnerabilities with "Open" status in latest Status column, grouped by Risk Factor
    Returns: (categories, first_audit_values, followup_audit_values)
    """
    try:
        wb = load_workbook(excel_file)
        
        # Find "API VAPT" worksheet
        ws = None
        for sheet_name in wb.sheetnames:
            if "api vapt" in sheet_name.lower() or "api" in sheet_name.lower():
                ws = wb[sheet_name]
                break
        
        if not ws:
            print("❌ 'API VAPT' worksheet not found")
            wb.close()
            return None, None, None
        
        print(f"📊 Reading from worksheet: {ws.title}")
        
        categories = ['Critical', 'High', 'Medium', 'Low']
        first_audit_counts = {'Critical': 0, 'High': 0, 'Medium': 0, 'Low': 0}
        followup_audit_counts = {'Critical': 0, 'High': 0, 'Medium': 0, 'Low': 0}
        
        # Find header row and locate columns
        header_row = 1
        risk_factor_col = None
        status_columns = []
        
        # Find Risk Factor column
        for col in range(1, ws.max_column + 1):
            cell_value = ws.cell(row=header_row, column=col).value
            if cell_value:
                cell_str = str(cell_value).strip().lower()
                if 'risk' in cell_str and 'factor' in cell_str:
                    risk_factor_col = col
                    print(f"📊 Found Risk Factor column at column {col}")
                    break
        
        if risk_factor_col is None:
            # Default to column D (4)
            risk_factor_col = 4
            print(f"📊 Using default Risk Factor column D ({risk_factor_col})")
        
        # Find all Status columns with dates
        for col in range(1, ws.max_column + 1):
            cell_value = ws.cell(row=header_row, column=col).value
            if cell_value:
                cell_str = str(cell_value).strip()
                if cell_str.lower().startswith('status'):
                    # Try to extract date from "Status - DD.MM.YYYY" format
                    date_match = re.search(r'(\d{1,2})\.(\d{1,2})\.(\d{4})', cell_str)
                    if date_match:
                        day = int(date_match.group(1))
                        month = int(date_match.group(2))
                        year = int(date_match.group(3))
                        date_obj = datetime(year, month, day)
                        status_columns.append({
                            'col': col,
                            'header': cell_str,
                            'date': date_obj
                        })
        
        # Find latest Status column
        latest_status_col = None
        if status_columns:
            status_columns.sort(key=lambda x: x['date'], reverse=True)
            latest_status_col = status_columns[0]
            print(f"📊 Found latest Status column: {latest_status_col['header']} at column {latest_status_col['col']}")
        else:
            print("⚠️ No Status columns with dates found")
        
        # Count vulnerabilities starting from row 2 (assuming row 1 is header)
        for row in range(2, ws.max_row + 1):
            # Get Risk Factor
            risk_factor_cell = ws.cell(row=row, column=risk_factor_col)
            risk_factor = None
            if risk_factor_cell.value:
                risk_factor = str(risk_factor_cell.value).strip()
            
            # Skip if no risk factor
            if not risk_factor:
                continue
            
            # Normalize risk factor
            risk_factor_upper = risk_factor.upper()
            if 'CRITICAL' in risk_factor_upper:
                risk_factor = 'Critical'
            elif 'HIGH' in risk_factor_upper:
                risk_factor = 'High'
            elif 'MEDIUM' in risk_factor_upper:
                risk_factor = 'Medium'
            elif 'LOW' in risk_factor_upper:
                risk_factor = 'Low'
            else:
                continue  # Skip unknown risk factors
            
            # Count for First Audit (all vulnerabilities)
            if risk_factor in first_audit_counts:
                first_audit_counts[risk_factor] += 1
            
            # Count for Follow-up Audit (only if Status is "Open" in latest Status column)
            if latest_status_col:
                status_cell = ws.cell(row=row, column=latest_status_col['col'])
                status_value = None
                if status_cell.value:
                    status_value = str(status_cell.value).strip().upper()
                
                if status_value == 'OPEN':
                    if risk_factor in followup_audit_counts:
                        followup_audit_counts[risk_factor] += 1
        
        wb.close()
        
        # Convert to lists in order
        first_audit_values = [first_audit_counts[cat] for cat in categories]
        followup_audit_values = [followup_audit_counts[cat] for cat in categories]
        
        print(f"📊 First Audit counts: {first_audit_values}")
        print(f"📊 Follow-up Audit counts: {followup_audit_values}")
        
        return categories, first_audit_values, followup_audit_values
        
    except Exception as e:
        print(f"❌ Error extracting vulnerability counts: {e}")
        traceback.print_exc()
        return None, None, None


def update_followup_chart_in_docx(docx_path, excel_file, chart_file=None):
    """Update the chart in the Follow-Up Audit DOCX file with two data series"""
    try:
        # Step 1: Get vulnerability data from Excel
        categories, first_audit_values, followup_audit_values = get_followup_vulnerability_counts(excel_file)
        if not categories or first_audit_values is None or followup_audit_values is None:
            print("❌ Failed to extract vulnerability data from Excel")
            return False

        print(f"📊 First Audit data: {dict(zip(categories, first_audit_values))}")
        print(f"📊 Follow-up Audit data: {dict(zip(categories, followup_audit_values))}")

        # Create a temporary copy to work with
        temp_dir = tempfile.mkdtemp()
        temp_docx = os.path.join(temp_dir, "temp.docx")
        shutil.copy2(docx_path, temp_docx)

        # Step 2: Read the DOCX as a ZIP file and find chart files
        with zipfile.ZipFile(temp_docx, 'r') as zin:
            file_list = zin.namelist()
            
            # Find chart files if not specified
            if chart_file is None:
                chart_files = [f for f in file_list if f.startswith('word/charts/') and f.endswith('.xml')]
                if not chart_files:
                    print("❌ No chart files found in Word document")
                    return False
                chart_file = chart_files[0]  # Use the first chart found
                print(f"📊 Using chart file: {chart_file}")
            
            try:
                with zin.open(chart_file) as chart_file_obj:
                    chart_xml = chart_file_obj.read()
            except KeyError:
                print(f"❌ Chart file not found: {chart_file}")
                return False

        # Step 3: Parse and update the chart XML
        tree = etree.fromstring(chart_xml)

        ns = {
            'c': 'http://schemas.openxmlformats.org/drawingml/2006/chart',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
        }

        # Find all series in the chart (should be 2: First Audit and Follow-up Audit)
        all_series = tree.findall('.//c:ser', namespaces=ns)
        print(f"📊 Found {len(all_series)} data series in chart")

        # Update each series with corresponding data
        for series_idx, series in enumerate(all_series[:2]):  # Limit to first 2 series
            if series_idx == 0:
                # First series: First Audit data
                values_to_use = first_audit_values
                series_name = "First Audit"
            elif series_idx == 1:
                # Second series: Follow-up Audit data
                values_to_use = followup_audit_values
                series_name = "Follow-up Audit"
            else:
                continue
            
            print(f"📊 Updating series {series_idx + 1} ({series_name}): {values_to_use}")
            
            # Find value points in this series
            val_pts = series.findall('.//c:val//c:numCache//c:pt', namespaces=ns)
            
            for i, val_pt in enumerate(val_pts):
                if i < len(values_to_use):
                    val_elem = val_pt.find('c:v', namespaces=ns)
                    if val_elem is not None:
                        val_elem.text = str(values_to_use[i])
                        print(f"  Updated {series_name} data point {i} ({categories[i]}): {values_to_use[i]}")

        # Y-axis scaling logic
        all_values = first_audit_values + followup_audit_values
        max_value = max(all_values) if all_values else 1
        
        if max_value == 0:
            max_value = 1
            print("📊 All values are 0, setting max to 1 for chart display")
        
        # Determine Y-axis max and interval
        if max_value <= 10:
            max_axis = 10
            interval = 1
        elif max_value <= 20:
            max_axis = 20
            interval = 2
        elif max_value <= 40:
            max_axis = 40
            interval = 4
        elif max_value <= 80:
            max_axis = 80
            interval = 8
        elif max_value <= 120:
            max_axis = 120
            interval = 12
        else:
            max_axis = 1
            while max_axis < max_value:
                max_axis *= 2
            interval = max_axis // 4

        # Update Y-axis scaling
        val_ax = tree.find('.//c:valAx', namespaces=ns)
        if val_ax is not None:
            scaling_elem = val_ax.find('.//c:scaling', namespaces=ns)
            if scaling_elem is not None:
                max_elem = scaling_elem.find('.//c:max', namespaces=ns)
                if max_elem is None:
                    max_elem = etree.SubElement(scaling_elem, '{http://schemas.openxmlformats.org/drawingml/2006/chart}max')
                max_elem.set('val', str(max_axis))
            
            # Update major unit
            major_unit_elem = val_ax.find('.//c:majorUnit', namespaces=ns)
            if major_unit_elem is None:
                major_unit_elem = etree.SubElement(val_ax, '{http://schemas.openxmlformats.org/drawingml/2006/chart}majorUnit')
            if major_unit_elem is not None:
                major_unit_elem.set('val', str(interval))

        # Serialize and update the DOCX
        updated_chart_xml = etree.tostring(tree, pretty_print=True, xml_declaration=True, encoding='UTF-8')

        with zipfile.ZipFile(temp_docx, 'w') as zout:
            for file in file_list:
                if file == chart_file:
                    zout.writestr(file, updated_chart_xml)
                else:
                    with zipfile.ZipFile(docx_path, 'r') as zin:
                        with zin.open(file) as src_file:
                            zout.writestr(file, src_file.read())

        shutil.move(temp_docx, docx_path)
        shutil.rmtree(temp_dir)

        print(f"✅ Chart updated successfully in: {docx_path}")
        print(f"📈 Y-axis scaled to max: {max_axis}, interval: {interval}")
        return True

    except Exception as e:
        print(f"❌ Error updating chart: {str(e)}")
        traceback.print_exc()
        if 'temp_dir' in locals():
            shutil.rmtree(temp_dir, ignore_errors=True)
        return False


def generate_follow_up_audit_word_report(excel_file_path, form_metadata=None):
    """
    Generate Follow-Up Audit Word report for API from Excel file
    form_metadata: Dictionary containing metadata from form (if provided, will be used instead of Excel Meta_Data)
    """
    try:
        template_path = os.path.join("static", "Formats_and_Catalog", "API_VAPT_Follow_Up_Audit_Report.docx")

        if not os.path.exists(template_path):
            raise Exception(f"Template file not found: {template_path}")

        print(f"📄 Using template: {template_path}")

        doc = Document(template_path)

        # Use form metadata if provided, otherwise extract from Excel
        if form_metadata:
            print("🔍 Using metadata from form data...")
            metadata = form_metadata
        else:
            print("🔍 Extracting metadata from Excel file...")
            metadata = find_and_print_metadata(excel_file_path)

        if metadata:
            print("📋 Using metadata:")
            for key, value in metadata.items():
                if value:
                    print(f"  {key}: {value}")

        if metadata:
            print("🔍 Replacing metadata placeholders...")
            replace_metadata_placeholders(doc, metadata)

            print("🔍 Processing auditing team members...")
            replace_auditor_placeholders_and_add_rows(doc, metadata)
        else:
            print("⚠️ No metadata found")

        print("🔍 Replacing scope placeholders with data from form or Meta_Data worksheet...")
        replace_scope_placeholders_with_data(doc, excel_file_path, form_metadata)

        print("🔍 Replacing vulnerability details with images...")
        replace_followup_vulnerability_details_with_images(doc, excel_file_path)

        print("🔍 Replacing POC placeholders with images...")
        replace_followup_poc_placeholders_with_images(doc, excel_file_path)

        print("🔍 Replacing OPEN_PORTSSSS with Nmap table...")
        replace_open_ports_with_nmap_table(doc, excel_file_path, "OPEN_PORTSSSS")

        print("🔍 Replacing a.1.1 table with vulnerability data...")
        replace_a11_row_with_vulnerabilities(doc, excel_file_path)

        # Replace timeline placeholders from Excel Meta_Data worksheet
        print("🔍 Reading timeline data from Excel Meta_Data worksheet...")
        replace_timeline_from_excel_metadata(doc, excel_file_path, metadata)
        
        # Replace dd.dd.dddd and ee.ee.eeee based on Finish Audit Date
        if metadata and metadata.get("end_date"):
            print("🔍 Replacing dd.dd.dddd and ee.ee.eeee placeholders...")
            replace_finish_date_placeholders(doc, metadata.get("end_date"))

        marker_to_remove_after = "CEH Cert. No: 12345678901234567890"
        print(f"🔍 Removing content after marker: '{marker_to_remove_after}'...")
        remove_content_after_marker(doc, marker_to_remove_after)

        temp_dir = tempfile.mkdtemp()
        temp_file_path = os.path.join(temp_dir, "API_VAPT_Follow_Up_Audit_Report.docx")
        doc.save(temp_file_path)

        print("🔍 Updating chart with vulnerability data...")
        update_followup_chart_in_docx(temp_file_path, excel_file_path)

        print(f"✅ Word report generated successfully: {temp_file_path}")
        return temp_file_path

    except Exception as e:
        print(f"❌ Error generating Word report: {e}")
        traceback.print_exc()
        raise Exception(f"Error generating Word report: {str(e)}")


@api_follow_up_word_report_bp.route('/process_api_vapt_follow_up_audit_word_report', methods=['POST'])
def process_api_vapt_follow_up_audit_word_report():
    """
    Process the uploaded Excel file and generate an API Follow-Up Audit Word report
    """
    try:
        if 'followUpExcelFile' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400

        file = request.files['followUpExcelFile']

        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not file.filename.lower().endswith('.xlsx'):
            return jsonify({'error': 'Please upload an Excel file (.xlsx)'}), 400

        # Extract form metadata
        organization = request.form.get('organization', '')
        if organization == 'other':
            organization = request.form.get('otherOrganization', '')
        
        city = request.form.get('city', '')
        if city == 'other':
            city = request.form.get('otherCity', '')
        
        state = request.form.get('state', '')
        first_audit_report_id = request.form.get('firstAuditReportId', '')
        first_audit_report_date = request.form.get('firstAuditReportDate', '')
        start_date = request.form.get('startDate', '')
        end_date = request.form.get('endDate', '')
        # Report Prepared By
        prepared_by_title = request.form.get('preparedByTitle', '')
        prepared_by_name = request.form.get('preparedByName', '')
        report_prepared_by = f"{prepared_by_title} {prepared_by_name}".strip() if prepared_by_title and prepared_by_name else ''
        
        # Auditee Details
        auditee_title = request.form.get('auditeeTitle', '')
        auditee_name = request.form.get('auditeeName', '')
        auditee_details_1 = f"{auditee_title} {auditee_name}".strip() if auditee_title and auditee_name else ''
        auditee_details_2 = request.form.get('designation', '')
        
        # Bank Email Addresses
        bank_email_addresses = request.form.getlist('bankEmail[]')
        bank_email_addresses = [email.strip() for email in bank_email_addresses if email.strip()]
        
        # Asset Details
        asset_descriptions = request.form.getlist('assetDescription[]')
        asset_criticalities = request.form.getlist('assetCriticality[]')
        asset_urls = request.form.getlist('assetUrl[]')
        asset_hashes = request.form.getlist('assetHash[]')
        asset_versions = request.form.getlist('assetVersion[]')
        
        # Build asset data list
        asset_data = []
        for idx in range(len(asset_descriptions)):
            if asset_descriptions[idx] and asset_criticalities[idx] and asset_urls[idx] and asset_hashes[idx] and asset_versions[idx]:
                asset_info = {
                    'description': asset_descriptions[idx],
                    'criticality': asset_criticalities[idx],
                    'url': asset_urls[idx],
                    'hash_value': asset_hashes[idx],
                    'version': asset_versions[idx],
                    'index': idx + 1
                }
                asset_data.append(asset_info)
        
        print(f"📋 Extracted {len(asset_data)} assets from form")
        
        # Auditing Team Members
        team_names = request.form.getlist('teamName[]')
        team_designations = request.form.getlist('teamDesignation[]')
        team_emails = request.form.getlist('teamEmail[]')
        team_qualifications = request.form.getlist('teamQualification[]')
        team_certified = []
        
        # Get team certified values (they come as teamCertified[0], teamCertified[1], etc.)
        i = 0
        while True:
            certified_value = request.form.get(f'teamCertified[{i}]')
            if certified_value is None:
                break
            team_certified.append(certified_value)
            i += 1
        
        auditing_team_members = []
        for idx in range(len(team_names)):
            if team_names[idx] and team_designations[idx] and team_emails[idx] and team_qualifications[idx]:
                member_data = {
                    f"Team Member {idx + 1} - Name": team_names[idx],
                    f"Team Member {idx + 1} - Designation": team_designations[idx],
                    f"Team Member {idx + 1} - Email": team_emails[idx],
                    f"Team Member {idx + 1} - Qualification": team_qualifications[idx],
                    f"Team Member {idx + 1} - Certified": team_certified[idx] if idx < len(team_certified) else 'no'
                }
                auditing_team_members.append(member_data)
        
        # Build metadata dictionary
        form_metadata = {
            "organization_name": organization,
            "city": city,
            "state": state,
            "start_date": start_date,
            "end_date": end_date,
            "first_audit_report_id": first_audit_report_id,
            "first_audit_report_date": first_audit_report_date,
            "report_prepared_by": report_prepared_by,
            "auditee_details_1": auditee_details_1,
            "auditee_details_2": auditee_details_2,
            "bank_email_addresses": bank_email_addresses,
            "auditing_team_members": auditing_team_members,
            "asset_data": asset_data  # Add asset data to metadata
        }
        
        print(f"📋 Form metadata extracted:")
        for key, value in form_metadata.items():
            if value:
                print(f"  {key}: {value}")

        temp_dir = tempfile.mkdtemp()
        excel_file_path = os.path.join(temp_dir, file.filename)
        file.save(excel_file_path)

        print(f"📁 Saved Excel file to: {excel_file_path}")

        try:
            word_file_path = generate_follow_up_audit_word_report(excel_file_path, form_metadata)

            return send_file(
                word_file_path,
                as_attachment=True,
                download_name='API_VAPT_Follow_Up_Audit_Report.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        finally:
            try:
                if os.path.exists(excel_file_path):
                    os.remove(excel_file_path)
                if 'word_file_path' in locals() and os.path.exists(word_file_path):
                    os.remove(word_file_path)
                    temp_word_dir = os.path.dirname(word_file_path)
                    if temp_word_dir != temp_dir and os.path.exists(temp_word_dir):
                        os.rmdir(temp_word_dir)
                os.rmdir(temp_dir)
            except Exception as cleanup_error:
                print(f"Warning: Error during cleanup: {cleanup_error}")

    except Exception as e:
        print(f"❌ Error processing api follow-up audit word report: {e}")
        traceback.print_exc()
        return jsonify({'error': f'Error generating report: {str(e)}'}), 500

