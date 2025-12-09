import os
import tempfile
import traceback
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from flask import Blueprint, request, send_file, jsonify
import re

# Create blueprint for certificate generation
vapt_first_audit_certificate_bp = Blueprint('vapt_first_audit_certificate', __name__)

def replace_text_in_paragraph(paragraph, old_text, new_text, force_font=None):
    """Replace specific text in a paragraph while preserving formatting and keeping the rest of the paragraph intact"""
    if old_text in paragraph.text:
        # Find the run that contains the old text
        for run in paragraph.runs:
            if old_text in run.text:
                # Store the formatting of the current run
                font_name = run.font.name
                font_size = run.font.size
                bold = run.bold
                italic = run.italic
                underline = run.underline
                font_color = None
                try:
                    if run.font.color.rgb:
                        font_color = run.font.color.rgb
                except:
                    pass
                
                # Replace the text in this run
                run.text = run.text.replace(old_text, new_text)
                
                # Restore the formatting
                if force_font:
                    run.font.name = force_font
                elif font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = font_size
                if bold is not None:
                    run.bold = bold
                if italic is not None:
                    run.italic = italic
                if underline is not None:
                    run.underline = underline
                if font_color:
                    try:
                        run.font.color.rgb = font_color
                    except:
                        pass
                return True
        return False
    return False

def replace_text_in_document(doc, replacements):
    """Replace text throughout the document"""
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                # Check if this is a date field that needs Times New Roman
                force_font = None
                if old_text in ["11.11.1111", "22.22.2222"]:
                    force_font = "Times New Roman"
                
                # Try simple replacement first
                if not replace_text_in_paragraph(paragraph, old_text, new_text, force_font):
                    # If simple replacement fails, try multi-run replacement
                    replace_text_in_runs(paragraph, old_text, new_text, force_font)
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            # Check if this is a date field that needs Times New Roman
                            force_font = None
                            if old_text in ["11.11.1111", "22.22.2222"]:
                                force_font = "Times New Roman"
                            
                            # Try simple replacement first
                            if not replace_text_in_paragraph(paragraph, old_text, new_text, force_font):
                                # If simple replacement fails, try multi-run replacement
                                replace_text_in_runs(paragraph, old_text, new_text, force_font)

def replace_text_in_runs(paragraph, old_text, new_text, force_font=None):
    """Handle text replacement when text spans multiple runs"""
    # Combine all run texts to check if old_text exists
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return False
    
    # Find the position of old_text in the combined text
    start_pos = full_text.find(old_text)
    end_pos = start_pos + len(old_text)
    
    # Get the formatting of the first run
    first_run = paragraph.runs[0] if paragraph.runs else None
    if not first_run:
        return False
    
    # Store formatting
    font_name = first_run.font.name
    font_size = first_run.font.size
    bold = first_run.bold
    italic = first_run.italic
    underline = first_run.underline
    font_color = None
    try:
        if first_run.font.color.rgb:
            font_color = first_run.font.color.rgb
    except:
        pass
    
    # Clear all runs
    for run in paragraph.runs:
        run.clear()
    
    # Rebuild the paragraph with replacement
    before_text = full_text[:start_pos]
    after_text = full_text[end_pos:]
    
    # Add text before the replacement
    if before_text:
        before_run = paragraph.add_run(before_text)
        before_run.font.name = font_name
        before_run.font.size = font_size
        before_run.bold = bold
        before_run.italic = italic
        before_run.underline = underline
        if font_color:
            try:
                before_run.font.color.rgb = font_color
            except:
                pass
    
    # Add the replacement text
    new_run = paragraph.add_run(new_text)
    if force_font:
        new_run.font.name = force_font
    else:
        new_run.font.name = font_name
    new_run.font.size = font_size
    new_run.bold = bold
    new_run.italic = italic
    new_run.underline = underline
    if font_color:
        try:
            new_run.font.color.rgb = font_color
        except:
            pass
    
    # Add text after the replacement
    if after_text:
        after_run = paragraph.add_run(after_text)
        after_run.font.name = font_name
        after_run.font.size = font_size
        after_run.bold = bold
        after_run.italic = italic
        after_run.underline = underline
        if font_color:
            try:
                after_run.font.color.rgb = font_color
            except:
                pass
    
    return True

def format_date_to_dd_mm_yyyy(date_str):
    """Convert date from YYYY-MM-DD to DD.MM.YYYY format"""
    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        return date_obj.strftime('%d.%m.%Y')
    except ValueError:
        return date_str

def get_current_financial_year():
    """Get current financial year in YYYY-YYYY format"""
    current_year = datetime.now().year
    current_month = datetime.now().month
    
    # Financial year starts from April (month 4)
    if current_month >= 4:
        return f"{current_year}-{current_year + 1}"
    else:
        return f"{current_year - 1}-{current_year}"

def generate_vapt_first_audit_certificate(organization_name, city, state, auditee_designation, scope,
                                   start_audit_date, end_audit_date, report_id):
    """
    Generate a Infrastructure VAPT First Audit Completion Certificate document by copying the template and replacing placeholders
    """
    try:
        # Path to the template file
        template_path = os.path.join('static', 'Formats_and_Catalog', 'Infrastructure_VAPT_First_Audit_Conduct_Certificate.docx')
        
        # Check if template file exists
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        # Create temp directory
        temp_dir = tempfile.mkdtemp()
        
        # Copy the template file to temp directory
        import shutil
        certificate_filename = "Infrastrcture_VAPT_Certificate.docx"
        certificate_path = os.path.join(temp_dir, certificate_filename)
        shutil.copy2(template_path, certificate_path)
        
        print(f"✅ Certificate template copied successfully: {certificate_path}")
        
        # Load the copied document for text replacement
        doc = Document(certificate_path)
        
        # Format dates
        start_date_formatted = format_date_to_dd_mm_yyyy(start_audit_date)
        end_date_formatted = format_date_to_dd_mm_yyyy(end_audit_date)
        current_financial_year = get_current_financial_year()
        
        # Define replacements
        replacements = {
            "Designationnnn": auditee_designation,
            "Organization_Nameeee": organization_name,
            "Cityyy": city,
            "Stateeee": state,
            "YYYY1–YYYY2": current_financial_year,
            "Report_IDDD": report_id,
            "Scoppeee": scope,
            "11.11.1111": start_date_formatted,
            "22.22.2222": end_date_formatted
        }
        
        # Perform text replacements
        print("🔄 Replacing placeholders in certificate...")
        replace_text_in_document(doc, replacements)
        
        # Save the updated document
        doc.save(certificate_path)
        
        print(f"✅ Certificate generated successfully with replaced data: {certificate_path}")
        print(f"📋 Certificate Data Applied:")
        print(f"   Organization: {organization_name}")
        print(f"   City: {city}")
        print(f"   State: {state}")
        print(f"   Auditee Designation: {auditee_designation}")
        print(f"   Scope: {scope}")
        print(f"   Financial Year: {current_financial_year}")
        print(f"   Start Date: {start_date_formatted}")
        print(f"   End Date: {end_date_formatted}")
        print(f"   Report ID: {report_id}")
        
        return certificate_path
        
    except Exception as e:
        print(f"❌ Error generating certificate: {e}")
        traceback.print_exc()
        raise e

@vapt_first_audit_certificate_bp.route('/process_vapt_first_audit_certificate', methods=['POST'])
def process_vapt_first_audit_certificate():
    """
    Process the certificate form data and generate a Infrastructure VAPT First Audit Completion Certificate
    """
    try:
        # Get form data
        organization_name = request.form.get('organizationName', '').strip()
        city = request.form.get('city', '').strip()
        state = request.form.get('state', '').strip()
        auditee_designation = request.form.get('auditeeDesignation', '').strip()
        scope = request.form.get('scope', '').strip()
        start_audit_date = request.form.get('startAuditDate', '').strip()
        end_audit_date = request.form.get('endAuditDate', '').strip()
        report_id = request.form.get('reportId', '').strip()
        
        # Handle "Other (Please specify)" fields
        if organization_name == 'other':
            organization_name = request.form.get('otherOrganizationName', '').strip()
        if city == 'other':
            city = request.form.get('otherCity', '').strip()
        if state == 'other':
            state = request.form.get('otherState', '').strip()
        
        # Validate required fields
        if not all([organization_name, city, state, auditee_designation, scope, start_audit_date, end_audit_date, report_id]):
            return jsonify({'error': 'All fields are required'}), 400
        
        print(f"📋 Certificate Data Received:")
        print(f"   Organization: {organization_name}")
        print(f"   City: {city}")
        print(f"   State: {state}")
        print(f"   Auditee Designation: {auditee_designation}")
        print(f"   Scope: {scope}")
        print(f"   Start Date: {start_audit_date}")
        print(f"   End Date: {end_audit_date}")
        print(f"   Report ID: {report_id}")
        
        # Generate the certificate
        certificate_path = generate_vapt_first_audit_certificate(
            organization_name, city, state, auditee_designation, scope,
            start_audit_date, end_audit_date, report_id
        )
        
        # Send the generated certificate file
        return send_file(
            certificate_path,
            as_attachment=True,
            download_name="Infrastrcture_VAPT_Certificate.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"❌ Error processing VAPT first audit certificate: {e}")
        traceback.print_exc()
        return jsonify({'error': f'Error generating certificate: {str(e)}'}), 500

