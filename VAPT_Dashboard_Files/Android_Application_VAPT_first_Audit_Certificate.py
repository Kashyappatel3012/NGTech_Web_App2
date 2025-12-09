import os
import tempfile
import traceback
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401
from docx.shared import Inches, Pt, RGBColor  # noqa: F401
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL  # noqa: F401
from flask import Blueprint, request, send_file, jsonify
import re  # noqa: F401

# Create blueprint for Android Application VAPT First Audit Certificate generation
android_app_vapt_first_audit_certificate_bp = Blueprint('android_app_vapt_first_audit_certificate', __name__)


def replace_text_in_paragraph(paragraph, old_text, new_text, force_font=None):
    """Replace specific text in a paragraph while preserving formatting and keeping the rest of the paragraph intact"""
    if old_text in paragraph.text:
        for run in paragraph.runs:
            if old_text in run.text:
                font_name = run.font.name
                font_size = run.font.size
                bold = run.bold
                italic = run.italic
                underline = run.underline
                font_color = None
                try:
                    if run.font.color.rgb:
                        font_color = run.font.color.rgb
                except:  # noqa: E722
                    pass

                run.text = run.text.replace(old_text, new_text)

                if force_font:
                    run.font.name = force_font
                elif font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = font_size
                if bold is not None:
                    run.bold = bold
                if italic is not None:
                    run.italic = italic
                if underline is not None:
                    run.underline = underline
                if font_color:
                    try:
                        run.font.color.rgb = font_color
                    except:  # noqa: E722
                        pass
                return True
        return False
    return False


def replace_text_in_document(doc, replacements):
    """Replace text throughout the document"""
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                force_font = None
                if old_text in ["11.11.1111", "22.22.2222"]:
                    force_font = "Times New Roman"

                if not replace_text_in_paragraph(paragraph, old_text, new_text, force_font):
                    replace_text_in_runs(paragraph, old_text, new_text, force_font)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            force_font = None
                            if old_text in ["11.11.1111", "22.22.2222"]:
                                force_font = "Times New Roman"

                            if not replace_text_in_paragraph(paragraph, old_text, new_text, force_font):
                                replace_text_in_runs(paragraph, old_text, new_text, force_font)


def replace_text_in_runs(paragraph, old_text, new_text, force_font=None):
    """Handle text replacement when text spans multiple runs"""
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return False

    start_pos = full_text.find(old_text)
    end_pos = start_pos + len(old_text)

    first_run = paragraph.runs[0] if paragraph.runs else None
    if not first_run:
        return False

    font_name = first_run.font.name
    font_size = first_run.font.size
    bold = first_run.bold
    italic = first_run.italic
    underline = first_run.underline
    font_color = None
    try:
        if first_run.font.color.rgb:
            font_color = first_run.font.color.rgb
    except:  # noqa: E722
        pass

    for run in paragraph.runs:
        run.clear()

    before_text = full_text[:start_pos]
    after_text = full_text[end_pos:]

    if before_text:
        before_run = paragraph.add_run(before_text)
        before_run.font.name = font_name
        before_run.font.size = font_size
        before_run.bold = bold
        before_run.italic = italic
        before_run.underline = underline
        if font_color:
            try:
                before_run.font.color.rgb = font_color
            except:  # noqa: E722
                pass

    new_run = paragraph.add_run(new_text)
    if force_font:
        new_run.font.name = force_font
    else:
        new_run.font.name = font_name
    new_run.font.size = font_size
    new_run.bold = bold
    new_run.italic = italic
    new_run.underline = underline
    if font_color:
        try:
            new_run.font.color.rgb = font_color
        except:  # noqa: E722
            pass

    if after_text:
        after_run = paragraph.add_run(after_text)
        after_run.font.name = font_name
        after_run.font.size = font_size
        after_run.bold = bold
        after_run.italic = italic
        after_run.underline = underline
        if font_color:
            try:
                after_run.font.color.rgb = font_color
            except:  # noqa: E722
                pass

    return True


def format_date_to_dd_mm_yyyy(date_str):
    """Convert date from YYYY-MM-DD to DD.MM.YYYY format"""
    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        return date_obj.strftime('%d.%m.%Y')
    except ValueError:
        return date_str


def get_current_financial_year():
    """Get current financial year in YYYY-YYYY format"""
    current_year = datetime.now().year
    current_month = datetime.now().month

    if current_month >= 4:
        return f"{current_year}-{current_year + 1}"
    else:
        return f"{current_year - 1}-{current_year}"


def generate_android_app_vapt_first_audit_certificate(organization_name, city, state, auditee_designation, scope,
                                                      start_audit_date, end_audit_date, report_id):
    """
    Generate an Android Mobile Application VAPT First Audit Completion Certificate document by copying the template and replacing placeholders
    """
    try:
        template_path = os.path.join('static', 'Formats_and_Catalog', 'Android_Application_VAPT_First_Audit_Conduct_Certificate.docx')

        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")

        temp_dir = tempfile.mkdtemp()

        import shutil
        certificate_filename = "Android_Application_VAPT_Certificate.docx"
        certificate_path = os.path.join(temp_dir, certificate_filename)
        shutil.copy2(template_path, certificate_path)

        print(f"✅ Android Application certificate template copied successfully: {certificate_path}")

        doc = Document(certificate_path)

        start_date_formatted = format_date_to_dd_mm_yyyy(start_audit_date)
        end_date_formatted = format_date_to_dd_mm_yyyy(end_audit_date)
        current_financial_year = get_current_financial_year()

        replacements = {
            "Designationnnn": auditee_designation,
            "Organization_Nameeee": organization_name,
            "Cityyy": city,
            "Stateeee": state,
            "YYYY1–YYYY2": current_financial_year,
            "Report_IDDD": report_id,
            "Scoppeee": scope,
            "11.11.1111": start_date_formatted,
            "22.22.2222": end_date_formatted
        }

        print("🔄 Replacing placeholders in Android Application certificate...")
        replace_text_in_document(doc, replacements)

        doc.save(certificate_path)

        print(f"✅ Android Application certificate generated successfully with replaced data: {certificate_path}")
        print("📋 Android Application Certificate Data Applied:")
        print(f"   Organization: {organization_name}")
        print(f"   City: {city}")
        print(f"   State: {state}")
        print(f"   Auditee Designation: {auditee_designation}")
        print(f"   Scope: {scope}")
        print(f"   Financial Year: {current_financial_year}")
        print(f"   Start Date: {start_date_formatted}")
        print(f"   End Date: {end_date_formatted}")
        print(f"   Report ID: {report_id}")

        return certificate_path

    except Exception as e:
        print(f"❌ Error generating Android Application certificate: {e}")
        traceback.print_exc()
        raise e


@android_app_vapt_first_audit_certificate_bp.route('/process_android_app_vapt_first_audit_certificate', methods=['POST'])
def process_android_app_vapt_first_audit_certificate():
    """
    Process the certificate form data and generate an Android Mobile Application VAPT First Audit Completion Certificate
    """
    try:
        organization_name = request.form.get('organizationName', '').strip()
        city = request.form.get('city', '').strip()
        state = request.form.get('state', '').strip()
        auditee_designation = request.form.get('auditeeDesignation', '').strip()
        scope = request.form.get('scope', '').strip()
        start_audit_date = request.form.get('startAuditDate', '').strip()
        end_audit_date = request.form.get('endAuditDate', '').strip()
        report_id = request.form.get('reportId', '').strip()

        if organization_name == 'other':
            organization_name = request.form.get('otherOrganizationName', '').strip()
        if city == 'other':
            city = request.form.get('otherCity', '').strip()
        if state == 'other':
            state = request.form.get('otherState', '').strip()

        if not all([organization_name, city, state, auditee_designation, scope, start_audit_date, end_audit_date, report_id]):
            return jsonify({'error': 'All fields are required'}), 400

        print("📋 Android Application Certificate Data Received:")
        print(f"   Organization: {organization_name}")
        print(f"   City: {city}")
        print(f"   State: {state}")
        print(f"   Auditee Designation: {auditee_designation}")
        print(f"   Scope: {scope}")
        print(f"   Start Date: {start_audit_date}")
        print(f"   End Date: {end_audit_date}")
        print(f"   Report ID: {report_id}")

        certificate_path = generate_android_app_vapt_first_audit_certificate(
            organization_name, city, state, auditee_designation, scope,
            start_audit_date, end_audit_date, report_id
        )

        return send_file(
            certificate_path,
            as_attachment=True,
            download_name="Android_Application_VAPT_Certificate.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        print(f"❌ Error processing Android Application first audit certificate: {e}")
        traceback.print_exc()
        return jsonify({'error': f'Error generating certificate: {str(e)}'}), 500


