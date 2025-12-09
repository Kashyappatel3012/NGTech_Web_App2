import os
import re
import tempfile
import traceback
from flask import Blueprint, request, send_file, jsonify
from docx import Document
from openpyxl import load_workbook

from VAPT_Dashboard_Files.Web_Application_First_Audit_Word_Report import (
    clean_value,
    _replace_text_in_runs,
    find_and_print_metadata,
    replace_metadata_placeholders,
    replace_auditor_placeholders_and_add_rows,
    replace_vulnerability_details_with_images,
    replace_poc_placeholders_with_images,
    replace_open_ports_with_nmap_table,
    replace_scope_placeholders_with_data,
    replace_a11_row_with_vulnerabilities,
    remove_content_after_marker,
    update_chart_in_docx
)
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import RGBColor, Pt
import pandas as pd

ios_app_word_report_bp = Blueprint('ios_app_word_report', __name__)


def _extract_ios_asset_hash_values(excel_file_path):
    """Read Meta_Data worksheet and collect Asset Hash Value entries in order."""
    hash_values = []
    try:
        wb = load_workbook(excel_file_path, data_only=True)
        if 'Meta_Data' not in wb.sheetnames:
            print("⚠️ 'Meta_Data' worksheet not found while extracting hash values.")
            wb.close()
            return hash_values

        ws = wb['Meta_Data']
        pattern = re.compile(r'Asset\s+(\d+)\s*-\s*Hash Value', re.IGNORECASE)

        for row in ws.iter_rows(values_only=True):
            if not row or row[0] is None:
                continue

            key = str(row[0]).strip()
            if pattern.match(key):
                value = clean_value(row[1]) if len(row) > 1 else "NA"
                hash_values.append(value)

        wb.close()
    except Exception as e:
        print(f"❌ Error extracting asset hash values: {e}")
        traceback.print_exc()

    return hash_values


def _replace_hash_placeholders_in_doc(doc: Document, hash_values):
    """Replace sequential HAAAASH placeholders in the document with provided hash values."""
    index = 0

    def next_hash_value():
        nonlocal index
        if index < len(hash_values):
            value = hash_values[index]
        else:
            value = "NA"
        index += 1
        return value

    def process_paragraph(paragraph):
        while "HAAAASH" in paragraph.text:
            replacement = next_hash_value()
            _replace_text_in_runs(paragraph, "HAAAASH", replacement)

    def process_table(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
                for nested_table in cell.tables:
                    process_table(nested_table)

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    for table in doc.tables:
        process_table(table)


def _prepare_ios_excel_for_web_logic(original_excel_path):
    """
    Create a temporary copy of the IOS Excel file with worksheet names tailored for
    the web application Word generation logic.
    """
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
    temp_file.close()

    try:
        wb = load_workbook(original_excel_path)

        if 'IOS Mobile Application VAPT' in wb.sheetnames and 'Web Application VAPT' not in wb.sheetnames:
            ws = wb['IOS Mobile Application VAPT']
            ws.title = 'Web Application VAPT'

        wb.save(temp_file.name)
        wb.close()
    except Exception as e:
        print(f"⚠️ Error preparing IOS Excel for Word logic: {e}")
        traceback.print_exc()
        # Fallback: copy original file as-is
        try:
            wb.close()
        except Exception:
            pass
        import shutil
        shutil.copy2(original_excel_path, temp_file.name)

    return temp_file.name


def _check_empty_a2_and_handle_ios(doc, original_excel_file):
    """Check if A2 is empty in IOS Mobile Application VAPT worksheet and handle accordingly"""
    try:
        wb = load_workbook(original_excel_file, data_only=True)
        if 'IOS Mobile Application VAPT' in wb.sheetnames:
            ws = wb['IOS Mobile Application VAPT']
            a2_value = ws.cell(row=2, column=1).value  # A2 is row 2, column 1 (1-indexed)
            # Check if A2 is empty or None
            if a2_value is None or str(a2_value).strip() == '':
                print("A2 cell is empty - removing table with a.1.1 and Vulnerability___Name, replacing Tables_Heree with message")
                
                # Find the target table (the one containing "a.1.1")
                target_table = None
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if "a.1.1" in cell.text.lower():
                                target_table = table
                                break
                        if target_table:
                            break
                    if target_table:
                        break
                
                if target_table:
                    # Remove the entire table
                    target_table._element.getparent().remove(target_table._element)
                    print("✅ Removed a.1.1 table")
                else:
                    print("⚠️ Could not find table containing 'a.1.1'")
                
                # Remove all tables containing "Recommendation___" (in case they're on different pages)
                tables_to_remove = []
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if "Recommendation___" in cell.text:
                                tables_to_remove.append(table)
                                break
                        if table in tables_to_remove:
                            break
                
                for table in tables_to_remove:
                    try:
                        table._element.getparent().remove(table._element)
                        print(f"✅ Removed table containing 'Recommendation___'")
                    except Exception as e:
                        print(f"⚠️ Could not remove table: {e}")
                
                # Replace "Tables_Heree" with green bold message
                for paragraph in doc.paragraphs:
                    if "Tables_Heree" in paragraph.text:
                        # Clear existing runs
                        paragraph.clear()
                        # Add line break before the message
                        break_run = paragraph.add_run()
                        break_run.add_break(WD_BREAK.LINE)
                        # Add new run with green bold text
                        run = paragraph.add_run("Note: No vulnerabilities were identified by the auditor during the audit.")
                        run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
                        run.font.bold = True
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print("✅ Replaced 'Tables_Heree' with green bold message")
                        break
                
                # Also check in runs
                for paragraph in doc.paragraphs:
                    for idx, run in enumerate(paragraph.runs):
                        if "Tables_Heree" in run.text:
                            # Create a new run with line break and insert it before the current run
                            break_run = paragraph.add_run()
                            break_run.add_break(WD_BREAK.LINE)
                            # Get the XML element for the break run and move it before the current run
                            run_element = run._element
                            break_element = break_run._element
                            run_element.getparent().remove(break_element)
                            run_element.getparent().insert(run_element.getparent().index(run_element), break_element)
                            
                            run.text = "Note: No vulnerabilities were identified by the auditor during the audit."
                            run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
                            run.font.bold = True
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            print("✅ Replaced 'Tables_Heree' in run with green bold message")
                            break
                
                wb.close()
                return True
        wb.close()
    except Exception as e:
        print(f"⚠️ Error checking A2 cell: {e}")
    return False


def _check_empty_a2_for_a11_ios(doc, original_excel_file):
    """Check if A2 is empty in IOS Mobile Application VAPT worksheet for a11 function"""
    try:
        df = pd.read_excel(original_excel_file, sheet_name='IOS Mobile Application VAPT')
        
        # Check if A2 cell is empty (row 2, column 1, index 0 for row, index 0 for column)
        a2_empty = False
        if not df.empty and len(df) >= 1:
            first_cell_value = df.iloc[0, 0] if len(df.columns) > 0 else None
            note_message = "Note: No vulnerabilities were identified by the auditor during the audit."
            if pd.isna(first_cell_value) or str(first_cell_value).strip() == '' or note_message in str(first_cell_value):
                a2_empty = True
        elif df.empty:
            a2_empty = True
        
        if a2_empty:
            # Find the target table (the one containing "a.1.1")
            target_table = None
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "a.1.1" in cell.text.lower():
                            target_table = table
                            break
                    if target_table:
                        break
                if target_table:
                    break

            if target_table:
                print("A2 is empty - removing a.1.1 table...")
                target_table._element.getparent().remove(target_table._element)
                print("✅ Removed a.1.1 table")
                
                # Remove all tables containing "Recommendation___"
                tables_to_remove = []
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if "Recommendation___" in cell.text:
                                tables_to_remove.append(table)
                                break
                        if table in tables_to_remove:
                            break
                
                for table in tables_to_remove:
                    try:
                        table._element.getparent().remove(table._element)
                        print(f"✅ Removed table containing 'Recommendation___'")
                    except Exception as e:
                        print(f"⚠️ Could not remove table: {e}")
                
                # Replace "Tables_Heree" with note message
                for paragraph in doc.paragraphs:
                    if "Tables_Heree" in paragraph.text:
                        paragraph.clear()
                        break_run = paragraph.add_run()
                        break_run.add_break(WD_BREAK.LINE)
                        run = paragraph.add_run("Note: No vulnerabilities were identified by the auditor during the audit.")
                        run.font.name = 'Calibri'
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 128, 0)
                        run.font.bold = True
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print("✅ Replaced 'Tables_Heree' with note message")
                        break
                
                # Also check in runs
                for paragraph in doc.paragraphs:
                    for idx, run in enumerate(paragraph.runs):
                        if "Tables_Heree" in run.text:
                            break_run = paragraph.add_run()
                            break_run.add_break(WD_BREAK.LINE)
                            run_element = run._element
                            break_element = break_run._element
                            run_element.getparent().remove(break_element)
                            run_element.getparent().insert(run_element.getparent().index(run_element), break_element)
                            
                            run.text = "Note: No vulnerabilities were identified by the auditor during the audit."
                            run.font.name = 'Calibri'
                            run.font.size = Pt(12)
                            run.font.color.rgb = RGBColor(0, 128, 0)
                            run.font.bold = True
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            print("✅ Replaced 'Tables_Heree' in run with note message")
                            break
                
                return True
    except Exception as e:
        print(f"⚠️ Error checking A2 cell for a11: {e}")
    return False


def generate_ios_word_report_from_excel(excel_file_path):
    """
    Generate IOS Mobile Application First Audit Word report by leveraging the
    Web Application Word report logic and applying IOS-specific placeholder replacements.
    """
    temp_excel_for_web = _prepare_ios_excel_for_web_logic(excel_file_path)

    try:
        template_path = os.path.join('static', 'Formats_and_Catalog', 'IOS_Mobile_Application_VAPT_First_Audit_Report.docx')
        if not os.path.exists(template_path):
            raise Exception(f"Template file not found at: {template_path}")

        doc = Document(template_path)

        print("🔍 Extracting metadata from Excel file...")
        metadata = find_and_print_metadata(temp_excel_for_web)

        if metadata:
            print("📋 Extracted metadata:")
            for key, value in metadata.items():
                if value:
                    print(f"  {key}: {value}")

        if metadata:
            print("🔍 Replacing metadata placeholders...")
            replace_metadata_placeholders(doc, metadata)

            print("🔍 Processing auditing team members...")
            replace_auditor_placeholders_and_add_rows(doc, metadata)
        else:
            print("⚠️ No metadata found in Excel file")

        # Check original Excel file for empty A2 before processing
        if _check_empty_a2_and_handle_ios(doc, excel_file_path):
            # A2 is empty, tables already removed and message added, skip vulnerability processing
            pass
        else:
            replace_vulnerability_details_with_images(doc, temp_excel_for_web)
        
        replace_poc_placeholders_with_images(doc, temp_excel_for_web)
        replace_open_ports_with_nmap_table(doc, temp_excel_for_web, "OPEN_PORTSSSS")

        print("🔍 Replacing asset placeholders with data from Meta_Data worksheet...")
        replace_scope_placeholders_with_data(doc, temp_excel_for_web)

        # Check original Excel file for empty A2 before calling a11 function
        if not _check_empty_a2_for_a11_ios(doc, excel_file_path):
            replace_a11_row_with_vulnerabilities(doc, temp_excel_for_web)

        marker_to_remove_after = "CEH Cert. No: 12345678901234567890"
        remove_content_after_marker(doc, marker_to_remove_after)

        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_docx:
            doc.save(temp_docx.name)
            temp_docx_path = temp_docx.name

        try:
            print("Updating chart in Word document with vulnerability data...")
            chart_update_success = update_chart_in_docx(temp_docx_path, temp_excel_for_web)
            if chart_update_success:
                print("✅ Chart updated successfully in Word document")
            else:
                print("⚠️ Chart update failed, but continuing with document generation")
        except Exception as e:
            print(f"❌ Error updating chart: {e}")
            print("⚠️ Continuing with document generation without chart update")

        doc = Document(temp_docx_path)
        os.unlink(temp_docx_path)

        hash_values = _extract_ios_asset_hash_values(excel_file_path)
        _replace_hash_placeholders_in_doc(doc, hash_values)

        print("✅ IOS Mobile Application Word report generated successfully from Excel file")
        return doc

    finally:
        try:
            if os.path.exists(temp_excel_for_web):
                os.unlink(temp_excel_for_web)
        except Exception:
            pass


@ios_app_word_report_bp.route('/process_ios_app_vapt_first_audit_word_report', methods=['POST'])
def process_ios_app_first_audit_word_report():
    """Process the uploaded Excel file and generate an IOS Mobile Application Word report."""
    try:
        if 'excelFile' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400

        excel_file = request.files['excelFile']
        if excel_file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not excel_file.filename.lower().endswith('.xlsx'):
            return jsonify({'error': 'Please upload a valid Excel file (.xlsx)'}), 400

        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as temp_excel:
            excel_file.save(temp_excel.name)
            temp_excel_path = temp_excel.name

        try:
            print("🔍 Starting IOS Mobile Application Word report generation...")
            doc = generate_ios_word_report_from_excel(temp_excel_path)

            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_doc:
                doc.save(temp_doc.name)
                temp_doc_path = temp_doc.name

            response = send_file(
                temp_doc_path,
                as_attachment=True,
                download_name='IOS_Mobile_Application_VAPT_First_Audit_Report.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

            os.unlink(temp_excel_path)

            import atexit

            def cleanup_generated_doc():
                try:
                    if os.path.exists(temp_doc_path):
                        os.unlink(temp_doc_path)
                        print(f"🧹 Cleaned up Word file: {temp_doc_path}")
                except Exception:
                    pass

            atexit.register(cleanup_generated_doc)

            print("✅ IOS Mobile Application Word report generated successfully.")
            return response

        except Exception as processing_error:
            print(f"❌ Error generating IOS Word report: {processing_error}")
            traceback.print_exc()
            if os.path.exists(temp_excel_path):
                os.unlink(temp_excel_path)
            return jsonify({'error': str(processing_error)}), 500

    except Exception as e:
        print(f"❌ Error in IOS Word report route: {e}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

