import copy
import io
import os
import re
import tempfile
import traceback
import zipfile
import shutil
from datetime import datetime, timedelta
from lxml import etree
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor
import pandas as pd
from openpyxl import load_workbook
import openpyxl
from PIL import Image as PILImage
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import qn
from docx.oxml.ns import nsdecls
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from flask import Blueprint, request, send_file, jsonify


vapt_follow_up_word_report_bp = Blueprint('vapt_follow_up_word_report', __name__)

def apply_1pt_border_to_picture(picture):
    """
    Apply a 1pt solid black border to a picture object.
    Works for all images added via add_picture().
    
    Args:
        picture: The picture object returned by run.add_picture()
    """
    try:
        # Access the picture element
        pic = picture._inline.graphic.graphicData.pic
        spPr = pic.spPr
        
        # Remove any existing line borders
        for ln in spPr.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}ln'):
            spPr.remove(ln)
        
        # Create 1pt border with the width attribute directly in the tag
        from docx.oxml import parse_xml
        border = parse_xml(
            '<a:ln xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" w="12700">'
            '<a:solidFill><a:srgbClr val="000000"/></a:solidFill>'
            '<a:prstDash val="solid"/>'
            '</a:ln>'
        )
        spPr.append(border)
        
        print(f"✅ Applied 1pt border (12700 EMUs)")
        
    except Exception as e:
        print(f"❌ Error applying border: {e}")
        import traceback
        traceback.print_exc()

# Helper function to convert NaN values to "NA"
def clean_value(value):
    """Convert NaN, None, empty strings to 'NA'"""
    if pd.isna(value) or value is None or str(value).lower() in ['nan', 'none', '']:
        return "NA"
    return str(value)

# Helper to convert openpyxl color to RGB
def get_rgb(color):
    if color is None or color.type != "rgb":
        return None
    if len(color.rgb) > 6:
        return color.rgb[2:]
    return color.rgb

def add_borders_to_cell(cell):
    """Add borders to a table cell"""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Create border elements
        borders = ['top', 'left', 'bottom', 'right']
        for border_name in borders:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcPr.append(border)
    except Exception as e:
        print(f"Warning: Could not add borders to cell: {e}")

def add_borders_to_row(row):
    """Add borders to all cells in a table row"""
    for cell in row.cells:
        add_borders_to_cell(cell)

from datetime import datetime, timedelta

def add_ordinal_suffix(day):
    """Add ordinal suffix to day with leading zero and superscript (01ˢᵗ, 02ⁿᵈ, 03ʳᵈ, 04ᵗʰ, etc.)"""
    if 10 <= day % 100 <= 20:
        suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    
    # Convert suffix to superscript using Unicode characters
    superscript_map = {
        's': 'ˢ', 't': 'ᵗ', 'n': 'ⁿ', 'd': 'ᵈ', 'r': 'ʳ', 'h': 'ʰ'
    }
    superscript_suffix = ''.join(superscript_map.get(c, c) for c in suffix)
    
    # Add leading zero for single digit days
    day_str = f"{day:02d}"
    return f"{day_str}{superscript_suffix}"

def convert_to_dd_mm_yyyy(date_str):
    """Convert various date formats to DD.MM.YYYY format"""
    date_formats = ["%d.%m.%Y", "%Y-%m-%d", "%m/%d/%Y", "%d-%m-%Y", "%B %Y", "%b %Y"]
    for fmt in date_formats:
        try:
            date_obj = datetime.strptime(date_str, fmt)
            return date_obj.strftime("%d.%m.%Y")
        except ValueError:
            continue
    try:
        year = int(date_str)
        return f"01.01.{year}"
    except ValueError:
        raise ValueError(f"Could not parse date: {date_str}")


def format_date_for_month_year(date_str):
    """Format date string to 'Month YYYY' format"""
    try:
        if not date_str:
            return ""
        
        # Try different date formats
        date_formats = ["%d.%m.%Y", "%Y-%m-%d", "%m/%d/%Y", "%d-%m-%Y"]
        for fmt in date_formats:
            try:
                date_obj = datetime.strptime(str(date_str), fmt)
                return date_obj.strftime("%B %Y")
            except ValueError:
                continue
        
        # If all formats fail, try to extract year
        year_match = re.search(r'\b(20\d{2})\b', str(date_str))
        if year_match:
            year = year_match.group(1)
            return f"January {year}"
        
        return str(date_str)
    except Exception as e:
        print(f"Error formatting date: {e}")
        return str(date_str)

def format_date_for_range(start_date_str, end_date_str):
    """Format date range to 'DD Month YYYY to DD Month YYYY' format"""
    try:
        if not start_date_str or not end_date_str:
            return ""
        
        start_formatted = format_date_for_dd_month_yyyy(start_date_str)
        end_formatted = format_date_for_dd_month_yyyy(end_date_str)
        
        return f"{start_formatted} to {end_formatted}"
    except Exception as e:
        print(f"Error formatting date range: {e}")
        return ""

def format_date_for_dd_month_yyyy(date_str):
    """Format date string to 'DD Month YYYY' format with ordinal suffix"""
    try:
        if not date_str:
            return ""
        
        # Try different date formats
        date_formats = ["%d.%m.%Y", "%Y-%m-%d", "%m/%d/%Y", "%d-%m-%Y"]
        for fmt in date_formats:
            try:
                date_obj = datetime.strptime(str(date_str), fmt)
                day = date_obj.day
                month = date_obj.strftime("%B")
                year = date_obj.year
                ordinal_day = add_ordinal_suffix(day)
                return f"{ordinal_day} {month} {year}"
            except ValueError:
                continue
        
        return str(date_str)
    except Exception as e:
        print(f"Error formatting date: {e}")
        return str(date_str)

def format_audit_date_period(start_date_str, end_date_str):
    """Format audit date period"""
    try:
        if not start_date_str or not end_date_str:
            return ""
        
        start_formatted = convert_to_dd_mm_yyyy(start_date_str)
        end_formatted = convert_to_dd_mm_yyyy(end_date_str)
        
        return f"{start_formatted} - {end_formatted}"
    except Exception as e:
        print(f"Error formatting audit date period: {e}")
        return ""

def generate_followup_vapt_timeline(start_date_str, end_date_str):
    """
    Generate Follow-Up Audit VAPT timeline with 2 phases
    - Phase 1: 85-90% of time (VAPT Execution)
    - Phase 2: 10-15% of time (Reporting)
    """
    
    BANK_HOLIDAYS = [
        "01.01.2025",  # New Year's Day
        "15.08.2025",  # Independence Day
        "02.10.2025",  # Gandhi Jayanti
    ]
    
    def is_working_day(date):
        date_str = date.strftime("%d.%m.%Y")
        return date.weekday() < 5 and date_str not in BANK_HOLIDAYS
    
    # Parse dates
    try:
        start_date = datetime.strptime(convert_to_dd_mm_yyyy(start_date_str), "%d.%m.%Y")
        end_date = datetime.strptime(convert_to_dd_mm_yyyy(end_date_str), "%d.%m.%Y")
    except Exception as e:
        print(f"Error parsing dates: {e}")
        return []
    
    # Collect working days
    all_dates = [d for d in (start_date + timedelta(days=i) for i in range((end_date - start_date).days + 1)) if is_working_day(d)]
    total_working_days = len(all_dates)
    
    if total_working_days == 0:
        return []
    
    timeline = []
    date_index = 0

    if total_working_days < 2:
        # Less than 2 days: assign all days to Phase 1
        if total_working_days > 0:
            assigned_day = all_dates[0]
            start = assigned_day.strftime("%d.%m.%Y")
            end = assigned_day.strftime("%d.%m.%Y")
            timeline.append({
                "Phase": "VAPT Execution",
                "Dates": f"{start} - {end}",
                "Start": start,
                "End": end
            })
        return timeline
    
    # Calculate days for each phase
    # Phase 1: 85-90% of time (using 87.5% as middle)
    # Phase 2: 10-15% of time (using 12.5% as middle)
    phase1_percentage = 0.875  # 87.5% for Phase 1
    phase2_percentage = 0.125  # 12.5% for Phase 2
    
    phase1_days = max(1, int(total_working_days * phase1_percentage))
    phase2_days = total_working_days - phase1_days
    
    # Ensure Phase 2 has at least 1 day if we have enough total days
    if phase2_days < 1 and total_working_days > 1:
        phase1_days = total_working_days - 1
        phase2_days = 1
    
    # Phase 1: VAPT Execution (85-90% of time)
    if phase1_days > 0 and date_index < len(all_dates):
        end_index = min(date_index + phase1_days, len(all_dates))
        phase_dates = all_dates[date_index:end_index]
        if phase_dates:
            start = phase_dates[0].strftime("%d.%m.%Y")
            end = phase_dates[-1].strftime("%d.%m.%Y")
            timeline.append({
                "Phase": "VAPT Execution",
                "Dates": f"{start} - {end}",
                "Start": start,
                "End": end
            })
            date_index = end_index
    
    # Phase 2: Reporting (10-15% of time)
    if phase2_days > 0 and date_index < len(all_dates):
        end_index = min(date_index + phase2_days, len(all_dates))
        phase_dates = all_dates[date_index:end_index]
        if phase_dates:
            start = phase_dates[0].strftime("%d.%m.%Y")
            end = phase_dates[-1].strftime("%d.%m.%Y")
            timeline.append({
                "Phase": "Reporting",
                "Dates": f"{start} - {end}",
                "Start": start,
                "End": end
            })
            date_index = end_index

    return timeline

def get_followup_vulnerability_counts(excel_file):
    """
    Extract vulnerability counts for both First Audit and Follow-up Audit from Meta_Data worksheet
    Returns: (categories, first_audit_values, followup_audit_values)
    """
    try:
        # Load the workbook
        wb = load_workbook(excel_file)
        
        # Check if 'Meta_Data' worksheet exists
        if 'Meta_Data' not in wb.sheetnames:
            print("❌ 'Meta_Data' worksheet not found")
            return None, None, None
        
        sheet = wb['Meta_Data']
        
        # Initialize data structures
        categories = ['Critical', 'High', 'Medium', 'Low']
        first_audit_counts = {'Critical': 0, 'High': 0, 'Medium': 0, 'Low': 0}
        followup_audit_counts = {'Critical': 0, 'High': 0, 'Medium': 0, 'Low': 0}
        
        # Search for "FIRST AUDIT" and "FOLLOW UP AUDIT" sections
        first_audit_found = False
        followup_audit_found = False
        
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value is not None:
                    cell_value = str(cell.value).strip().upper()
                    
                    # Check for "FIRST AUDIT" section
                    if "FIRST AUDIT" in cell_value and not "FOLLOW" in cell_value:
                        first_audit_found = True
                        print(f"📊 Found 'FIRST AUDIT' at row {cell.row}")
                        # Read the next 4 rows for Critical, High, Medium, Low
                        for i, severity in enumerate(categories, start=1):
                            severity_cell = sheet.cell(row=cell.row + i, column=cell.column)
                            count_cell = sheet.cell(row=cell.row + i, column=cell.column + 1)
                            
                            if severity_cell.value and str(severity_cell.value).strip() == severity:
                                try:
                                    count = int(count_cell.value) if count_cell.value else 0
                                    first_audit_counts[severity] = count
                                    print(f"  {severity}: {count}")
                                except (ValueError, TypeError):
                                    print(f"  ⚠️ Invalid count for {severity}: {count_cell.value}")
                    
                    # Check for "FOLLOW UP AUDIT" section
                    if "FOLLOW UP AUDIT" in cell_value or "FOLLOW-UP AUDIT" in cell_value:
                        followup_audit_found = True
                        print(f"📊 Found 'FOLLOW UP AUDIT' at row {cell.row}")
                        # Read the next 4 rows for Critical, High, Medium, Low
                        for i, severity in enumerate(categories, start=1):
                            severity_cell = sheet.cell(row=cell.row + i, column=cell.column)
                            count_cell = sheet.cell(row=cell.row + i, column=cell.column + 1)
                            
                            if severity_cell.value and str(severity_cell.value).strip() == severity:
                                try:
                                    count = int(count_cell.value) if count_cell.value else 0
                                    followup_audit_counts[severity] = count
                                    print(f"  {severity}: {count}")
                                except (ValueError, TypeError):
                                    print(f"  ⚠️ Invalid count for {severity}: {count_cell.value}")
        
        wb.close()
        
        if not first_audit_found:
            print("⚠️ 'FIRST AUDIT' section not found in Meta_Data")
        if not followup_audit_found:
            print("⚠️ 'FOLLOW UP AUDIT' section not found in Meta_Data")
        
        # Convert to lists in order
        first_audit_values = [first_audit_counts[cat] for cat in categories]
        followup_audit_values = [followup_audit_counts[cat] for cat in categories]
        
        print(f"📊 First Audit counts: {first_audit_values}")
        print(f"📊 Follow-up Audit counts: {followup_audit_values}")
        
        return categories, first_audit_values, followup_audit_values
        
    except Exception as e:
        print(f"❌ Error extracting vulnerability counts: {e}")
        traceback.print_exc()
        return None, None, None

def update_followup_chart_in_docx(docx_path, excel_file, chart_file=None):
    """Update the chart in the Follow-Up Audit DOCX file with two data series"""
    try:
        # Step 1: Get vulnerability data from Excel
        categories, first_audit_values, followup_audit_values = get_followup_vulnerability_counts(excel_file)
        if not categories or first_audit_values is None or followup_audit_values is None:
            print("❌ Failed to extract vulnerability data from Excel")
            return False

        print(f"📊 First Audit data: {dict(zip(categories, first_audit_values))}")
        print(f"📊 Follow-up Audit data: {dict(zip(categories, followup_audit_values))}")

        # Create a temporary copy to work with
        temp_dir = tempfile.mkdtemp()
        temp_docx = os.path.join(temp_dir, "temp.docx")
        shutil.copy2(docx_path, temp_docx)

        # Step 2: Read the DOCX as a ZIP file and find chart files
        with zipfile.ZipFile(temp_docx, 'r') as zin:
            file_list = zin.namelist()
            
            # Find chart files if not specified
            if chart_file is None:
                chart_files = [f for f in file_list if f.startswith('word/charts/') and f.endswith('.xml')]
                if not chart_files:
                    print("❌ No chart files found in Word document")
                    return False
                chart_file = chart_files[0]  # Use the first chart found
                print(f"📊 Using chart file: {chart_file}")
            
            try:
                with zin.open(chart_file) as chart_file_obj:
                    chart_xml = chart_file_obj.read()
            except KeyError:
                print(f"❌ Chart file not found: {chart_file}")
                return False

        # Step 3: Parse and update the chart XML
        tree = etree.fromstring(chart_xml)

        ns = {
            'c': 'http://schemas.openxmlformats.org/drawingml/2006/chart',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
        }

        # Find all series in the chart (should be 2: First Audit and Follow-up Audit)
        all_series = tree.findall('.//c:ser', namespaces=ns)
        print(f"📊 Found {len(all_series)} data series in chart")

        # Update each series with corresponding data
        for series_idx, series in enumerate(all_series[:2]):  # Limit to first 2 series
            if series_idx == 0:
                # First series: First Audit data
                values_to_use = first_audit_values
                series_name = "First Audit"
            elif series_idx == 1:
                # Second series: Follow-up Audit data
                values_to_use = followup_audit_values
                series_name = "Follow-up Audit"
            else:
                continue
            
            print(f"📊 Updating series {series_idx + 1} ({series_name}): {values_to_use}")
            
            # Find value points in this series
            val_pts = series.findall('.//c:val//c:numCache//c:pt', namespaces=ns)
            
            for i, val_pt in enumerate(val_pts):
                if i < len(values_to_use):
                    val_elem = val_pt.find('c:v', namespaces=ns)
                    if val_elem is not None:
                        val_elem.text = str(values_to_use[i])
                        print(f"  Updated {series_name} data point {i} ({categories[i]}): {values_to_use[i]}")

        # Y-axis scaling logic
        all_values = first_audit_values + followup_audit_values
        max_value = max(all_values) if all_values else 1
        
        if max_value == 0:
            max_value = 1
            print("📊 All values are 0, setting max to 1 for chart display")
        
        # Determine Y-axis max and interval
        if max_value <= 10:
            max_axis = 10
            interval = 1
        elif max_value <= 20:
            max_axis = 20
            interval = 2
        elif max_value <= 40:
            max_axis = 40
            interval = 4
        elif max_value <= 80:
            max_axis = 80
            interval = 8
        elif max_value <= 120:
            max_axis = 120
            interval = 12
        else:
            max_axis = 1
            while max_axis < max_value:
                max_axis *= 2
            interval = max_axis // 4

        # Update Y-axis scaling
        val_ax = tree.find('.//c:valAx', namespaces=ns)
        if val_ax is not None:
            scaling_elem = val_ax.find('.//c:scaling', namespaces=ns)
            if scaling_elem is not None:
                max_elem = scaling_elem.find('.//c:max', namespaces=ns)
                if max_elem is None:
                    max_elem = etree.SubElement(scaling_elem, '{http://schemas.openxmlformats.org/drawingml/2006/chart}max')
                max_elem.set('val', str(max_axis))
            
            # Update major unit
            major_unit_elem = val_ax.find('.//c:majorUnit', namespaces=ns)
            if major_unit_elem is None:
                major_unit_elem = etree.SubElement(val_ax, '{http://schemas.openxmlformats.org/drawingml/2006/chart}majorUnit')
            if major_unit_elem is not None:
                major_unit_elem.set('val', str(interval))

        # Serialize and update the DOCX
        updated_chart_xml = etree.tostring(tree, pretty_print=True, xml_declaration=True, encoding='UTF-8')

        with zipfile.ZipFile(temp_docx, 'w') as zout:
            for file in file_list:
                if file == chart_file:
                    zout.writestr(file, updated_chart_xml)
                else:
                    with zipfile.ZipFile(docx_path, 'r') as zin:
                        with zin.open(file) as src_file:
                            zout.writestr(file, src_file.read())

        shutil.move(temp_docx, docx_path)
        shutil.rmtree(temp_dir)

        print(f"✅ Chart updated successfully in: {docx_path}")
        print(f"📈 Y-axis scaled to max: {max_axis}, interval: {interval}")
        return True

    except Exception as e:
        print(f"❌ Error updating chart: {str(e)}")
        traceback.print_exc()
        if 'temp_dir' in locals():
            shutil.rmtree(temp_dir, ignore_errors=True)
        return False

def replace_followup_timeline_placeholders(doc, timeline):
    """
    Replace timeline placeholders in the Word document with generated Follow-Up Audit timeline data
    """
    try:
        # Create mapping of placeholders to phase names (2 phases only)
        phase_placeholders = {
            "VAPT Execution": {
                "start": "11.11.1111",
                "end": "22.22.2222"
            },
            "Reporting": {
                "start": "55.55.5555",
                "end": "66.66.6666"
            }
        }

        # Extract dates from timeline
        phase_dates = {}
        for phase in timeline:
            phase_name = phase['Phase']
            phase_dates[phase_name] = {
                'start': phase['Start'],
                'end': phase['End']
            }

        # Create replacements dictionary
        timeline_replacements = {}
        for phase_name, placeholders in phase_placeholders.items():
            if phase_name in phase_dates:
                timeline_replacements[placeholders['start']] = phase_dates[phase_name]['start']
                timeline_replacements[placeholders['end']] = phase_dates[phase_name]['end']
            else:
                # Use default values if phase not found in timeline
                timeline_replacements[placeholders['start']] = placeholders['start']
                timeline_replacements[placeholders['end']] = placeholders['end']

        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for old_text, new_text in timeline_replacements.items():
                if old_text in paragraph.text:
                    _replace_text_in_runs(paragraph, old_text, new_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in timeline_replacements.items():
                            if old_text in paragraph.text:
                                _replace_text_in_runs(paragraph, old_text, new_text)
        
        print("✅ Follow-Up Audit timeline placeholders replaced successfully")
        
    except Exception as e:
        print(f"❌ Error replacing timeline placeholders: {str(e)}")
        traceback.print_exc()

def find_and_print_metadata(file_path):
    """
    Extract metadata from Excel file's Meta_Data worksheet
    """
    try:
        # Load the workbook
        workbook = openpyxl.load_workbook(file_path)
        
        # Check if 'Meta_Data' worksheet exists
        if 'Meta_Data' not in workbook.sheetnames:
            print("Error: 'Meta_Data' worksheet not found!")
            return {}
        
        # Get the Meta_Data worksheet
        sheet = workbook['Meta_Data']
        
        # Dictionary to store the values we're looking for
        target_values = {
            "Organization Name": None,
            "City": None,
            "State": None,
            "Start Date": None,
            "End Date": None,
            "First Audit Report ID": None,
            "First Audit Report Date": None
        }
        
        # Variables to store the special values
        report_prepared_by_value = None
        auditee_details_value1 = None  # First value (diagonal cell)
        auditee_details_value2 = None  # Second value (two steps down and right)
        bank_email_addresses = []      # List to store all email addresses
        auditing_team_members = []     # List to store all auditing team member data
        
        # Iterate through all cells in the worksheet
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value is not None:
                    # Check if cell value matches any of our target keys
                    for key in target_values.keys():
                        if str(cell.value).strip() == key:
                            # Get the adjacent cell (next column, same row)
                            adjacent_cell = sheet.cell(row=cell.row, column=cell.column + 1)
                            if adjacent_cell.value is not None:
                                target_values[key] = adjacent_cell.value
                            break
                    
                    # Special handling for "REPORT PREPARED BY" (separate from other keys)
                    if str(cell.value).strip() == "REPORT PREPARED BY":
                        # Get the cell below and to the right (next column, next row)
                        diagonal_cell = sheet.cell(row=cell.row + 1, column=cell.column + 1)
                        if diagonal_cell.value is not None:
                            report_prepared_by_value = diagonal_cell.value
                    
                    # Special handling for "AUDITEE DETAILS" section
                    if str(cell.value).strip() == "AUDITEE DETAILS":
                        # Get the first value: cell below and to the right (next column, next row)
                        diagonal_cell1 = sheet.cell(row=cell.row + 1, column=cell.column + 1)
                        if diagonal_cell1.value is not None:
                            auditee_details_value1 = diagonal_cell1.value
                        
                        # Get the second value: two steps down and to the right (next column, row+2)
                        diagonal_cell2 = sheet.cell(row=cell.row + 2, column=cell.column + 1)
                        if diagonal_cell2.value is not None:
                            auditee_details_value2 = diagonal_cell2.value
                    
                    # Special handling for "BANK EMAIL ADDRESSES" section
                    if str(cell.value).strip() == "BANK EMAIL ADDRESSES":
                        # Start from the cell below and to the right (next column, next row)
                        current_row = cell.row + 1
                        current_col = cell.column + 1
                        
                        # Keep reading cells vertically until an empty cell is found
                        while True:
                            email_cell = sheet.cell(row=current_row, column=current_col)
                            if email_cell.value is not None and str(email_cell.value).strip() != "":
                                bank_email_addresses.append(str(email_cell.value).strip())
                                current_row += 1  # Move to next row
                            else:
                                break  # Stop when empty cell is found
                    
                    # Special handling for "AUDITING TEAM MEMBER" entries
                    if "AUDITING TEAM MEMBER" in str(cell.value).strip():
                        member_data = {}
                        member_number = str(cell.value).strip().split()[-1]  # Get the number (1, 2, 3, etc.)
                        
                        # Define the labels for each row
                        labels = [
                            f"Team Member {member_number} - Name",
                            f"Team Member {member_number} - Designation", 
                            f"Team Member {member_number} - Email",
                            f"Team Member {member_number} - Qualification",
                            f"Team Member {member_number} - Certified"
                        ]
                        
                        # Start from the cell below and to the right (next column, next row)
                        current_row = cell.row + 1
                        current_col = cell.column + 1
                        label_index = 0
                        
                        # Read data for each label
                        for label in labels:
                            if current_row <= sheet.max_row and label_index < len(labels):
                                # Get the value from the cell below and to the right
                                value_cell = sheet.cell(row=current_row, column=current_col)
                                if value_cell.value is not None:
                                    member_data[label] = str(value_cell.value).strip()
                                
                                # Move to next row for next label
                                current_row += 1
                                label_index += 1
                        
                        # Only add if we found some data
                        if member_data:
                            auditing_team_members.append(member_data)
        
        # Create the metadata dictionary
        metadata = {
            "organization_name": target_values.get("Organization Name", ""),
            "city": target_values.get("City", ""),
            "state": target_values.get("State", ""),
            "start_date": target_values.get("Start Date", ""),
            "end_date": target_values.get("End Date", ""),
            "first_audit_report_id": target_values.get("First Audit Report ID", ""),
            "first_audit_report_date": target_values.get("First Audit Report Date", ""),
            "report_prepared_by": report_prepared_by_value or "",
            "auditee_details_1": auditee_details_value1 or "",
            "auditee_details_2": auditee_details_value2 or "",
            "bank_emails": bank_email_addresses,
            "auditing_team": auditing_team_members
        }
        
        return metadata
        
    except Exception as e:
        print(f"Error extracting metadata: {e}")
        return {}

def _replace_text_in_runs(paragraph, old_text, new_text):
    """Replace text in paragraph runs while preserving formatting"""
    if old_text not in paragraph.text:
        return
    
    # Get all runs in the paragraph
    runs = paragraph.runs
    
    # Find runs that contain the old text
    for i, run in enumerate(runs):
        if old_text in run.text:
            # Replace the text in this run
            run.text = run.text.replace(old_text, new_text)
            break
    else:
        # If old_text spans multiple runs, we need to handle it differently
        full_text = paragraph.text
        if old_text in full_text:
            # Clear all runs and create a new one with the replaced text
            for run in runs:
                run.clear()
            if runs:
                runs[0].text = full_text.replace(old_text, new_text)

def _replace_text_in_runs_with_bold(paragraph, old_text, new_text, make_bold=True):
    """Replace text in paragraph runs and make only the replacement text bold"""
    if old_text not in paragraph.text:
        return
    
    # Get all runs in the paragraph
    runs = paragraph.runs
    
    # Find runs that contain the old text
    for i, run in enumerate(runs):
        if old_text in run.text:
            # Replace the text and make it bold with Calibri font
            run.text = run.text.replace(old_text, new_text)
            run.bold = make_bold
            run.font.name = 'Calibri'
            break
    else:
        # If old_text spans multiple runs, handle it more carefully
        full_text = paragraph.text
        if old_text in full_text:
            # Find the position of old_text in the full text
            start_pos = full_text.find(old_text)
            end_pos = start_pos + len(old_text)
            
            # Split the text into parts
            before_text = full_text[:start_pos]
            after_text = full_text[end_pos:]
            
            # Clear all runs
            for run in runs:
                run.clear()
            
            # Rebuild the paragraph with proper positioning
            if before_text:
                runs[0].text = before_text
                runs[0].bold = False
                runs[0].font.name = 'Calibri'
            
            # Add the bold replacement text
            if before_text:
                # Add new run for the bold text
                bold_run = paragraph.add_run(new_text)
                bold_run.bold = make_bold
                bold_run.font.name = 'Calibri'
            else:
                # Use first run for bold text
                runs[0].text = new_text
                runs[0].bold = make_bold
                runs[0].font.name = 'Calibri'
            
            # Add text after the replacement
            if after_text:
                if before_text:
                    # Add new run for text after
                    after_run = paragraph.add_run(after_text)
                    after_run.bold = False
                    after_run.font.name = 'Calibri'
                else:
                    # Add to second run if it exists, otherwise create new run
                    if len(runs) > 1:
                        runs[1].text = after_text
                        runs[1].bold = False
                        runs[1].font.name = 'Calibri'
                    else:
                        after_run = paragraph.add_run(after_text)
                        after_run.bold = False
                        after_run.font.name = 'Calibri'

def replace_metadata_placeholders(doc, metadata):
    """Replace metadata placeholders in the document"""
    try:
        # Prepare bank email addresses as a single string
        bank_emails = ""
        if metadata.get("bank_emails"):
            bank_emails = "\n".join(metadata.get("bank_emails", []))
        
        # Format dates
        month_year = format_date_for_month_year(metadata.get("end_date", ""))
        date_range = format_date_for_range(metadata.get("start_date", ""), metadata.get("end_date", ""))
        
        # Define placeholder mappings
        replacements = {
            "ORGANIZATION_NAMEE": metadata.get("organization_name", ""),
            "Organization_Namee": metadata.get("organization_name", ""),
            "CITYY": metadata.get("city", ""),
            "Stateee": metadata.get("state", ""),
            "First_Audit_Report_Idd": metadata.get("first_audit_report_id", ""),
            "First_Audit_Report_Datee": metadata.get("first_audit_report_date", ""),
            "Audit_Date_Period": format_audit_date_period(metadata.get('start_date', ''), metadata.get('end_date', '')),
            "Maker_Name_R": metadata.get("report_prepared_by", ""),
            "Organization_Personn": metadata.get("auditee_details_1", ""),
            "Designationn": metadata.get("auditee_details_2", ""),
            "Auditee_Email_Adresss": bank_emails,
            "Monthh Yearr": month_year,
            "00rd Month Year to 00th Month Year": date_range
        }
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    # Special handling for date_range to make it bold
                    if old_text == "00rd Month Year to 00th Month Year":
                        _replace_text_in_runs_with_bold(paragraph, old_text, new_text, make_bold=True)
                    else:
                        _replace_text_in_runs(paragraph, old_text, new_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            if old_text in paragraph.text:
                                # Special handling for date_range to make it bold
                                if old_text == "00rd Month Year to 00th Month Year":
                                    _replace_text_in_runs_with_bold(paragraph, old_text, new_text, make_bold=True)
                                else:
                                    _replace_text_in_runs(paragraph, old_text, new_text)
        
        
        print("✅ Metadata placeholders replaced successfully")
        
    except Exception as e:
        print(f"❌ Error replacing metadata placeholders: {e}")
        raise

def read_scope_worksheet_data(excel_file):
    """
    Read data from the Scope worksheet in Excel file
    Handles cases where branches may have multiple IP addresses
    
    Args:
        excel_file (str): Path to Excel file
        
    Returns:
        list: List of dictionaries containing branch names and IP addresses
    """
    try:
        # Load Excel workbook
        wb = load_workbook(excel_file)
        
        # Check if 'Scope' worksheet exists
        if 'Scope' not in wb.sheetnames:
            print("Warning: 'Scope' worksheet not found in Excel file")
            return []
        
        # Get the Scope worksheet
        ws = wb['Scope']
        
        scope_data = []
        current_branch = None
        
        # Read data from the worksheet (assuming data starts from row 2, with headers in row 1)
        for row in ws.iter_rows(min_row=2, values_only=True):
            sr_no = row[0] if row[0] is not None else ""
            branch_name = str(row[1]).strip() if row[1] else ""
            host_ip = str(row[2]).strip() if row[2] else ""
            
            # If we have both branch name and IP, add to scope data
            if branch_name and host_ip:
                current_branch = branch_name
                scope_data.append({
                    'branch_name': branch_name,
                    'host_ip': host_ip
                })
            # If branch name is empty but we have an IP and a current branch, it's an additional IP
            elif not branch_name and host_ip and current_branch:
                scope_data.append({
                    'branch_name': current_branch,
                    'host_ip': host_ip
                })
        
        wb.close()
        print(f"✅ Read {len(scope_data)} entries from Scope worksheet")
        return scope_data
        
    except Exception as e:
        print(f"❌ Error reading Scope worksheet: {e}")
        return []

def replace_scope_placeholders_with_data(doc, excel_file):
    """
    Find and replace 'Branch Namee or Server Namee' and 'IP Addressesss' placeholders
    with data from the Scope worksheet, creating new rows for each entry
    """
    try:
        # Read scope data from Excel
        scope_data = read_scope_worksheet_data(excel_file)
        
        if not scope_data:
            print("No scope data found, skipping placeholder replacement")
            return
        
        # Find tables containing the placeholders
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        if "Branch Namee or Server Namee" in paragraph.text or "IP Addressesss" in paragraph.text:
                            # Found the target row, now replace and add new rows
                            replace_scope_placeholders_in_table(table, row_idx, cell_idx, scope_data)
                            return
        
        print("⚠️ Scope placeholders not found in any table")
        
    except Exception as e:
        print(f"❌ Error replacing scope placeholders: {e}")
        traceback.print_exc()

def replace_scope_placeholders_in_table(table, target_row_idx, target_cell_idx, scope_data):
    """
    Replace scope placeholders in a specific table row and add new rows for additional data
    Groups IP addresses by branch name for better organization
    """
    try:
        target_row = table.rows[target_row_idx]
        
        # Group scope data by branch name
        grouped_data = {}
        for entry in scope_data:
            branch_name = entry['branch_name']
            if branch_name not in grouped_data:
                grouped_data[branch_name] = []
            grouped_data[branch_name].append(entry['host_ip'])
        
        # Convert to list of entries with combined IP addresses
        processed_data = []
        for branch_name, ip_list in grouped_data.items():
            # Combine multiple IPs with line breaks
            combined_ips = '\n'.join(ip_list)
            processed_data.append({
                'branch_name': branch_name,
                'host_ip': combined_ips
            })
        
        # Replace placeholders in the first row with first processed entry
        if processed_data:
            first_entry = processed_data[0]
            
            # Replace Branch Namee or Server Namee
            for paragraph in target_row.cells[target_cell_idx].paragraphs:
                if "Branch Namee or Server Namee" in paragraph.text:
                    _replace_text_in_runs(paragraph, "Branch Namee or Server Namee", first_entry['branch_name'])
            
            # Replace IP Addressesss (two columns after Branch Name column - A to C)
            ip_cell_idx = target_cell_idx + 2
            if ip_cell_idx < len(target_row.cells):
                for paragraph in target_row.cells[ip_cell_idx].paragraphs:
                    if "IP Addressesss" in paragraph.text:
                        _replace_text_in_runs(paragraph, "IP Addressesss", first_entry['host_ip'])
            
            # Add borders to the target row
            add_borders_to_row(target_row)
        
        # Add new rows for remaining processed entries
        for i in range(1, len(processed_data)):
            entry = processed_data[i]
            
            # Create a new row by copying the target row structure
            new_row = table.add_row()
            
            # Copy the structure and content from the target row
            for j, cell in enumerate(new_row.cells):
                if j < len(target_row.cells):
                    # Clear the new cell
                    cell._element.clear()
                    
                    # Copy formatting from target cell
                    target_cell = target_row.cells[j]
                    
                    # Determine content based on column position
                    if j == target_cell_idx:  # Branch Name column
                        new_content = entry['branch_name']
                    elif j == target_cell_idx + 2:  # IP Address column (two columns after Branch Name - A to C)
                        new_content = entry['host_ip']
                    elif j == 0:  # Sr.No column
                        new_content = str(i + 1)  # Continue numbering
                    else:
                        # Keep original content for other columns
                        new_content = target_cell.text
                    
                    # Add content with proper formatting
                    if new_content:
                        new_paragraph = cell.add_paragraph()
                        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Set consistent formatting
                        new_run = new_paragraph.add_run(new_content)
                        new_run.font.name = 'Calibri (Body)'
                        new_run.font.size = Pt(12)
                        new_run.font.bold = False
                        new_run.font.italic = False
                        new_run.font.underline = False
                        
                        # Set cell vertical alignment
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Add borders to the new row
            add_borders_to_row(new_row)
        
        print(f"✅ Replaced scope placeholders and added {len(processed_data)} rows of data")
        
    except Exception as e:
        print(f"❌ Error replacing scope placeholders in table: {e}")
        traceback.print_exc()

def replace_auditor_placeholders_and_add_rows(doc, metadata):
    """Replace auditor placeholders and add additional rows for multiple team members"""
    try:
        auditing_team_members = metadata.get("auditing_team", [])
        
        print(f"🔍 Debug: Found {len(auditing_team_members)} auditing team members")
        for i, member in enumerate(auditing_team_members):
            print(f"  Member {i+1}: {member}")
        
        if not auditing_team_members:
            print("No auditing team members found")
            return
        
        # Find the table containing auditor placeholders
        target_table = None
        print(f"🔍 Debug: Searching through {len(doc.tables)} tables for auditor placeholders...")
        
        for table_idx, table in enumerate(doc.tables):
            print(f"  Checking table {table_idx + 1} with {len(table.rows)} rows...")
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    if "Auditorrrrr" in cell.text:
                        print(f"    Found 'Auditorrrrr' in table {table_idx + 1}, row {row_idx + 1}, cell {cell_idx + 1}")
                        target_table = table
                        break
                if target_table:
                    break
            if target_table:
                break
        
        if not target_table:
            print("❌ Table with auditor placeholders not found")
            print("Available table contents:")
            for table_idx, table in enumerate(doc.tables):
                print(f"  Table {table_idx + 1}:")
                for row_idx, row in enumerate(table.rows):
                    row_text = " | ".join([cell.text[:50] for cell in row.cells])
                    print(f"    Row {row_idx + 1}: {row_text}")
            return
        
        # Find the row with auditor placeholders
        target_row = None
        column_mapping = {}  # Store which column contains which placeholder
        
        print(f"🔍 Debug: Searching for auditor placeholders in target table...")
        for row_idx, row in enumerate(target_table.rows):
            # Check if this row contains auditor placeholders
            has_auditor_placeholders = False
            for cell_idx, cell in enumerate(row.cells):
                if "Auditorrrrr" in cell.text:
                    print(f"    Found 'Auditorrrrr' in row {row_idx + 1}, cell {cell_idx + 1}")
                    has_auditor_placeholders = True
                    target_row = row
                    break
            
            if has_auditor_placeholders:
                print(f"    Mapping columns in row {row_idx + 1}...")
                # Map all columns in this row
                for j, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        if "Auditorrrrr_Desi" in paragraph.text:
                            column_mapping[j] = "designation"
                            print(f"      Column {j}: designation")
                        elif "Auditorrrrr_email" in paragraph.text:
                            column_mapping[j] = "email"
                            print(f"      Column {j}: email")
                        elif "Auditorrrrr_Qua" in paragraph.text:
                            column_mapping[j] = "qualification"
                            print(f"      Column {j}: qualification")
                        elif "Auditorrrrr_Cert" in paragraph.text:
                            column_mapping[j] = "certified"
                            print(f"      Column {j}: certified")
                        elif "Auditorrrrr" in paragraph.text and "Desi" not in paragraph.text and "email" not in paragraph.text and "Qua" not in paragraph.text and "Cert" not in paragraph.text:
                            column_mapping[j] = "name"
                            print(f"      Column {j}: name")
                break
        
        if not target_row:
            print("❌ Row with auditor placeholders not found")
            return
        
        print(f"📋 Column mapping detected: {column_mapping}")
        
        # Replace placeholders in the first row with fixed values
        fixed_auditor_data = {
            "name": "Niraj Goyal",
            "designation": "CEO / Director", 
            "email": "admin@ngtech.co.in",
            "qualification": "CA, CEH, DISA",
            "certified": "Yes"
        }
        
        for j, cell in enumerate(target_row.cells):
            if j in column_mapping:
                column_type = column_mapping[j]
                for paragraph in cell.paragraphs:
                    # Replace Team Member 1 placeholders with fixed values
                    if column_type == "name":
                        _replace_text_in_runs(paragraph, "Auditorrrrr", fixed_auditor_data["name"])
                    elif column_type == "designation":
                        _replace_text_in_runs(paragraph, "Auditorrrrr_Desi", fixed_auditor_data["designation"])
                    elif column_type == "email":
                        _replace_text_in_runs(paragraph, "Auditorrrrr_email", fixed_auditor_data["email"])
                    elif column_type == "qualification":
                        _replace_text_in_runs(paragraph, "Auditorrrrr_Qua", fixed_auditor_data["qualification"])
                    elif column_type == "certified":
                        _replace_text_in_runs(paragraph, "Auditorrrrr_Cert", fixed_auditor_data["certified"])
        
        # Add additional rows for Team Member 2, 3, etc. (limit to 8 additional rows)
        # Since first row is fixed, we start from index 0 of auditing_team_members for Team Member 2
        max_additional_rows = min(8, len(auditing_team_members))
        
        for i in range(max_additional_rows):
            member_data = auditing_team_members[i]
            member_number = i + 2  # Start from Team Member 2
            
            print(f"🔍 Processing Team Member {member_number} (Excel index {i})")
            print(f"📊 Available keys in member_data: {list(member_data.keys())}")
            
            # Create a new row by copying the target row
            new_row = target_table.add_row()
            
            # Copy the structure from the target row
            for j, cell in enumerate(new_row.cells):
                if j < len(target_row.cells):
                    # Safely clear the new cell content
                    try:
                        # Remove all existing paragraphs to avoid extra spacing
                        while len(cell.paragraphs) > 0:
                            p = cell.paragraphs[0]
                            p._element.getparent().remove(p._element)
                    except Exception as e:
                        print(f"Warning: Could not clear cell content: {e}")
                        # Fallback: clear the entire cell element
                        cell._element.clear()
                    
                    # Copy formatting and add content
                    original_cell = target_row.cells[j]
                    
                    # Determine content based on column mapping
                    new_content = ""
                    if j in column_mapping:
                        column_type = column_mapping[j]
                        # The Excel data is stored with the actual member number from Excel (1, 2, 3, etc.)
                        # We need to find the correct key in the member_data
                        excel_member_number = None
                        for key in member_data.keys():
                            if key.startswith("Team Member ") and key.endswith(" - Name"):
                                # Extract the number from the key
                                try:
                                    excel_member_number = key.split("Team Member ")[1].split(" - Name")[0]
                                    break
                                except:
                                    continue
                        
                        if excel_member_number:
                            if column_type == "name":
                                new_content = member_data.get(f"Team Member {excel_member_number} - Name", "")
                            elif column_type == "designation":
                                new_content = member_data.get(f"Team Member {excel_member_number} - Designation", "")
                            elif column_type == "email":
                                new_content = member_data.get(f"Team Member {excel_member_number} - Email", "")
                            elif column_type == "qualification":
                                new_content = member_data.get(f"Team Member {excel_member_number} - Qualification", "")
                            elif column_type == "certified":
                                certified_value = member_data.get(f"Team Member {excel_member_number} - Certified", "")
                                # Convert to proper case (Yes/No)
                                if certified_value.lower() == "yes":
                                    new_content = "Yes"
                                elif certified_value.lower() == "no":
                                    new_content = "No"
                                else:
                                    new_content = certified_value  # Keep original if not yes/no
                            
                            print(f"📝 {column_type}: '{new_content}' (from key: Team Member {excel_member_number} - {column_type.title()})")
                        else:
                            print(f"⚠️  No Excel member number found for Team Member {member_number}")
                    else:
                        # Check if this is the first column (Sr. No. column)
                        if j == 0 and original_cell.text.strip().isdigit():
                            # This is the Sr. No. column, update the number
                            new_content = str(member_number)
                        else:
                            # Keep original text for other non-auditor columns
                            new_content = original_cell.text
                    
                    # Add content with clean formatting
                    if new_content:
                        new_paragraph = cell.add_paragraph()
                        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Remove any extra spacing - set to exactly 0
                        new_paragraph.paragraph_format.space_before = Pt(0)
                        new_paragraph.paragraph_format.space_after = Pt(0)
                        new_paragraph.paragraph_format.line_spacing = 1.0
                        new_paragraph.paragraph_format.first_line_indent = Pt(0)
                        new_paragraph.paragraph_format.left_indent = Pt(0)
                        new_paragraph.paragraph_format.right_indent = Pt(0)
                        
                        # Ensure no extra spacing is applied
                        new_paragraph.paragraph_format.space_before_auto = False
                        new_paragraph.paragraph_format.space_after_auto = False
                        
                        # Set consistent formatting for all new rows
                        new_run = new_paragraph.add_run(new_content)
                        new_run.font.name = 'Calibri (Body)'
                        new_run.font.size = Pt(12)
                        new_run.font.bold = False
                        new_run.font.italic = False
                        new_run.font.underline = False
                        
                        # Set cell vertical alignment
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        
                        # Set cell padding (top and bottom 1px equivalent)
                        try:
                            from docx.oxml.shared import qn
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            
                            # Set cell margins for padding - use minimal margins to avoid extra spacing
                            from docx.oxml import parse_xml
                            cell_margin_xml = f'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:w="0" w:type="dxa"/><w:bottom w:w="0" w:type="dxa"/></w:tcMar>'
                            cell_margin = parse_xml(cell_margin_xml)
                            tcPr.append(cell_margin)
                        except Exception as e:
                            print(f"Warning: Could not set cell margins: {e}")
                            # Continue without setting margins to avoid corruption
        
        print(f"✅ Added {max_additional_rows + 1} auditing team member(s) to table (max 8 additional rows)")
        
    except Exception as e:
        print(f"❌ Error processing auditing team: {e}")
        traceback.print_exc()
        raise

def extract_followup_poc_images_from_excel(excel_file):
    """
    Extract POC and Old POC images from the Excel file for Follow-Up Audit.
    Returns: (poc_images, old_poc_images, df)
    """
    try:
        # Load workbook and worksheet
        wb = load_workbook(excel_file)
        ws = wb['Infra_VAPT']
        
        # Read the dataframe
        df = pd.read_excel(excel_file, sheet_name='Infra_VAPT')
        
        # Initialize dictionaries to store images
        poc_images = {}  # New POC images (from POC column)
        old_poc_images = {}  # Old POC images (from Old POC column)
        
        # Find column ranges for POC (T-Z) and Old POC (M-S) by looking at merged cells
        poc_col_start = None
        poc_col_end = None
        old_poc_col_start = None
        old_poc_col_end = None
        
        # Check merged cells in row 1 for POC and Old POC headers
        for merged_range in ws.merged_cells.ranges:
            if merged_range.min_row == 1 and merged_range.max_row == 1:
                first_cell = ws.cell(row=1, column=merged_range.min_col)
                if first_cell.value and str(first_cell.value).strip() == "POC":
                    poc_col_start = merged_range.min_col
                    poc_col_end = merged_range.max_col
                    print(f"📸 Found POC columns from {poc_col_start} to {poc_col_end}")
                elif first_cell.value and str(first_cell.value).strip() == "Old POC":
                    old_poc_col_start = merged_range.min_col
                    old_poc_col_end = merged_range.max_col
                    print(f"📸 Found Old POC columns from {old_poc_col_start} to {old_poc_col_end}")
        
        # Extract images from Excel if they exist
        if hasattr(ws, "_images"):
            print(f"📸 Found {len(ws._images)} images in Excel")
            
            for img in ws._images:
                try:
                    # Get the row and column of the image
                    row_idx = img.anchor._from.row + 1  # Excel row (1-based)
                    col_idx = img.anchor._from.col  # Excel column (0-based)
                    
                    # Get image data
                    img_data = img._data() if callable(img._data) else img._data
                    
                    # Get vulnerability name for this row
                    if row_idx <= len(df) + 1:  # +1 for header
                        excel_row_idx = row_idx - 2  # Convert to DataFrame index (0-based, excluding header)
                        if 0 <= excel_row_idx < len(df):
                            vuln_name = str(df.iloc[excel_row_idx]['Name of Vulnerability'])
                            
                            # Determine if this image is in POC (T-Z) or Old POC (M-S) columns
                            if poc_col_start is not None and poc_col_end is not None:
                                # Check if column is within POC range (T-Z)
                                if poc_col_start <= col_idx + 1 <= poc_col_end:
                                    # This is a POC (new) image
                                    if vuln_name not in poc_images:
                                        poc_images[vuln_name] = []
                                    poc_images[vuln_name].append(img_data)
                                    print(f"  ✅ Added POC image for '{vuln_name}' (row {row_idx}, col {col_idx})")
                            
                            if old_poc_col_start is not None and old_poc_col_end is not None:
                                # Check if column is within Old POC range (M-S)
                                if old_poc_col_start <= col_idx + 1 <= old_poc_col_end:
                                    # This is an Old POC image
                                    if vuln_name not in old_poc_images:
                                        old_poc_images[vuln_name] = []
                                    old_poc_images[vuln_name].append(img_data)
                                    print(f"  ✅ Added Old POC image for '{vuln_name}' (row {row_idx}, col {col_idx})")
                
                except Exception as e:
                    print(f"  ⚠️ Error processing image: {e}")
                    continue
        
        wb.close()
        
        print(f"📊 Extracted {len(poc_images)} vulnerabilities with POC images")
        print(f"📊 Extracted {len(old_poc_images)} vulnerabilities with Old POC images")
        
        return poc_images, old_poc_images, df
    
    except Exception as e:
        print(f"❌ Error extracting images: {e}")
        traceback.print_exc()
        return {}, {}, pd.DataFrame()

def add_followup_images_at_end_of_document(doc, poc_images, old_poc_images, df):
    """
    Add POC and Old POC images at the end of the document for Follow-Up Audit.
    """
    try:
        # Check if there are any images to add
        has_poc = any(poc_images.values())
        has_old_poc = any(old_poc_images.values())
        
        if not has_poc and not has_old_poc:
            print("⚠️ No images to add at end of document")
            return
        
        # Add page break and title
        doc.add_page_break()
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run("Proof of Concept (POC) Images")
        title_run.font.size = Pt(16)
        title_run.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        image_count = 0
        
        # Iterate through each vulnerability in the DataFrame
        for idx, row in enumerate(df.iterrows(), start=1):
            index, row_data = row
            vuln_name = str(row_data['Name of Vulnerability'])
            
            # Check if this vulnerability has any images
            has_images = (vuln_name in poc_images and poc_images[vuln_name]) or \
                        (vuln_name in old_poc_images and old_poc_images[vuln_name])
            
            if has_images:
                image_count += 1
                
                # Add vulnerability title
                vuln_title = doc.add_paragraph()
                vuln_run = vuln_title.add_run(f"{image_count}. {vuln_name}")
                vuln_run.font.size = Pt(14)
                vuln_run.bold = True
                
                # Add Old POC images first
                if vuln_name in old_poc_images and old_poc_images[vuln_name]:
                    old_poc_subtitle = doc.add_paragraph()
                    old_poc_subtitle_run = old_poc_subtitle.add_run("Old POC Images:")
                    old_poc_subtitle_run.font.size = Pt(12)
                    old_poc_subtitle_run.bold = True
                    
                    for i, image_data in enumerate(old_poc_images[vuln_name], start=1):
                        try:
                            pil_image = PILImage.open(io.BytesIO(image_data))
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                                pil_image.save(tmp, format='PNG')
                                temp_image_path = tmp.name
                            
                            caption_para = doc.add_paragraph()
                            caption_run = caption_para.add_run(f"Old POC Image {i}")
                            caption_run.italic = True
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            image_para = doc.add_paragraph()
                            image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = image_para.add_run()
                            picture = run.add_picture(temp_image_path, width=Inches(6.0))
                            apply_1pt_border_to_picture(picture)
                            os.unlink(temp_image_path)
                            doc.add_paragraph()
                        except Exception as e:
                            error_para = doc.add_paragraph()
                            error_para.add_run(f"Error loading Old POC image {i}: {str(e)}")
                
                # Add POC (new) images
                if vuln_name in poc_images and poc_images[vuln_name]:
                    poc_subtitle = doc.add_paragraph()
                    poc_subtitle_run = poc_subtitle.add_run("New POC Images:")
                    poc_subtitle_run.font.size = Pt(12)
                    poc_subtitle_run.bold = True
                    
                    for i, image_data in enumerate(poc_images[vuln_name], start=1):
                        try:
                            pil_image = PILImage.open(io.BytesIO(image_data))
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                                pil_image.save(tmp, format='PNG')
                                temp_image_path = tmp.name
                            
                            caption_para = doc.add_paragraph()
                            caption_run = caption_para.add_run(f"New POC Image {i}")
                            caption_run.italic = True
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            image_para = doc.add_paragraph()
                            image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = image_para.add_run()
                            picture = run.add_picture(temp_image_path, width=Inches(6.0))
                            apply_1pt_border_to_picture(picture)
                            os.unlink(temp_image_path)
                            doc.add_paragraph()
                        except Exception as e:
                            error_para = doc.add_paragraph()
                            error_para.add_run(f"Error loading POC image {i}: {str(e)}")
                
                # Add page break between vulnerabilities (except for the last one)
                total_vulns_with_images = sum(1 for v in [poc_images, old_poc_images] 
                                              for name in df['Name of Vulnerability'] 
                                              if str(name) in v and v[str(name)])
                if image_count < total_vulns_with_images:
                    doc.add_page_break()
        
        print(f"✅ Added {image_count} vulnerability sections with images at end of document")
    
    except Exception as e:
        print(f"❌ Error adding images at end of document: {e}")
        traceback.print_exc()

def _replace_text_in_runs(paragraph, old_text, new_text):
    """
    Helper function to replace text in a paragraph while preserving formatting.
    This handles cases where a single placeholder is split across multiple runs.
    """
    # Combine all run texts to a single string for easy replacement
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return False

    # Get the formatting of the runs before we delete them
    run_formats = []
    for run in paragraph.runs:
        # Safely get font color to avoid None reference errors
        font_color = None
        try:
            if run.font.color.rgb is not None:
                font_color = run.font.color.rgb
        except:
            font_color = None
            
        run_formats.append({
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': 'calibri',
            'font_size': run.font.size,
            'font_color': font_color
        })

    # Clear the paragraph safely
    try:
        # Clear runs one by one to avoid corruption
        runs_to_clear = list(paragraph.runs)  # Create a copy to avoid modification during iteration
        for run in runs_to_clear:
            run.clear()
    except Exception as e:
        print(f"Warning: Error clearing paragraph runs: {e}")
        # Fallback: clear all runs at once
        for run in paragraph.runs:
            run.clear()
    
    # Replace the text
    new_full_text = full_text.replace(old_text, new_text)

    # Rebuild the runs with the new text and old formatting
    new_runs_texts = re.split(r'(' + re.escape(new_text) + r')', new_full_text)
    
    # Apply the formatting of the first run to all new segments
    first_run_format = run_formats[0] if run_formats else {}
    
    for segment in new_runs_texts:
        if segment:
            new_run = paragraph.add_run(segment)
            new_run.bold = first_run_format.get('bold')
            new_run.italic = first_run_format.get('italic')
            new_run.underline = first_run_format.get('underline')
            new_run.font.name = first_run_format.get('font_name')
            new_run.font.size = first_run_format.get('font_size')
            # Safely set font color
            if first_run_format.get('font_color') is not None:
                try:
                    new_run.font.color.rgb = first_run_format.get('font_color')
                except:
                    pass  # Skip color setting if it fails

    return True

def replace_text_in_table(table, old_text, new_text):
    """
    Replace text in table cells, handling placeholders that might be split across runs.
    Returns True if at least one replacement was made.
    """
    replaced = False
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if _replace_text_in_runs(paragraph, old_text, new_text):
                    replaced = True
    return replaced

def replace_text_in_table_with_risk_colors(table, old_text, risk_factor):
    """
    Replace text in a table while preserving formatting and applying risk-based background colors
    
    Args:
        table: The table object from python-docx
        old_text: Text to be replaced
        risk_factor: Text to replace with (should be the risk factor)
    """
    
    # Define risk level colors (using hex values for better compatibility)
    risk_colors = {
        'CRITICAL': '8B0000',      # Dark red
        'HIGH': 'FF0000',          # Red  
        'MEDIUM': 'FFA500',        # Orange
        'LOW': '008000'            # Green
    }
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # Use _replace_text_in_runs to handle text split across runs
                full_text = "".join(run.text for run in paragraph.runs)
                if old_text in full_text:
                    # Replace the text first
                    _replace_text_in_runs(paragraph, old_text, risk_factor)
                    
                    # Make all runs in this paragraph bold and white
                    for run in paragraph.runs:
                        run.bold = True
                        try:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                        except:
                            pass
                    
                    # Apply background color based on risk
                    risk_upper = risk_factor.upper().strip()
                    if 'CRITICAL' in risk_upper:
                        cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), risk_colors['CRITICAL'])))
                    elif 'HIGH' in risk_upper:
                        cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), risk_colors['HIGH'])))
                    elif 'MEDIUM' in risk_upper:
                        cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), risk_colors['MEDIUM'])))
                    elif 'LOW' in risk_upper:
                        cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), risk_colors['LOW'])))

def copy_table_with_formatting(table):
    """Create a deep copy of a table with all formatting"""
    return copy.deepcopy(table)

def is_ip_address(text):
    """Check if text looks like an IP address"""
    import re
    ip_pattern = r'^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}$'
    return bool(re.match(ip_pattern, text.strip()))

def replace_affected_systems_with_table(table, placeholder, affected_systems_text):
    """Replace affected systems placeholder with a formatted table"""
    target_cell = None
    for row in table.rows:
        for cell in row.cells:
            if placeholder in cell.text:
                target_cell = cell
                break
            if target_cell:
                        break
                
    if not target_cell:
        return
    
    if not affected_systems_text or affected_systems_text.strip() == '':
        target_cell.text = "No affected systems"
        return
    
    lines = affected_systems_text.strip().split('\n')
    table_data = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        if 'Branch' in line or not is_ip_address(line):
            branch_name = line
            ip_addresses = []
            
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if not next_line:
                    j += 1
                    continue
                elif is_ip_address(next_line):
                    ip_addresses.append(next_line)
                    j += 1
                else:
                    break
            
            if ip_addresses:
                if len(ip_addresses) == 1:
                    table_data.append([branch_name, ip_addresses[0]])
                else:
                    table_data.append([branch_name, ip_addresses[0]])
                    for ip in ip_addresses[1:]:
                        table_data.append(["", ip])
            
            i = j
        else:
            table_data.append(["", line])
            i += 1
    
    if not table_data:
        target_cell.text = "No affected systems"
        return
    
    target_cell.text = ""
    
    for paragraph in target_cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    
    new_table = target_cell.add_table(rows=len(table_data), cols=2)
    new_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    for row in new_table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(2.0)
    
    for row in new_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            for border_name in ['top', 'left', 'bottom', 'right']:
                existing_border = tcPr.find(qn(f'w:{border_name}'))
                if existing_border is not None:
                    tcPr.remove(existing_border)
            
            for border_name in ['top', 'left', 'bottom', 'right']:
                border_xml = f'<w:{border_name} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                border = parse_xml(border_xml)
                tcPr.append(border)
    
    for i, (branch_name, ip_address) in enumerate(table_data):
        row = new_table.rows[i]
        row.cells[0].text = branch_name
        row.cells[1].text = ip_address
        
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 0
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Inches(0.17)

def replace_followup_vulnerability_details_with_images(doc, excel_file):
    """
    Replaces vulnerability details and creates new tables with POC placeholders for Follow-Up Audit.
    Includes Status-based logic for Statussssss, Newwww_or_Repeatttt, and Follow_up_Remarksssss.
    """
    try:
        poc_images, old_poc_images, df = extract_followup_poc_images_from_excel(excel_file)
        
        if df.empty:
            print("⚠️ No vulnerability data found")
            return

        # Find the template table
        template_table = None
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Vulnerability___Name" in cell.text:
                        template_table = table
                        break
                if template_table:
                    break
            if template_table:
                break
        
        if not template_table:
            print("⚠️ Template table with 'Vulnerability___Name' not found")
            return

        # Find the marker paragraph
        marker_paragraph = None
        for paragraph in doc.paragraphs:
            if "Tables_Heree" in paragraph.text:
                marker_paragraph = paragraph
                break
        
        if not marker_paragraph:
            print("⚠️ Marker 'Tables_Heree' not found")
            return

        parent = marker_paragraph._element.getparent()
        marker_index = parent.index(marker_paragraph._element)
        parent.remove(template_table._element)
        
        for run in marker_paragraph.runs:
            if "Tables_Heree" in run.text:
                run.text = run.text.replace("Tables_Heree", "")

        # Create table for each vulnerability
        for idx, row in enumerate(df.iterrows(), start=1):
            index, row_data = row
            new_table = copy_table_with_formatting(template_table)
            
            vuln_name = str(row_data['Name of Vulnerability'])
            status_value = clean_value(row_data['Status']) if 'Status' in row_data else "New"
            status_lower = status_value.lower().strip()
            
            # Replace basic fields
            replace_text_in_table(new_table, "Vulnerability___Name", vuln_name)
            
            risk_factor = str(row_data['Risk Factor'])
            replace_text_in_table_with_risk_colors(new_table, "Riskkkk", risk_factor)
            
            # Handle both old and new column names for backward compatibility
            cve_id_value = row_data.get('CVE/CWE ID', row_data.get('CVE ID', 'NA'))
            replace_text_in_table(new_table, "CWE_ID__", str(cve_id_value) if pd.notna(cve_id_value) else "NA")
            replace_text_in_table(new_table, "CVSS__", str(row_data['CVSS']) if pd.notna(row_data['CVSS']) else "NA")
            
            # Affected systems table
            replace_affected_systems_with_table(new_table, "Affected_URL___", str(row_data['Affected Systems']))
            
            # Simple text replacements
            replace_text_in_table(new_table, "Observation___", str(row_data['Audit Observation']) if pd.notna(row_data['Audit Observation']) else "NA")
            replace_text_in_table(new_table, "Impact___", str(row_data['Impact']) if pd.notna(row_data['Impact']) else "NA")
            replace_text_in_table(new_table, "Recommendation___", str(row_data['Recommendation/Countermeasure']) if pd.notna(row_data['Recommendation/Countermeasure']) else "NA")
            replace_text_in_table(new_table, "Reference_Link___", str(row_data['Reference Link']) if pd.notna(row_data['Reference Link']) else "NA")
            
            # STATUS-BASED REPLACEMENTS
            
            # 1. Replace "Statussssss" with Status value
            replace_text_in_table(new_table, "Statussssss", status_value)
            
            # 2. Replace "Newwww_or_Repeatttt" based on Status
            if "closed" in status_lower or "close" in status_lower:
                newwww_or_repeatttt_value = "-"
            elif "new" in status_lower:
                newwww_or_repeatttt_value = "New Observation"
            elif "open" in status_lower:
                newwww_or_repeatttt_value = "Repeat Observation"
            else:
                newwww_or_repeatttt_value = "New Observation"  # Default
            replace_text_in_table(new_table, "Newwww_or_Repeatttt", newwww_or_repeatttt_value)
            
            # 3. Replace "Follow_up_Remarksssss" based on Status
            if "closed" in status_lower or "close" in status_lower:
                followup_remarks_value = "This vulnerability has been mitigated by the Bank."
            elif "open" in status_lower or "new" in status_lower:
                followup_remarks_value = "This vulnerability has not been mitigated by the Bank."
            else:
                followup_remarks_value = "This vulnerability has not been mitigated by the Bank."  # Default
            replace_text_in_table(new_table, "Follow_up_Remarksssss", followup_remarks_value)
            
            # 4. Replace "Not Applicableeee"
            replace_text_in_table(new_table, "Not Applicableeee", "Not Applicable")
            
            # 5. POC placeholders - POCsss for Old POC, New_POCCsssss for new POC
            # Use unique separator to make extraction easier
            print(f"🔍 Debug: Attempting to replace 'POCsss' with 'POCsss_{idx}_11' in table {idx}")
            replaced_old = replace_text_in_table(new_table, "POCsss", f"POCsss_{idx}_11")  # Old POC
            print(f"  {'✅' if replaced_old else '⚠️'} Old POC replacement {'succeeded' if replaced_old else 'failed'}")
            
            print(f"🔍 Debug: Attempting to replace 'New_POCCsssss' with 'New_POCCsssss_{idx}_11' in table {idx}")
            replaced_new = replace_text_in_table(new_table, "New_POCCsssss", f"New_POCCsssss_{idx}_11")  # New POC
            print(f"  {'✅' if replaced_new else '⚠️'} New POC replacement {'succeeded' if replaced_new else 'failed'}")
            
            # Update numbering (replace 1.X with idx.X)
            for table_row in new_table.rows:
                for cell in table_row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        # Updated pattern to match 1.1, 1.2, ... 1.9, 1.10, 1.11, etc.
                        pattern = r'\b1\.(\d+)\b'
                        
                        def replace_numbering(match):
                            return f"{idx}.{match.group(1)}"
                        
                        new_text = re.sub(pattern, replace_numbering, original_text)
                        
                        if new_text != original_text:
                            font_props = {}
                            if paragraph.runs:
                                run = paragraph.runs[0]
                                font_props['name'] = 'calibri'
                                font_props['size'] = run.font.size
                                font_props['color_rgb'] = run.font.color.rgb if run.font.color.rgb else None
                                font_props['bold'] = run.bold
                                font_props['italic'] = run.italic
                                font_props['underline'] = run.underline

                            paragraph.text = ""
                            new_run = paragraph.add_run(new_text)
                            if font_props.get('name') is not None:
                                new_run.font.name = font_props['name']
                            if font_props.get('size') is not None:
                                new_run.font.size = font_props['size']
                            if font_props.get('color_rgb') is not None:
                                new_run.font.color.rgb = font_props['color_rgb']
                            if font_props.get('bold') is not None:
                                new_run.bold = font_props['bold']
                            if font_props.get('italic') is not None:
                                new_run.italic = font_props['italic']
                            if font_props.get('underline') is not None:
                                new_run.underline = font_props['underline']

            parent.insert(marker_index + 1, new_table._element)
            marker_index += 1
            
            # Add page break between vulnerabilities
            if idx < len(df):
                p = doc.add_paragraph()
                p.add_run().add_break(WD_BREAK.PAGE)
                parent.insert(marker_index + 1, p._element)
                marker_index += 1
        
        # Add images at the end of document
        add_followup_images_at_end_of_document(doc, poc_images, old_poc_images, df)

        print("✅ Vulnerability details with POC links updated successfully")
        
    except Exception as e:
        print(f"❌ Error replacing vulnerability details: {e}")
        traceback.print_exc()

def replace_followup_poc_placeholders_with_images(doc, excel_file):
    """
    Finds POC and New_POC placeholders and replaces them with images from Excel.
    POCsss<number> = Old POC images
    New_POCCsssss<number> = New POC images
    If no images are found, replaces the placeholder with "NIL".
    """
    try:
        poc_images, old_poc_images, df = extract_followup_poc_images_from_excel(excel_file)
        
        # Create a mapping of vulnerability index to images
        poc_images_by_index = {}
        old_poc_images_by_index = {}
        
        for idx, row in enumerate(df.iterrows(), start=1):
            index, row_data = row
            vuln_name = str(row_data['Name of Vulnerability'])
            
            if vuln_name in poc_images and poc_images[vuln_name]:
                poc_images_by_index[idx] = poc_images[vuln_name]
                print(f"📸 Mapped POC images for vulnerability #{idx}: {vuln_name} ({len(poc_images[vuln_name])} images)")
            
            if vuln_name in old_poc_images and old_poc_images[vuln_name]:
                old_poc_images_by_index[idx] = old_poc_images[vuln_name]
                print(f"📸 Mapped Old POC images for vulnerability #{idx}: {vuln_name} ({len(old_poc_images[vuln_name])} images)")
        
        print(f"🔍 Debug: POC images mapped for indices: {list(poc_images_by_index.keys())}")
        print(f"🔍 Debug: Old POC images mapped for indices: {list(old_poc_images_by_index.keys())}")
        
        old_poc_found_count = 0
        new_poc_found_count = 0
        
        # First, scan all tables to see what placeholders exist
        print(f"\n🔍 Scanning document for all POC-related text...")
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        text_content = "".join(run.text for run in paragraph.runs)
                        if "POC" in text_content and ("_" in text_content or "sss" in text_content):
                            print(f"  Table {table_idx+1}, Row {row_idx+1}, Cell {cell_idx+1}, Para {para_idx+1}: '{text_content[:100]}'")
        
        # Process all tables in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Check for both placeholders in this cell
                    for paragraph in list(cell.paragraphs):
                        # Combine all runs to get complete text (handles split placeholders)
                        text_content = "".join(run.text for run in paragraph.runs)
                        
                        # Match POCsss_<number>_11 format
                        old_poc_match = re.search(r'(POCsss_(\d+)_11)', text_content)
                        if old_poc_match:
                            old_poc_found_count += 1
                            full_placeholder = old_poc_match.group(1)
                            vuln_index = int(old_poc_match.group(2))
                            
                            try:
                                print(f"🔍 Found Old POC placeholder: {full_placeholder} → vulnerability index {vuln_index}")
                                
                                if vuln_index in old_poc_images_by_index:
                                    image_data_list = old_poc_images_by_index[vuln_index]
                                    print(f"  ✅ Replacing with {len(image_data_list)} Old POC images")
                                    
                                    # Clear the placeholder paragraph
                                    combined_text = "".join(run.text for run in paragraph.runs)
                                    if "POCsss" in combined_text:
                                        paragraph._element.getparent().remove(paragraph._element)
                                    
                                    # Add images
                                    for i, image_data in enumerate(image_data_list):
                                        try:
                                            image_paragraph = cell.add_paragraph()
                                            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            
                                            image_stream = io.BytesIO(image_data)
                                            run = image_paragraph.add_run()
                                            picture = run.add_picture(image_stream, width=Inches(6.5))
                                            
                                            # Apply 1pt border to image
                                            apply_1pt_border_to_picture(picture)
                                            
                                            # Add space after every image
                                            cell.add_paragraph()
                                        except Exception as e:
                                            print(f"Error adding Old POC image {i+1}: {e}")
                                else:
                                    # No images found, clear entire cell and add NIL
                                    print(f"  ⚠️ No Old POC images found for vulnerability #{vuln_index}, replacing with NIL")
                                    # Clear all paragraphs in the cell
                                    for p in list(cell.paragraphs):
                                        p._element.getparent().remove(p._element)
                                    # Add new paragraph with NIL
                                    nil_paragraph = cell.add_paragraph()
                                    nil_paragraph.add_run("NIL")
                                    nil_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception as e:
                                print(f"❌ Error processing Old POC placeholder: {e}")
                                traceback.print_exc()
                        
                        # Match New_POCCsssss_<number>_11 format
                        new_poc_match = re.search(r'(New_POCCsssss_(\d+)_11)', text_content)
                        if new_poc_match:
                            new_poc_found_count += 1
                            full_placeholder = new_poc_match.group(1)
                            vuln_index = int(new_poc_match.group(2))
                            
                            try:
                                print(f"🔍 Found New POC placeholder: {full_placeholder} → vulnerability index {vuln_index}")
                                
                                if vuln_index in poc_images_by_index:
                                    image_data_list = poc_images_by_index[vuln_index]
                                    print(f"  ✅ Replacing with {len(image_data_list)} New POC images")
                                    
                                    # Clear the placeholder paragraph
                                    combined_text = "".join(run.text for run in paragraph.runs)
                                    if "New_POCCsssss" in combined_text:
                                        paragraph._element.getparent().remove(paragraph._element)
                                    
                                    # Add images
                                    for i, image_data in enumerate(image_data_list):
                                        try:
                                            image_paragraph = cell.add_paragraph()
                                            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            
                                            image_stream = io.BytesIO(image_data)
                                            run = image_paragraph.add_run()
                                            picture = run.add_picture(image_stream, width=Inches(6.5))
                                            
                                            # Apply 1pt border to image
                                            apply_1pt_border_to_picture(picture)
                                            
                                            # Add space after every image
                                            cell.add_paragraph()
                                        except Exception as e:
                                            print(f"Error adding New POC image {i+1}: {e}")
                                else:
                                    # No images found, clear entire cell and add NIL
                                    print(f"  ⚠️ No New POC images found for vulnerability #{vuln_index}, replacing with NIL")
                                    # Clear all paragraphs in the cell
                                    for p in list(cell.paragraphs):
                                        p._element.getparent().remove(p._element)
                                    # Add new paragraph with NIL
                                    nil_paragraph = cell.add_paragraph()
                                    nil_paragraph.add_run("NIL")
                                    nil_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception as e:
                                print(f"❌ Error processing New POC placeholder: {e}")
                                traceback.print_exc()
        
        print(f"\n📊 Summary:")
        print(f"  - Found {old_poc_found_count} Old POC placeholders (POCsss_X_11)")
        print(f"  - Found {new_poc_found_count} New POC placeholders (New_POCCsssss_X_11)")
        print(f"  - Available Old POC images for indices: {list(old_poc_images_by_index.keys())}")
        print(f"  - Available New POC images for indices: {list(poc_images_by_index.keys())}")
        
        if new_poc_found_count == 0 and len(poc_images_by_index) > 0:
            print(f"\n⚠️ WARNING: Found {len(poc_images_by_index)} vulnerabilities with New POC images but no 'New_POCCsssss' placeholders in document!")
            print("  → Check if your Word template has 'New_POCCsssss' placeholder or if it uses a different name")
        
        print("✅ POC placeholders replaced with images successfully")
    
    except Exception as e:
        print(f"❌ Error replacing POC placeholders: {e}")
        traceback.print_exc()

def highlight_keywords(word_cell, value):
    """
    Checks if a cell's text contains specific keywords and changes its background color to blue.
    """
    if value.strip().upper() in ["HOST", "PORT", "SERVICE"]:
        blue_rgb = "1376d1"  # Hex code for blue
        tc = word_cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), blue_rgb)
        tcPr.append(shd)

def create_nmap_table_from_excel(doc, excel_file):
    """
    Create a formatted table from the 'Nmap Files' worksheet in Excel
    Returns the table object without adding it to document
    """
    try:
        # Load Excel workbook and get the 'Nmap Files' worksheet
        wb = load_workbook(excel_file)
        
        # Try to get the 'Nmap Files' worksheet, fallback to active sheet if not found
        try:
            ws = wb['Nmap Files']
        except KeyError:
            print("Worksheet 'Nmap Files' not found, using active sheet")
            ws = wb.active
        
        # Create a table with Excel size (but don't add it to doc yet)
        table = doc.add_table(rows=ws.max_row, cols=ws.max_column)
        table.style = 'Table Grid'

        # A set to keep track of merged cells to avoid overwriting them
        merged_cells_to_skip = set()

        # First, handle merged cells
        for merged_range in ws.merged_cells:
            merged_range_coords = list(merged_range.cells)
            
            top_left_cell_ref = merged_range_coords[0]
            start_row = top_left_cell_ref[0] - 1
            start_col = top_left_cell_ref[1] - 1
            
            bottom_right_cell_ref = merged_range_coords[-1]
            end_row = bottom_right_cell_ref[0] - 1
            end_col = bottom_right_cell_ref[1] - 1
            
            start_word_cell = table.cell(start_row, start_col)
            end_word_cell = table.cell(end_row, end_col)
            
            start_word_cell.merge(end_word_cell)
            
            for r_idx in range(start_row, end_row + 1):
                for c_idx in range(start_col, end_col + 1):
                    if (r_idx, c_idx) != (start_row, start_col):
                        merged_cells_to_skip.add((r_idx, c_idx))

        # Fill Word table with Excel data and formatting
        for i, row in enumerate(ws.iter_rows(values_only=False)):
            for j, cell in enumerate(row):
                if (i, j) in merged_cells_to_skip:
                    continue
                    
                word_cell = table.cell(i, j)
                value = "" if cell.value is None else str(cell.value)
                word_cell.text = value

                # Format font
                if cell.font:
                    # Ensure we have at least one run in the paragraph
                    if not word_cell.paragraphs[0].runs:
                        word_cell.paragraphs[0].add_run(value)
                    
                    run = word_cell.paragraphs[0].runs[0]
                    run.font.name = cell.font.name or "Calibri"
                    run.font.size = Pt(12)
                    run.font.bold = cell.font.b
                    run.font.italic = cell.font.i
                    if cell.font.color:
                        rgb = get_rgb(cell.font.color)
                        if rgb:
                            run.font.color.rgb = RGBColor.from_string(rgb)

                # Set horizontal and vertical alignment
                word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Call the function to highlight keywords
                highlight_keywords(word_cell, value)

        print("✅ Nmap table created successfully from Excel")
        wb.close()
        return table

    except Exception as e:
        print(f"❌ Error creating Nmap table: {e}")
        traceback.print_exc()
        return None

def replace_open_ports_with_nmap_table(doc, excel_file, placeholder="OPEN_PORTSSSS"):
    """
    Replace OPEN_PORTSSSS placeholder with Nmap table from Excel
    """
    try:
        # Create the Nmap table
        nmap_table = create_nmap_table_from_excel(doc, excel_file)
        if not nmap_table:
            print("❌ Failed to create Nmap table")
            return False
        
        # Search for the placeholder in paragraphs
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                # Store the parent and index for insertion
                parent = paragraph._element.getparent()
                index = parent.index(paragraph._element)
                
                # Remove the placeholder paragraph
                parent.remove(paragraph._element)
                
                # Insert the table at the same position
                parent.insert(index, nmap_table._element)
                print(f"✅ Replaced '{placeholder}' with Nmap table")
                return True
        
        # If not found in paragraphs, check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            # Clear the cell and add the table
                            cell._element.clear()
                            cell._element.append(nmap_table._element)
                            print(f"✅ Replaced '{placeholder}' with Nmap table in table cell")
                            return True
        
        print(f"❌ Placeholder '{placeholder}' not found in document")
        return False
        
    except Exception as e:
        print(f"Error replacing OPEN_PORTSSSS with Nmap table: {e}")
        traceback.print_exc()
        return False

def remove_content_after_marker(doc, marker_text):
    """
    Safely removes all content (paragraphs, tables, etc.) after a specific marker text in a Word document,
    including the marker text itself.

    Args:
        doc: The python-docx Document object.
        marker_text: The string to search for as the end marker (will be removed).
    """
    marker_found = False

    # Iterate over paragraphs and tables in body in document order
    body_elements = list(doc.element.body)

    for element in body_elements:
        if element.tag.endswith('p'):  # Paragraph
            para = next((p for p in doc.paragraphs if p._element == element), None)
            if para and marker_text in para.text:
                marker_found = True
                # Find text before the marker and keep only that
                para_text = para.text
                marker_index = para_text.find(marker_text)
                
                if marker_index > 0:
                    # Keep text before the marker
                    text_to_keep = para_text[:marker_index].strip()
                    para.clear()
                    if text_to_keep:
                        para.add_run(text_to_keep)
                else:
                    # Marker is at beginning or no text before it, remove the entire paragraph
                    element.getparent().remove(element)
                continue

            if marker_found:
                element.getparent().remove(element)

        elif element.tag.endswith('tbl'):  # Table
            tbl = next((t for t in doc.tables if t._element == element), None)
            if tbl:
                marker_in_table = False
                for row in tbl.rows:
                    for cell in row.cells:
                        if marker_text in cell.text:
                            marker_in_table = True
                            # Remove marker text from cell
                            cell_text = cell.text
                            marker_index = cell_text.find(marker_text)
                            if marker_index > 0:
                                cell.text = cell_text[:marker_index].strip()
                            else:
                                cell.text = ""
                            break
                    if marker_in_table:
                        break
                
                if marker_in_table:
                    marker_found = True
                    continue

            if marker_found:
                element.getparent().remove(element)

    return doc

def replace_a11_row_with_vulnerabilities(doc, excel_file):
    """Find the table containing a.1.1 and add Excel data with proper columns including Status column"""
    try:
        # Read Excel data
        df = pd.read_excel(excel_file, sheet_name='Infra_VAPT')
        
        if df.empty:
            print("No data found in 'Infra_VAPT' worksheet")
            return

        # Find the target table (the one containing "a.1.1")
        target_table = None
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "a.1.1" in cell.text.lower():
                        target_table = table
                        break
                if target_table:
                    break
            if target_table:
                break
        
        if not target_table:
            print("Could not find table containing 'a.1.1'")
            return

        print("Found target table - updating content...")

        # Clear all existing rows except header (assuming first row is header)
        while len(target_table.rows) > 1:
            target_table._tbl.remove(target_table.rows[1]._tr)

        # Add Excel data with proper columns
        for idx in range(len(df)):
            new_row = target_table.add_row()
            
            # Ensure we have at least 12 columns for Follow-Up Audit (includes Status column)
            if len(new_row.cells) < 12:
                continue

            for cell in new_row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER    
            
            # Column 1: Sr.No (index + 1)
            first_cell = new_row.cells[0]
            for paragraph in first_cell.paragraphs:
                paragraph.clear()
            p = first_cell.paragraphs[0]
            p.text = str(idx + 1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 2: Affected System (Excel column J, index 9)
            second_cell = new_row.cells[1]
            for paragraph in second_cell.paragraphs:
                paragraph.clear()
            affected_system = str(df.iloc[idx, 9]) if pd.notna(df.iloc[idx, 9]) else "NA"
            p = second_cell.paragraphs[0]
            p.text = affected_system
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 3: Vulnerability Description (Excel column C, index 2)
            third_cell = new_row.cells[2]
            for paragraph in third_cell.paragraphs:
                paragraph.clear()
            vulnerability = str(df.iloc[idx, 2]) if pd.notna(df.iloc[idx, 2]) else "NA"
            p = third_cell.paragraphs[0]
            p.text = vulnerability
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 4: CVE/CWE ID (Excel column E, index 4)
            fourth_cell = new_row.cells[3]
            for paragraph in fourth_cell.paragraphs:
                paragraph.clear()
            # Get the value and handle both NaN and string "nan"
            raw_value = df.iloc[idx, 4]
            if pd.isna(raw_value) or str(raw_value).lower() == 'nan':
                cve_id = "NA"
            else:
                cve_id = str(raw_value)
            p = fourth_cell.paragraphs[0]
            p.text = cve_id
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)
                
            # Column 5: Risk Factor (Excel column D, index 3)
            fifth_cell = new_row.cells[7]
            for paragraph in fifth_cell.paragraphs:
                paragraph.clear()
            risk_factor = str(df.iloc[idx, 3]) if pd.notna(df.iloc[idx, 3]) else "NA"
            p = fifth_cell.paragraphs[0]
            p.text = risk_factor
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(255, 255, 255) 
                run.bold = True

            # Apply risk-based background colors
            risk_factor_lower = risk_factor.lower()
            if 'critical' in risk_factor_lower:
                fifth_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="8B0000"/>'.format(nsdecls('w'))))  # Dark Red
            elif 'high' in risk_factor_lower:
                fifth_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w'))))  # Red
            elif 'medium' in risk_factor_lower:
                fifth_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FFA500"/>'.format(nsdecls('w'))))  # Orange
            elif 'low' in risk_factor_lower:
                fifth_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="008000"/>'.format(nsdecls('w'))))  # Green
            
            # Column 6: CVSS (Excel column F, index 5)
            sixth_cell = new_row.cells[4]
            for paragraph in sixth_cell.paragraphs:
                paragraph.clear()
            p = sixth_cell.paragraphs[0]
            p.text = "NA"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 7: Audit Observation (Excel column G, index 6)
            seventh_cell = new_row.cells[5]
            for paragraph in seventh_cell.paragraphs:
                paragraph.clear()
            p = seventh_cell.paragraphs[0]
            p.text = "NA"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 8: Impact (Excel column H, index 7)
            eighth_cell = new_row.cells[6]
            for paragraph in eighth_cell.paragraphs:
                paragraph.clear()
            p = eighth_cell.paragraphs[0]
            p.text = "NA"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 9: Recommendation (Excel column I, index 8)
            ninth_cell = new_row.cells[8]
            for paragraph in ninth_cell.paragraphs:
                paragraph.clear()
            recommendation_w = clean_value(df.iloc[idx, 8])
            p = ninth_cell.paragraphs[0]
            p.text = recommendation_w
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 10: Reference Link (Excel column K, index 10)
            tenth_cell = new_row.cells[9]
            for paragraph in tenth_cell.paragraphs:
                paragraph.clear()
            reference = clean_value(df.iloc[idx, 10])
            p = tenth_cell.paragraphs[0]
            p.text = reference
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 12: Status from Excel (Excel column M, index 12)
            twelve_cell = new_row.cells[11]
            for paragraph in twelve_cell.paragraphs:
                paragraph.clear()
            status_value = clean_value(df.iloc[idx, 11])
            p = twelve_cell.paragraphs[0]
            p.text = status_value
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

            # Column 11: Status (New/Repeat/-) based on Column 12
            eleven_cell = new_row.cells[10]
            for paragraph in eleven_cell.paragraphs:
                paragraph.clear()
            
            # Determine eleven_cell content based on twelve_cell status
            status_lower = status_value.lower().strip()
            if "new" in status_lower:
                eleven_cell_value = "New"
            elif "open" in status_lower:
                eleven_cell_value = "Repeat"
            elif "closed" in status_lower or "close" in status_lower:
                eleven_cell_value = "-"
            else:
                eleven_cell_value = "New"  # Default to "New" if status is unclear
            
            p = eleven_cell.paragraphs[0]
            p.text = eleven_cell_value
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(12)

        print(f"✅ Successfully updated a.1.1 table with {len(df)} rows of data")
        
    except Exception as e:
        print(f"Error updating a.1.1 table: {str(e)}")
        traceback.print_exc()

def generate_follow_up_audit_word_report(excel_file_path):
    """
    Generate Follow-Up Audit Word report from Excel file
    """
    try:
        # Template path for follow-up audit report
        template_path = os.path.join("static", "Formats_and_Catalog", "Infrastructure_VAPT_Follow_Up_Audit_Report.docx")
        
        if not os.path.exists(template_path):
            raise Exception(f"Template file not found: {template_path}")
        
        print(f"📄 Using template: {template_path}")
        
        # Load the Word template
        doc = Document(template_path)
        
        # Extract metadata from Excel file
        print("🔍 Extracting metadata from Excel file...")
        metadata = find_and_print_metadata(excel_file_path)
        
        # Print extracted metadata for debugging
        if metadata:
            print("📋 Extracted metadata:")
            for key, value in metadata.items():
                if value:
                    print(f"  {key}: {value}")
        
        # Replace metadata placeholders in the document
        if metadata:
            print("🔍 Replacing metadata placeholders...")
            replace_metadata_placeholders(doc, metadata)
            
            # Replace auditor placeholders and add additional rows for multiple team members
            print("🔍 Processing auditing team members...")
            replace_auditor_placeholders_and_add_rows(doc, metadata)
        else:
            print("⚠️ No metadata found in Excel file")
        
        # Replace scope placeholders with data from Scope worksheet
        print("🔍 Replacing scope placeholders with data from Scope worksheet...")
        replace_scope_placeholders_with_data(doc, excel_file_path)
        
        # Replace vulnerability details with images and POC placeholders
        print("🔍 Replacing vulnerability details with images...")
        replace_followup_vulnerability_details_with_images(doc, excel_file_path)
        
        # Replace POC placeholders with actual images (must be after vulnerability details)
        print("🔍 Replacing POC placeholders with images...")
        replace_followup_poc_placeholders_with_images(doc, excel_file_path)
        
        # Replace OPEN_PORTSSSS placeholder with Nmap table
        print("🔍 Replacing OPEN_PORTSSSS with Nmap table...")
        replace_open_ports_with_nmap_table(doc, excel_file_path, "OPEN_PORTSSSS")
        
        # Replace a.1.1 table with vulnerabilities data
        print("🔍 Replacing a.1.1 table with vulnerability data...")
        replace_a11_row_with_vulnerabilities(doc, excel_file_path)
        
        # Generate and replace timeline placeholders
        if metadata and metadata.get("start_date") and metadata.get("end_date"):
            print("🔍 Generating Follow-Up Audit VAPT timeline...")
            timeline = generate_followup_vapt_timeline(metadata.get("start_date"), metadata.get("end_date"))
            if timeline:
                print(f"📅 Generated {len(timeline)} timeline phases")
                for phase in timeline:
                    print(f"  - {phase['Phase']}: {phase['Dates']}")
                print("🔍 Replacing timeline placeholders...")
                replace_followup_timeline_placeholders(doc, timeline)
            else:
                print("⚠️ No timeline generated")
        else:
            print("⚠️ Start date or end date not found, skipping timeline generation")
        
        # Remove content after marker
        marker_to_remove_after = "CEH Cert. No: 12345678901234567890"
        print(f"🔍 Removing content after marker: '{marker_to_remove_after}'...")
        remove_content_after_marker(doc, marker_to_remove_after)
        
        # Save the document to a temporary file
        temp_dir = tempfile.mkdtemp()
        temp_file_path = os.path.join(temp_dir, "Infrastructure_VAPT_Follow_Up_Audit_Report.docx")
        doc.save(temp_file_path)
        
        # Update chart with vulnerability data
        print("🔍 Updating chart with vulnerability data...")
        update_followup_chart_in_docx(temp_file_path, excel_file_path)
        
        print(f"✅ Word report generated successfully: {temp_file_path}")
        return temp_file_path
        
    except Exception as e:
        print(f"❌ Error generating Word report: {e}")
        traceback.print_exc()
        raise Exception(f"Error generating Word report: {str(e)}")

@vapt_follow_up_word_report_bp.route('/process_vapt_follow_up_audit_word_report', methods=['POST'])
def process_follow_up_audit_word_report():
    """
    Process the uploaded Excel file and generate a Follow-Up Audit Word report
    """
    try:
        # Check if file was uploaded
        if 'followUpExcelFile' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['followUpExcelFile']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not file.filename.lower().endswith('.xlsx'):
            return jsonify({'error': 'Please upload an Excel file (.xlsx)'}), 400
        
        # Save uploaded file to temporary location
        temp_dir = tempfile.mkdtemp()
        excel_file_path = os.path.join(temp_dir, file.filename)
        file.save(excel_file_path)
        
        print(f"📁 Saved Excel file to: {excel_file_path}")
        
        try:
            # Generate the Word report
            word_file_path = generate_follow_up_audit_word_report(excel_file_path)
            
            # Send the generated Word file
            return send_file(
                word_file_path,
                as_attachment=True,
                download_name='Infrastructure_VAPT_Follow_Up_Audit_Report.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
        finally:
            # Clean up temporary files
            try:
                if os.path.exists(excel_file_path):
                    os.remove(excel_file_path)
                if 'word_file_path' in locals() and os.path.exists(word_file_path):
                    os.remove(word_file_path)
                    # Also remove the temp directory if it's empty
                    temp_word_dir = os.path.dirname(word_file_path)
                    if temp_word_dir != temp_dir and os.path.exists(temp_word_dir):
                        os.rmdir(temp_word_dir)
                os.rmdir(temp_dir)
            except Exception as cleanup_error:
                print(f"Warning: Error during cleanup: {cleanup_error}")
        
    except Exception as e:
        print(f"❌ Error processing follow-up audit word report: {e}")
        traceback.print_exc()
        return jsonify({'error': f'Error generating report: {str(e)}'}), 500
